# -*- coding: utf-8 -*-
"""
PORTAL HR PRO - Versión Ultra Mejorada
Listo para Render / GitHub / uso local.

Usuarios demo:
- Administrador: admin / admin123
- Trabajador: DNI 74324033 / correo omar@demo.com

Variables Render opcionales:
- SECRET_KEY
- PERSIST_DIR=/data
- APP_TIMEZONE=America/Lima
"""

import os
import re
import sqlite3
import html
import base64
import hashlib
import uuid
import json
from datetime import datetime
from pathlib import Path
from copy import copy
from functools import wraps
from zoneinfo import ZoneInfo

from flask import Flask, request, redirect, url_for, session, send_file, render_template_string, flash, jsonify, abort
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
try:
    from docx import Document
except Exception:
    Document = None
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

BASE_DIR = Path(__file__).resolve().parent
PERSIST_DIR = Path(os.getenv("PERSIST_DIR", "/data" if Path("/data").is_dir() else str(BASE_DIR)))
STATIC_DIR = BASE_DIR / "static"
UPLOAD_DIR = PERSIST_DIR / "uploads"
EXCEL_LOCAL_DIR = PERSIST_DIR / "REGISTROS_EXCEL_LOCAL"
DB_PATH = PERSIST_DIR / "boletas_prize.db"
APP_TZ = ZoneInfo(os.getenv("APP_TIMEZONE", "America/Lima"))

for d in (PERSIST_DIR, STATIC_DIR, UPLOAD_DIR, EXCEL_LOCAL_DIR):
    d.mkdir(parents=True, exist_ok=True)

app = Flask(__name__)

# Módulo de contratación deshabilitado por solicitud: se oculta del menú y se bloquea el acceso directo.
@app.before_request
def bloquear_modulo_contratacion_portal_hr():
    ruta = request.path.lower()
    if ruta.startswith('/admin/contratacion') or ruta.startswith('/contratacion/') or ruta.startswith('/admin/firma') or ruta.startswith('/firma/') or 'contratacion' in ruta:
        flash('El módulo de contratación fue retirado de PORTAL HR PRO.', 'ok')
        return redirect(url_for('admin') if session.get('admin_id') else url_for('panel'))


@app.after_request
def permitir_camara_firma(response):
    # Permite cámara/micrófono en la propia página del sistema.
    # IMPORTANTE: Chrome/Edge/Safari solo permiten cámara en HTTPS o en localhost/127.0.0.1.
    # Para celular, publicar en Render/HTTPS o ejecutar local con APP_SSL=1 y abrir https://IP-DE-TU-PC:5000.
    response.headers['Permissions-Policy'] = 'camera=(self), microphone=(self), fullscreen=(self)'
    response.headers['Feature-Policy'] = "camera 'self'; microphone 'self'"
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    return response


def contexto_camara_seguro_texto():
    return (
        'La cámara web del navegador solo funciona en HTTPS o en localhost/127.0.0.1. '
        'En PC local usa http://127.0.0.1:5000. En celular usa el enlace HTTPS de Render, '
        'o ejecuta local con APP_SSL=1 y abre https://IP-DE-TU-PC:5000 aceptando el certificado.'
    )
app.secret_key = os.getenv("SECRET_KEY", "prize_documentos_ultra_2026")
app.config["MAX_CONTENT_LENGTH"] = 80 * 1024 * 1024
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"

# =============================
# CONFIGURACIÓN FUNCIONAL
# =============================
TIPOS_PAGO = [
    ("Utilidad", "Boletas utilidades", "<i class=\'bi bi-file-earmark-text\'></i>"),
    ("Vacaciones", "Boletas vacaciones", "<i class=\'bi bi-file-earmark-text\'></i>"),
    ("Normal", "Boletas normal", "<i class=\'bi bi-file-earmark-text\'></i>"),
    ("CTS", "Boletas CTS", "<i class=\'bi bi-file-earmark-text\'></i>"),
    ("Liquidación", "Boletas liquidación", "<i class=\'bi bi-file-earmark-text\'></i>"),
    ("Gratificación", "Boletas gratificación", "<i class=\'bi bi-file-earmark-text\'></i>"),
]
TIPOS_EMPRESA = [
    ("Contrato de Trabajo", "Contrato de Trabajo", "<i class=\'bi bi-file-earmark-richtext\'></i>"),
    ("Reglamento Interno", "Reglamento Interno", "<i class=\'bi bi-journal-text\'></i>"),
    ("Reglamento de SST", "Reglamento de SST", "<i class=\'bi bi-shield-check\'></i>"),
    ("Código de Conducta", "Código de Conducta", "<i class=\'bi bi-file-lock\'></i>"),
    ("Políticas", "Políticas", "<i class=\'bi bi-pin-angle\'></i>"),
    ("Comunicados", "Comunicados", "<i class=\'bi bi-megaphone\'></i>"),
    ("Formatos", "Formatos", "<i class=\'bi bi-receipt\'></i>"),
]
TIPOS_PERSONALES = [
    ("Otros", "Otros documentos", "<i class=\'bi bi-folder2\'></i>"),
    ("Contrato Personal", "Contrato de Trabajo", "<i class=\'bi bi-person-vcard\'></i>"),
]
ALL_TIPOS = {k: (label, icon, "pago") for k, label, icon in TIPOS_PAGO}
ALL_TIPOS.update({k: (label, icon, "empresa") for k, label, icon in TIPOS_EMPRESA})
ALL_TIPOS.update({k: (label, icon, "personal") for k, label, icon in TIPOS_PERSONALES})
EXT_ALLOWED = {".pdf", ".png", ".jpg", ".jpeg", ".webp", ".doc", ".docx", ".xls", ".xlsx"}


# =============================
# TRABAJADORES OBSERVADOS - CATÁLOGOS GENÉRICOS
# =============================
MOTIVOS_TRABAJADOR_OBSERVADO = [
    'AGRESIÓN',
    'ANTECEDENTES GRADO I',
    'ANTECEDENTES GRADO II',
    'BAJO RENDIMIENTO COSECHA',
    'CASO INTERNO',
    'DAÑO DE ACTIVO',
    'DENUNCIA SUNAFIL',
    'DROGA',
    'ESTADO DE EBRIEDAD',
    'FALSIFICACIÓN DE DM',
    'FALTA DE COMPROMISO',
    'FALTA DE RESPETO AL EVALUADOR',
    'HURTO EN LA EMPRESA',
    'HURTO Y CONSUMO DE DROGA',
    'INCUMPLIMIENTO DE PROCEDIMIENTO DE INGRESO',
    'INCUMPLIMIENTO DEL RIT',
    'MAL COMPORTAMIENTO',
    'EXTORSION',
    'RESTRICCIÓN FÍSICA',
    'SINDICALISTA',
    'SUSTRACCION DE ARANDANO',
    'SUSTRAER BIENES',
    'HOSTIGAMIENTO',
    'OTROS MOTIVOS'
]
NIVELES_RESTRICCION_OBSERVADO = ['NIVEL 1', 'NIVEL 2', 'NIVEL 3']


# Campos disponibles para correspondencia de contratos (plantillas Word/PDF).
# El texto de "nombre" es el que verá el usuario; "origen" sirve como llave para anexar al documento.
CONTRATACION_CAMPOS_CORRESPONDENCIA = [
    ('Indeterminado','Indeterminado','Text'),
    ('Centro Costo','CentroCosto','Text'),
    ('Tipo Contrato','TipoContrato','Text'),
    ('Fecha Fin Contrato','FechaFinContrato','DateTime'),
    ('Fecha Inicio Contrato','FechaInicioContrato','DateTime'),
    ('Dirección Simple','DireccionActual','Text'),
    ('Dni','Dni','Text'),
    ('Nombre Trabajador','NombreCompletoTrabajador','Text'),
    ('Nombre Moneda','NombreMoneda','Text'),
    ('Remuneración Básica','RemuneracionBasica','Number'),
    ('Estado','Estado','Text'),
    ('Comentario','Comentario','Text'),
    ('Nivel Jerárquico','NivelJerarquico','Text'),
    ('Distrito','Distrito','Text'),
    ('Provincia','Provincia','Text'),
    ('Departamento','Departamento','Text'),
    ('Area','Area','Text'),
    ('Puesto','Puesto','Text'),
    ('Número Celular','NroTelefonoMovil','Text'),
    ('Email','Email','Text'),
    ('Símbolo Moneda','SimboloMoneda','Text'),
    ('Situacion Especial','SituacionEspecial','Text'),
    ('Planilla','Planilla','Text'),
    ('Condición','Condicion','Text'),
    ('Gerencia','Gerencia','Text'),
    ('Cargo','Cargo','Text'),
    ('Ocupación','Ocupacion','Text'),
    ('Tipo Trabajador','TipoTrabajador','Text'),
    ('Zona','Zona','Text'),
    ('Categoría Ocupacional','CategoriaOcupacional','Text'),
    ('Grupo Impresión','GrupoImpresion','Text'),
    ('Nivel Educativo','NivelEducativo','Text'),
    ('Sede','Sede','Text'),
    ('Producto','Producto','Text'),
    ('Actividad','Actividad','Text'),
    ('Remuneración Letra','RemuneracionLetra','Text'),
    ('Fecha Inicio Prueba','FechaInicioPrueba','DateTime'),
    ('Fecha Fin Prueba','FechaFinPrueba','DateTime'),
    ('Meses Renovación','MesesRenovacion','Number'),
    ('Fecha Firma','FechaFirma','DateTime'),
    ('Fecha Inicio Contrato Origen','FechaInicioContratoOrigen','DateTime'),
    ('Fecha Fin Contrato Origen','FechaFinContratoOrigen','DateTime'),
    ('Fecha de Nacimiento ISO','FechaNacimientoISO','DateTime'),
    ('Fecha de Nacimiento Guion','FechaNacimientoGuion','DateTime'),
    ('Fecha de Nacimiento Barra','FechaNacimientoBarra','DateTime'),
    ('Fecha de Nacimiento Mayuscula','FechaNacimientoMayuscula','DateTime'),
    ('Fecha de Nacimiento Minuscula','FechaNacimientoMinuscula','DateTime'),
    ('Fecha Fin Contrato ISO','FechaFinContratoISO','DateTime'),
    ('Fecha Fin Contrato Guion','FechaFinContratoGuion','DateTime'),
    ('Fecha Fin Contrato Barra','FechaFinContratoBarra','DateTime'),
]
CONDICION_OPERADORES = ['=', '<>', 'CONTIENE', 'NO CONTIENE', '>', '<', '>=', '<=']
VALORES_CONDICION = {
    'Planilla': ['EMPLEADOS RÉGIMEN GENERAL','EMPLEADOS RÉGIMEN AGRÍCOLA','OBREROS RÉGIMEN AGRÍCOLA','OBREROS REGIMEN GENERAL','OBREROS REGIMEN PACKING','PRACTICANTES'],
    'Tipo Contrato': ['INTERMITENTE OBRERO','INTERMITENTE EMPLEADO','INDETERMINADO','TEMPORAL','RENOVACIÓN'],
    'TipoContrato': ['INTERMITENTE OBRERO','INTERMITENTE EMPLEADO','INDETERMINADO','TEMPORAL','RENOVACIÓN'],
    'Condición': ['ACTIVO','INACTIVO'],
    'Estado': ['ACTIVO','INACTIVO'],
    'Tipo Trabajador': ['OBRERO','EMPLEADO','PRACTICANTE'],
    'Zona': ['PLANTA','CAMPO','PACKING','OFICINA'],
    'Area': ['RECURSOS HUMANOS','OPERACIONES','CAMPO','PACKING','ADMINISTRACIÓN'],
    'Cargo': ['OBRERO','OPERARIO','AUXILIAR','ASISTENTE','ANALISTA','SUPERVISOR','JEFE'],
    'Puesto': ['OBRERO','OPERARIO','AUXILIAR','ASISTENTE','ANALISTA','SUPERVISOR','JEFE'],
}


# Campos oficiales del esquema "Trabajador Contrato Laboral".
# Estos son los campos que se muestran con la lupita y también se exportan a Excel.
CAMPOS_ESQUEMA_TRABAJADOR_CONTRATO_LABORAL = [
    ('Indeterminado', 'SI/NO'), ('CentroCosto', 'CC101'), ('TipoContrato', 'Indefinido'),
    ('FechaFinContrato', '31/12/2025'), ('FechaIniContrato', '1/01/2023'),
    ('DireccionActual', 'Av. Secundaria 456'), ('Dni', '12345678'),
    ('NombreCompletoTrabajador', 'Juan Antonio Pérez Garcia'), ('NombreMoneda', 'Sol Peruano'),
    ('RemunBasica', '2500'), ('Estado', 'Activo'), ('Comentario', 'Sin observaciones'),
    ('NivelJerarquico', 'Jefe'), ('Distrito', 'Miraflores'), ('Provincia', 'Lima'),
    ('Departamento', 'Lima'), ('Area', 'Recursos Humanos'), ('Puesto', 'Analista'),
    ('NroTelefonoMovil', '987654321'), ('Email', 'juan.perez@empresa.com'), ('SimboloMoneda', 'S/'),
    ('SituacionEspecial', 'Ninguna'), ('Planilla', 'Planilla General'), ('Condicion', 'Permanente'),
    ('Gerencia', 'Gerencia General'), ('Cargo', 'Contador'), ('Ocupacion', 'Ingeniero'),
    ('TipoTrabajador', 'Empleado'), ('Zona', 'Zona Norte'), ('CategoriaOcupacional', 'Profesional'),
    ('GrupoImpresion', 'Grupo A'), ('NivelEducativo', 'Universitario'), ('Sede', 'Sede Central'),
    ('Producto', 'Producto A'), ('Actividad', 'Ventas'), ('RemuneracionLetra', 'Dos Mil Quinientos Soles'),
    ('FechaIniPrueba', '1/06/2023'), ('FechaFinPrueba', '1/12/2023'), ('MesesRenovacion', '12'),
    ('FechaFirma', '15/01/2023'), ('FechaInicioContratoOrigen', '1/01/2023'),
    ('FechaFinContratoOrigen', '31/12/2025'), ('FechaNacimientoISO', '15/05/1980'),
    ('FechaNacimientoGuion', '15/05/1980'), ('FechaNacimientoBarra', '15/05/1980'),
    ('FechaNacimientoTextoMayuscula', '15 DE MAYO DE 1980'),
    ('FechaNacimientoTextoMinuscula', '15 de mayo de 1980'), ('FechaFinContratoISO', '31/12/2025'),
    ('FechaFinContratoGuion', '31/12/2025'), ('FechaFinContratoBarra', '31/12/2025'),
    ('FechaFinContratoTextoMayuscula', '31 DE DICIEMBRE DE 2025'),
    ('FechaFinContratoTextoMinuscula', '31 de diciembre de 2025'), ('FechaIniContratoISO', '0001-01-01'),
    ('FechaIniContratoGuion', '01-01-0001'), ('FechaIniContratoBarra', '01/01/0001'),
    ('FechaIniContratoTextoMayuscula', '01 DE ENERO DEL 0001'),
    ('FechaIniContratoTextoMinuscula', '01 de enero del 0001'), ('DireccionDNI', 'Calle de Ejemplo'),
    ('RemunBasicaAgraria', '2500'), ('ApellidoPaternoTrabajador', 'Pérez'),
    ('ApellidoMaternoTrabajador', 'Garcia'), ('NombreTrabajador', 'Juan Antonio'),
    ('NombreTipoDocumentoIdentidad', 'DOC. NACIONAL DE IDENTIDAD'), ('NombreCortoTipoDocumentoIdentidad', 'DNI'),
    ('PeriodicidadPago', 'Quincenal'), ('TipoPago', 'Efectivo'), ('Cuspp', 'XXXXX'),
    ('RegimenLaboral', 'AGRARIO'), ('Nacionalidad', 'Peruana'), ('EstadoCivil', 'SOLTERO/A'),
    ('SistemaPensionario', 'ONP'), ('TipoVia', 'Calle'), ('PaisNacimiento', 'Perú'),
    ('Sexo', 'MASCULINO'), ('Funciones', 'Funciones para el cargo'), ('MesesContrato', 'TRES'),
    ('NumeroMesesContrato', '12'), ('DuracionContratoTexto', '1 mes y 2 días'),
]


def tipo_dato_esquema(campo):
    c = (campo or '').lower()
    if 'fecha' in c:
        return 'DateTime'
    if 'remun' in c or 'meses' in c or 'numero' in c:
        return 'Numeric'
    return 'Text'


def nombre_legible_campo(campo):
    return re.sub(r'(?<!^)(?=[A-ZÁÉÍÓÚÑ])', ' ', campo or '').strip() or campo

# Mantiene compatibilidad con plantillas existentes, pero ahora incluye TODOS los campos del Excel modelo.
for _campo, _desc in CAMPOS_ESQUEMA_TRABAJADOR_CONTRATO_LABORAL:
    if not any(x[1] == _campo for x in CONTRATACION_CAMPOS_CORRESPONDENCIA):
        CONTRATACION_CAMPOS_CORRESPONDENCIA.append((nombre_legible_campo(_campo), _campo, tipo_dato_esquema(_campo)))
# Carpeta documental LOCAL y dinámica:
# Por defecto se crea al costado de app.py, en la misma carpeta donde tienes tus archivos.
# Ejemplo Windows: si app.py está en D:\MiSistema\, se crea D:\MiSistema\DOCUMENTOS_PRIZE_AUTO\
# En Render también se crea dentro del proyecto, pero para carga real masiva se recomienda uso local.
DOCUMENTOS_BASE_DIR = Path(os.getenv("DOCUMENTOS_BASE_DIR", str(BASE_DIR / "DOCUMENTOS_PRIZE_AUTO")))
DOCUMENTOS_BASE_DIR.mkdir(parents=True, exist_ok=True)

def slug_folder(v):
    v = clean(v).upper() if 'clean' in globals() else str(v or '').strip().upper()
    v = (v.replace('Á','A').replace('É','E').replace('Í','I').replace('Ó','O').replace('Ú','U').replace('Ñ','N'))
    return re.sub(r"[^A-Z0-9]+", " ", v).strip() or "GENERAL"

def asegurar_carpetas_documentales(tipo=None):
    """Crea automáticamente la estructura física al entrar/click en pestañas."""
    grupos = []
    if tipo and tipo in ALL_TIPOS:
        label, icon, cat = ALL_TIPOS[tipo]
        grupos = [(tipo, label, cat)]
    else:
        grupos = [(k,l,'pago') for k,l,i in TIPOS_PAGO] + [(k,l,'empresa') for k,l,i in TIPOS_EMPRESA] + [(k,l,'personal') for k,l,i in TIPOS_PERSONALES]
    for k, label, cat in grupos:
        base = DOCUMENTOS_BASE_DIR / ({'pago':'DOCUMENTOS DE PAGO','empresa':'DOCUMENTOS DE LA EMPRESA','personal':'DOCUMENTOS PERSONALES'}.get(cat,'DOCUMENTOS')) / slug_folder(label)
        base.mkdir(parents=True, exist_ok=True)
        if k == 'Normal':
            (base / 'MENSUAL').mkdir(parents=True, exist_ok=True)
            (base / 'SEMANAL').mkdir(parents=True, exist_ok=True)
        if cat == 'pago':
            for y in range(datetime.now(APP_TZ).year-1, datetime.now(APP_TZ).year+2):
                (base / str(y)).mkdir(parents=True, exist_ok=True)
    return True


def now_txt():
    return datetime.now(APP_TZ).strftime("%d/%m/%Y %I:%M %p")


def now_file():
    return datetime.now(APP_TZ).strftime("%Y%m%d_%H%M%S")


def clean(v):
    return str(v or "").strip()


def h(v):
    return html.escape(str(v or ''))



def fecha_sin_hora(v):
    """Muestra fechas sin 00:00:00, aceptando Excel datetime, ISO y texto dd/mm/aaaa."""
    if v is None:
        return ''
    if hasattr(v, 'strftime'):
        return v.strftime('%d/%m/%Y')
    txt = clean(v)
    if not txt:
        return ''
    txt = re.sub(r'\s+00:00:00$', '', txt)
    for fmt in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%d', '%d/%m/%Y %H:%M:%S', '%d/%m/%Y'):
        try:
            return datetime.strptime(txt, fmt).strftime('%d/%m/%Y')
        except Exception:
            pass
    return txt.split()[0] if '00:00:00' in txt else txt


def excel_cell_fecha(v):
    return fecha_sin_hora(v)


def exportar_tabla_excel(nombre_archivo, tabla, columnas):
    """Respaldo local en Excel para que la información sobreviva reinicios y se pueda auditar."""
    try:
        EXCEL_LOCAL_DIR.mkdir(parents=True, exist_ok=True)
        path = EXCEL_LOCAL_DIR / nombre_archivo
        wb = Workbook(); ws = wb.active; ws.title = tabla[:31]
        ws.append([titulo for titulo, campo in columnas])
        with db() as con:
            rows = con.execute(f"SELECT * FROM {tabla}").fetchall()
        for r in rows:
            ws.append([fecha_sin_hora(r[campo]) if 'fecha' in campo.lower() else (r[campo] if campo in r.keys() else '') for titulo, campo in columnas])
        for i, _ in enumerate(columnas, 1):
            ws.column_dimensions[chr(64+i) if i <= 26 else 'A'].width = 24
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill('solid', fgColor='1F2937')
            cell.alignment = Alignment(horizontal='center')
        ws.freeze_panes = 'A2'
        wb.save(path)
        return path
    except Exception as e:
        print('No se pudo exportar Excel local', tabla, e)
        return None


def respaldar_exceles_locales():
    exportar_tabla_excel('01_TRABAJADORES_LOCAL.xlsx', 'trabajadores', [
        ('EMPRESA','empresa'),('DNI','dni'),('TRABAJADOR','nombre'),('CARGO','cargo'),('AREA','area'),('JEFE INMEDIATO','jefe_dni'),('JEFE NOMBRE','jefe_nombre'),('PLANILLA','planilla'),('CORREO','correo'),('FECHA NACIMIENTO','fecha_nacimiento'),('FECHA INGRESO','fecha_ingreso'),('CELULAR','celular'),('CONTACTO EMERGENCIA','contacto_emergencia'),('TELEFONO EMERGENCIA','telefono_emergencia'),('TIPO CONTRATO','tipo_contrato'),('FECHA FIN CONTRATO','fecha_fin_contrato'),('REMUNERACION BASICA','remuneracion_basica'),('DIRECCION','direccion'),('DEPARTAMENTO','departamento'),('PROVINCIA','provincia'),('DISTRITO','distrito'),('NIVEL EDUCATIVO','nivel_educativo'),('PROCEDENCIA','procedencia'),('INDUMENTARIA','indumentaria'),('CARNET CONADIS','carnet_conadis'),('OBSERVACION','observacion'),('USUARIO','usuario_portal'),('CLAVE','clave_portal'),('ACTIVO','activo'),('FECHA REGISTRO','fecha_registro')])
    exportar_tabla_excel('02_VACACIONES_SALDOS_LOCAL.xlsx', 'vacaciones_saldos', [
        ('DNI','dni'),('TRABAJADOR','trabajador'),('EMPRESA','empresa'),('AREA','area'),('JEFE','jefe'),('JEFE DNI','jefe_dni'),('FECHA INGRESO','fecha_ingreso'),('I_PERIODO','periodo_inicio'),('F_PERIODO','periodo_fin'),('DIAS GANADOS','dias_ganados'),('DIAS GOZADOS','dias_gozados'),('SALDO','saldo'),('PERIODO','periodo'),('FECHA CARGA','fecha_carga')])
    exportar_tabla_excel('03_VACACIONES_SOLICITUDES_LOCAL.xlsx', 'vacaciones_solicitudes', [
        ('ID','id'),('DNI','dni'),('TRABAJADOR','trabajador'),('JEFE DNI','jefe_dni'),('FECHA INICIO','fecha_inicio'),('FECHA FIN','fecha_fin'),('DIAS','dias'),('MOTIVO','motivo'),('ESTADO','estado'),('FECHA SOLICITUD','fecha_solicitud'),('PERIODO DETALLE','periodo_detalle'),('PERIODO IDS','periodo_ids'),('COMENTARIO JEFE','comentario_jefe'),('COMENTARIO GH','comentario_gh')])


def restaurar_trabajadores_desde_excel_si_db_vacia():
    """Si Render/local reinicia con BD vacía, recupera trabajadores desde el Excel local."""
    path = EXCEL_LOCAL_DIR / '01_TRABAJADORES_LOCAL.xlsx'
    if not path.exists():
        return
    try:
        with db() as con:
            total = con.execute('SELECT COUNT(*) FROM trabajadores').fetchone()[0]
            if total > 1:
                return
            wb = load_workbook(path, data_only=True); ws = wb.active
            headers=[clean(c.value).upper() for c in ws[1]]
            def idx(n): return headers.index(n) if n in headers else -1
            for row in ws.iter_rows(min_row=2, values_only=True):
                dni=normalizar_dni(row[idx('DNI')] if idx('DNI')>=0 else '')
                if not dni: continue
                fecha_nac=excel_cell_fecha(row[idx('FECHA NACIMIENTO')] if idx('FECHA NACIMIENTO')>=0 else '')
                fecha_ing=excel_cell_fecha(row[idx('FECHA INGRESO')] if idx('FECHA INGRESO')>=0 else '')
                clave=clean(row[idx('CLAVE')] if idx('CLAVE')>=0 else '') or (re.sub(r'\D','', fecha_nac) or dni)
                con.execute("INSERT OR REPLACE INTO trabajadores(dni,nombre,correo,cargo,area,jefe_dni,jefe_nombre,empresa,planilla,fecha_nacimiento,fecha_ingreso,usuario_portal,clave_portal,activo,fecha_registro) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,1,?)", (dni,clean(row[idx('TRABAJADOR')] if idx('TRABAJADOR')>=0 else ''),clean(row[idx('CORREO')] if idx('CORREO')>=0 else '').lower(),clean(row[idx('CARGO')] if idx('CARGO')>=0 else ''),clean(row[idx('AREA')] if idx('AREA')>=0 else ''),normalizar_dni(row[idx('JEFE INMEDIATO')] if idx('JEFE INMEDIATO')>=0 else row[idx('JEFE DNI')] if idx('JEFE DNI')>=0 else ''),clean(row[idx('JEFE NOMBRE')] if idx('JEFE NOMBRE')>=0 else ''),clean(row[idx('EMPRESA')] if idx('EMPRESA')>=0 else 'AQUANQA') or 'AQUANQA',clean(row[idx('PLANILLA')] if idx('PLANILLA')>=0 else ''),fecha_nac,fecha_ing,dni,clave,now_txt()))
            con.commit()
    except Exception as e:
        print('No se pudo restaurar trabajadores desde Excel local', e)


def restaurar_vacaciones_desde_excel_si_db_vacia():
    """Recupera saldos y solicitudes si la BD se reinicia o Render pierde memoria temporal."""
    try:
        with db() as con:
            total_saldos = con.execute('SELECT COUNT(*) FROM vacaciones_saldos').fetchone()[0]
            total_sol = con.execute('SELECT COUNT(*) FROM vacaciones_solicitudes').fetchone()[0]

            saldos_path = EXCEL_LOCAL_DIR / '02_VACACIONES_SALDOS_LOCAL.xlsx'
            if total_saldos == 0 and saldos_path.exists():
                wb = load_workbook(saldos_path, data_only=True); ws = wb.active
                headers=[clean(c.value).upper() for c in ws[1]]
                def idx(n): return headers.index(n) if n in headers else -1

                for row in ws.iter_rows(min_row=2, values_only=True):
                    dni = normalizar_dni(row[idx('DNI')] if idx('DNI')>=0 else '')
                    if not dni: 
                        continue

                    periodo_inicio = periodo_year_value(row[idx('I_PERIODO')] if idx('I_PERIODO')>=0 else '')
                    periodo_fin = periodo_year_value(row[idx('F_PERIODO')] if idx('F_PERIODO')>=0 else '')
                    trabajador = clean(row[idx('TRABAJADOR')] if idx('TRABAJADOR')>=0 else '')
                    jefe_dni = normalizar_dni(row[idx('JEFE DNI')] if idx('JEFE DNI')>=0 else row[idx('JEFE')] if idx('JEFE')>=0 else '')

                    con.execute("""INSERT OR REPLACE INTO vacaciones_saldos
                    (dni,trabajador,empresa,area,jefe,jefe_dni,fecha_ingreso,periodo_inicio,periodo_fin,dias_ganados,dias_gozados,saldo,periodo,fecha_carga,uploaded_by)
                    VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (
                        dni,
                        trabajador,
                        clean(row[idx('EMPRESA')] if idx('EMPRESA')>=0 else ''),
                        clean(row[idx('AREA')] if idx('AREA')>=0 else ''),
                        jefe_dni,
                        jefe_dni,
                        excel_cell_fecha(row[idx('FECHA INGRESO')] if idx('FECHA INGRESO')>=0 else ''),
                        periodo_inicio,
                        periodo_fin,
                        float(row[idx('DIAS GANADOS')] if idx('DIAS GANADOS')>=0 and row[idx('DIAS GANADOS')] not in (None,'') else 0),
                        float(row[idx('DIAS GOZADOS')] if idx('DIAS GOZADOS')>=0 and row[idx('DIAS GOZADOS')] not in (None,'') else 0),
                        float(row[idx('SALDO')] if idx('SALDO')>=0 and row[idx('SALDO')] not in (None,'') else 0),
                        f'{periodo_inicio}/{periodo_fin}',
                        now_txt(),
                        'AUTO-RESTORE'
                    ))

            solicitudes_path = EXCEL_LOCAL_DIR / '03_VACACIONES_SOLICITUDES_LOCAL.xlsx'
            if total_sol == 0 and solicitudes_path.exists():
                wb = load_workbook(solicitudes_path, data_only=True); ws = wb.active
                headers=[clean(c.value).upper() for c in ws[1]]
                def idx2(n): return headers.index(n) if n in headers else -1

                for row in ws.iter_rows(min_row=2, values_only=True):
                    dni = normalizar_dni(row[idx2('DNI')] if idx2('DNI')>=0 else '')
                    if not dni:
                        continue

                    con.execute("""INSERT INTO vacaciones_solicitudes
                    (dni,trabajador,jefe_dni,fecha_inicio,fecha_fin,dias,motivo,estado,fecha_solicitud,periodo_detalle,periodo_ids,comentario_jefe,comentario_gh)
                    VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""", (
                        dni,
                        clean(row[idx2('TRABAJADOR')] if idx2('TRABAJADOR')>=0 else ''),
                        normalizar_dni(row[idx2('JEFE DNI')] if idx2('JEFE DNI')>=0 else ''),
                        clean(row[idx2('FECHA INICIO')] if idx2('FECHA INICIO')>=0 else ''),
                        clean(row[idx2('FECHA FIN')] if idx2('FECHA FIN')>=0 else ''),
                        float(row[idx2('DIAS')] if idx2('DIAS')>=0 and row[idx2('DIAS')] not in (None,'') else 0),
                        clean(row[idx2('MOTIVO')] if idx2('MOTIVO')>=0 else ''),
                        clean(row[idx2('ESTADO')] if idx2('ESTADO')>=0 else 'Pendiente jefe'),
                        clean(row[idx2('FECHA SOLICITUD')] if idx2('FECHA SOLICITUD')>=0 else now_txt()),
                        clean(row[idx2('PERIODO DETALLE')] if idx2('PERIODO DETALLE')>=0 else ''),
                        clean(row[idx2('PERIODO IDS')] if idx2('PERIODO IDS')>=0 else ''),
                        clean(row[idx2('COMENTARIO JEFE')] if idx2('COMENTARIO JEFE')>=0 else ''),
                        clean(row[idx2('COMENTARIO GH')] if idx2('COMENTARIO GH')>=0 else '')
                    ))
            con.commit()
    except Exception as e:
        print('No se pudo restaurar vacaciones desde Excel local', e)


def normalizar_dni(v):
    d = re.sub(r"\D", "", str(v or ""))
    return d[-8:].zfill(8) if d else ""


def safe_periodo(p):
    return re.sub(r"[^A-Za-z0-9_\- ]", "", clean(p))[:50] or datetime.now(APP_TZ).strftime("%Y-%m")





def periodo_year_value(v):
    """Normaliza I_PERIODO/F_PERIODO para trabajar solo con años: 2025, 2026, etc."""
    if v is None:
        return ''
    if hasattr(v, 'year'):
        return str(v.year)
    txt = clean(v)
    m = re.search(r'(20\d{2}|19\d{2})', txt)
    return m.group(1) if m else txt

def periodo_anual_texto(inicio='', fin=''):
    """Devuelve periodo en formato 2025/2026 usando fechas de inicio y fin."""
    def year(v):
        if v is None:
            return ''
        if hasattr(v, 'year'):
            return str(v.year)
        txt = clean(v)
        m = re.search(r'(20\d{2}|19\d{2})', txt)
        return m.group(1) if m else ''
    yi, yf = year(inicio), year(fin)
    if yi and yf and yi != yf:
        return f'{yi}/{yf}'
    return yi or yf or clean(inicio) or clean(fin) or datetime.now(APP_TZ).strftime('%Y')

def logo_url():
    # Reconoce logo en la carpeta raíz o static: logo_prize.png, logo.png, prize.png, etc.
    nombres = ["logo_prize.png", "logo.png", "prize.png", "LOGO.png", "Logo.png", "logo_prize.jpg", "logo_prize.jpeg"]
    for folder in (BASE_DIR, STATIC_DIR):
        for name in nombres:
            p = folder / name
            if p.exists():
                if folder == STATIC_DIR:
                    return url_for("static", filename=name)
                return url_for("logo_file", filename=name)
    return url_for("logo_svg")


def db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.create_function('normalizar_dni_sql', 1, normalizar_dni)
    return conn



# =============================
# IA HR PRO - DOCUMENTAL / VACACIONAL / LEGAL PERÚ
# =============================
def ia_sql_count(con, sql, params=()):
    try:
        return int(con.execute(sql, params).fetchone()[0] or 0)
    except Exception:
        return 0


def ia_sql_rows(con, sql, params=()):
    try:
        return con.execute(sql, params).fetchall()
    except Exception:
        return []


def ia_norm(txt):
    txt = clean(txt).lower()
    repl = {'á':'a','é':'e','í':'i','ó':'o','ú':'u','ñ':'n'}
    for a,b in repl.items():
        txt = txt.replace(a,b)
    return re.sub(r'[^a-z0-9 ]+', ' ', txt)


def ia_score(pregunta, row):
    q = set(ia_norm(pregunta).split())
    base = ' '.join([clean(row[k]) for k in row.keys() if k in ['pregunta_clave','palabras_clave','tema','modulo','respuesta']])
    b = set(ia_norm(base).split())
    return len(q & b)


def asegurar_ia_hr_pro():
    """Crea y precarga la IA del Portal HR. No borra respuestas editadas por el administrador."""
    semillas_sistema = [
        ('documental','admin','Dónde cargar boletas de pago','Para cargar boletas ingresa a Gestión Documental > Documentos de pago. Selecciona el tipo: Normal, CTS, Gratificación, Utilidad, Vacaciones o Liquidación. También puedes entrar a Subir / gestionar documentos para cargas individuales o masivas.','/admin/documentos','boleta boletas pago cargar subir normal cts gratificacion utilidad liquidacion documentos pago'),
        ('documental','admin','Dónde sincronizar o detectar PDFs por DNI','Ingresa a Gestión Documental > Detectar PDFs o a Sincronizar carpetas. El sistema revisa DOCUMENTOS_PRIZE_AUTO, identifica DNI y registra documentos vinculados al trabajador.','/admin/sincronizar','detectar pdf dni sincronizar carpeta documentos prize auto'),
        ('documental','admin','Dónde crear carpetas documentales','Ingresa a Gestión Documental > Crear carpetas. El sistema arma la estructura local para documentos de pago, empresa y personales.','/admin/crear_carpetas','crear carpetas documental estructura documentos'),
        ('documental','admin','Dónde descargar plantilla documental','Ingresa a Gestión Documental > Plantilla Documental. Descarga el Excel base para organizar cargas de documentos.','/admin/plantilla_gestion/documental','plantilla documental excel descargar'),
        ('documental','trabajador','Dónde veo mis boletas','Ingresa a Gestión Documental > Documentos de pago. Ahí podrás ver tus boletas normales, CTS, gratificaciones, utilidades, vacaciones o liquidaciones que RR.HH. haya cargado para tu DNI.','/panel','mis boletas ver descargar pago cts gratificacion utilidad liquidacion'),
        ('documental','trabajador','No veo mi boleta','Si no visualizas tu boleta, puede que RR.HH. aún no la haya cargado o que el archivo no esté vinculado a tu DNI. Revisa el tipo de documento y periodo; si no aparece, comunícate con RR.HH.','/panel','no veo mi boleta no aparece documento'),
        ('documental','ambos','Qué significa documento pendiente','Un documento pendiente es aquel que fue cargado, pero aún no registra lectura, aceptación, firma o aprobación final, según el flujo configurado por RR.HH.','','documento pendiente estado lectura firma aprobacion'),
        ('vacacional','admin','Dónde cargar saldos vacacionales','Ingresa a Gestión Vacacional > Saldos Vacacionales o al Dashboard Vacacional. Desde ahí puedes cargar el Excel de saldos por DNI, periodo, días ganados, gozados y saldo.','/admin/vacaciones#cargar-saldos','cargar saldos vacaciones vacacionales excel saldo'),
        ('vacacional','admin','Dónde aprobar vacaciones','Ingresa a Gestión Vacacional > Aprobaciones. Allí se revisan solicitudes pendientes de jefe y de Gestión Humana, según el flujo.','/admin/vacaciones#aprobaciones','aprobar vacaciones aprobaciones jefe gestion humana'),
        ('vacacional','admin','Dónde ver reportes vacacionales','Ingresa a Gestión Vacacional > Reportes. Podrás revisar solicitudes, saldos, aprobadas, rechazadas y pendientes.','/admin/vacaciones#reportes','reportes vacaciones vacacional indicadores'),
        ('vacacional','trabajador','Cómo solicito vacaciones','Ingresa a Gestión Vacacional > Saldo y solicitud. Selecciona el periodo con saldo, registra fechas y envía la solicitud para aprobación.','/vacaciones/mi_solicitud#solicitar','solicitar vacaciones pedido solicitud saldo'),
        ('vacacional','trabajador','Dónde veo mi saldo vacacional','Ingresa a Gestión Vacacional > Dashboard vacacional. Ahí verás tus periodos, días ganados, gozados y saldo disponible.','/vacaciones/mi_solicitud','mi saldo vacaciones cuantos dias tengo'),
        ('vacacional','trabajador','Dónde veo el estado de mi solicitud','Ingresa a Gestión Vacacional > Dashboard vacacional. En Mis solicitudes verás si está pendiente, aprobada, rechazada o anulada.','/vacaciones/mi_solicitud','estado solicitud vacaciones pendiente aprobada rechazada'),
        ('portal','ambos','Qué puede hacer la IA HR','La IA HR ayuda con dudas del sistema, gestión documental, boletas, vacaciones y conceptos laborales peruanos. Las respuestas se filtran por rol: administrador o trabajador.','','ia asistente ayuda portal rrhh'),
        ('portal','trabajador','Por qué no puedo cargar base de trabajadores','Esa opción es solo para administradores. Como trabajador puedes consultar tus documentos, boletas, vacaciones y estado de tus solicitudes.','','cargar base trabajadores no puedo'),
    ]
    semillas_legal = [
        ('Boleta de pago','Qué es una boleta de pago','La boleta de pago es el documento que informa al trabajador el detalle de su remuneración, descuentos, aportes y otros conceptos registrados en planilla. Sirve como constancia del pago realizado.','D.S. N.° 001-98-TR y modificatorias','El D.S. N.° 001-98-TR regula la obligación de llevar planillas y entregar boletas; el D.S. N.° 009-2011-TR modificó reglas de entrega.','boleta pago remuneracion descuentos aportes planilla'),
        ('CTS','Qué es CTS','La Compensación por Tiempo de Servicios (CTS) es un beneficio social de previsión frente al cese. En términos simples, funciona como un fondo de protección para el trabajador cuando termina el vínculo laboral.','TUO de la Ley de CTS - D.S. N.° 001-97-TR','La CTS se regula por el TUO aprobado por D.S. N.° 001-97-TR y su reglamento.','cts compensacion tiempo servicios deposito mayo noviembre'),
        ('Gratificación','Qué es gratificación','La gratificación legal es un beneficio que corresponde en Fiestas Patrias y Navidad, siempre que se cumplan los requisitos legales. Se paga en julio y diciembre según corresponda.','Ley N.° 27735 y D.S. N.° 005-2002-TR','La Ley N.° 27735 regula las gratificaciones de Fiestas Patrias y Navidad; su reglamento precisa oportunidad y requisitos.','gratificacion julio diciembre fiestas patrias navidad'),
        ('Vacaciones','Cuándo gano vacaciones','En el régimen laboral privado general, el trabajador adquiere derecho a 30 días calendario de descanso vacacional por cada año completo de servicios, sujeto al cumplimiento del récord vacacional aplicable.','D. Leg. N.° 713 y reglamento','El D. Leg. N.° 713 regula descansos remunerados; el derecho vacacional se genera por año completo y récord.','vacaciones cuando gano derecho record vacacional 30 dias'),
        ('Vacaciones','Qué es récord vacacional','El récord vacacional es el requisito de días efectivamente laborados que debe cumplir el trabajador dentro del año de servicios para acceder al descanso vacacional.','D. Leg. N.° 713','El récord depende de la jornada; por ejemplo, para jornadas de 6 días se exige un mínimo mayor que para jornadas de 5 días.','record vacacional requisito dias laborados'),
        ('Vacaciones','Qué son vacaciones truncas','Las vacaciones truncas son el pago proporcional que corresponde cuando el trabajador cesa antes de cumplir un nuevo año completo, siempre que haya generado el derecho proporcional conforme a ley.','D. Leg. N.° 713 y normas complementarias','Se calculan de forma proporcional al tiempo laborado pendiente.','vacaciones truncas cese liquidacion proporcional'),
        ('Liquidación','Qué es liquidación','La liquidación de beneficios sociales es el cálculo de conceptos pendientes al cese, como remuneraciones, vacaciones truncas o pendientes, CTS, gratificación trunca u otros conceptos aplicables.','Normativa laboral peruana según beneficio','Depende del tipo de cese, régimen laboral y conceptos generados.','liquidacion beneficios sociales renuncia cese'),
        ('Utilidades','Qué son utilidades','La participación en utilidades es un beneficio que puede corresponder a trabajadores de empresas que generan renta empresarial y cumplen condiciones legales. Su cálculo depende de la actividad y resultados de la empresa.','D. Leg. N.° 892 y normas complementarias','La obligación depende de la actividad empresarial, renta y número de trabajadores, entre otros factores.','utilidades participacion trabajadores empresa'),
        ('AFP ONP','Qué es AFP y ONP','AFP y ONP son sistemas de pensiones. La AFP pertenece al sistema privado de pensiones; la ONP corresponde al sistema nacional de pensiones. El trabajador aporta según el sistema elegido.','Sistema pensionario peruano','La afiliación y aportes se rigen por normas del sistema pensionario aplicable.','afp onp pension aporte descuento'),
        ('Vida Ley','Qué es Vida Ley','El Seguro Vida Ley es un seguro obligatorio a favor del trabajador, conforme a los supuestos y reglas de la normativa laboral peruana.','D. Leg. N.° 688 y normas complementarias','Tiene finalidad de protección ante contingencias cubiertas por la póliza.','vida ley seguro trabajador'),
        ('SCTR','Qué es SCTR','El Seguro Complementario de Trabajo de Riesgo cubre actividades de alto riesgo cuando corresponda según la normativa aplicable.','Ley N.° 26790 y normas complementarias','Aplica según actividad de riesgo y cobertura definida por ley.','sctr seguro complementario trabajo riesgo'),
        ('Contratos', 'Qué es un contrato de trabajo', 'El contrato de trabajo es el acuerdo por el cual una persona presta servicios personales, remunerados y subordinados para un empleador. En el régimen privado, si existe prestación personal, remuneración y subordinación, se presume relación laboral.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'Base referencial: artículos iniciales del TUO de la Ley de Productividad y Competitividad Laboral.', 'contrato trabajo relacion laboral subordinacion remuneracion servicios'),
        ('Contratos', 'Cuántos tipos de contrato hay en Perú', 'En el régimen laboral privado se reconocen contratos a plazo indeterminado, contratos sujetos a modalidad o plazo fijo, contratos a tiempo parcial y otras formas especiales según norma aplicable. Los contratos sujetos a modalidad requieren causa objetiva y formalidad escrita.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'El contrato indeterminado es la regla general; los contratos modales proceden solo en supuestos regulados.', 'tipos contrato indeterminado plazo fijo sujeto modalidad parcial'),
        ('Contratos', 'Qué es contrato indeterminado', 'Es el contrato sin fecha de término. Puede ser verbal o escrito y se presume cuando existen servicios personales, remunerados y subordinados sin causa temporal válida.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'En el régimen privado, la contratación indeterminada es la regla general.', 'contrato indeterminado indefinido estable sin fecha fin'),
        ('Contratos', 'Qué es contrato sujeto a modalidad', 'Es un contrato temporal con fecha de inicio y fin, sustentado en una causa objetiva prevista por ley, como inicio o incremento de actividad, necesidad de mercado, obra determinada, suplencia, emergencia, intermitente o temporada.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'Debe constar por escrito e indicar la causa objetiva que justifica la temporalidad.', 'contrato sujeto modalidad temporal plazo fijo causa objetiva'),
        ('Contratos', 'Qué es contrato intermitente', 'Es un contrato sujeto a modalidad usado para labores permanentes pero discontinuas, cuando la prestación se ejecuta por periodos alternados. Debe sustentarse en la necesidad real de la actividad.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'Aplica cuando hay prestación intercalada por la naturaleza de la labor.', 'contrato intermitente discontinuo temporada periodo alternado'),
        ('Contratos', 'Qué es periodo de prueba', 'El periodo de prueba es el tiempo inicial de la relación laboral en el que el empleador evalúa la adaptación del trabajador. En el régimen general suele ser de tres meses, salvo ampliaciones válidas para ciertos cargos según ley.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'La ampliación debe respetar límites y justificación según el tipo de puesto.', 'periodo prueba tres meses contrato inicio laboral'),
        ('Contratos', 'Qué es contrato a tiempo parcial', 'Es el contrato en el que la jornada ordinaria es menor a cuatro horas diarias en promedio. Debe documentarse conforme a las reglas aplicables.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'Algunos beneficios pueden variar según el cumplimiento de jornada mínima y régimen aplicable.', 'contrato tiempo parcial part time cuatro horas'),
        ('Licencias', 'Qué es licencia por paternidad', 'La licencia por paternidad es el permiso remunerado que corresponde al trabajador padre por nacimiento de hijo. La regla general vigente es 10 días calendario consecutivos, con ampliaciones en supuestos especiales como parto múltiple o nacimiento prematuro.', 'Ley N.° 29409 modificada por Ley N.° 30807', 'Debe validarse el supuesto concreto y la documentación sustentatoria.', 'licencia paternidad nacimiento hijo padre diez dias'),
        ('Licencias', 'Cuántos días son de licencia por paternidad', 'La licencia por paternidad es de 10 días calendario consecutivos en parto natural o cesárea. Puede ampliarse en casos especiales, como nacimiento prematuro, parto múltiple o complicaciones graves, según la norma.', 'Ley N.° 29409 y Ley N.° 30807', 'RR.HH. debe revisar el certificado o documento sustentatorio para aplicar la duración correcta.', 'dias licencia paternidad 10 20 30 parto multiple prematuro'),
        ('Licencias', 'Qué es licencia por fallecimiento o luto', 'Es la licencia remunerada por fallecimiento de familiares directos. En el sector privado, comprende 5 días calendario por fallecimiento de cónyuge, padres, hijos o hermanos, con reglas de ampliación por traslado cuando corresponda.', 'Ley N.° 31602 y D.S. N.° 013-2023-TR', 'El trabajador debe sustentar el vínculo y fallecimiento según el reglamento.', 'licencia luto fallecimiento familiar padres hijos hermanos conyuge'),
        ('Licencias', 'Cuántos días son de licencia por fallecimiento', 'La licencia por fallecimiento en el sector privado es de 5 días calendario con goce de remuneración por fallecimiento de cónyuge, padres, hijos o hermanos. Puede extenderse por traslado si el fallecimiento ocurre en lugar distinto al centro de trabajo.', 'Ley N.° 31602 y D.S. N.° 013-2023-TR', 'La extensión depende de la vía de transporte y la distancia, conforme al reglamento.', 'dias licencia fallecimiento luto 5 dias calendario traslado'),
        ('Licencias', 'Qué es licencia por maternidad', 'Es el descanso remunerado relacionado con el embarazo y nacimiento. En Perú, el descanso pre y postnatal suma 98 días calendario, usualmente 49 días antes y 49 días después del parto, con reglas especiales según el caso.', 'Ley N.° 26644 y normas modificatorias', 'Puede diferirse o ampliarse según supuestos regulados, previa documentación médica.', 'licencia maternidad pre natal post natal 98 dias'),
        ('Licencias', 'Qué es hora de lactancia', 'Es el derecho de la madre trabajadora a una hora diaria de permiso por lactancia materna hasta que el hijo cumpla un año, conforme a la normativa aplicable.', 'Ley N.° 27240 y normas modificatorias', 'El ejercicio puede coordinarse con el empleador sin afectar el derecho.', 'lactancia hora diaria madre trabajadora hijo un año'),
        ('Licencias', 'Qué es descanso médico', 'El descanso médico es la indicación de reposo emitida por profesional autorizado por una incapacidad temporal. Sirve para justificar inasistencia y, según el caso, gestionar subsidios o canjes ante EsSalud.', 'Normas de seguridad social en salud y procedimientos EsSalud', 'RR.HH. debe validar certificado, días, diagnóstico reservado y procedimiento de subsidio.', 'descanso medico certificado incapacidad temporal essalud'),
        ('Subsidios', 'Qué es subsidio por incapacidad temporal', 'Es la prestación económica vinculada a incapacidad temporal para el trabajo, conforme a reglas de EsSalud y normativa de seguridad social. Usualmente se gestiona cuando se cumplen requisitos y días correspondientes.', 'Ley N.° 26790 y normas EsSalud', 'La aplicación depende de aportes, acreditación y documentación médica.', 'subsidio incapacidad temporal descanso medico essalud'),
        ('Subsidios', 'Qué es subsidio por maternidad', 'Es la prestación económica que cubre el periodo de descanso por maternidad cuando se cumplen requisitos de aseguramiento y aportes ante EsSalud.', 'Ley N.° 26790 y normas EsSalud', 'Requiere validación de vínculo, aportes y documentación médica.', 'subsidio maternidad essalud descanso prenatal postnatal'),
        ('Planillas', 'Qué es asignación familiar', 'La asignación familiar es un beneficio remunerativo equivalente al 10% de la Remuneración Mínima Vital, para trabajadores del régimen privado que acrediten tener hijos menores de edad o mayores en estudios conforme a ley.', 'Ley N.° 25129 y D.S. N.° 035-90-TR', 'Tiene naturaleza remunerativa y no varía por número de hijos.', 'asignacion familiar 10 rmv hijos menor estudios'),
        ('Planillas', 'Cuánto es la asignación familiar', 'La asignación familiar equivale al 10% de la Remuneración Mínima Vital vigente. Para calcularla se debe verificar la RMV vigente al periodo de pago.', 'Ley N.° 25129 y D.S. N.° 035-90-TR', 'El monto cambia si cambia la RMV.', 'monto asignacion familiar 10 rmv cuanto'),
        ('Jornada', 'Cuál es la jornada máxima de trabajo', 'La jornada máxima ordinaria en Perú es de 8 horas diarias o 48 horas semanales, salvo jornadas especiales, acumulativas o atípicas válidamente implementadas.', 'Constitución Política del Perú y D. Leg. N.° 854', 'RR.HH. debe revisar régimen, jornada, horario y descansos.', 'jornada maxima 8 horas 48 semanales horario trabajo'),
        ('Jornada', 'Qué son horas extras', 'Las horas extras son trabajo realizado por encima de la jornada ordinaria. Deben ser voluntarias, registradas y pagadas con sobretasa legal o compensadas conforme a acuerdo válido.', 'D. Leg. N.° 854 y D.S. N.° 007-2002-TR', 'La sobretasa mínima suele ser 25% para las dos primeras horas y 35% para las siguientes.', 'horas extras sobretiempo sobretasa 25 35'),
        ('Jornada', 'Qué es descanso semanal obligatorio', 'Es el derecho a un descanso remunerado semanal, usualmente de 24 horas consecutivas, conforme a la jornada y organización del trabajo.', 'D. Leg. N.° 713', 'Debe coordinarse con la programación de turnos y actividad empresarial.', 'descanso semanal obligatorio dso 24 horas'),
        ('Jornada', 'Qué pasa si trabajo feriado', 'El trabajo en feriado no laborable puede generar pago adicional o descanso sustitutorio, según corresponda. RR.HH. debe verificar si hubo descanso compensatorio.', 'D. Leg. N.° 713', 'El tratamiento depende de si el feriado fue laborado y si existió sustitución.', 'feriado trabajado pago triple descanso sustitutorio'),
        ('Remuneración', 'Qué es remuneración', 'La remuneración es todo lo que el trabajador recibe por sus servicios, en dinero o especie, siempre que sea de libre disposición, salvo conceptos excluidos por ley.', 'TUO del D. Leg. N.° 728 y normas complementarias', 'Es base para varios beneficios, según la naturaleza del concepto.', 'remuneracion sueldo salario conceptos libre disposicion'),
        ('Planillas', 'Qué es planilla electrónica', 'La planilla electrónica es el registro declarado por el empleador con información laboral, previsional y de seguridad social de trabajadores y prestadores, según normas de SUNAT y MTPE.', 'Normas de Planilla Electrónica - SUNAT/MTPE', 'Incluye T-Registro y PLAME, según corresponda.', 'planilla electronica t registro plame sunat mtpe'),
        ('Beneficios', 'Qué son beneficios sociales', 'Son derechos económicos laborales como CTS, gratificaciones, vacaciones, asignación familiar, utilidades, seguros u otros que correspondan según régimen y condiciones del trabajador.', 'Normativa laboral peruana según cada beneficio', 'No todos aplican igual en todos los regímenes.', 'beneficios sociales cts gratificacion vacaciones asignacion utilidades'),
        ('Beneficios', 'Qué son beneficios truncos', 'Son beneficios proporcionales generados al cese antes de completar el periodo completo, como CTS trunca, gratificación trunca o vacaciones truncas, según corresponda.', 'Normativa laboral peruana según cada beneficio', 'Se calculan según tiempo laborado y conceptos computables.', 'beneficios truncos cese liquidacion proporcional'),
        ('Cese', 'Qué es renuncia', 'La renuncia es la decisión voluntaria del trabajador de terminar el vínculo laboral. En el régimen privado se comunica por escrito y, salvo exoneración, debe observar el plazo legal de preaviso.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'El empleador puede aceptar o exonerar el plazo de preaviso.', 'renuncia carta trabajador cese voluntario preaviso'),
        ('Cese', 'Qué es despido arbitrario', 'Es el despido sin causa justa o sin cumplir el procedimiento legal. Puede generar indemnización u otras consecuencias según el tipo de trabajador y régimen aplicable.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'La evaluación requiere revisar causa, procedimiento y pruebas.', 'despido arbitrario causa justa indemnizacion'),
        ('Cese', 'Qué es abandono de trabajo', 'Es una falta relacionada con inasistencias injustificadas bajo condiciones previstas por la ley. Debe evaluarse con cuidado y seguirse el procedimiento disciplinario correspondiente.', 'TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR', 'RR.HH. debe validar días, comunicaciones y descargos.', 'abandono trabajo inasistencia injustificada falta grave'),
        ('Regímenes', 'Qué es régimen laboral general', 'Es el régimen laboral común de la actividad privada, regulado principalmente por el D. Leg. 728 y normas complementarias. Reconoce beneficios como CTS, gratificaciones, vacaciones, entre otros, según requisitos.', 'D. Leg. N.° 728 y normas complementarias', 'Es la base general, salvo regímenes especiales.', 'regimen general actividad privada beneficios'),
        ('Regímenes', 'Qué es régimen laboral agrario', 'Es un régimen especial para trabajadores de actividades agrarias, con reglas propias sobre remuneración, beneficios y condiciones laborales según la Ley del régimen agrario vigente.', 'Ley N.° 31110 y reglamento', 'Debe revisarse actividad, puesto, jornada y conceptos aplicables.', 'regimen agrario ley 31110 trabajadores agrarios'),
        ('Regímenes', 'Qué es régimen MYPE', 'Es el régimen especial aplicable a micro y pequeñas empresas inscritas y calificadas conforme a ley, con beneficios laborales diferenciados según tipo de empresa.', 'TUO de la Ley MYPE y normas complementarias', 'La aplicación depende de inscripción, condición MYPE y fecha correspondiente.', 'regimen mype microempresa pequena empresa beneficios'),
        ('Regímenes', 'Qué es CAS', 'El Contrato Administrativo de Servicios es un régimen especial del sector público. No corresponde al régimen privado común y tiene reglas propias.', 'D. Leg. N.° 1057 y normas complementarias', 'Aplicable principalmente en entidades públicas.', 'cas contrato administrativo servicios sector publico'),
        ('Regímenes', 'Qué son practicantes', 'Los practicantes se vinculan mediante modalidades formativas laborales, no como contrato laboral común. Existen prácticas preprofesionales y profesionales con subvención y condiciones reguladas.', 'Ley N.° 28518', 'No deben usarse para encubrir relación laboral ordinaria.', 'practicantes practicas preprofesionales profesionales modalidades formativas'),
        ('SST', 'Qué es SST', 'La Seguridad y Salud en el Trabajo es el sistema de prevención de riesgos laborales destinado a proteger la vida, salud e integridad de los trabajadores.', 'Ley N.° 29783 y reglamento', 'Incluye política, IPERC, capacitaciones, investigación de accidentes y comité/supervisor.', 'sst seguridad salud trabajo ley 29783'),
        ('SST', 'Qué es IPERC', 'IPERC significa Identificación de Peligros, Evaluación de Riesgos y Controles. Es una herramienta preventiva para gestionar riesgos en el trabajo.', 'Ley N.° 29783 y D.S. N.° 005-2012-TR', 'Debe actualizarse según cambios en procesos, puestos o incidentes.', 'iperc peligros riesgos controles seguridad salud'),
        ('SST', 'Qué es accidente de trabajo', 'Es todo suceso repentino relacionado con el trabajo que produce lesión, perturbación funcional, invalidez o muerte, según la normativa de SST.', 'Ley N.° 29783 y reglamento', 'Debe investigarse, registrarse y reportarse cuando corresponda.', 'accidente trabajo lesion investigacion reporte sst'),
        ('SST', 'Qué es incidente de trabajo', 'Es un suceso relacionado con el trabajo que no necesariamente causa lesión, pero pudo ocasionarla. Sirve para prevenir accidentes futuros.', 'Ley N.° 29783 y reglamento', 'Debe registrarse y analizarse como parte de la prevención.', 'incidente trabajo casi accidente sst'),
        ('SST', 'Qué es comité de SST', 'Es el órgano paritario que participa en la gestión de seguridad y salud en el trabajo cuando la empresa cumple los supuestos legales. En empresas más pequeñas puede corresponder supervisor de SST.', 'Ley N.° 29783 y reglamento', 'Depende del número de trabajadores y estructura de la empresa.', 'comite sst supervisor seguridad salud trabajo'),
        ('Hostigamiento', 'Qué es hostigamiento sexual laboral', 'Es una conducta de naturaleza o connotación sexual o sexista no deseada que afecta la dignidad o genera un ambiente intimidatorio, hostil o humillante. El empleador debe prevenir, investigar y sancionar.', 'Ley N.° 27942 y reglamento', 'Debe existir canal de denuncia y procedimiento de atención.', 'hostigamiento sexual laboral denuncia investigacion'),
        ('Teletrabajo', 'Qué es teletrabajo', 'El teletrabajo es una modalidad especial de prestación de servicios usando tecnologías de información, fuera del centro de trabajo o en forma híbrida, conforme a acuerdo y reglas legales.', 'Ley N.° 31572 y reglamento', 'Debe regularse con condiciones, jornada, equipos, seguridad y desconexión digital.', 'teletrabajo remoto home office desconexion digital'),
        ('Teletrabajo', 'Qué es desconexión digital', 'Es el derecho del trabajador a no responder comunicaciones o requerimientos laborales fuera de su jornada o durante descansos, salvo excepciones justificadas.', 'Ley N.° 31572 y normas sobre teletrabajo', 'Debe respetarse especialmente en teletrabajo y trabajo remoto/híbrido.', 'desconexion digital fuera jornada teletrabajo'),
        ('Vacaciones', 'Qué es fraccionamiento de vacaciones', 'Es la posibilidad de dividir el descanso vacacional conforme a reglas legales y acuerdo entre trabajador y empleador, respetando mínimos y formalidades.', 'D. Leg. N.° 713 y D. Leg. N.° 1405', 'Debe documentarse correctamente y respetar límites.', 'fraccionamiento vacaciones partir dias descanso'),
        ('Vacaciones', 'Qué es adelanto de vacaciones', 'Es el goce anticipado de vacaciones antes de cumplir el año y récord, sujeto a acuerdo escrito entre trabajador y empleador según la normativa vigente.', 'D. Leg. N.° 713 y D. Leg. N.° 1405', 'RR.HH. debe controlar el saldo para evitar inconsistencias al cese.', 'adelanto vacaciones anticipadas acuerdo trabajador empleador'),
        ('Vacaciones', 'Qué es indemnización vacacional', 'Es un pago que puede generarse cuando el trabajador no goza el descanso vacacional oportunamente, conforme a las reglas legales aplicables.', 'D. Leg. N.° 713', 'Requiere revisar periodo, vencimiento y si hubo goce efectivo.', 'indemnizacion vacacional vacaciones vencidas triple'),
        ('Documentos laborales', 'Qué es reglamento interno de trabajo', 'El Reglamento Interno de Trabajo contiene reglas internas de orden laboral, deberes, derechos, medidas disciplinarias y procedimientos de la empresa, cuando corresponde por obligación legal.', 'D.S. N.° 039-91-TR y normas complementarias', 'Debe ser conocido por los trabajadores y alineado a la normativa vigente.', 'reglamento interno trabajo rit reglas empresa'),
        ('Documentos laborales', 'Qué es certificado de trabajo', 'Es el documento que acredita la prestación de servicios del trabajador, usualmente emitido al cese o cuando corresponde según solicitud y política interna.', 'Normativa laboral peruana y buenas prácticas de RR.HH.', 'Debe contener información laboral veraz y no discriminatoria.', 'certificado trabajo constancia laboral cese'),
        ('Documentos laborales', 'Qué es constancia de trabajo', 'Es una constancia emitida por el empleador que acredita vínculo laboral, cargo, fecha de ingreso u otra información laboral permitida.', 'Buenas prácticas laborales y normativa aplicable', 'Debe emitirse con información objetiva y autorizada.', 'constancia trabajo laboral vinculo cargo fecha ingreso'),
        ('Disciplina', 'Qué es medida disciplinaria', 'Es una acción del empleador frente a incumplimientos laborales, como amonestación, suspensión o despido, respetando proporcionalidad, razonabilidad y derecho de defensa.', 'TUO del D. Leg. N.° 728 y principios laborales', 'Debe estar sustentada, documentada y seguir procedimiento cuando corresponda.', 'medida disciplinaria amonestacion suspension descargo'),
        ('RRHH', 'Qué es inducción laboral', 'La inducción laboral es el proceso de bienvenida y orientación al trabajador sobre la empresa, funciones, normas, SST, políticas y cultura organizacional.', 'Buenas prácticas de RR.HH. y obligaciones de información/capacitación según materia', 'Debe documentarse especialmente en temas de SST y políticas internas.', 'induccion laboral bienvenida trabajador capacitacion'),
        ('RRHH', 'Qué es clima laboral', 'El clima laboral es la percepción de los trabajadores sobre ambiente, liderazgo, comunicación, reconocimiento y condiciones de trabajo. Sirve para detectar oportunidades de mejora.', 'Buenas prácticas de gestión humana', 'No es un beneficio legal específico, pero impacta gestión y retención.', 'clima laboral ambiente encuesta bienestar'),
        ('RRHH', 'Qué es evaluación de desempeño', 'Es un proceso para medir cumplimiento de objetivos, competencias y resultados del trabajador. Debe aplicarse con criterios objetivos y comunicados.', 'Buenas prácticas de RR.HH.', 'Puede vincularse a capacitación, promoción o mejora.', 'evaluacion desempeño objetivos competencias resultados'),
    ]
    with db() as con:
        con.execute("""CREATE TABLE IF NOT EXISTS ia_base_conocimiento_hr(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            modulo TEXT,
            rol_destino TEXT DEFAULT 'ambos',
            pregunta_clave TEXT,
            respuesta TEXT,
            ruta TEXT,
            palabras_clave TEXT,
            activo INTEGER DEFAULT 1,
            fecha_registro TEXT,
            registrado_por TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS ia_legislacion_peru(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tema TEXT,
            pregunta_clave TEXT,
            respuesta TEXT,
            base_legal TEXT,
            detalle_legal TEXT,
            palabras_clave TEXT,
            activo INTEGER DEFAULT 1,
            fecha_actualizacion TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS ia_hr_log(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rol TEXT,
            modulo TEXT,
            dni TEXT,
            pregunta TEXT,
            respuesta_tipo TEXT,
            fecha TEXT,
            usuario TEXT
        )""")
        for modulo, rol, preg, resp, ruta, keys in semillas_sistema:
            existe = con.execute('SELECT id FROM ia_base_conocimiento_hr WHERE UPPER(pregunta_clave)=UPPER(?) AND rol_destino=? LIMIT 1', (preg, rol)).fetchone()
            if not existe:
                con.execute("""INSERT INTO ia_base_conocimiento_hr(modulo,rol_destino,pregunta_clave,respuesta,ruta,palabras_clave,activo,fecha_registro,registrado_por)
                               VALUES(?,?,?,?,?,?,1,?,?)""", (modulo,rol,preg,resp,ruta,keys,now_txt(),'SISTEMA'))
        for tema, preg, resp, base, det, keys in semillas_legal:
            existe = con.execute('SELECT id FROM ia_legislacion_peru WHERE UPPER(pregunta_clave)=UPPER(?) LIMIT 1', (preg,)).fetchone()
            if not existe:
                con.execute("""INSERT INTO ia_legislacion_peru(tema,pregunta_clave,respuesta,base_legal,detalle_legal,palabras_clave,activo,fecha_actualizacion)
                               VALUES(?,?,?,?,?,?,1,?)""", (tema,preg,resp,base,det,keys,now_txt()))
        con.commit()


def ia_registrar_log(rol, modulo, dni, pregunta, tipo):
    try:
        with db() as con:
            con.execute('INSERT INTO ia_hr_log(rol,modulo,dni,pregunta,respuesta_tipo,fecha,usuario) VALUES(?,?,?,?,?,?,?)',
                        (rol, modulo, dni, pregunta, tipo, now_txt(), session.get('admin_nombre') or session.get('nombre') or session.get('dni') or 'Sistema'))
            con.commit()
    except Exception:
        pass


def ia_buscar_base_conocimiento(pregunta, rol, modulo=''):
    try:
        with db() as con:
            if rol == 'admin':
                # Administrador: acceso total a respuestas de admin, trabajador y ambos roles.
                rows = con.execute("""SELECT * FROM ia_base_conocimiento_hr
                                      WHERE activo=1""").fetchall()
            else:
                rows = con.execute("""SELECT * FROM ia_base_conocimiento_hr
                                      WHERE activo=1 AND rol_destino IN ('ambos', ?)""", (rol,)).fetchall()
        if not rows:
            return None
        ranked = sorted(rows, key=lambda r: ia_score(pregunta + ' ' + modulo, r), reverse=True)
        if ranked and ia_score(pregunta + ' ' + modulo, ranked[0]) >= 1:
            r = ranked[0]
            ruta = f"<p><b>Ruta sugerida:</b> <a href='{h(r['ruta'])}'>{h(r['ruta'])}</a></p>" if clean(r['ruta']) else ''
            return f"<div class='ia-result ok'><h3>{h(r['pregunta_clave'])}</h3><p>{h(r['respuesta'])}</p>{ruta}</div>"
    except Exception:
        pass
    return None


def ia_buscar_legal(pregunta):
    try:
        with db() as con:
            rows = con.execute('SELECT * FROM ia_legislacion_peru WHERE activo=1').fetchall()
        ranked = sorted(rows, key=lambda r: ia_score(pregunta, r), reverse=True)
        if ranked and ia_score(pregunta, ranked[0]) >= 1:
            r = ranked[0]
            return f"""
            <div class='ia-result legal'>
              <h3>{h(r['pregunta_clave'])}</h3>
              <p>{h(r['respuesta'])}</p>
              <p><b>Base legal referencial:</b> {h(r['base_legal'])}</p>
              <p class='muted'>{h(r['detalle_legal'])}</p>
              <small>Respuesta informativa. Para casos especiales, RR.HH. debe validar régimen laboral, jornada y situación concreta.</small>
            </div>"""
    except Exception:
        pass
    return None


def ia_resumen_documental(rol, dni=''):
    with db() as con:
        if rol == 'admin':
            total = ia_sql_count(con, 'SELECT COUNT(*) FROM documentos')
            pendientes = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE COALESCE(estado,'Pendiente') IN ('Pendiente','Aceptado','Firmado')")
            rechazados = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE estado='Rechazado'")
            leidos = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE COALESCE(fecha_lectura,'')<>''")
            rows = ia_sql_rows(con, 'SELECT tipo, COUNT(*) c FROM documentos GROUP BY tipo ORDER BY c DESC LIMIT 6')
            detalle = ''.join([f"<li>{h(r['tipo'])}: <b>{r['c']}</b></li>" for r in rows]) or '<li>Sin documentos cargados.</li>'
            return f"<div class='ia-result ok'><h3>Resumen documental administrador</h3><p>Total: <b>{total}</b> · Pendientes: <b>{pendientes}</b> · Leídos: <b>{leidos}</b> · Rechazados: <b>{rechazados}</b></p><ul>{detalle}</ul><p>Ruta: Gestión Documental > Subir / gestionar documentos.</p></div>"
        dni = normalizar_dni(dni)
        total = ia_sql_count(con, 'SELECT COUNT(*) FROM documentos WHERE dni=?', (dni,))
        leidos = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE dni=? AND COALESCE(fecha_lectura,'')<>''", (dni,))
        rows = ia_sql_rows(con, 'SELECT tipo, periodo, estado, fecha_subida FROM documentos WHERE dni=? ORDER BY id DESC LIMIT 6', (dni,))
        detalle = ''.join([f"<li>{h(r['tipo'])} · {h(r['periodo'])} · {h(r['estado'] or 'Pendiente')}</li>" for r in rows]) or '<li>Aún no tienes documentos cargados para tu DNI.</li>'
        return f"<div class='ia-result ok'><h3>Mis documentos</h3><p>Tienes <b>{total}</b> documento(s) cargado(s). Leídos: <b>{leidos}</b>.</p><ul>{detalle}</ul><p>Ruta: Gestión Documental > Documentos de pago / empresa / personales.</p></div>"


def ia_resumen_vacacional(rol, dni=''):
    with db() as con:
        if rol == 'admin':
            saldos = ia_sql_count(con, 'SELECT COUNT(*) FROM vacaciones_saldos')
            solicitudes = ia_sql_count(con, 'SELECT COUNT(*) FROM vacaciones_solicitudes')
            pendientes = ia_sql_count(con, "SELECT COUNT(*) FROM vacaciones_solicitudes WHERE estado LIKE 'Pendiente%'")
            aprobadas = ia_sql_count(con, "SELECT COUNT(*) FROM vacaciones_solicitudes WHERE estado LIKE 'Aprobado%'")
            return f"<div class='ia-result warn'><h3>Resumen vacacional administrador</h3><p>Saldos registrados: <b>{saldos}</b> · Solicitudes: <b>{solicitudes}</b> · Pendientes: <b>{pendientes}</b> · Aprobadas: <b>{aprobadas}</b>.</p><p>Ruta: Gestión Vacacional > Dashboard / Saldos / Aprobaciones / Reportes.</p></div>"
        dni = normalizar_dni(dni)
        rows = ia_sql_rows(con, 'SELECT periodo_inicio, periodo_fin, dias_ganados, dias_gozados, saldo FROM vacaciones_saldos WHERE dni=? ORDER BY periodo_inicio DESC, periodo_fin DESC', (dni,))
        total_saldo = sum(float(r['saldo'] or 0) for r in rows)
        sol = ia_sql_rows(con, 'SELECT fecha_inicio, fecha_fin, dias, estado FROM vacaciones_solicitudes WHERE dni=? ORDER BY id DESC LIMIT 5', (dni,))
        detalle = ''.join([f"<li>Periodo {h(r['periodo_inicio'])}/{h(r['periodo_fin'])}: saldo <b>{r['saldo']}</b> día(s)</li>" for r in rows[:5]]) or '<li>No se encontraron saldos cargados para tu DNI.</li>'
        solicitudes = ''.join([f"<li>{h(s['fecha_inicio'])} al {h(s['fecha_fin'])}: {h(s['estado'])}</li>" for s in sol]) or '<li>No tienes solicitudes registradas.</li>'
        return f"<div class='ia-result warn'><h3>Mis vacaciones</h3><p>Saldo total referencial: <b>{total_saldo:g}</b> día(s).</p><h4>Saldos</h4><ul>{detalle}</ul><h4>Solicitudes recientes</h4><ul>{solicitudes}</ul><p>Ruta: Gestión Vacacional > Saldo y solicitud.</p></div>"


def ia_hr_responder(pregunta, modulo='', rol=None, dni=None):
    pregunta = clean(pregunta)
    rol = rol or ('admin' if session.get('admin_id') else 'trabajador')
    dni = normalizar_dni(dni or session.get('dni') or '')
    q = ia_norm(pregunta)
    if not pregunta:
        return "<div class='ia-result warn'><h3>IA HR</h3><p>Escribe una pregunta. Ejemplo: ¿qué es CTS?, ¿dónde veo mis boletas?, ¿cuánto saldo tengo?</p></div>", 'vacio'
    admin_terms = ['cargar base','cargar trabajadores','subir saldos','cargar saldos','subir boletas masiva','sincronizar pdf','crear carpetas','reporte general','todos los trabajadores']
    if rol != 'admin' and any(t in q for t in admin_terms):
        return "<div class='ia-result bad'><h3>Acceso limitado</h3><p>Esa consulta corresponde a opciones administrativas. Desde tu portal puedes consultar tus boletas, documentos, saldo vacacional, solicitudes y conceptos laborales.</p></div>", 'bloqueo_rol'
    if any(x in q for x in ['mis boletas','mis documentos','documentos','boleta','cts','gratificacion','utilidad','liquidacion']) and any(x in q for x in ['tengo','ver','donde','mis','cargadas','aparece','pendiente','resumen']):
        return ia_resumen_documental(rol, dni), 'datos_documental'
    if any(x in q for x in ['vacacion','saldo','solicitud','aprobacion','dias']) and any(x in q for x in ['tengo','mis','estado','pendiente','resumen','cuanto','cuantos']):
        return ia_resumen_vacacional(rol, dni), 'datos_vacacional'
    r = ia_buscar_base_conocimiento(pregunta, rol, modulo)
    if r:
        return r, 'base_conocimiento'
    r = ia_buscar_legal(pregunta)
    if r:
        return r, 'legal_peru'
    if 'vac' in ia_norm(modulo):
        return ia_resumen_vacacional(rol, dni), 'fallback_vacacional'
    if 'doc' in ia_norm(modulo) or 'panel' in ia_norm(modulo):
        return ia_resumen_documental(rol, dni), 'fallback_documental'
    if rol == 'admin':
        return "<div class='ia-result ok'><h3>IA HR Administrador</h3><p>Puedo ayudarte con Gestión Documental, boletas, carga de saldos, vacaciones, reportes e información legal laboral peruana. Prueba: ¿dónde cargo boletas?, ¿qué documentos están pendientes?, ¿qué es CTS?</p></div>", 'fallback'
    return "<div class='ia-result ok'><h3>IA HR Trabajador</h3><p>Puedo ayudarte a revisar tus boletas, documentos, vacaciones, solicitudes y conceptos laborales. Prueba: ¿qué es CTS?, ¿cuánto saldo tengo?, ¿dónde veo mi boleta?</p></div>", 'fallback'


def ia_widget_hr(active=''):
    if not (session.get('admin_id') or session.get('dni')):
        return ''
    rol = 'Administrador' if session.get('admin_id') else 'Trabajador'
    ejemplos = '¿Dónde cargo boletas? | ¿Qué documentos están pendientes? | ¿Qué es CTS?' if session.get('admin_id') else '¿Dónde veo mi boleta? | ¿Cuánto saldo tengo? | ¿Qué es gratificación?'
    return f"""
    <div class='iahr-fab' onclick='iahrToggle()'><i class='bi bi-robot'></i><span>IA HR</span></div>
    <div class='iahr-panel' id='iahrPanel'>
      <div class='iahr-head'><div><b>🤖 IA HR PRO</b><small>{h(rol)} · {h(active)}</small></div><button type='button' onclick='iahrToggle()'>×</button></div>
      <div class='iahr-body'>
        <div class='iahr-ex'>{h(ejemplos)}</div>
        <input id='iahrModulo' type='hidden' value='{h(active)}'>
        <textarea id='iahrPregunta' placeholder='Escribe tu pregunta aquí...'></textarea>
        <button type='button' class='btn-green iahr-send' onclick='iahrAsk(event)'>Consultar IA</button>
        <div id='iahrRespuesta' class='iahr-respuesta'><div class='muted'>La respuesta se filtrará según tu rol.</div></div>
      </div>
    </div>"""

def sincronizar_jefes_vacaciones(con):
    """Sincroniza JEFE INMEDIATO por DNI entre saldos, trabajadores y solicitudes."""
    try:
        for r in con.execute("SELECT id, jefe_dni FROM vacaciones_saldos WHERE COALESCE(jefe_dni,'')<>''").fetchall():
            nd = normalizar_dni(r['jefe_dni'])
            if nd and nd != (r['jefe_dni'] or ''):
                con.execute("UPDATE vacaciones_saldos SET jefe_dni=? WHERE id=?", (nd, r['id']))
    except Exception:
        pass
    try:
        for r in con.execute("SELECT id, jefe_dni FROM vacaciones_solicitudes WHERE COALESCE(jefe_dni,'')<>''").fetchall():
            nd = normalizar_dni(r['jefe_dni'])
            if nd and nd != (r['jefe_dni'] or ''):
                con.execute("UPDATE vacaciones_solicitudes SET jefe_dni=? WHERE id=?", (nd, r['id']))
    except Exception:
        pass
    try:
        con.execute("""
            UPDATE trabajadores
               SET jefe_dni = COALESCE((SELECT s.jefe_dni FROM vacaciones_saldos s
                        WHERE s.dni = trabajadores.dni AND COALESCE(s.jefe_dni,'')<>''
                        ORDER BY s.periodo_inicio DESC, s.periodo_fin DESC, s.id DESC LIMIT 1), jefe_dni),
                   jefe_nombre = COALESCE((SELECT tj.nombre FROM vacaciones_saldos s
                        LEFT JOIN trabajadores tj ON tj.dni = s.jefe_dni
                        WHERE s.dni = trabajadores.dni AND COALESCE(s.jefe_dni,'')<>''
                        ORDER BY s.periodo_inicio DESC, s.periodo_fin DESC, s.id DESC LIMIT 1), jefe_nombre)
             WHERE EXISTS (SELECT 1 FROM vacaciones_saldos s WHERE s.dni = trabajadores.dni AND COALESCE(s.jefe_dni,'')<>'')
        """)
    except Exception:
        pass
    try:
        con.execute("""
            UPDATE vacaciones_solicitudes
               SET jefe_dni = COALESCE((SELECT s.jefe_dni FROM vacaciones_saldos s
                        WHERE s.dni = vacaciones_solicitudes.dni AND COALESCE(s.jefe_dni,'')<>''
                        ORDER BY s.periodo_inicio DESC, s.periodo_fin DESC, s.id DESC LIMIT 1),
                        (SELECT tr.jefe_dni FROM trabajadores tr
                        WHERE tr.dni = vacaciones_solicitudes.dni AND COALESCE(tr.jefe_dni,'')<>'' LIMIT 1), jefe_dni)
             WHERE estado IN ('Pendiente jefe','Pendiente Jefe','Pendiente') OR COALESCE(jefe_dni,'')=''
        """)
    except Exception:
        pass

def init_db():
    with db() as con:
        con.execute("""
        CREATE TABLE IF NOT EXISTS usuarios_admin(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT UNIQUE,
            clave_hash TEXT,
            nombre TEXT,
            rol TEXT DEFAULT 'admin',
            activo INTEGER DEFAULT 1
        )""")
        con.execute("""
        CREATE TABLE IF NOT EXISTS trabajadores(
            dni TEXT PRIMARY KEY,
            nombre TEXT,
            correo TEXT,
            cargo TEXT,
            area TEXT,
            empresa TEXT,
            activo INTEGER DEFAULT 1,
            fecha_registro TEXT
        )""")
        con.execute("""
        CREATE TABLE IF NOT EXISTS documentos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            dni TEXT,
            categoria TEXT,
            tipo TEXT,
            periodo TEXT,
            detalle TEXT,
            observacion TEXT,
            archivo_nombre TEXT,
            ruta_archivo TEXT,
            extension TEXT,
            fecha_subida TEXT,
            uploaded_by TEXT
        )""")
        # Migraciones livianas para versiones anteriores
        for col, ddl in [
            ('planilla', 'ALTER TABLE trabajadores ADD COLUMN planilla TEXT'),
            ('foto_ruta', 'ALTER TABLE trabajadores ADD COLUMN foto_ruta TEXT'),
            ('fecha_nacimiento', 'ALTER TABLE trabajadores ADD COLUMN fecha_nacimiento TEXT'),
            ('fecha_ingreso', 'ALTER TABLE trabajadores ADD COLUMN fecha_ingreso TEXT'),
            ('empresa_login', 'ALTER TABLE trabajadores ADD COLUMN empresa_login TEXT'),
            ('usuario_portal', 'ALTER TABLE trabajadores ADD COLUMN usuario_portal TEXT'),
            ('clave_portal', 'ALTER TABLE trabajadores ADD COLUMN clave_portal TEXT'),
            ('jefe_dni', 'ALTER TABLE trabajadores ADD COLUMN jefe_dni TEXT'),
            ('jefe_nombre', 'ALTER TABLE trabajadores ADD COLUMN jefe_nombre TEXT'),
            ('celular', 'ALTER TABLE trabajadores ADD COLUMN celular TEXT'),
            ('contacto_emergencia', 'ALTER TABLE trabajadores ADD COLUMN contacto_emergencia TEXT'),
            ('telefono_emergencia', 'ALTER TABLE trabajadores ADD COLUMN telefono_emergencia TEXT'),
            ('tipo_contrato', 'ALTER TABLE trabajadores ADD COLUMN tipo_contrato TEXT'),
            ('fecha_fin_contrato', 'ALTER TABLE trabajadores ADD COLUMN fecha_fin_contrato TEXT'),
            ('remuneracion_basica', 'ALTER TABLE trabajadores ADD COLUMN remuneracion_basica TEXT'),
            ('direccion', 'ALTER TABLE trabajadores ADD COLUMN direccion TEXT'),
            ('departamento', 'ALTER TABLE trabajadores ADD COLUMN departamento TEXT'),
            ('provincia', 'ALTER TABLE trabajadores ADD COLUMN provincia TEXT'),
            ('distrito', 'ALTER TABLE trabajadores ADD COLUMN distrito TEXT'),
            ('nivel_educativo', 'ALTER TABLE trabajadores ADD COLUMN nivel_educativo TEXT'),
            ('procedencia', 'ALTER TABLE trabajadores ADD COLUMN procedencia TEXT'),
            ('indumentaria', 'ALTER TABLE trabajadores ADD COLUMN indumentaria TEXT'),
            ('carnet_conadis', 'ALTER TABLE trabajadores ADD COLUMN carnet_conadis TEXT'),
            ('observacion', 'ALTER TABLE trabajadores ADD COLUMN observacion TEXT'),
        ]:
            try: con.execute(ddl)
            except Exception: pass
        for col, ddl in [
            ('fecha_ingreso', 'ALTER TABLE vacaciones_saldos ADD COLUMN fecha_ingreso TEXT'),
            ('periodo_inicio', 'ALTER TABLE vacaciones_saldos ADD COLUMN periodo_inicio TEXT'),
            ('periodo_fin', 'ALTER TABLE vacaciones_saldos ADD COLUMN periodo_fin TEXT'),
            ('jefe_dni', 'ALTER TABLE vacaciones_saldos ADD COLUMN jefe_dni TEXT'),
        ]:
            try: con.execute(ddl)
            except Exception: pass
        for col, ddl in [
            ('estado', "ALTER TABLE documentos ADD COLUMN estado TEXT DEFAULT 'Pendiente'"),
            ('comentario_rechazo', 'ALTER TABLE documentos ADD COLUMN comentario_rechazo TEXT'),
            ('fecha_aceptacion', 'ALTER TABLE documentos ADD COLUMN fecha_aceptacion TEXT'),
            ('fecha_firma', 'ALTER TABLE documentos ADD COLUMN fecha_firma TEXT'),
            ('fecha_aprobacion', 'ALTER TABLE documentos ADD COLUMN fecha_aprobacion TEXT'),
            ('fecha_lectura', 'ALTER TABLE documentos ADD COLUMN fecha_lectura TEXT'),
            ('leido_por', 'ALTER TABLE documentos ADD COLUMN leido_por TEXT'),
        ]:
            try: con.execute(ddl)
            except Exception: pass
        con.execute('''
        CREATE TABLE IF NOT EXISTS eventos_documento(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            documento_id INTEGER,
            dni TEXT,
            evento TEXT,
            fecha TEXT,
            detalle TEXT
        )''')
        con.execute('''
        CREATE TABLE IF NOT EXISTS app_config(
            clave TEXT PRIMARY KEY,
            valor TEXT
        )''')
        con.execute('''
        CREATE TABLE IF NOT EXISTS login_intentos(
            dni TEXT PRIMARY KEY,
            intentos INTEGER DEFAULT 0,
            bloqueado INTEGER DEFAULT 0,
            ultima_fecha TEXT
        )''')

        con.execute('''
        CREATE TABLE IF NOT EXISTS vacaciones_saldos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            dni TEXT,
            trabajador TEXT,
            empresa TEXT,
            area TEXT,
            jefe TEXT,
            jefe_dni TEXT,
            fecha_ingreso TEXT,
            periodo_inicio TEXT,
            periodo_fin TEXT,
            dias_ganados REAL DEFAULT 0,
            dias_gozados REAL DEFAULT 0,
            saldo REAL DEFAULT 0,
            periodo TEXT,
            fecha_carga TEXT,
            uploaded_by TEXT
        )''')
        try:
            cols_info = con.execute("PRAGMA table_info(vacaciones_saldos)").fetchall()
            has_id = any(c[1] == 'id' for c in cols_info)
            dni_is_pk = any(c[1] == 'dni' and c[5] == 1 for c in cols_info)
            if (not has_id) or dni_is_pk:
                con.execute('ALTER TABLE vacaciones_saldos RENAME TO vacaciones_saldos_old')
                con.execute('''CREATE TABLE vacaciones_saldos(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    dni TEXT, trabajador TEXT, empresa TEXT, area TEXT, jefe TEXT, jefe_dni TEXT,
                    fecha_ingreso TEXT, periodo_inicio TEXT, periodo_fin TEXT,
                    dias_ganados REAL DEFAULT 0, dias_gozados REAL DEFAULT 0, saldo REAL DEFAULT 0,
                    periodo TEXT, fecha_carga TEXT, uploaded_by TEXT
                )''')
                con.execute('''INSERT INTO vacaciones_saldos(dni,trabajador,empresa,area,jefe,jefe_dni,fecha_ingreso,periodo_inicio,periodo_fin,dias_ganados,dias_gozados,saldo,periodo,fecha_carga,uploaded_by)
                    SELECT dni,trabajador,empresa,area,jefe,jefe_dni,fecha_ingreso,periodo_inicio,periodo_fin,dias_ganados,dias_gozados,saldo,periodo,fecha_carga,uploaded_by FROM vacaciones_saldos_old''')
                con.execute('DROP TABLE vacaciones_saldos_old')
        except Exception:
            pass
        try:
            con.execute('CREATE UNIQUE INDEX IF NOT EXISTS idx_vac_saldos_dni_periodo ON vacaciones_saldos(dni, periodo_inicio, periodo_fin)')
        except Exception:
            pass
        for col, ddl in [
            ('fecha_ingreso', 'ALTER TABLE vacaciones_saldos ADD COLUMN fecha_ingreso TEXT'),
            ('periodo_inicio', 'ALTER TABLE vacaciones_saldos ADD COLUMN periodo_inicio TEXT'),
            ('periodo_fin', 'ALTER TABLE vacaciones_saldos ADD COLUMN periodo_fin TEXT'),
            ('jefe_dni', 'ALTER TABLE vacaciones_saldos ADD COLUMN jefe_dni TEXT'),
        ]:
            try: con.execute(ddl)
            except Exception: pass

        con.execute('''
        CREATE TABLE IF NOT EXISTS vacaciones_solicitudes(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            dni TEXT,
            trabajador TEXT,
            jefe_dni TEXT,
            fecha_inicio TEXT,
            fecha_fin TEXT,
            dias REAL,
            motivo TEXT,
            estado TEXT DEFAULT 'Pendiente jefe',
            comentario_jefe TEXT,
            comentario_gh TEXT,
            fecha_solicitud TEXT,
            fecha_jefe TEXT,
            fecha_gh TEXT
        )''')
        for col, ddl in [
            ('jefe_dni', 'ALTER TABLE vacaciones_solicitudes ADD COLUMN jefe_dni TEXT'),
            ('periodo_detalle', 'ALTER TABLE vacaciones_solicitudes ADD COLUMN periodo_detalle TEXT'),
            ('periodo_ids', 'ALTER TABLE vacaciones_solicitudes ADD COLUMN periodo_ids TEXT'),
        ]:
            try: con.execute(ddl)
            except Exception: pass

        con.execute('''
        CREATE TABLE IF NOT EXISTS contratacion_docs(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            dni TEXT,
            trabajador TEXT,
            empresa TEXT,
            etapa TEXT,
            tipo_doc TEXT,
            estado TEXT DEFAULT 'Generado',
            archivo_nombre TEXT,
            ruta_archivo TEXT,
            fecha_registro TEXT,
            uploaded_by TEXT
        )''')
        con.execute('''
        CREATE TABLE IF NOT EXISTS contratacion_tipos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codigo TEXT,
            descripcion TEXT,
            etapa TEXT,
            obligatorio INTEGER DEFAULT 1,
            activo INTEGER DEFAULT 1,
            requiere_firma_digital INTEGER DEFAULT 1,
            requiere_firma_electronica INTEGER DEFAULT 1,
            requiere_confirmacion_recepcion INTEGER DEFAULT 0
        )''')
        # Migración PRO: campos solicitados para Tipos de Documento por Etapa.
        for col, ddl in [
            ('requiere_firma_digital', 'ALTER TABLE contratacion_tipos ADD COLUMN requiere_firma_digital INTEGER DEFAULT 1'),
            ('requiere_firma_electronica', 'ALTER TABLE contratacion_tipos ADD COLUMN requiere_firma_electronica INTEGER DEFAULT 1'),
            ('requiere_confirmacion_recepcion', 'ALTER TABLE contratacion_tipos ADD COLUMN requiere_confirmacion_recepcion INTEGER DEFAULT 0'),
        ]:
            try: con.execute(ddl)
            except Exception: pass
        base_tipos=[
            ('468','ANEXO DE RIESGOS','Incorporacion',1,1,1,1,0),
            ('890','BOLETIN SIS. PENSIONARIO','Incorporacion',1,1,1,1,0),
            ('323','CARGO DE ENTREGA','Incorporacion',1,1,1,1,0),
            ('706','CARTA DE COMPROMISO','Incorporacion',1,1,1,1,0),
            ('400','ELECCIÓN DE BENEFICIOS SOCIALES','Incorporacion',1,1,1,1,0),
            ('214','AUTODECLARACION BUENAS PRACTICAS','Incorporacion',1,1,1,1,0),
            ('340','ACUERDO PREFERENCIAL','Incorporacion',1,1,1,1,0),
            ('159','NOTA DE CARGO','Incorporacion',1,1,1,1,0),
            ('929','CONTRATO TRABAJADOR(RENOVACIÓN)','Renovacion',1,1,1,1,0),
            ('057','CARGO ENTREGA RENOVACION','Renovacion',1,1,1,1,0),
        ]
        for codigo, desc, etapa, obligatorio, activo, fd, fe, cr in base_tipos:
            existe = con.execute('SELECT id FROM contratacion_tipos WHERE codigo=? OR UPPER(descripcion)=UPPER(?) LIMIT 1', (codigo, desc)).fetchone()
            if not existe:
                con.execute('''INSERT INTO contratacion_tipos(codigo,descripcion,etapa,obligatorio,activo,requiere_firma_digital,requiere_firma_electronica,requiere_confirmacion_recepcion)
                               VALUES(?,?,?,?,?,?,?,?)''', (codigo, desc, etapa, obligatorio, activo, fd, fe, cr))

        # PRO: Maestro Tipo Documento de Empleado (similar a Adapta)
        con.execute('''
        CREATE TABLE IF NOT EXISTS contratacion_tipo_empleado(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codigo TEXT,
            descripcion TEXT,
            nombre_corto TEXT,
            grupo TEXT,
            activo INTEGER DEFAULT 1,
            permite_doc_sin_cargo INTEGER DEFAULT 0,
            requiere_validacion INTEGER DEFAULT 0,
            tiene_fecha_vencimiento INTEGER DEFAULT 0,
            es_contrato INTEGER DEFAULT 0
        )''')
        base_tipo_empleado=[
            ('089','CONTRATO TRABAJADOR','CONTRATO','Contratación Trabajadores',1,0,0,0,1),
            ('929','CONTRATO TRABAJADOR(RENOVACIÓN)','CONTRATO RENOVACIÓN','Renovaciones de Contrato',1,0,0,0,1),
            ('468','ANEXO DE RIESGOS','ANEXO DE RIESGOS','Contratación Trabajadores',1,0,0,0,0),
            ('890','BOLETIN SIS. PENSIONARIO','BOLETIN PENSIONARIO','Contratación Trabajadores',1,0,0,0,0),
            ('323','CARGO DE ENTREGA','CARGO DE ENTREGA','Contratación Trabajadores',1,0,0,0,0),
            ('706','CARTA DE COMPROMISO','CARTA COMPROMISO','Contratación Trabajadores',1,0,0,0,0),
            ('159','NOTA DE CARGO','NOTA DE CARGO','Contratación Trabajadores',1,0,0,0,0),
            ('057','CARGO ENTREGA RENOVACION','CARGO RENOVACION','Renovaciones de Contrato',1,0,0,0,0),
            ('400','ELECCIÓN DE BENEFICIOS SOCIALES','ELECCIÓN BENEFICIOS','Contratación Trabajadores',1,0,0,0,0),
            ('214','AUTODECLARACION BUENAS PRACTICAS','BUENAS PRACTICAS','Contratación Trabajadores',1,0,0,0,0),
        ]
        for codigo, descripcion, nombre_corto, grupo, activo, sin_cargo, validacion, vencimiento, contrato in base_tipo_empleado:
            existe = con.execute('SELECT id FROM contratacion_tipo_empleado WHERE codigo=? OR UPPER(descripcion)=UPPER(?) LIMIT 1', (codigo, descripcion)).fetchone()
            if not existe:
                con.execute('''INSERT INTO contratacion_tipo_empleado(codigo,descripcion,nombre_corto,grupo,activo,permite_doc_sin_cargo,requiere_validacion,tiene_fecha_vencimiento,es_contrato)
                               VALUES(?,?,?,?,?,?,?,?,?)''', (codigo, descripcion, nombre_corto, grupo, activo, sin_cargo, validacion, vencimiento, contrato))

        # PRO: Maestro Cargos (con estado Activo/Inactivo y ficha editable)
        con.execute('''
        CREATE TABLE IF NOT EXISTS contratacion_cargos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codigo TEXT,
            nombre TEXT,
            nombre_corto TEXT,
            resumen TEXT,
            funciones TEXT,
            activo INTEGER DEFAULT 1
        )''')
        base_cargos=[
            ('803','AUXILIAR DE TRANSPORTE Y COMEDORES','AUXILIAR DE TRANSPORTE Y COMEDORES','', 'Coordinar con transportistas para asegurar el recojo y llegada del personal a tiempo en fundo.\\nRecepcionar y archivar manifiestos y facturas de los proveedores de transporte.'),
            ('023','AUXILIAR DOCUMENTARIO DE COMPRAS','AUXILIAR DOCUMENTARIO DE COMPRAS','', 'Gestionar documentos de compras, validar sustentos y apoyar el archivo documentario del área.'),
            ('741','OPERARIO DE ALMACEN AGRICOLA','OPERARIO DE ALMACEN AGRICOLA','', 'Recepcionar, ordenar y controlar materiales agrícolas según procedimientos internos.'),
            ('764','AUXILIAR DE COMERCIO EXTERIOR','AUXILIAR DE COMERCIO EXTERIOR','', 'Apoyar en documentación, seguimiento y control de operaciones de comercio exterior.'),
            ('357','AUXILIAR DE PRODUCCION AGRICOLA','AUXILIAR DE PRODUCCION AGRICOLA','', 'Apoyar las actividades operativas de producción agrícola y control de avances diarios.'),
            ('2950','AUXILIAR DE FERTIRIEGO','AUXILIAR DE FERTIRIEGO','', 'Apoyar la programación, control y ejecución de labores de fertirriego.'),
            ('2949','ASISTENTE DE FERTIRIEGO Y OSMOSIS','ASISTENTE DE FERTIRIEGO Y OSMOSIS','', 'Registrar, monitorear y reportar actividades de fertirriego y sistemas de osmosis.'),
            ('2948','SUB GERENTE AGRÍCOLA','SUB GERENTE AGRÍCOLA','', 'Supervisar la gestión agrícola, coordinar equipos y controlar indicadores productivos.'),
            ('2947','ANALISTA DE GESTIÓN DE LA INFORMACIÓN AGRÍCOLA','ANALISTA DE GESTIÓN DE LA INFORMACIÓN AGRÍCOLA','', 'Analizar información agrícola, generar reportes y proponer mejoras de gestión.'),
            ('2946','AUXILIAR DE DIGITACIÓN - DESPACHO','AUXILIAR DE DIGITACIÓN - DESPACHO','', 'Registrar información de despacho, validar datos y mantener reportes actualizados.'),
        ]
        for codigo, nombre, nombre_corto, resumen, funciones in base_cargos:
            existe = con.execute('SELECT id FROM contratacion_cargos WHERE codigo=? OR UPPER(nombre)=UPPER(?) LIMIT 1', (codigo, nombre)).fetchone()
            if not existe:
                con.execute('''INSERT INTO contratacion_cargos(codigo,nombre,nombre_corto,resumen,funciones,activo) VALUES(?,?,?,?,?,1)''', (codigo, nombre, nombre_corto, resumen, funciones))
        con.execute('''
        CREATE TABLE IF NOT EXISTS contratacion_plantillas(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre_plantilla TEXT,
            descripcion TEXT,
            tipo_documento TEXT,
            esquema TEXT DEFAULT 'Trabajador Contrato Laboral',
            condicion TEXT DEFAULT 'SIN CONDICIONES',
            version TEXT DEFAULT 'Version 01',
            activo INTEGER DEFAULT 1,
            archivo_nombre TEXT,
            ruta_archivo TEXT,
            fecha_creacion TEXT,
            fecha_actualizacion TEXT,
            creado_por TEXT
        )''')
        con.execute('''
        CREATE TABLE IF NOT EXISTS contratacion_plantilla_campos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plantilla_id INTEGER,
            descripcion TEXT,
            tipo_campo TEXT DEFAULT 'Origen de Datos',
            nombre_campo TEXT,
            campo_origen TEXT,
            tipo_dato TEXT DEFAULT 'Text',
            requerido TEXT DEFAULT 'SI',
            activo INTEGER DEFAULT 1
        )''')
        for col, ddl in [
            ('valor_default', 'ALTER TABLE contratacion_plantilla_campos ADD COLUMN valor_default TEXT'),
            ('opciones', 'ALTER TABLE contratacion_plantilla_campos ADD COLUMN opciones TEXT'),
            ('editable_admin', 'ALTER TABLE contratacion_plantilla_campos ADD COLUMN editable_admin INTEGER DEFAULT 1'),
        ]:
            try: con.execute(ddl)
            except Exception: pass
        con.execute('''
        CREATE TABLE IF NOT EXISTS contratacion_plantilla_condiciones(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plantilla_id INTEGER,
            nombre_campo TEXT,
            condicion TEXT DEFAULT '=',
            valor TEXT,
            activo INTEGER DEFAULT 1,
            fecha_registro TEXT,
            creado_por TEXT
        )''')
        con.execute('''
        CREATE TABLE IF NOT EXISTS firma_configuracion(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            proveedor TEXT DEFAULT 'INTERNO',
            modo TEXT DEFAULT 'RECONOCIMIENTO FACIAL / FIRMA DIGITAL',
            reniec_activo INTEGER DEFAULT 0,
            firma_digital_activo INTEGER DEFAULT 0,
            url_api TEXT,
            token_ref TEXT,
            observacion TEXT,
            fecha_registro TEXT
        )''')
        con.execute('''
        CREATE TABLE IF NOT EXISTS firma_solicitudes(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            documento_id INTEGER,
            dni TEXT,
            trabajador TEXT,
            metodo TEXT DEFAULT 'PENDIENTE',
            estado TEXT DEFAULT 'Pendiente',
            evidencia_ref TEXT,
            fecha_envio TEXT,
            fecha_firma TEXT,
            observacion TEXT
        )''')
        # Migración segura: solicitudes de firma con cámara laptop/celular y trazabilidad.
        for col, ddl in [
            ('firma_token', 'ALTER TABLE firma_solicitudes ADD COLUMN firma_token TEXT'),
            ('selfie_path', 'ALTER TABLE firma_solicitudes ADD COLUMN selfie_path TEXT'),
            ('documento_firmado_path', 'ALTER TABLE firma_solicitudes ADD COLUMN documento_firmado_path TEXT'),
            ('camara_origen', 'ALTER TABLE firma_solicitudes ADD COLUMN camara_origen TEXT'),
            ('acepta_terminos', 'ALTER TABLE firma_solicitudes ADD COLUMN acepta_terminos INTEGER DEFAULT 0'),
            ('ip_registro', 'ALTER TABLE firma_solicitudes ADD COLUMN ip_registro TEXT'),
            ('user_agent', 'ALTER TABLE firma_solicitudes ADD COLUMN user_agent TEXT'),
            ('hash_evidencia', 'ALTER TABLE firma_solicitudes ADD COLUMN hash_evidencia TEXT'),
            ('fecha_captura', 'ALTER TABLE firma_solicitudes ADD COLUMN fecha_captura TEXT'),
            ('validacion_estado', 'ALTER TABLE firma_solicitudes ADD COLUMN validacion_estado TEXT'),
            ('proveedor_respuesta', 'ALTER TABLE firma_solicitudes ADD COLUMN proveedor_respuesta TEXT'),
        ]:
            try: con.execute(ddl)
            except Exception: pass

        con.execute('''
        CREATE TABLE IF NOT EXISTS trabajadores_observados(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo_persona TEXT DEFAULT 'Trabajador',
            tipo_documento TEXT DEFAULT 'DNI',
            numero_documento TEXT,
            nombre TEXT,
            motivo TEXT,
            nivel_restriccion TEXT,
            comentario TEXT,
            estado TEXT DEFAULT 'Active',
            creado_por TEXT,
            fecha_creacion TEXT,
            editor_por TEXT,
            fecha_edicion TEXT
        )''')
        for col, ddl in [
            ('tipo_persona', "ALTER TABLE trabajadores_observados ADD COLUMN tipo_persona TEXT DEFAULT 'Trabajador'"),
            ('tipo_documento', "ALTER TABLE trabajadores_observados ADD COLUMN tipo_documento TEXT DEFAULT 'DNI'"),
            ('editor_por', 'ALTER TABLE trabajadores_observados ADD COLUMN editor_por TEXT'),
            ('fecha_edicion', 'ALTER TABLE trabajadores_observados ADD COLUMN fecha_edicion TEXT'),
        ]:
            try: con.execute(ddl)
            except Exception: pass
        if not con.execute("SELECT 1 FROM contratacion_plantillas LIMIT 1").fetchone():
            semillas=[
                ('ACUERDO PREFERENCIAL','ACUERDO PREFERENCIAL','ACUERDO PREFERENCIAL','CONTRATACION PREFERENCIAL AQUII'),
                ('AUTODECLARACION BUENAS PRACTICAS','AUTODECLARACION BUENAS PRACTICAS','AUTODECLARACION BUENAS PRACTICAS','AUTODECLARACION AQU ANQA II'),
                ('CARGO ENTREGA RENOVACION','CARGO ENTREGA RENOVACION','CARGO ENTREGA RENOVACION','CARGO DE ENTREGA - RENOVACION - AQII'),
                ('ELECCIÓN DE BENEFICIOS SOCIALES','ELECCIÓN DE BENEFICIOS SOCIALES','ELECCIÓN DE BENEFICIOS SOCIALES','ELECCION DE BENEFICIOS SOCIALES AQII'),
                ('NOTA DE CARGO','NOTA DE CARGO','NOTA DE CARGO','NOTA DE CARGO AQII'),
                ('CARTA DE COMPROMISO','CARTA DE COMPROMISO','CARTA DE COMPROMISO','CARTA DE COMPROMISO - AQ II'),
                ('CARGO DE ENTREGA','CARGO DE ENTREGA','CARGO DE ENTREGA','CARGO DE ENTREGA AQII')]
            for nom, desc, tipo, arch in semillas:
                con.execute("INSERT INTO contratacion_plantillas(nombre_plantilla,descripcion,tipo_documento,archivo_nombre,fecha_creacion,fecha_actualizacion,creado_por) VALUES(?,?,?,?,?,?,?)", (nom,desc,tipo,arch,now_txt(),now_txt(),'Admin_AQUA'))
            campos=CONTRATACION_CAMPOS_CORRESPONDENCIA
            for pid in [r['id'] for r in con.execute('SELECT id FROM contratacion_plantillas').fetchall()]:
                for nom, origen, td in campos:
                    con.execute("INSERT INTO contratacion_plantilla_campos(plantilla_id,descripcion,nombre_campo,campo_origen,tipo_dato) VALUES(?,?,?,?,?)", (pid,'',nom,origen,td))
        asegurar_carpetas_documentales()
        # Datos demo seguros
        if not con.execute("SELECT 1 FROM usuarios_admin WHERE usuario='admin'").fetchone():
            con.execute("INSERT INTO usuarios_admin(usuario,clave_hash,nombre,rol) VALUES(?,?,?,?)",
                        ("admin", generate_password_hash("admin123"), "PORTAL HR PRO", "admin"))
        if not con.execute("SELECT 1 FROM trabajadores WHERE dni='74324033'").fetchone():
            con.execute("INSERT INTO trabajadores(dni,nombre,correo,cargo,area,empresa,activo,fecha_registro) VALUES(?,?,?,?,?,?,?,?)",
                        ("74324033", "AZABACHE LUJAN, OMAR EDUARDO", "omar@demo.com", "Analista", "RR.HH.", "AQUANQA", 1, now_txt()))
        # Normaliza registros antiguos de demostración para que no aparezca PORTAL HR PRO SUPERFRUITS al trabajador.
        try:
            con.execute("UPDATE trabajadores SET empresa='AQUANQA' WHERE UPPER(COALESCE(empresa,''))='PORTAL HR PRO SUPERFRUITS'")
        except Exception:
            pass
        # Repara solicitudes antiguas que quedaron sin jefe_dni: toma el jefe desde saldos o ficha trabajadores.
        try:
            con.execute("""
                UPDATE vacaciones_solicitudes
                   SET jefe_dni = COALESCE(
                       (SELECT s.jefe_dni FROM vacaciones_saldos s
                         WHERE s.dni = vacaciones_solicitudes.dni
                           AND COALESCE(s.jefe_dni,'')<>''
                         ORDER BY s.periodo_inicio, s.periodo_fin LIMIT 1),
                       (SELECT tr.jefe_dni FROM trabajadores tr
                         WHERE tr.dni = vacaciones_solicitudes.dni
                           AND COALESCE(tr.jefe_dni,'')<>'' LIMIT 1),
                       jefe_dni
                   )
                 WHERE COALESCE(jefe_dni,'')=''
            """)
        except Exception:
            pass
        # Seguridad: cualquier solicitud registrada con inicio anterior a hoy queda anulada para no consumir saldo.
        try:
            hoy_txt = hoy_lima().isoformat()
            con.execute("""UPDATE vacaciones_solicitudes
                           SET estado='Anulado - fecha anterior a hoy',
                               comentario_gh=COALESCE(comentario_gh,'') || ' | Anulado automáticamente por fecha anterior a hoy.'
                         WHERE date(fecha_inicio) < date(?)
                           AND estado NOT LIKE 'Rechazado%'
                           AND estado NOT LIKE 'Anulado%'""", (hoy_txt,))
        except Exception:
            pass
        try:
            sincronizar_jefes_vacaciones(con)
        except Exception:
            pass
        con.commit()



def asegurar_plantillas_contratacion_base():
    """Carga/actualiza el catálogo base de Plantilla Documentos sin borrar las plantillas existentes."""
    plantillas_base = [
        ('CONTRATO INTERMITENTE PLANTA RENOVACION', 'CONTRATO INTERMITENTE PLANTA RENOVACION', 'CONTRATO TRABAJADOR(RENOVACIÓN)', 'CONDICIONES', 'CONTRATO INTEMRITENTE - PLANTA - RENOVACION AQI (OBP)'),
        ('CONTRATO INTERMITENTE RENOVACION', 'CONTRATO INTERMITENTE RENOVACION', 'CONTRATO TRABAJADOR(RENOVACIÓN)', 'CONDICIONES', 'CONTRATO INTERMITENTE - RENOVACION AQI (OBR)'),
        ('CONTRATO INTERMITENTE PLANTA', 'CONTRATO INTERMITENTE PLANTA', 'CONTRATO TRABAJADOR', 'CONDICIONES', 'CONTRATO INTEMRITENTE - PLANTA AQI (OBP)'),
        ('CONTRATO INTERMITENTE', 'CONTRATO INTERMITENTE', 'CONTRATO TRABAJADOR', 'CONDICIONES', 'CONTRATO INTEMRITENTE - PLANTA AQI (OBP)'),
        ('ACUERDO PREFERENCIAL', 'ACUERDO PREFERENCIAL', 'ACUERDO PREFERENCIAL', 'SIN CONDICIONES', 'CONTRATACION PREFERENCIAL AQUI'),
        ('AUTODECLARACION BUENAS PRACTICAS', 'AUTODECLARACION BUENAS PRACTICAS', 'AUTODECLARACION BUENAS PRACTICAS', 'SIN CONDICIONES', 'AUTODECLARACION BUENAS PRACTICAS 2025 AQU I'),
        ('CARGO ENTREGA RENOVACION', 'CARGO ENTREGA RENOVACION', 'CARGO ENTREGA RENOVACION', 'SIN CONDICIONES', 'CARGO DE ENTREGA - RENOVACION - AQI'),
        ('ELECCIÓN DE BENEFICIOS SOCIALES', 'ELECCIÓN DE BENEFICIOS SOCIALES', 'ELECCIÓN DE BENEFICIOS SOCIALES', 'SIN CONDICIONES', 'ELECCION DE BENEFICIOS SOCIALES AQI'),
        ('NOTA DE CARGO', 'NOTA DE CARGO', 'NOTA DE CARGO', 'SIN CONDICIONES', 'NOTA DE CARGO AQI'),
        ('CARTA DE COMPROMISO', 'CARTA DE COMPROMISO', 'CARTA DE COMPROMISO', 'SIN CONDICIONES', 'CARTA DE COMPROMISO - AQ I'),
        ('CARGO DE ENTREGA', 'CARGO DE ENTREGA', 'CARGO DE ENTREGA', 'SIN CONDICIONES', 'CARGO DE ENTREGA AQI'),
    ]
    campos_base=CONTRATACION_CAMPOS_CORRESPONDENCIA
    with db() as con:
        for nombre, desc, tipo_doc, condicion, archivo in plantillas_base:
            row = con.execute('SELECT id FROM contratacion_plantillas WHERE UPPER(nombre_plantilla)=UPPER(?) LIMIT 1', (nombre,)).fetchone()
            if row:
                pid = row['id']
                con.execute("UPDATE contratacion_plantillas SET descripcion=?, tipo_documento=?, esquema=?, condicion=?, version=COALESCE(NULLIF(version,''),'Version 01'), activo=COALESCE(activo,1), archivo_nombre=COALESCE(NULLIF(archivo_nombre,''),?), fecha_actualizacion=? WHERE id=?", (desc, tipo_doc, 'Trabajador Contrato Laboral', condicion, archivo, now_txt(), pid))
            else:
                cur = con.execute("INSERT INTO contratacion_plantillas(nombre_plantilla,descripcion,tipo_documento,esquema,condicion,version,activo,archivo_nombre,fecha_creacion,fecha_actualizacion,creado_por) VALUES(?,?,?,?,?,'Version 01',1,?,?,?,?)", (nombre, desc, tipo_doc, 'Trabajador Contrato Laboral', condicion, archivo, now_txt(), now_txt(), 'Admin_AQUA'))
                pid = cur.lastrowid
            for nom, origen, td in campos_base:
                existe = con.execute('SELECT 1 FROM contratacion_plantilla_campos WHERE plantilla_id=? AND nombre_campo=?', (pid, nom)).fetchone()
                if not existe:
                    con.execute("INSERT INTO contratacion_plantilla_campos(plantilla_id,descripcion,nombre_campo,campo_origen,tipo_dato,activo,requerido) VALUES(?,?,?,?,?,1,'SI')", (pid, '', nom, origen, td))
            if condicion == 'CONDICIONES':
                cond_count = con.execute('SELECT COUNT(*) FROM contratacion_plantilla_condiciones WHERE plantilla_id=?', (pid,)).fetchone()[0]
                if cond_count == 0:
                    con.execute("INSERT INTO contratacion_plantilla_condiciones(plantilla_id,nombre_campo,condicion,valor,activo,fecha_registro,creado_por) VALUES(?,?,?,?,?,?,?)", (pid,'Planilla','=','OBREROS RÉGIMEN AGRÍCOLA',1,now_txt(),'Admin_AQUA'))
                    con.execute("INSERT INTO contratacion_plantilla_condiciones(plantilla_id,nombre_campo,condicion,valor,activo,fecha_registro,creado_por) VALUES(?,?,?,?,?,?,?)", (pid,'Tipo Contrato','=','INTERMITENTE OBRERO',1,now_txt(),'Admin_AQUA'))
        tipos_doc = sorted({x[2] for x in plantillas_base})
        for tipo_doc in tipos_doc:
            if not con.execute('SELECT 1 FROM contratacion_tipos WHERE UPPER(descripcion)=UPPER(?) LIMIT 1', (tipo_doc,)).fetchone():
                con.execute('INSERT INTO contratacion_tipos(codigo,descripcion,etapa,obligatorio,activo) VALUES(?,?,?,?,1)', ('AUTO', tipo_doc, 'Plantilla Documentos', 1))
        con.commit()

init_db()
try:
    asegurar_ia_hr_pro()
except Exception as e:
    print('No se pudo inicializar IA HR PRO:', e)
try:
    with db() as con:
        temas_extra_ia = [
            ('Utilidades','Qué es utilidades','La participación en utilidades es un beneficio laboral que corresponde cuando la empresa genera renta y cumple las condiciones legales. Su pago depende de la actividad económica, el número de trabajadores y los resultados de la empresa.','D. Leg. N.° 892 y normas complementarias','La participación se calcula según porcentajes por actividad y reglas legales aplicables.','utilidades utilidad participacion trabajadores renta empresa pago'),
            ('Licencia por paternidad','Qué es licencia por paternidad','Es el permiso remunerado que corresponde al trabajador por el nacimiento de su hijo o hija, conforme a la ley aplicable.','Ley N.° 29409 y modificatorias','La duración puede variar según supuestos especiales previstos por la normativa.','licencia paternidad nacimiento hijo padre'),
            ('Licencia por fallecimiento','Qué es licencia por luto','Es la licencia otorgada al trabajador por fallecimiento de familiares directos, conforme a las reglas legales vigentes y políticas internas aplicables.','Ley N.° 31602 y normas complementarias','Aplica según parentesco y condiciones establecidas por la norma.','licencia luto fallecimiento familiar muerte'),
            ('Contratos','Qué es un contrato de trabajo','Es el acuerdo por el cual una persona presta servicios personales, remunerados y subordinados a favor de un empleador.','TUO del D. Leg. N.° 728 y normas complementarias','Los elementos principales son prestación personal, remuneración y subordinación.','contrato trabajo laboral tipos indeterminado plazo fijo sujeto modalidad')
        ]
        for tema,preg,resp,base,det,keys in temas_extra_ia:
            ex=con.execute('SELECT id FROM ia_legislacion_peru WHERE UPPER(pregunta_clave)=UPPER(?) LIMIT 1',(preg,)).fetchone()
            if not ex:
                con.execute('INSERT INTO ia_legislacion_peru(tema,pregunta_clave,respuesta,base_legal,detalle_legal,palabras_clave,activo,fecha_actualizacion) VALUES(?,?,?,?,?,?,1,?)',(tema,preg,resp,base,det,keys,now_txt()))
        con.commit()
except Exception as e:
    print('No se pudo cargar base legal adicional IA:', e)
# asegurar_plantillas_contratacion_base()  # módulo retirado
restaurar_trabajadores_desde_excel_si_db_vacia()
restaurar_vacaciones_desde_excel_si_db_vacia()
try:
    with db() as con:
        sincronizar_jefes_vacaciones(con)
        con.commit()
except Exception:
    pass
respaldar_exceles_locales()

def get_config(clave, default=''):
    with db() as con:
        r = con.execute('SELECT valor FROM app_config WHERE clave=?', (clave,)).fetchone()
    return r['valor'] if r else default

def set_config(clave, valor):
    with db() as con:
        con.execute('INSERT INTO app_config(clave,valor) VALUES(?,?) ON CONFLICT(clave) DO UPDATE SET valor=excluded.valor', (clave, str(valor)))
        con.commit()

def modo_prueba_activo():
    return get_config('modo_prueba', '0') == '1'

def marca_carga(usuario='sistema'):
    usuario = clean(usuario or 'sistema')
    return usuario + (' [MODO PRUEBA]' if modo_prueba_activo() else '')

def reset_intentos_login(dni):
    with db() as con:
        con.execute('DELETE FROM login_intentos WHERE dni=?', (normalizar_dni(dni),))
        con.commit()

def registrar_intento_fallido(dni):
    dni = normalizar_dni(dni)
    with db() as con:
        r = con.execute('SELECT intentos FROM login_intentos WHERE dni=?', (dni,)).fetchone()
        n = (int(r['intentos']) if r else 0) + 1
        bloqueado = 1 if n >= 3 else 0
        con.execute('INSERT INTO login_intentos(dni,intentos,bloqueado,ultima_fecha) VALUES(?,?,?,?) ON CONFLICT(dni) DO UPDATE SET intentos=?, bloqueado=?, ultima_fecha=?', (dni,n,bloqueado,now_txt(),n,bloqueado,now_txt()))
        con.commit()
    return n, bloqueado

def esta_bloqueado(dni):
    with db() as con:
        r = con.execute('SELECT bloqueado,intentos FROM login_intentos WHERE dni=?', (normalizar_dni(dni),)).fetchone()
    return bool(r and int(r['bloqueado'] or 0)==1), (int(r['intentos']) if r else 0)

# =============================
# SEGURIDAD / DECORADORES
# =============================
def admin_required(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not session.get("admin_id"):
            return redirect(url_for("admin_login"))
        return fn(*args, **kwargs)
    return wrapper


def worker_required(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not session.get("dni"):
            return redirect(url_for("login"))
        return fn(*args, **kwargs)
    return wrapper


def portal_required(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not session.get("dni") and not session.get("admin_id"):
            return redirect(url_for("login"))
        return fn(*args, **kwargs)
    return wrapper


# =============================
# IA HR PRO PLUS - RR.HH. PERÚ / SISTEMA / DATOS REALES
# =============================
# Mejora: administrador puede consultar TODO (admin + trabajador + legal + datos).
# Mejora: trabajador queda limitado a su información personal y consultas legales generales.
# Mejora: preguntas tipo "qué es..." priorizan base legal antes que datos documentales.
# Mejora: mayor base inicial de RR.HH. Perú, régimen privado y régimen agrario.

def ia_keywords_hit(q_norm, terms):
    return any(ia_norm(t) in q_norm for t in terms)


def ia_score(pregunta, row):
    """Score tolerante: coincidencia exacta + tokens + plurales simples."""
    qn = ia_norm(pregunta)
    q_tokens = [x for x in qn.split() if len(x) > 2]
    q = set(q_tokens)
    try:
        base = ' '.join([clean(row[k]) for k in row.keys() if k in ['pregunta_clave','palabras_clave','tema','modulo','respuesta','base_legal']])
    except Exception:
        base = str(row)
    bn = ia_norm(base)
    b = set([x for x in bn.split() if len(x) > 2])
    score = len(q & b)
    if qn and qn in bn:
        score += 12
    # boost por raíz simple para plurales: utilidad/utilidades, boleta/boletas, vacacion/vacaciones
    for tok in q:
        raiz = tok[:-2] if tok.endswith('es') else tok[:-1] if tok.endswith('s') else tok
        if len(raiz) >= 4 and raiz in bn:
            score += 1
    return score


def ia_hr_semillas_legales_plus():
    return [
        ('Boleta de pago','Qué es una boleta de pago','La boleta de pago es el documento que muestra al trabajador el detalle de su remuneración, descuentos, aportes, periodo pagado y otros conceptos de planilla. Sirve como constancia del pago efectuado.','D.S. N.° 001-98-TR y normas de planilla electrónica','Debe contener información suficiente para identificar al empleador, trabajador, periodo y conceptos pagados/descontados.','boleta boletas pago remuneracion sueldo descuentos aportes planilla recibo haberes'),
        ('Boleta de utilidades','Qué es boleta de utilidades','La boleta de utilidades es el documento o constancia donde se informa el pago de participación en utilidades cuando la empresa está obligada a distribuirlas. Muestra el importe pagado y, según el formato interno, puede detallar criterios como días laborados y remuneraciones computables.','D. Leg. N.° 892, D.S. N.° 009-98-TR y normas complementarias','En empresas obligadas, la participación se distribuye conforme a la actividad económica y reglas legales; en régimen agrario se consideran reglas especiales de Ley N.° 31110.','boleta utilidades utilidad participacion trabajadores renta empresa pago reparto'),
        ('Utilidades','Qué son utilidades','La participación en utilidades es un beneficio por el cual ciertos trabajadores participan en un porcentaje de la renta anual antes de impuestos de empresas que generan rentas de tercera categoría y cumplen condiciones legales.','D. Leg. N.° 892, D.S. N.° 009-98-TR y modificatorias','El porcentaje depende de la actividad económica. La distribución se realiza, de forma general, en función a días laborados y remuneraciones.','utilidades utilidad participacion trabajadores renta tercera categoria porcentaje reparto'),
        ('CTS','Qué es CTS','La Compensación por Tiempo de Servicios (CTS) es un beneficio social de previsión frente al cese. Funciona como un fondo de protección económica para el trabajador cuando termina la relación laboral.','TUO de la Ley de CTS - D.S. N.° 001-97-TR y D.S. N.° 004-97-TR','Generalmente se deposita en mayo y noviembre cuando corresponde, según régimen y cumplimiento de requisitos.','cts compensacion tiempo servicios deposito mayo noviembre cese'),
        ('CTS','Cuándo se deposita la CTS','En el régimen privado general, la CTS se deposita semestralmente, normalmente hasta la primera quincena de mayo y noviembre, siempre que el trabajador cumpla los requisitos legales.','TUO de la Ley de CTS - D.S. N.° 001-97-TR','El tratamiento puede variar por régimen especial, tiempo parcial, microempresa o remuneración integral anual.','cuando depositan cts mayo noviembre fecha deposito'),
        ('Gratificación','Qué es gratificación','La gratificación legal es un beneficio social que se paga por Fiestas Patrias y Navidad, usualmente en julio y diciembre, cuando el trabajador cumple los requisitos legales.','Ley N.° 27735 y D.S. N.° 005-2002-TR','Puede existir gratificación trunca si cesa antes del semestre completo y cumple condiciones.','gratificacion gratificaciones julio diciembre fiestas patrias navidad'),
        ('Gratificación','Qué es gratificación trunca','La gratificación trunca es el pago proporcional que corresponde cuando el trabajador cesa antes de julio o diciembre, siempre que haya laborado al menos un mes completo dentro del semestre computable.','Ley N.° 27735 y D.S. N.° 005-2002-TR','Debe revisarse fecha de ingreso, cese y remuneración computable.','gratificacion trunca proporcional cese semestre'),
        ('Vacaciones','Cuándo gano vacaciones','En el régimen laboral privado general, el trabajador adquiere derecho a 30 días calendario de descanso vacacional por cada año completo de servicios, sujeto al récord vacacional.','D. Leg. N.° 713 y D.S. N.° 012-92-TR','El récord depende de la jornada semanal y días efectivamente laborados.','vacaciones cuando gano derecho record vacacional 30 dias'),
        ('Vacaciones','Qué es récord vacacional','El récord vacacional es el requisito de asistencia o días efectivos de labor dentro del año de servicios para acceder al descanso vacacional.','D. Leg. N.° 713 y D.S. N.° 012-92-TR','Se evalúa según jornada y supuestos legalmente considerados como días efectivos.','record vacacional dias laborados requisito asistencia'),
        ('Vacaciones','Qué son vacaciones truncas','Las vacaciones truncas son el pago proporcional por vacaciones generado al cese antes de cumplir un nuevo año completo, conforme al tiempo laborado y reglas aplicables.','D. Leg. N.° 713 y D.S. N.° 012-92-TR','Se calculan en liquidación de beneficios sociales.','vacaciones truncas cese liquidacion proporcional'),
        ('Vacaciones','Qué es indemnización vacacional','La indemnización vacacional puede corresponder cuando el trabajador no goza oportunamente su descanso vacacional dentro del plazo legal, según las reglas del régimen privado.','D. Leg. N.° 713','Debe evaluarse periodo vencido, oportunidad de goce y responsabilidad del empleador.','indemnizacion vacacional triple vacaciones vencidas no gozadas'),
        ('Liquidación','Qué es liquidación de beneficios sociales','La liquidación de beneficios sociales es el cálculo de conceptos pendientes al cese: remuneraciones, CTS, gratificación trunca, vacaciones truncas o pendientes y otros pagos que correspondan.','Normativa laboral peruana según cada beneficio','Depende de régimen laboral, motivo de cese, fecha de ingreso, fecha de cese y pagos previos.','liquidacion beneficios sociales renuncia cese despido calculo'),
        ('Asignación familiar','Qué es asignación familiar','La asignación familiar es un beneficio equivalente al 10% de la Remuneración Mínima Vital para trabajadores del régimen privado que cumplen los requisitos legales y tienen hijos a cargo según la norma.','Ley N.° 25129 y D.S. N.° 035-90-TR','El monto cambia cuando cambia la RMV. Debe sustentarse la condición familiar.','asignacion familiar hijo 10 rmv beneficio'),
        ('Jornada','Cuál es la jornada máxima de trabajo','La jornada máxima ordinaria en Perú es de 8 horas diarias o 48 horas semanales, salvo jornadas especiales, acumulativas o atípicas válidamente implementadas.','Constitución Política del Perú y D.S. N.° 007-2002-TR','RR.HH. debe revisar régimen, jornada, horario, asistencia y descansos.','jornada maxima 8 horas 48 semanales horario trabajo'),
        ('Horas extras','Qué son horas extras','Las horas extras son labores realizadas por encima de la jornada ordinaria. Deben ser voluntarias, registradas y pagadas con sobretasa legal o compensadas conforme a acuerdo válido.','D.S. N.° 007-2002-TR y reglamento','La sobretasa mínima usual es 25% para las dos primeras horas y 35% para las siguientes.','horas extras sobretiempo sobretasa 25 35'),
        ('Feriados','Qué pasa si trabajo feriado','Si se trabaja en feriado no laborable, puede corresponder descanso sustitutorio o pago adicional según la norma. RR.HH. debe verificar si existió compensación.','D. Leg. N.° 713','El tratamiento depende del feriado, programación y descanso sustitutorio.','feriado trabajado pago descanso sustitutorio triple'),
        ('Descanso semanal','Qué es descanso semanal obligatorio','Es el descanso remunerado semanal que, de forma general, debe ser de al menos 24 horas consecutivas, según la organización de la jornada.','D. Leg. N.° 713','Puede programarse según operación, turnos y régimen laboral.','descanso semanal obligatorio dso 24 horas'),
        ('Contrato','Qué es un contrato de trabajo','Es el acuerdo por el cual una persona presta servicios personales, remunerados y subordinados para un empleador. Si existen esos elementos, se configura relación laboral.','TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR','Los elementos clave son prestación personal, remuneración y subordinación.','contrato trabajo relacion laboral subordinacion remuneracion servicios'),
        ('Contrato','Cuántos tipos de contrato hay','En el régimen privado se suele distinguir contratos a plazo indeterminado y contratos sujetos a modalidad o plazo fijo. También existen contratos a tiempo parcial y modalidades especiales según actividad.','TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR','Los contratos sujetos a modalidad requieren causa objetiva y formalidades.','tipos contrato indeterminado plazo fijo sujeto modalidad parcial'),
        ('Contrato','Qué es contrato intermitente','Es un contrato sujeto a modalidad para actividades permanentes pero discontinuas, donde la prestación se activa según necesidad de la actividad.','TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR','Debe existir causa objetiva y respetarse la formalidad contractual.','contrato intermitente discontinuo campaña agro actividad'),
        ('Contrato','Qué es periodo de prueba','El periodo de prueba es la etapa inicial de la relación laboral durante la cual se evalúa la adaptación del trabajador al puesto. En general es de 3 meses, con reglas especiales para cargos de confianza o dirección.','TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR','Debe verificarse puesto, convenio y condiciones pactadas.','periodo prueba tres meses confianza direccion'),
        ('Licencia por paternidad','Qué es licencia por paternidad','Es el permiso remunerado que corresponde al trabajador padre por nacimiento de hijo o hija. La regla general vigente es 10 días calendario consecutivos, con ampliaciones en supuestos especiales.','Ley N.° 29409 modificada por Ley N.° 30807','Puede ser 20 o 30 días en supuestos especiales como nacimiento prematuro, parto múltiple, enfermedad congénita terminal, discapacidad severa o complicaciones graves de la madre.','licencia paternidad nacimiento hijo padre 10 dias parto multiple prematuro'),
        ('Licencia por paternidad','Cuántos días son de licencia por paternidad','La licencia por paternidad es de 10 días calendario consecutivos en parto natural o cesárea. Puede ampliarse a 20 o 30 días en casos especiales regulados.','Ley N.° 29409 modificada por Ley N.° 30807','RR.HH. debe validar el certificado y el supuesto aplicable.','dias licencia paternidad 10 20 30 parto cesarea'),
        ('Licencia por maternidad','Qué es licencia por maternidad','Es el descanso pre y postnatal de la trabajadora gestante. En Perú suma 98 días naturales, usualmente 49 días antes y 49 días después del parto, con ampliaciones en casos especiales.','Ley N.° 26644 y normas modificatorias','Puede ampliarse por nacimiento múltiple o niño con discapacidad, entre otros supuestos.','licencia maternidad prenatal postnatal 98 dias gestante'),
        ('Lactancia','Qué es permiso por lactancia','Es el derecho de la madre trabajadora a una hora diaria de permiso por lactancia materna hasta que su hijo cumpla un año, conforme a la ley.','Ley N.° 27240 y normas modificatorias','El horario debe coordinarse sin afectar el derecho.','lactancia permiso hora diaria madre hijo un año'),
        ('Licencia por fallecimiento','Qué es licencia por luto','Es la licencia remunerada por fallecimiento de familiares directos. En el sector privado comprende 5 días calendario por fallecimiento de cónyuge, padres, hijos o hermanos, con reglas de ampliación por traslado.','Ley N.° 31602 y D.S. N.° 013-2023-TR','El trabajador debe sustentar parentesco y fallecimiento; la extensión depende del traslado.','licencia luto fallecimiento muerte familiar conyuge padres hijos hermanos 5 dias'),
        ('Licencia por adopción','Qué es licencia por adopción','Es una licencia laboral vinculada al proceso de adopción, según los supuestos y requisitos establecidos por la ley.','Ley N.° 27409','RR.HH. debe validar resolución o documento sustentatorio aplicable.','licencia adopcion trabajador hijo adoptivo'),
        ('Licencia familiar grave','Qué es licencia por familiar grave','Es la licencia para trabajadores con familiares directos con enfermedad grave o terminal o que sufran accidente grave, bajo requisitos legales.','Ley N.° 30012 y D.S. N.° 008-2017-TR','Debe acreditarse el vínculo y la condición médica según reglamento.','licencia familiar enfermedad grave terminal accidente grave'),
        ('Descanso médico','Qué es descanso médico','Es la indicación de reposo emitida por profesional autorizado por incapacidad temporal. Sirve para justificar inasistencia y, según corresponda, gestionar subsidios.','Ley N.° 26790, normas EsSalud y procedimientos aplicables','RR.HH. debe validar certificado, CITT o canje, fechas y requisitos.','descanso medico certificado incapacidad temporal citt essalud'),
        ('Subsidio','Qué es subsidio por incapacidad temporal','Es una prestación económica vinculada a incapacidad temporal para el trabajo, conforme a reglas de EsSalud y requisitos de aportación/acreditación.','Ley N.° 26790 y normas EsSalud','El procedimiento depende de días, acreditación y documentación médica.','subsidio incapacidad temporal essalud descanso medico'),
        ('Subsidio','Qué es subsidio por maternidad','Es la prestación económica asociada al descanso por maternidad, sujeta a requisitos de acreditación y aportes ante EsSalud.','Ley N.° 26790 y normas EsSalud','Debe revisarse acreditación, vínculo laboral y descanso pre/postnatal.','subsidio maternidad essalud descanso prenatal postnatal'),
        ('AFP ONP','Qué es AFP y ONP','AFP y ONP son sistemas de pensiones. AFP corresponde al Sistema Privado de Pensiones y ONP al Sistema Nacional de Pensiones. El trabajador aporta según su afiliación.','Normativa del sistema pensionario peruano','RR.HH. debe validar afiliación, comisión, aportes y declaración en planilla.','afp onp pension aporte descuento sistema previsional'),
        ('Seguro Vida Ley','Qué es Vida Ley','El Seguro Vida Ley es un seguro obligatorio de vida a favor del trabajador, contratado por el empleador según los supuestos y coberturas reguladas.','D. Leg. N.° 688 y normas complementarias','Otorga coberturas por fallecimiento o invalidez según la póliza y norma.','vida ley seguro vida obligatorio trabajador'),
        ('SCTR','Qué es SCTR','El Seguro Complementario de Trabajo de Riesgo cubre prestaciones de salud y pensiones por accidentes de trabajo o enfermedades profesionales en actividades de alto riesgo.','Ley N.° 26790 y normas complementarias','Aplica según actividad económica y riesgo.','sctr seguro complementario trabajo riesgo salud pension'),
        ('SST','Qué es SST','SST significa Seguridad y Salud en el Trabajo. Es el sistema de prevención de riesgos laborales para proteger la vida, salud e integridad de los trabajadores.','Ley N.° 29783 y reglamento','Incluye IPERC, capacitaciones, comité/supervisor, investigación de accidentes y medidas preventivas.','sst seguridad salud trabajo prevencion riesgos'),
        ('IPERC','Qué es IPERC','IPERC es la Identificación de Peligros, Evaluación de Riesgos y Controles. Sirve para reconocer peligros, valorar riesgos y definir controles preventivos.','Ley N.° 29783 y D.S. N.° 005-2012-TR','Debe actualizarse según cambios de proceso, incidentes o condiciones de trabajo.','iperc peligros riesgos controles matriz sst'),
        ('Accidente de trabajo','Qué es accidente de trabajo','Es todo suceso repentino que sobreviene por causa o con ocasión del trabajo y que produce lesión, perturbación funcional, invalidez o muerte, según la normativa de SST.','Ley N.° 29783 y reglamento','Debe investigarse, registrarse y reportarse según corresponda.','accidente trabajo lesion investigacion sst'),
        ('Incidente','Qué es incidente de trabajo','Es un suceso relacionado con el trabajo que pudo causar daño o lesión, aunque no necesariamente haya producido una lesión. Sirve para prevenir accidentes.','Ley N.° 29783 y reglamento','Debe analizarse para tomar acciones preventivas.','incidente trabajo casi accidente sst'),
        ('Hostigamiento sexual','Qué es hostigamiento sexual laboral','Es una forma de violencia que puede presentarse en el ámbito laboral mediante conductas de naturaleza sexual o sexista no deseadas. Debe atenderse mediante procedimiento interno y medidas de protección.','Ley N.° 27942 y normas complementarias','El empleador debe prevenir, investigar y sancionar conforme al procedimiento aplicable.','hostigamiento sexual laboral acoso procedimiento comite'),
        ('Teletrabajo','Qué es teletrabajo','El teletrabajo es una modalidad especial de prestación de servicios usando tecnologías de la información, fuera del centro de trabajo o con esquema híbrido, según acuerdo y normativa.','Ley N.° 31572 y reglamento','Debe regularse mediante acuerdo, condiciones, seguridad de información y SST.','teletrabajo trabajo remoto hibrido ley 31572'),
        ('Régimen general','Qué es régimen laboral general privado','Es el régimen laboral aplicable a trabajadores de la actividad privada bajo el marco del D. Leg. N.° 728 y normas complementarias, salvo regímenes especiales.','TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR','Incluye reglas generales sobre contratación, beneficios, jornada y extinción.','regimen general privado 728 actividad privada'),
        ('Régimen MYPE','Qué es régimen MYPE','Es un régimen especial para micro y pequeña empresa inscrita según la normativa aplicable. Los beneficios laborales pueden variar respecto del régimen general.','TUO de la Ley MYPE y normas complementarias','Debe verificarse inscripción REMYPE, fecha y categoría de empresa.','regimen mype microempresa pequena empresa remype beneficios'),
        ('Régimen agrario','Qué es régimen laboral agrario','Es un régimen especial para actividades agrarias y agroindustriales comprendidas en Ley N.° 31110. Tiene reglas sobre remuneración básica, remuneración diaria, CTS/gratificaciones, BETA, jornada y otros derechos.','Ley N.° 31110 y D.S. N.° 005-2021-MIDAGRI','No todo personal está comprendido; debe revisarse actividad, área y alcance de la ley.','regimen agrario ley 31110 trabajador agrario agroindustrial beta rd rb'),
        ('Régimen agrario','Qué es BETA agrario','La BETA es la Bonificación Especial por Trabajo Agrario equivalente al 30% de la RMV, con carácter no remunerativo y no pensionable, según Ley N.° 31110.','Ley N.° 31110 y D.S. N.° 005-2021-MIDAGRI','Puede pagarse mensual o proporcionalmente según días laborados, conforme a la norma.','beta bonificacion especial trabajo agrario 30 rmv'),
        ('Régimen agrario','Cómo se paga CTS y gratificación en régimen agrario','En régimen agrario, la RB no puede ser menor a la RMV; las gratificaciones equivalen a 16.66% de la RB y la CTS a 9.72% de la RB, pudiendo integrarse en la remuneración diaria o pagarse en oportunidades generales si el trabajador lo elige conforme a la norma.','Ley N.° 31110 y D.S. N.° 005-2021-MIDAGRI','La decisión debe comunicarse por escrito dentro del plazo establecido en el reglamento.','cts gratificacion regimen agrario 9.72 16.66 remuneracion diaria'),
        ('CAS','Qué es régimen CAS','El CAS es un régimen de contratación administrativa de servicios usado en el sector público, con reglas propias y diferentes al régimen laboral privado.','D. Leg. N.° 1057 y normas complementarias','No debe confundirse con el régimen privado D. Leg. 728.','cas contrato administrativo servicios sector publico'),
        ('Practicantes','Qué son modalidades formativas laborales','Son convenios formativos como prácticas preprofesionales o profesionales, orientados al aprendizaje y formación, no equivalentes automáticamente a contrato laboral si cumplen la ley.','Ley N.° 28518 y D.S. N.° 007-2005-TR','Deben respetarse subvención, jornada formativa, convenio y plan de aprendizaje.','practicantes practicas preprofesionales profesionales modalidades formativas'),
        ('Intermediación','Qué es intermediación laboral','Es la provisión de personal mediante empresas autorizadas para servicios temporales, complementarios o especializados, bajo reglas legales específicas.','Ley N.° 27626 y D.S. N.° 003-2002-TR','Debe verificarse registro y límites legales.','intermediacion laboral services cooperativas tercerizacion'),
        ('Tercerización','Qué es tercerización laboral','Es la contratación de una empresa para realizar una actividad con autonomía empresarial, recursos propios y asunción de riesgos, conforme a la normativa.','Ley N.° 29245, D. Leg. N.° 1038 y reglamento','No debe usarse para simple provisión de personal.','tercerizacion outsourcing autonomia empresarial'),
        ('Remuneración','Qué es remuneración computable','Es la remuneración que se toma como base para calcular determinados beneficios, según la naturaleza del concepto y la norma de cada beneficio.','Normativa laboral peruana según beneficio','No todos los conceptos pagados son computables para todos los beneficios.','remuneracion computable calculo beneficios cts gratificacion'),
        ('Conceptos no remunerativos','Qué conceptos no son remunerativos','Algunos conceptos pueden no tener carácter remunerativo por ley o por su naturaleza, como ciertas condiciones de trabajo, movilidad supeditada, gratificaciones extraordinarias, entre otros, según evaluación.','TUO del D. Leg. N.° 728 y normas complementarias','Debe analizarse regularidad, libre disposición y finalidad del pago.','conceptos no remunerativos movilidad condiciones trabajo'),
        ('Despido','Qué es despido arbitrario','Es la extinción unilateral del vínculo por el empleador sin causa justa o sin cumplir el procedimiento legal. Puede generar indemnización según el régimen aplicable.','TUO del D. Leg. N.° 728 - D.S. N.° 003-97-TR','Debe revisarse tipo de contrato, antigüedad, remuneración y causal.','despido arbitrario indemnizacion causa justa'),
        ('Renuncia','Qué pasa si renuncio','La renuncia es la decisión del trabajador de terminar el vínculo laboral. Usualmente debe comunicarse con anticipación o solicitar exoneración del plazo. Corresponde liquidación de beneficios generados.','TUO del D. Leg. N.° 728 y normas complementarias','RR.HH. debe calcular beneficios pendientes hasta la fecha de cese.','renuncia voluntaria cese carta liquidacion 30 dias'),
        ('Suspensión perfecta','Qué es suspensión perfecta','Es la suspensión temporal de la obligación de trabajar y pagar remuneración, bajo supuestos legales específicos. Debe estar sustentada y cumplir requisitos.','TUO del D. Leg. N.° 728 y normas especiales','Debe evaluarse caso concreto y formalidades.','suspension perfecta contrato remuneracion temporal'),
        ('Reglamento interno','Qué es reglamento interno de trabajo','Es el documento interno que contiene reglas de orden laboral, deberes, derechos, disciplina, horarios y procedimientos dentro de la empresa, cuando corresponde.','D.S. N.° 039-91-TR','Debe difundirse y aplicarse respetando la ley.','reglamento interno trabajo rit normas empresa'),
        ('Sindicato','Qué es libertad sindical','Es el derecho de trabajadores a organizarse, afiliarse, constituir sindicatos y realizar actividad sindical conforme a la ley.','Constitución, convenios OIT y normas laborales peruanas','El empleador debe respetar la actividad sindical lícita.','sindicato libertad sindical afiliacion negociacion colectiva'),
        ('Inspección laboral','Qué es SUNAFIL','SUNAFIL es la entidad inspectiva encargada de fiscalizar el cumplimiento de normas sociolaborales y de seguridad y salud en el trabajo.','Ley N.° 29981 y normas inspectivas','Puede realizar actuaciones inspectivas y proponer sanciones ante incumplimientos.','sunafil inspeccion laboral fiscalizacion multa'),
        ('Boleta vacaciones','Qué es boleta de vacaciones','Es el documento de pago asociado a remuneración vacacional o concepto vinculado al descanso vacacional, según el proceso de planilla de la empresa.','D. Leg. N.° 713 y normas de planillas','Debe corresponder al periodo, monto y trabajador correcto.','boleta vacaciones pago vacacional remuneracion vacacional'),
        ('Boleta CTS','Qué es boleta CTS','Es el documento o constancia asociada al depósito o cálculo de CTS. Informa el beneficio calculado para el periodo correspondiente.','D.S. N.° 001-97-TR y normas de planilla','Puede estar vinculada al depósito semestral o liquidación, según el caso.','boleta cts deposito compensacion tiempo servicios'),
        ('Boleta gratificación','Qué es boleta de gratificación','Es el documento que informa el pago de la gratificación legal de julio o diciembre, o gratificación trunca si corresponde.','Ley N.° 27735 y normas de planilla','Debe reflejar monto, periodo y bonificación extraordinaria cuando aplique.','boleta gratificacion julio diciembre bonificacion extraordinaria'),
    ]


def ia_hr_semillas_sistema_plus():
    return [
        ('documental','admin','Dónde cargar boletas de pago','Para cargar boletas ingresa a Gestión Documental > Documentos de pago o Subir / gestionar documentos. Selecciona tipo: Normal, CTS, Gratificación, Utilidad, Vacaciones o Liquidación, periodo y archivo.','/admin/documentos','boleta boletas pago cargar subir normal cts gratificacion utilidad liquidacion documentos pago'),
        ('documental','admin','Dónde cargar boletas utilidades','Ingresa a Gestión Documental > Documentos de pago > Utilidad. También puedes cargar desde Subir / gestionar documentos seleccionando tipo Utilidad y periodo correspondiente.','/admin/documentos?tipo=Utilidad','cargar boletas utilidades utilidad subir documentos pago'),
        ('documental','admin','Dónde sincronizar PDFs por DNI','Ingresa a Gestión Documental > Detectar PDFs o Sincronizar carpetas. El sistema revisa la carpeta DOCUMENTOS_PRIZE_AUTO, identifica DNI y registra documentos vinculados al trabajador.','/admin/sincronizar','detectar pdf dni sincronizar carpeta documentos prize auto'),
        ('documental','admin','Dónde crear carpetas documentales','Ingresa a Gestión Documental > Crear carpetas. El sistema arma la estructura local para documentos de pago, empresa y personales.','/admin/crear_carpetas','crear carpetas documental estructura documentos'),
        ('documental','admin','Dónde descargar plantilla documental','Ingresa a Gestión Documental > Plantilla Documental. Descarga el Excel base para organizar cargas de documentos.','/admin/plantilla_gestion/documental','plantilla documental excel descargar'),
        ('documental','admin','Qué documentos están pendientes','Consulta Gestión Documental > Subir / gestionar documentos o usa la IA preguntando por resumen documental. Verás pendientes, leídos, aceptados, firmados y rechazados.','/admin/documentos','documentos pendientes lectura aceptacion firma rechazo'),
        ('documental','admin','Cómo revisar documentos rechazados','Ingresa a Gestión Documental > Subir / gestionar documentos y filtra por estado Rechazado. Revisa comentario, trabajador, tipo y periodo.','/admin/documentos','documentos rechazados comentario estado filtro'),
        ('documental','trabajador','Dónde veo mis boletas','Ingresa a Gestión Documental > Documentos de pago. Ahí verás tus boletas normales, CTS, gratificaciones, utilidades, vacaciones o liquidaciones cargadas para tu DNI.','/panel','mis boletas ver descargar pago cts gratificacion utilidad liquidacion'),
        ('documental','trabajador','No veo mi boleta','Si no visualizas tu boleta, puede que RR.HH. aún no la haya cargado o que el archivo no esté vinculado a tu DNI. Revisa tipo y periodo; si no aparece, consulta con RR.HH.','/panel','no veo mi boleta no aparece documento'),
        ('documental','ambos','Qué significa documento pendiente','Un documento pendiente es aquel que fue cargado, pero aún no registra lectura, aceptación, firma o aprobación final, según el flujo configurado por RR.HH.','','documento pendiente estado lectura firma aprobacion'),
        ('vacacional','admin','Dónde cargar saldos vacacionales','Ingresa a Gestión Vacacional > Saldos Vacacionales o Dashboard Vacacional. Desde ahí cargas Excel con DNI, periodo, días ganados, gozados y saldo.','/admin/vacaciones#cargar-saldos','cargar saldos vacaciones vacacionales excel saldo'),
        ('vacacional','admin','Dónde aprobar vacaciones','Ingresa a Gestión Vacacional > Aprobaciones. Revisa solicitudes pendientes de jefe y Gestión Humana, según el flujo.','/admin/vacaciones#aprobaciones','aprobar vacaciones aprobaciones jefe gestion humana'),
        ('vacacional','admin','Dónde ver reportes vacacionales','Ingresa a Gestión Vacacional > Reportes. Revisa solicitudes, saldos, aprobadas, rechazadas y pendientes.','/admin/vacaciones#reportes','reportes vacaciones vacacional indicadores'),
        ('vacacional','admin','Cómo anular solicitud de vacaciones','Ingresa a Gestión Vacacional > Solicitudes o Aprobaciones, ubica el registro, revisa estado y usa la acción de anulación si está habilitada para administrador.','/admin/vacaciones','anular solicitud vacaciones eliminar rechazar'),
        ('vacacional','trabajador','Cómo solicito vacaciones','Ingresa a Gestión Vacacional > Saldo y solicitud. Selecciona periodo con saldo, registra fechas y envía la solicitud para aprobación.','/vacaciones/mi_solicitud#solicitar','solicitar vacaciones pedido solicitud saldo'),
        ('vacacional','trabajador','Dónde veo mi saldo vacacional','Ingresa a Gestión Vacacional > Dashboard vacacional. Ahí verás periodos, días ganados, gozados y saldo disponible.','/vacaciones/mi_solicitud','mi saldo vacaciones cuantos dias tengo'),
        ('vacacional','trabajador','Dónde veo el estado de mi solicitud','Ingresa a Gestión Vacacional > Dashboard vacacional. En Mis solicitudes verás si está pendiente, aprobada, rechazada o anulada.','/vacaciones/mi_solicitud','estado solicitud vacaciones pendiente aprobada rechazada'),
        ('portal','ambos','Qué puede hacer la IA HR','La IA HR ayuda con dudas del sistema, gestión documental, boletas, vacaciones, datos reales y conceptos laborales peruanos. La respuesta se filtra por rol.','','ia asistente ayuda portal rrhh'),
        ('portal','trabajador','Por qué no puedo cargar base de trabajadores','Esa opción es solo para administradores. Como trabajador puedes consultar tus documentos, boletas, vacaciones y estado de tus solicitudes.','','cargar base trabajadores no puedo'),
        ('portal','admin','Qué puede consultar el administrador','El administrador puede consultar todo: manual del sistema, datos documentales, datos vacacionales, consultas legales y preguntas orientadas al trabajador para validar cómo respondería el portal.','','administrador consultar todo roles trabajador legal datos'),
    ]


def asegurar_ia_hr_pro():
    """Crea y precarga IA HR. No borra respuestas editadas; agrega/actualiza catálogo base amplio."""
    with db() as con:
        con.execute("""CREATE TABLE IF NOT EXISTS ia_base_conocimiento_hr(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            modulo TEXT,
            rol_destino TEXT DEFAULT 'ambos',
            pregunta_clave TEXT,
            respuesta TEXT,
            ruta TEXT,
            palabras_clave TEXT,
            activo INTEGER DEFAULT 1,
            fecha_registro TEXT,
            registrado_por TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS ia_legislacion_peru(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tema TEXT,
            pregunta_clave TEXT,
            respuesta TEXT,
            base_legal TEXT,
            detalle_legal TEXT,
            palabras_clave TEXT,
            activo INTEGER DEFAULT 1,
            fecha_actualizacion TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS ia_hr_log(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rol TEXT,
            modulo TEXT,
            dni TEXT,
            pregunta TEXT,
            respuesta_tipo TEXT,
            fecha TEXT,
            usuario TEXT
        )""")
        # Índices útiles para evitar duplicados y acelerar búsqueda
        try: con.execute('CREATE INDEX IF NOT EXISTS idx_ia_leg_activo ON ia_legislacion_peru(activo, tema)')
        except Exception: pass
        try: con.execute('CREATE INDEX IF NOT EXISTS idx_ia_base_activo ON ia_base_conocimiento_hr(activo, rol_destino, modulo)')
        except Exception: pass
        for modulo, rol, preg, resp, ruta, keys in ia_hr_semillas_sistema_plus():
            existe = con.execute('SELECT id FROM ia_base_conocimiento_hr WHERE UPPER(pregunta_clave)=UPPER(?) AND rol_destino=? LIMIT 1', (preg, rol)).fetchone()
            if not existe:
                con.execute("""INSERT INTO ia_base_conocimiento_hr(modulo,rol_destino,pregunta_clave,respuesta,ruta,palabras_clave,activo,fecha_registro,registrado_por)
                               VALUES(?,?,?,?,?,?,1,?,?)""", (modulo,rol,preg,resp,ruta,keys,now_txt(),'SISTEMA'))
        for tema, preg, resp, base, det, keys in ia_hr_semillas_legales_plus():
            existe = con.execute('SELECT id FROM ia_legislacion_peru WHERE UPPER(pregunta_clave)=UPPER(?) LIMIT 1', (preg,)).fetchone()
            if not existe:
                con.execute("""INSERT INTO ia_legislacion_peru(tema,pregunta_clave,respuesta,base_legal,detalle_legal,palabras_clave,activo,fecha_actualizacion)
                               VALUES(?,?,?,?,?,?,1,?)""", (tema,preg,resp,base,det,keys,now_txt()))
            else:
                # actualiza claves y respuesta base solo si mantiene activo, para mejorar coincidencias de versiones previas
                con.execute("""UPDATE ia_legislacion_peru SET tema=?, respuesta=?, base_legal=?, detalle_legal=?, palabras_clave=?, fecha_actualizacion=? WHERE id=?""",
                            (tema, resp, base, det, keys, now_txt(), existe['id']))
        con.commit()


def ia_buscar_base_conocimiento(pregunta, rol, modulo=''):
    try:
        with db() as con:
            if rol == 'admin':
                rows = con.execute("SELECT * FROM ia_base_conocimiento_hr WHERE activo=1").fetchall()
            else:
                rows = con.execute("""SELECT * FROM ia_base_conocimiento_hr
                                      WHERE activo=1 AND rol_destino IN ('ambos', ?)""", (rol,)).fetchall()
        if not rows:
            return None
        ranked = sorted(rows, key=lambda r: ia_score(pregunta + ' ' + modulo, r), reverse=True)
        top_score = ia_score(pregunta + ' ' + modulo, ranked[0]) if ranked else 0
        if ranked and top_score >= 2:
            r = ranked[0]
            ruta = f"<p><b>Ruta sugerida:</b> <a href='{h(r['ruta'])}'>{h(r['ruta'])}</a></p>" if clean(r['ruta']) else ''
            visible_rol = 'Administrador / Trabajador' if r['rol_destino']=='ambos' else r['rol_destino'].capitalize()
            return f"<div class='ia-result ok'><h3>{h(r['pregunta_clave'])}</h3><p>{h(r['respuesta'])}</p>{ruta}<small>Base sistema · Visible para: {h(visible_rol)}</small></div>"
    except Exception as e:
        return f"<div class='ia-result bad'><h3>Error base IA</h3><p>{h(e)}</p></div>"
    return None


def ia_buscar_legal(pregunta):
    try:
        with db() as con:
            rows = con.execute('SELECT * FROM ia_legislacion_peru WHERE activo=1').fetchall()
        if not rows:
            return None
        ranked = sorted(rows, key=lambda r: ia_score(pregunta, r), reverse=True)
        top_score = ia_score(pregunta, ranked[0]) if ranked else 0
        # Umbral bajo, pero con boost de frases y tokens evita fallback genérico.
        if ranked and top_score >= 1:
            r = ranked[0]
            return f"""
            <div class='ia-result legal'>
              <h3>{h(r['pregunta_clave'])}</h3>
              <p>{h(r['respuesta'])}</p>
              <p><b>Base legal referencial:</b> {h(r['base_legal'])}</p>
              <p class='muted'>{h(r['detalle_legal'])}</p>
              <small>Respuesta informativa. RR.HH. debe validar régimen, jornada, fecha, documentos y caso concreto antes de aplicar.</small>
            </div>"""
    except Exception as e:
        return f"<div class='ia-result bad'><h3>Error legal IA</h3><p>{h(e)}</p></div>"
    return None


def ia_resumen_documental(rol, dni=''):
    with db() as con:
        if rol == 'admin':
            total = ia_sql_count(con, 'SELECT COUNT(*) FROM documentos')
            pendientes = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE COALESCE(estado,'Pendiente') IN ('Pendiente','Aceptado','Firmado')")
            rechazados = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE estado='Rechazado'")
            leidos = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE COALESCE(fecha_lectura,'')<>' '")
            rows = ia_sql_rows(con, 'SELECT tipo, COUNT(*) c FROM documentos GROUP BY tipo ORDER BY c DESC LIMIT 6')
            detalle = ''.join([f"<li>{h(r['tipo'])}: <b>{r['c']}</b></li>" for r in rows]) or '<li>Sin documentos cargados.</li>'
            return f"<div class='ia-result ok'><h3>Resumen documental administrador</h3><p>Total: <b>{total}</b> · Pendientes: <b>{pendientes}</b> · Leídos: <b>{leidos}</b> · Rechazados: <b>{rechazados}</b></p><ul>{detalle}</ul><p>Ruta: Gestión Documental > Subir / gestionar documentos.</p></div>"
        dni = normalizar_dni(dni)
        total = ia_sql_count(con, 'SELECT COUNT(*) FROM documentos WHERE dni=?', (dni,))
        leidos = ia_sql_count(con, "SELECT COUNT(*) FROM documentos WHERE dni=? AND COALESCE(fecha_lectura,'')<>''", (dni,))
        rows = ia_sql_rows(con, 'SELECT tipo, periodo, estado, fecha_subida FROM documentos WHERE dni=? ORDER BY id DESC LIMIT 6', (dni,))
        detalle = ''.join([f"<li>{h(r['tipo'])} · {h(r['periodo'])} · {h(r['estado'] or 'Pendiente')}</li>" for r in rows]) or '<li>Aún no tienes documentos cargados para tu DNI.</li>'
        return f"<div class='ia-result ok'><h3>Mis documentos</h3><p>Tienes <b>{total}</b> documento(s) cargado(s). Leídos: <b>{leidos}</b>.</p><ul>{detalle}</ul><p>Ruta: Gestión Documental > Documentos de pago / empresa / personales.</p></div>"


def ia_hr_es_pregunta_conceptual(q):
    return ia_keywords_hit(q, ['que es','qué es','que son','qué son','concepto','definicion','definición','explicame','explícame','cuantos tipos','cuántos tipos','cuando gano','cuándo gano','cuantos dias','cuántos días','cual es','cuál es'])


def ia_hr_responder(pregunta, modulo='', rol=None, dni=None):
    pregunta = clean(pregunta)
    rol = rol or ('admin' if session.get('admin_id') else 'trabajador')
    dni = normalizar_dni(dni or session.get('dni') or '')
    q = ia_norm(pregunta)
    if not pregunta:
        return "<div class='ia-result warn'><h3>IA HR</h3><p>Escribe una pregunta. Ejemplo: ¿qué es CTS?, ¿dónde veo mis boletas?, ¿cuánto saldo tengo?</p></div>", 'vacio'

    # Trabajador: bloquear consultas administrativas. Admin: acceso total.
    admin_terms = ['cargar base','cargar trabajadores','subir saldos','cargar saldos','subir boletas masiva','sincronizar pdf','crear carpetas','reporte general','todos los trabajadores','eliminar trabajador','base trabajadores']
    if rol != 'admin' and any(t in q for t in admin_terms):
        return "<div class='ia-result bad'><h3>Acceso limitado</h3><p>Esa consulta corresponde a opciones administrativas. Desde tu portal puedes consultar tus boletas, documentos, saldo vacacional, solicitudes y conceptos laborales.</p></div>", 'bloqueo_rol'

    # 1) Conceptos legales primero para preguntas tipo "qué es...". Evita que "qué es boletas utilidades" caiga en resumen documental.
    if ia_hr_es_pregunta_conceptual(q):
        r = ia_buscar_legal(pregunta)
        if r:
            return r, 'legal_peru'

    # 2) Sistema/manual: dónde cargo, dónde veo, cómo hago.
    r = ia_buscar_base_conocimiento(pregunta, rol, modulo)
    if r:
        return r, 'base_conocimiento'

    # 3) Legal general aunque no pregunte "qué es".
    legal_terms = ['cts','gratificacion','gratificaciones','utilidad','utilidades','vacacion','vacaciones','liquidacion','boleta','paternidad','maternidad','luto','fallecimiento','lactancia','subsidio','afp','onp','vida ley','sctr','jornada','horas extras','feriado','contrato','agrario','beta','mype','cas','practicante','sst','iperc','hostigamiento','teletrabajo','asignacion familiar']
    if ia_keywords_hit(q, legal_terms):
        r = ia_buscar_legal(pregunta)
        if r:
            return r, 'legal_peru'

    # 4) Datos reales del sistema.
    if any(x in q for x in ['mis boletas','mis documentos','documentos','boletas cargadas','documentos pendientes','documentos rechazados','resumen documental']) and any(x in q for x in ['tengo','ver','donde','mis','cargadas','aparece','pendiente','resumen','cuantos','cuántos']):
        return ia_resumen_documental(rol, dni), 'datos_documental'
    if any(x in q for x in ['vacacion','saldo','solicitud','aprobacion','dias']) and any(x in q for x in ['tengo','mis','estado','pendiente','resumen','cuanto','cuantos','cuánto','cuántos']):
        return ia_resumen_vacacional(rol, dni), 'datos_vacacional'

    # 5) Fallback por módulo.
    if 'vac' in ia_norm(modulo):
        return ia_resumen_vacacional(rol, dni), 'fallback_vacacional'
    if 'doc' in ia_norm(modulo) or 'panel' in ia_norm(modulo):
        return ia_resumen_documental(rol, dni), 'fallback_documental'
    if rol == 'admin':
        return "<div class='ia-result ok'><h3>IA HR Administrador</h3><p>Tienes acceso a preguntas de administrador, trabajador, datos reales y conceptos legales. Prueba: ¿qué es boleta de utilidades?, ¿dónde cargo saldos?, ¿qué es licencia por paternidad?, ¿cuántos documentos están pendientes?</p></div>", 'fallback'
    return "<div class='ia-result ok'><h3>IA HR Trabajador</h3><p>Puedo ayudarte a revisar tus boletas, documentos, vacaciones, solicitudes y conceptos laborales. Prueba: ¿qué es CTS?, ¿cuánto saldo tengo?, ¿dónde veo mi boleta?</p></div>", 'fallback'

# Inicialización PRO PLUS después de redefinir funciones.
try:
    asegurar_ia_hr_pro()
except Exception as e:
    print('No se pudo inicializar IA HR PRO PLUS:', e)


@app.route('/ia_hr/api', methods=['POST'])
@app.route('/ia_hr_ask', methods=['POST'])
@app.route('/api/ia_hr', methods=['POST'])
@portal_required
def ia_hr_api():
    try:
        asegurar_ia_hr_pro()
    except Exception:
        pass
    data = request.get_json(silent=True) or request.form or {}
    pregunta = clean(data.get('pregunta'))
    modulo = clean(data.get('modulo'))
    rol = 'admin' if session.get('admin_id') else 'trabajador'
    dni = session.get('dni') or clean(data.get('dni'))
    try:
        html_resp, tipo = ia_hr_responder(pregunta, modulo, rol, dni)
    except Exception as e:
        html_resp = f"<div class='ia-result bad'><h3>Error IA HR</h3><p>No se pudo procesar la consulta.</p><small>{h(e)}</small></div>"
        tipo = 'error'
    ia_registrar_log(rol, modulo, normalizar_dni(dni), pregunta, tipo)
    return jsonify({'ok': True, 'html': html_resp, 'tipo': tipo, 'rol': rol})


@app.route('/admin/ia_hr', methods=['GET','POST'])
@admin_required
def admin_ia_hr():
    if request.method == 'POST':
        tipo = clean(request.form.get('tipo'))
        with db() as con:
            if tipo == 'legal':
                con.execute("""INSERT INTO ia_legislacion_peru(tema,pregunta_clave,respuesta,base_legal,detalle_legal,palabras_clave,activo,fecha_actualizacion)
                               VALUES(?,?,?,?,?,?,1,?)""", (clean(request.form.get('tema')), clean(request.form.get('pregunta_clave')), clean(request.form.get('respuesta')), clean(request.form.get('base_legal')), clean(request.form.get('detalle_legal')), clean(request.form.get('palabras_clave')), now_txt()))
            else:
                con.execute("""INSERT INTO ia_base_conocimiento_hr(modulo,rol_destino,pregunta_clave,respuesta,ruta,palabras_clave,activo,fecha_registro,registrado_por)
                               VALUES(?,?,?,?,?,?,1,?,?)""", (clean(request.form.get('modulo')), clean(request.form.get('rol_destino') or 'ambos'), clean(request.form.get('pregunta_clave')), clean(request.form.get('respuesta')), clean(request.form.get('ruta')), clean(request.form.get('palabras_clave')), now_txt(), session.get('admin_nombre','admin')))
            con.commit()
        flash('Respuesta IA registrada correctamente.', 'ok')
        return redirect(url_for('admin_ia_hr'))
    asegurar_ia_hr_pro()
    with db() as con:
        sistema = con.execute('SELECT * FROM ia_base_conocimiento_hr WHERE activo=1 ORDER BY modulo, rol_destino, id DESC').fetchall()
        legal = con.execute('SELECT * FROM ia_legislacion_peru WHERE activo=1 ORDER BY tema, id DESC').fetchall()
        logs = con.execute('SELECT * FROM ia_hr_log ORDER BY id DESC LIMIT 30').fetchall()
    rows_s = ''.join([f"<tr><td>{h(r['modulo'])}</td><td>{h(r['rol_destino'])}</td><td><b>{h(r['pregunta_clave'])}</b><br><small>{h(r['palabras_clave'])}</small></td><td>{h(r['ruta'])}</td><td><a class='btn-red mini-btn' href='/admin/ia_hr/eliminar/sistema/{r['id']}'>Eliminar</a></td></tr>" for r in sistema]) or "<tr><td colspan='5'>Sin respuestas.</td></tr>"
    rows_l = ''.join([f"<tr><td>{h(r['tema'])}</td><td><b>{h(r['pregunta_clave'])}</b><br><small>{h(r['base_legal'])}</small></td><td>{h(r['fecha_actualizacion'])}</td><td><a class='btn-red mini-btn' href='/admin/ia_hr/eliminar/legal/{r['id']}'>Eliminar</a></td></tr>" for r in legal]) or "<tr><td colspan='4'>Sin base legal.</td></tr>"
    rows_log = ''.join([f"<tr><td>{h(r['fecha'])}</td><td>{h(r['rol'])}</td><td>{h(r['modulo'])}</td><td>{h(r['pregunta'])}</td><td>{h(r['respuesta_tipo'])}</td></tr>" for r in logs]) or "<tr><td colspan='5'>Sin consultas.</td></tr>"
    content=f"""
    <div class='hero'><div class='topbar'><div><h1>Base <span class='accent'>IA HR</span></h1><div class='subtitle'>Administra el manual inteligente, la base legal Perú y revisa consultas realizadas.</div></div><a class='btn-green' href='/admin/ia_hr/cargar_base'>Cargar respuestas base IA</a></div></div>
    <section class='grid'>
      <div class='card span-12'><h2>Agregar respuesta del sistema</h2><form method='post' class='ia-admin-form'>
        <input type='hidden' name='tipo' value='sistema'><label>Módulo<input name='modulo' placeholder='documental / vacacional / portal'></label><label>Rol<select name='rol_destino'><option value='ambos'>Ambos</option><option value='admin'>Administrador</option><option value='trabajador'>Trabajador</option></select></label>
        <label class='full'>Pregunta clave<input name='pregunta_clave' required placeholder='¿Dónde cargo boletas?'></label><label class='full'>Respuesta<textarea name='respuesta' required></textarea></label><label>Ruta<input name='ruta' placeholder='/admin/documentos'></label><label>Palabras clave<input name='palabras_clave' placeholder='boletas pago cargar'></label><button class='btn-green full'>Guardar respuesta</button></form></div>
      <div class='card span-12'><h2>Agregar respuesta legal Perú</h2><form method='post' class='ia-admin-form'>
        <input type='hidden' name='tipo' value='legal'><label>Tema<input name='tema' placeholder='CTS'></label><label>Base legal<input name='base_legal' placeholder='D.S. N.° 001-97-TR'></label><label class='full'>Pregunta clave<input name='pregunta_clave' required placeholder='¿Qué es CTS?'></label><label class='full'>Respuesta<textarea name='respuesta' required></textarea></label><label class='full'>Detalle legal<textarea name='detalle_legal'></textarea></label><label class='full'>Palabras clave<input name='palabras_clave'></label><button class='btn-green full'>Guardar legal</button></form></div>
      <div class='card span-12'><h2>Respuestas del sistema</h2><div class='table-wrap'><table class='c-table ia-table'><tr><th>Módulo</th><th>Rol</th><th>Pregunta</th><th>Ruta</th><th></th></tr>{rows_s}</table></div></div>
      <div class='card span-12'><h2>Base legal Perú</h2><div class='table-wrap'><table class='c-table ia-table'><tr><th>Tema</th><th>Pregunta / Base</th><th>Actualización</th><th></th></tr>{rows_l}</table></div></div>
      <div class='card span-12'><h2>Últimas consultas IA</h2><div class='table-wrap'><table class='c-table'><tr><th>Fecha</th><th>Rol</th><th>Módulo</th><th>Pregunta</th><th>Tipo</th></tr>{rows_log}</table></div></div>
    </section>"""
    return render_page(content, active='IA HR')


@app.route('/admin/ia_hr/cargar_base')
@admin_required
def admin_ia_hr_cargar_base():
    asegurar_ia_hr_pro()
    flash('Base inicial IA HR cargada/actualizada.', 'ok')
    return redirect(url_for('admin_ia_hr'))


@app.route('/admin/ia_hr/eliminar/<tipo>/<int:item_id>')
@admin_required
def admin_ia_hr_eliminar(tipo, item_id):
    with db() as con:
        if tipo == 'legal':
            con.execute('UPDATE ia_legislacion_peru SET activo=0 WHERE id=?', (item_id,))
        else:
            con.execute('UPDATE ia_base_conocimiento_hr SET activo=0 WHERE id=?', (item_id,))
        con.commit()
    flash('Registro IA desactivado.', 'ok')
    return redirect(url_for('admin_ia_hr'))

# =============================
# DB HELPERS
# =============================
def get_trabajador(dni):
    dni = normalizar_dni(dni)
    with db() as con:
        return con.execute("SELECT * FROM trabajadores WHERE dni=?", (dni,)).fetchone()



def row_get(row, key, default=''):
    try:
        return row[key] if row and key in row.keys() and row[key] is not None else default
    except Exception:
        return default


def separar_nombres_apellidos(nombre):
    txt = clean(nombre)
    if ',' in txt:
        ap, nom = [x.strip() for x in txt.split(',', 1)]
        aps = ap.split()
        return (aps[0] if aps else '', ' '.join(aps[1:]) if len(aps) > 1 else '', nom)
    partes = txt.split()
    if len(partes) >= 3:
        return partes[0], partes[1], ' '.join(partes[2:])
    return '', '', txt


def formato_fecha_texto(fecha, mayus=False):
    d = parse_fecha_any(fecha)
    if not d:
        return fecha_sin_hora(fecha)
    meses = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']
    txt = f"{d.day:02d} de {meses[d.month-1]} de {d.year}"
    return txt.upper() if mayus else txt


def valores_esquema_desde_trabajador(trabajador=None):
    """Llena los campos de esquema desde la base de trabajadores cuando exista información."""
    t = trabajador
    nombre = clean(row_get(t, 'nombre'))
    ap_pat, ap_mat, nombres = separar_nombres_apellidos(nombre)
    fecha_ing = row_get(t, 'fecha_ingreso')
    fecha_nac = row_get(t, 'fecha_nacimiento')
    empresa = row_get(t, 'empresa')
    cargo = row_get(t, 'cargo')
    area = row_get(t, 'area')
    correo = row_get(t, 'correo')
    planilla = row_get(t, 'planilla')
    dni = row_get(t, 'dni')
    base = {
        'Dni': dni, 'NombreCompletoTrabajador': nombre, 'Email': correo, 'Cargo': cargo,
        'Puesto': cargo, 'Ocupacion': cargo, 'Area': area, 'Gerencia': area, 'Planilla': planilla,
        'FechaIniContrato': fecha_sin_hora(fecha_ing), 'FechaInicioContratoOrigen': fecha_sin_hora(fecha_ing),
        'FechaIniContratoBarra': fecha_sin_hora(fecha_ing), 'FechaIniContratoGuion': fecha_sin_hora(fecha_ing).replace('/','-'),
        'FechaIniContratoISO': fecha_iso_segura(fecha_ing), 'FechaIniContratoTextoMayuscula': formato_fecha_texto(fecha_ing, True),
        'FechaIniContratoTextoMinuscula': formato_fecha_texto(fecha_ing, False), 'FechaNacimientoISO': fecha_iso_segura(fecha_nac),
        'FechaNacimientoGuion': fecha_sin_hora(fecha_nac).replace('/','-'), 'FechaNacimientoBarra': fecha_sin_hora(fecha_nac),
        'FechaNacimientoTextoMayuscula': formato_fecha_texto(fecha_nac, True),
        'FechaNacimientoTextoMinuscula': formato_fecha_texto(fecha_nac, False), 'ApellidoPaternoTrabajador': ap_pat,
        'ApellidoMaternoTrabajador': ap_mat, 'NombreTrabajador': nombres, 'Estado': 'Activo',
        'Condicion': 'ACTIVO' if row_get(t, 'activo', 1) else 'INACTIVO', 'CentroCosto': empresa,
        'NombreMoneda': 'Sol Peruano', 'SimboloMoneda': 'S/', 'NombreTipoDocumentoIdentidad': 'DOC. NACIONAL DE IDENTIDAD',
        'NombreCortoTipoDocumentoIdentidad': 'DNI',
    }
    salida=[]
    for campo, ejemplo in CAMPOS_ESQUEMA_TRABAJADOR_CONTRATO_LABORAL:
        salida.append((campo, clean(base.get(campo)) or clean(ejemplo)))
    return salida




def mapa_campos_trabajador(trabajador=None):
    """Diccionario CampoOrigen -> valor real del trabajador para combinar correspondencia."""
    return {campo: valor for campo, valor in valores_esquema_desde_trabajador(trabajador)}


def evaluar_condicion_valor(actual, operador, esperado):
    """Evalúa condiciones de plantilla contra los datos del trabajador."""
    actual_txt = clean(actual).upper()
    esperado_txt = clean(esperado).upper()
    op = clean(operador).upper() or '='
    if op in ('=', '=='):
        return actual_txt == esperado_txt
    if op in ('<>', '!='):
        return actual_txt != esperado_txt
    if op == 'CONTIENE':
        return esperado_txt in actual_txt
    if op == 'NO CONTIENE':
        return esperado_txt not in actual_txt
    try:
        a = float(re.sub(r'[^0-9.-]', '', actual_txt) or 0)
        b = float(re.sub(r'[^0-9.-]', '', esperado_txt) or 0)
        if op == '>': return a > b
        if op == '<': return a < b
        if op == '>=': return a >= b
        if op == '<=': return a <= b
    except Exception:
        return False
    return False


def plantilla_cumple_condiciones(plantilla, condiciones, trabajador=None):
    """Retorna (cumple, detalle_html). Si la plantilla está SIN CONDICIONES, siempre cumple."""
    if not plantilla or (plantilla['condicion'] or '').upper() != 'CONDICIONES':
        return True, '<span class="cond-ok">Plantilla general: sin condiciones.</span>'
    valores = mapa_campos_trabajador(trabajador)
    if not condiciones:
        return False, '<span class="cond-bad">La plantilla exige condiciones, pero no tiene reglas registradas.</span>'
    partes = []
    cumple_todo = True
    for c in condiciones:
        if not int(c['activo'] or 0):
            continue
        campo_label = clean(c['nombre_campo'])
        campo_origen = None
        for nom, origen, td in CONTRATACION_CAMPOS_CORRESPONDENCIA:
            if nom == campo_label or origen == campo_label:
                campo_origen = origen
                break
        campo_origen = campo_origen or campo_label.replace(' ', '')
        actual = valores.get(campo_origen, '')
        ok = evaluar_condicion_valor(actual, c['condicion'], c['valor'])
        cumple_todo = cumple_todo and ok
        partes.append(f"<span class='{'cond-ok' if ok else 'cond-bad'}'>{html.escape(campo_label)} {html.escape(c['condicion'] or '=')} {html.escape(c['valor'] or '')} → dato trabajador: {html.escape(actual or 'VACÍO')}</span>")
    return cumple_todo, ' '.join(partes) if partes else '<span class="cond-bad">No hay condiciones activas.</span>'



def extraer_campos_word_docx(path):
    """Detecta campos de correspondencia escritos como «Campo» y {{Campo}} en un Word .docx."""
    campos = []
    if Document is None:
        return campos
    try:
        doc = Document(str(path))
    except Exception:
        return campos
    textos = []
    def add(txt):
        if txt:
            textos.append(str(txt))
    for p in doc.paragraphs:
        add(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    add(p.text)
    texto = "\n".join(textos)
    # Soporta campos Word tipo combinación: «NombreCompletoTrabajador»
    # y campos internos del sistema: {{NombreCompletoTrabajador}}
    patrones = [r'«\s*([^»\n\r]+?)\s*»', r'\{\{\s*([^}\n\r]+?)\s*\}\}']
    vistos = set()
    for pat in patrones:
        for m in re.finditer(pat, texto):
            campo = clean(m.group(1)).strip()
            campo = re.sub(r'\s+', '', campo)
            if campo and campo not in vistos:
                vistos.add(campo); campos.append(campo)
    return campos


def sincronizar_campos_desde_word(plantilla_id, ruta_archivo):
    """Carga automáticamente a la pestaña Campos los campos encontrados en Word."""
    if not ruta_archivo:
        return 0
    path = Path(ruta_archivo)
    if not path.exists() or path.suffix.lower() != '.docx':
        return 0
    encontrados = extraer_campos_word_docx(path)
    if not encontrados:
        return 0
    tipos = {origen: td for nom, origen, td in CONTRATACION_CAMPOS_CORRESPONDENCIA}
    nombres = {origen: nom for nom, origen, td in CONTRATACION_CAMPOS_CORRESPONDENCIA}
    n = 0
    with db() as con:
        for campo in encontrados:
            existe = con.execute('SELECT 1 FROM contratacion_plantilla_campos WHERE plantilla_id=? AND campo_origen=?', (plantilla_id, campo)).fetchone()
            if not existe:
                con.execute("INSERT INTO contratacion_plantilla_campos(plantilla_id,descripcion,tipo_campo,nombre_campo,campo_origen,tipo_dato,requerido,activo) VALUES(?,?,?,?,?,?,?,1)",
                            (plantilla_id, 'Detectado automáticamente desde Word', 'Origen de Datos', nombres.get(campo, nombre_legible_campo(campo)), campo, tipos.get(campo, tipo_dato_esquema(campo)), 'SI'))
                n += 1
        con.commit()
    return n

def reemplazar_texto_docx(doc, valores):
    """Reemplaza campos {{CampoOrigen}} y «CampoOrigen» en párrafos y tablas de un Word."""
    def repl(txt):
        if not txt:
            return txt
        for k, v in valores.items():
            val = clean(v)
            txt = txt.replace('{{' + k + '}}', val)
            txt = txt.replace('{{ ' + k + ' }}', val)
            txt = txt.replace('«' + k + '»', val)
            txt = txt.replace('« ' + k + ' »', val)
        return txt
    def proc_paragraph(p):
        full = ''.join(run.text for run in p.runs)
        new = repl(full)
        if new != full:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new
            else:
                p.add_run(new)
    for p in doc.paragraphs:
        proc_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    proc_paragraph(p)
    return doc


def generar_docx_desde_plantilla(pid, dni=''):
    """Genera un DOCX final combinado con datos del trabajador y valida condiciones."""
    if Document is None:
        raise RuntimeError('python-docx no está instalado.')
    dni = normalizar_dni(dni)
    with db() as con:
        pl = con.execute('SELECT * FROM contratacion_plantillas WHERE id=?', (pid,)).fetchone()
        campos = con.execute('SELECT * FROM contratacion_plantilla_campos WHERE plantilla_id=? AND activo=1 ORDER BY id', (pid,)).fetchall()
        condiciones = con.execute('SELECT * FROM contratacion_plantilla_condiciones WHERE plantilla_id=? ORDER BY id', (pid,)).fetchall()
        trabajador = con.execute('SELECT * FROM trabajadores WHERE dni=?', (dni,)).fetchone() if dni else con.execute('SELECT * FROM trabajadores ORDER BY nombre LIMIT 1').fetchone()
    if not pl:
        raise FileNotFoundError('Plantilla no encontrada.')
    valores = mapa_campos_trabajador(trabajador)
    # Asegura que todos los campos registrados existan, aunque no haya dato en trabajador.
    for c in campos:
        key = c['campo_origen'] or c['nombre_campo']
        try:
            tipo_campo = (c['tipo_campo'] or '').upper()
            valor_default = c['valor_default'] if 'valor_default' in c.keys() else ''
        except Exception:
            tipo_campo, valor_default = '', ''
        if tipo_campo in ('MANUAL','DESPLEGABLE') and valor_default:
            valores[key] = valor_default
        else:
            valores.setdefault(key, '')
    cumple, detalle = plantilla_cumple_condiciones(pl, condiciones, trabajador)
    ruta = Path(pl['ruta_archivo']) if pl['ruta_archivo'] else None
    if ruta and ruta.exists() and ruta.suffix.lower() == '.docx':
        doc = Document(str(ruta))
    else:
        doc = Document()
        doc.add_heading(pl['nombre_plantilla'] or 'Documento contractual', level=1)
        doc.add_paragraph('Plantilla generada automáticamente. Reemplace el contenido por su Word oficial o use los campos siguientes:')
        doc.add_paragraph('Trabajador: {{NombreCompletoTrabajador}}')
        doc.add_paragraph('DNI: {{Dni}}')
        doc.add_paragraph('Cargo: {{Cargo}}')
        doc.add_paragraph('Planilla: {{Planilla}}')
        doc.add_paragraph('Tipo Contrato: {{TipoContrato}}')
    reemplazar_texto_docx(doc, valores)
    safe_name = re.sub(r'[^A-Za-z0-9_ -]+', '', pl['nombre_plantilla'] or 'plantilla').strip() or 'plantilla'
    out_dir = UPLOAD_DIR / 'contratacion' / 'generados'
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{safe_name}_{dni or 'SIN_DNI'}_{now_file()}.docx"
    doc.save(out)
    return out, pl, trabajador, cumple, detalle


def docx_to_preview_html(path, valores=None):
    """Vista previa simple de Word en HTML: párrafos y tablas con campos reemplazados."""
    if Document is None:
        return '<div class="preview-empty"><b>No se puede previsualizar Word porque falta python-docx.</b><br>Solución: el ZIP ya incluye <code>python-docx==1.1.2</code> en requirements.txt. En Render usa <b>Clear build cache & deploy</b>; en local ejecuta <code>pip install -r requirements.txt</code>.</div>'
    doc = Document(str(path))
    if valores:
        reemplazar_texto_docx(doc, valores)
    parts = ["<div class='word-preview'>"]
    for p in doc.paragraphs:
        txt = html.escape(p.text or '').replace('\n','<br>')
        if txt.strip():
            parts.append(f"<p>{txt}</p>")
    for table in doc.tables:
        parts.append("<table class='word-table'>")
        for row in table.rows:
            parts.append('<tr>' + ''.join(f"<td>{html.escape(cell.text or '').replace(chr(10), '<br>')}</td>" for cell in row.cells) + '</tr>')
        parts.append('</table>')
    parts.append('</div>')
    return ''.join(parts)

def generar_clave_trabajador(dni, fecha_nac=''):
    """Clave del trabajador: fecha de nacimiento sin / ni guiones (ddmmaaaa)."""
    if hasattr(fecha_nac, 'strftime'):
        return fecha_nac.strftime('%d%m%Y')
    txt = str(fecha_nac or '').strip()
    nums = re.sub(r"\D", "", txt)
    if len(nums) >= 8:
        return nums[:8]
    dni = normalizar_dni(dni)
    return (dni[-4:] + "PORTAL HR PRO").upper()


def registrar_evento_documento(doc_id, dni, evento, detalle=''):
    with db() as con:
        con.execute("INSERT INTO eventos_documento(documento_id,dni,evento,fecha,detalle) VALUES(?,?,?,?,?)", (doc_id, normalizar_dni(dni), evento, now_txt(), clean(detalle)))
        if evento in ['Abierto','Leído','Recibido']:
            con.execute("UPDATE documentos SET fecha_lectura=?, leido_por=? WHERE id=? AND (fecha_lectura IS NULL OR fecha_lectura='')", (now_txt(), normalizar_dni(dni), doc_id))
        con.commit()


def listar_documentos(dni=None, tipo=None, categoria=None, periodo=None, buscar=None, limit=300):
    where, params = [], []
    if dni:
        where.append("(dni=? OR categoria='empresa')")
        params.append(normalizar_dni(dni))
    if tipo:
        where.append("tipo=?")
        params.append(tipo)
    if categoria:
        where.append("categoria=?")
        params.append(categoria)
    if periodo:
        where.append("periodo=?")
        params.append(periodo)
    if buscar:
        b = f"%{buscar}%"
        where.append("(dni LIKE ? OR tipo LIKE ? OR periodo LIKE ? OR detalle LIKE ? OR observacion LIKE ? OR archivo_nombre LIKE ?)")
        params += [b, b, b, b, b, b]
    sql = "SELECT * FROM documentos"
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY id DESC LIMIT ?"
    params.append(limit)
    with db() as con:
        return con.execute(sql, params).fetchall()


def periodos_disponibles(dni=None, tipo=None, categoria=None):
    where, params = ["periodo IS NOT NULL", "periodo<>''"], []
    if dni:
        where.append("(dni=? OR categoria='empresa')"); params.append(normalizar_dni(dni))
    if tipo:
        where.append("tipo=?"); params.append(tipo)
    if categoria:
        where.append("categoria=?"); params.append(categoria)
    sql = "SELECT DISTINCT periodo FROM documentos WHERE " + " AND ".join(where) + " ORDER BY periodo DESC LIMIT 80"
    with db() as con:
        return [r[0] for r in con.execute(sql, params).fetchall() if r[0]]


def parse_fecha_any(v):
    txt = clean(v)
    if not txt:
        return None
    if isinstance(v, datetime):
        return v.date()
    for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d"]:
        try:
            return datetime.strptime(txt.split()[0], fmt).date()
        except Exception:
            pass
    return None

def hoy_lima():
    """Fecha actual en Perú para validar vacaciones sin depender del navegador."""
    return datetime.now(APP_TZ).date()

def fecha_iso_segura(v):
    d = parse_fecha_any(v)
    return d.isoformat() if d else ''

def periodos_desde_ingreso(dni=None, tipo=None, max_meses=72):
    """Devuelve periodos desde la fecha de ingreso del trabajador hasta hoy."""
    inicio = None
    if dni:
        t = get_trabajador(dni)
        if t and 'fecha_ingreso' in t.keys():
            inicio = parse_fecha_any(t['fecha_ingreso'])
    hoy = datetime.now(APP_TZ).date()
    if not inicio:
        inicio = hoy.replace(day=1)
    inicio = inicio.replace(day=1)
    out=[]; y=inicio.year; m=inicio.month
    while (y < hoy.year or (y == hoy.year and m <= hoy.month)) and len(out) < max_meses:
        out.append(f"{y:04d}-{m:02d}")
        m += 1
        if m == 13:
            y += 1; m = 1
    docs = periodos_disponibles(dni=dni, tipo=tipo)
    return sorted(set(out + docs), reverse=True)


def guardar_documento(file_storage, dni, tipo, periodo, detalle="", observacion="", uploaded_by="sistema"):
    if not file_storage or not file_storage.filename:
        return None
    original = secure_filename(file_storage.filename)
    ext = Path(original).suffix.lower()
    if ext not in EXT_ALLOWED:
        raise ValueError(f"Extensión no permitida: {ext}")
    tipo_info = ALL_TIPOS.get(tipo, (tipo, "📄", "personal"))
    categoria = tipo_info[2]
    dni = normalizar_dni(dni) if categoria != "empresa" else ""
    periodo = safe_periodo(periodo)
    folder = UPLOAD_DIR / categoria / re.sub(r"[^A-Za-z0-9_\-]", "_", tipo) / periodo
    if dni:
        folder = folder / dni
    folder.mkdir(parents=True, exist_ok=True)
    tipo_file = re.sub(r"[^A-Za-z0-9_\-]+", "_", tipo)
    prefijo_dni = f"{dni}_" if dni else ""
    final = f"{prefijo_dni}{tipo_file}_{periodo}_{now_file()}_{original}"
    path = folder / final
    file_storage.save(path)
    # Copia automática a carpeta documental organizada para respaldo físico.
    try:
        label_auto, icon_auto, cat_auto = ALL_TIPOS.get(tipo, (tipo, '📄', categoria))
        root_name = {'pago':'DOCUMENTOS DE PAGO','empresa':'DOCUMENTOS DE LA EMPRESA','personal':'DOCUMENTOS PERSONALES'}.get(cat_auto,'DOCUMENTOS')
        auto_base = DOCUMENTOS_BASE_DIR / root_name / slug_folder(label_auto)
        if tipo == 'Normal' and 'seman' in clean(detalle).lower():
            auto_base = auto_base / 'SEMANAL'
        elif tipo == 'Normal':
            auto_base = auto_base / 'MENSUAL'
        auto_base = auto_base / periodo
        if dni: auto_base = auto_base / dni
        auto_base.mkdir(parents=True, exist_ok=True)
        auto_path = auto_base / final
        if str(auto_path) != str(path):
            auto_path.write_bytes(path.read_bytes())
    except Exception:
        pass
    with db() as con:
        con.execute("""
        INSERT INTO documentos(dni,categoria,tipo,periodo,detalle,observacion,archivo_nombre,ruta_archivo,extension,fecha_subida,uploaded_by)
        VALUES(?,?,?,?,?,?,?,?,?,?,?)
        """, (dni, categoria, tipo, periodo, clean(detalle), clean(observacion), original, str(path), ext, now_txt(), uploaded_by))
        con.commit()
    return str(path)


def inferir_tipo_desde_ruta(path: Path):
    texto = str(path).lower()
    reglas = [("util", "Utilidad"), ("vacacion", "Vacaciones"), ("normal", "Normal"), ("cts", "CTS"), ("liquid", "Liquidación"), ("grat", "Gratificación"), ("constancia", "Constancia Gratificación"), ("contrato", "Contrato de Trabajo"), ("sst", "Reglamento de SST"), ("interno", "Reglamento Interno"), ("conducta", "Código de Conducta"), ("politica", "Políticas"), ("políticas", "Políticas"), ("comunicado", "Comunicados")]
    for clave, tipo in reglas:
        if clave in texto:
            return tipo
    return "Otros"


def inferir_periodo_desde_ruta(path: Path):
    texto = str(path)
    m = re.search(r"(20\d{2})[-_ ]?(0[1-9]|1[0-2])", texto)
    if m: return f"{m.group(1)}-{m.group(2)}"
    m = re.search(r"(0[1-9]|1[0-2])[-_ ]?(20\d{2})", texto)
    if m: return f"{m.group(2)}-{m.group(1)}"
    m = re.search(r"(20\d{2})", texto)
    if m: return m.group(1)
    return datetime.now(APP_TZ).strftime("%Y-%m")


def documento_ya_indexado(path: Path):
    with db() as con:
        return con.execute("SELECT 1 FROM documentos WHERE ruta_archivo=?", (str(path),)).fetchone() is not None


def extraer_texto_pdf(path: Path, max_paginas=2):
    """Extrae texto de las primeras páginas para detectar DNI dentro de la boleta."""
    if path.suffix.lower() != '.pdf' or PdfReader is None:
        return ''
    try:
        reader = PdfReader(str(path))
        partes = []
        for page in reader.pages[:max_paginas]:
            try:
                partes.append(page.extract_text() or '')
            except Exception:
                pass
        return '\n'.join(partes)[:8000]
    except Exception:
        return ''


def detectar_dni_en_archivo(path: Path, dni_obj=''):
    """Busca DNI primero en ruta/nombre y, si es PDF, también dentro del contenido."""
    dni_obj = normalizar_dni(dni_obj) if dni_obj else ''
    texto_ruta = str(path)
    if dni_obj and dni_obj in texto_ruta:
        return dni_obj, 'ruta/nombre'
    m = re.search(r"(?<!\d)(\d{8})(?!\d)", texto_ruta)
    if m:
        return m.group(1), 'ruta/nombre'
    texto_pdf = extraer_texto_pdf(path)
    if texto_pdf:
        if dni_obj and dni_obj in re.sub(r'\D', '', texto_pdf):
            return dni_obj, 'contenido PDF'
        patrones = [
            r"(?:DNI|D\.?N\.?I\.?|DOC(?:UMENTO)?|COD(?:IGO)?|IDENTIDAD)\s*[:º°#-]?\s*(\d{8})",
            r"(?<!\d)(\d{8})(?!\d)",
        ]
        for pat in patrones:
            mm = re.search(pat, texto_pdf, flags=re.I)
            if mm:
                return normalizar_dni(mm.group(1)), 'contenido PDF'
    return '', ''


def detalle_auto_desde_ruta(path: Path):
    txt = str(path).lower()
    if 'semanal' in txt or 'semana' in txt:
        return 'Boleta semanal - Importado automáticamente desde carpeta'
    if 'mensual' in txt or 'mes' in txt:
        return 'Boleta mensual - Importado automáticamente desde carpeta'
    return 'Importado automáticamente desde carpeta'


def registrar_archivo_existente(path: Path, dni: str, tipo: str, uploaded_by="auto", fuente_dni='ruta/nombre'):
    uploaded_by = marca_carga(uploaded_by)
    if documento_ya_indexado(path): return False
    ext = path.suffix.lower()
    if ext not in EXT_ALLOWED: return False
    label, icon, categoria = ALL_TIPOS.get(tipo, (tipo, "📄", "personal"))
    dni = normalizar_dni(dni) if categoria != "empresa" else ""
    periodo = inferir_periodo_desde_ruta(path)
    detalle = detalle_auto_desde_ruta(path)
    obs = f"Detectado automáticamente por {fuente_dni}. Ruta: {path.parent}"
    with db() as con:
        con.execute("""
        INSERT INTO documentos(dni,categoria,tipo,periodo,detalle,observacion,archivo_nombre,ruta_archivo,extension,fecha_subida,uploaded_by)
        VALUES(?,?,?,?,?,?,?,?,?,?,?)
        """, (dni, categoria, tipo, periodo, detalle, obs, path.name, str(path), ext, now_txt(), uploaded_by))
        con.commit()
    return True


def sincronizar_documentos_carpeta(dni=None, devolver_resumen=False):
    # Detecta documentos desde DOCUMENTOS_PRIZE_AUTO, incluyendo:
    # DOCUMENTOS DE PAGO / BOLETAS NORMAL / SEMANAL.
    # Si el DNI no está en el nombre, lee el PDF e intenta encontrarlo dentro del texto.
    asegurar_carpetas_documentales()
    base_dirs = []
    for b in [DOCUMENTOS_BASE_DIR, BASE_DIR / "documentos_auto"]:
        if b.exists() and b.is_dir() and b not in base_dirs:
            base_dirs.append(b)
    total = 0
    revisados = omitidos = duplicados = sin_dni = sin_trabajador = 0
    dni_obj = normalizar_dni(dni) if dni else ""
    for base in base_dirs:
        for path in base.rglob("*"):
            if not path.is_file() or path.suffix.lower() not in EXT_ALLOWED:
                continue
            revisados += 1
            if documento_ya_indexado(path):
                duplicados += 1; continue
            tipo = inferir_tipo_desde_ruta(path)
            categoria = ALL_TIPOS.get(tipo, ("", "", "personal"))[2]
            dni_detectado, fuente = detectar_dni_en_archivo(path, dni_obj)
            if categoria != "empresa" and not dni_detectado:
                sin_dni += 1; omitidos += 1; continue
            if dni_obj and categoria != "empresa" and dni_detectado != dni_obj:
                omitidos += 1; continue
            if categoria != "empresa":
                trab = get_trabajador(dni_detectado)
                if not trab or int(trab['activo'] or 0) != 1:
                    sin_trabajador += 1; omitidos += 1; continue
            try:
                if registrar_archivo_existente(path, dni_detectado, tipo, uploaded_by="auto carpeta local", fuente_dni=fuente or 'carpeta'):
                    total += 1
            except Exception:
                omitidos += 1
    resumen = {'nuevos': total, 'revisados': revisados, 'duplicados': duplicados, 'omitidos': omitidos, 'sin_dni': sin_dni, 'sin_trabajador': sin_trabajador, 'base': str(DOCUMENTOS_BASE_DIR)}
    return resumen if devolver_resumen else total

# =============================
# ESTILOS Y LAYOUT
# =============================
BASE = r'''
<!doctype html><html lang="es"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>{{ title }}</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.css">
<style>
:root{--txt:#eef2f7;--mut:#a8b0bb;--yellow:#10b981;--yellow2:#16a34a;--dark:#15181d;--panel:#1e2025;--panel2:#171a20;--line:#343a43;--shadow:0 24px 60px rgba(0,0,0,.35)}
*{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;color:var(--txt);background:#15181d;font-weight:650}a{text-decoration:none;color:inherit}.hidden{display:none!important}.btn,.btn-blue,.btn-green,.btn-red{border:1px solid #3a414b;border-radius:14px;padding:12px 18px;background:#20242b;color:#eef2f7;font-weight:950;cursor:pointer;display:inline-flex;align-items:center;gap:8px;box-shadow:0 10px 22px rgba(0,0,0,.18);transition:.16s}.btn:hover,.btn-blue:hover,.btn-green:hover{transform:translateY(-1px);box-shadow:0 16px 34px rgba(0,0,0,.28)}.btn-blue{background:linear-gradient(135deg,#292e36,#1f232b);border-color:#4a525e;color:var(--yellow)}.btn-green{background:linear-gradient(135deg,var(--yellow2),var(--yellow));border:0;color:#1d1f24}.btn-red{background:#361a24;border-color:#7f1d1d;color:#fecaca}.flash{padding:13px 16px;border-radius:16px;margin:10px 0;background:#2a2a1b;border:1px solid #9a7b16;color:#ffefaa}.flash.err{background:#3a1720;border-color:#7f1d1d;color:#fecaca}.input,select,textarea{width:100%;border:1px solid #3a414b;border-radius:14px;padding:13px 14px;background:#111418;color:#f8fafc;font:inherit;outline:none}option{background:#111418;color:#f8fafc}.input:focus,select:focus,textarea:focus{border-color:var(--yellow);box-shadow:0 0 0 4px rgba(16,185,129,.12)}
/* LOGIN - estilo imagen negra/amarilla */
.login-body{min-height:100vh;display:grid;place-items:center;padding:20px;position:relative;overflow:hidden;background:linear-gradient(rgba(22,25,29,.86),rgba(22,25,29,.90)),radial-gradient(circle at 7% 4%,#10b981 0 23%,transparent 23.2%),radial-gradient(circle at 94% -2%,#10b981 0 11%,transparent 11.2%),radial-gradient(circle at 72% 112%,#10b981 0 20%,transparent 20.2%),linear-gradient(135deg,#2a2e33,#111418)}.login-card{width:min(92vw,500px);background:linear-gradient(180deg,rgba(25,28,33,.98),rgba(29,33,38,.95));border:1px solid rgba(255,255,255,.10);border-radius:18px;padding:38px 42px 0;box-shadow:0 38px 90px rgba(0,0,0,.52);overflow:hidden;position:relative}.login-card:before{content:"";position:absolute;left:-72px;bottom:-58px;width:365px;height:150px;background:linear-gradient(135deg,#2e4f86,#5d83e6);border-radius:50% 50% 0 0;transform:rotate(-8deg);opacity:.95}.login-card:after{content:"";position:absolute;right:-78px;bottom:-52px;width:350px;height:145px;background:linear-gradient(135deg,#253849,#475b6f);border-radius:50% 50% 0 0;transform:rotate(8deg);opacity:.92}.login-inner{position:relative;z-index:2}.login-logo{text-align:center}.login-logo img{max-width:145px;max-height:105px;object-fit:contain;background:rgba(255,255,255,.92);border-radius:10px;padding:7px;filter:drop-shadow(0 14px 24px rgba(0,0,0,.45))}.login-title{text-align:center;margin:20px 0 30px;color:#aeb7c3}.login-title h1{margin:0 0 7px;color:#fff;font-size:24px;letter-spacing:.5px;text-transform:uppercase}.login-title b{color:#98a4b3}.login-card .field label{display:none}.login-input{display:flex;align-items:center;gap:13px;background:transparent;border-bottom:1px solid rgba(226,232,240,.40);padding:0 6px;margin-bottom:22px;transition:.18s}.login-input:focus-within{border-bottom-color:var(--yellow);box-shadow:0 10px 0 -9px rgba(16,185,129,.9)}.login-input input{border:0;padding:15px 8px;width:100%;font:inherit;outline:none;background:transparent;color:#fff;font-weight:900}.login-input input::placeholder{color:#cbd5e1}.login-card .btn-green{width:auto;justify-content:center;font-size:15px;margin:8px 0 74px;padding:14px 34px;border-radius:28px;background:linear-gradient(135deg,var(--yellow2),var(--yellow));color:#212529;border:0;box-shadow:0 14px 30px rgba(22,163,74,.35)}.login-links{text-align:center;margin-top:-48px;padding-bottom:24px;position:relative;z-index:3}.login-links a{color:#dbeafe;font-size:13px;font-weight:900}
/* APP - dashboard ejecutivo */
.app{display:grid;grid-template-columns:320px 1fr;min-height:100vh;background:#15181d;transition:grid-template-columns .22s ease}.app.side-collapsed{grid-template-columns:86px 1fr}.side{background:linear-gradient(180deg,#1e2024,#171a1f 72%,#111318);color:#f2f4f8;position:sticky;top:0;height:100vh;overflow:auto;transition:.25s;width:320px;z-index:5;box-shadow:12px 0 35px rgba(0,0,0,.34);border-right:1px solid #33373d}.side.collapsed{width:86px}.side-top{height:54px;display:flex;align-items:center;justify-content:space-between;padding:0 14px;background:#17191e;border-bottom:1px solid rgba(255,255,255,.07);position:sticky;top:0;z-index:3}.toggle{cursor:pointer;background:transparent;border:0;color:white;font-size:21px}.brand{padding:28px 16px 22px;text-align:center}.brand img{max-width:150px;max-height:95px;background:rgba(255,255,255,.90);border-radius:16px;object-fit:contain;padding:8px;box-shadow:0 14px 30px rgba(0,0,0,.35);border:1px solid rgba(16,185,129,.28)}.brand p{color:#c8cdd6;font-size:14px;margin-top:18px}.side.collapsed .brand p,.side.collapsed .label,.side.collapsed .chev,.side.collapsed .subtxt,.side.collapsed .side-user{display:none}.side.collapsed .brand{padding:20px 8px}.side.collapsed .brand img{max-width:55px;max-height:55px;border-radius:14px;padding:4px}.menu-group{margin:10px 12px;border-radius:12px;overflow:hidden}.menu-title{width:100%;border:1px solid rgba(255,255,255,.06);display:flex;align-items:center;gap:12px;background:linear-gradient(135deg,#22252b,#1b1e24);color:#eef2f7;padding:15px 14px;font-size:15px;font-weight:1000;cursor:pointer;text-align:left;border-radius:12px}.menu-title:hover{background:linear-gradient(135deg,#2b2f36,#23272f)}.menu-group.force-open>.menu-title{background:linear-gradient(135deg,#22252b,#1b1e24);color:#eef2f7;box-shadow:none}.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,var(--yellow2),var(--yellow));color:#181a1f;box-shadow:0 14px 30px rgba(16,185,129,.20)}.menu-title .chev{margin-left:auto;transition:.18s}.menu-group.closed .chev{transform:rotate(-90deg)}.submenu{background:transparent;padding:9px 0;max-height:720px;transition:max-height .28s ease,padding .18s ease}.menu-group.closed .submenu{max-height:0;padding:0;overflow:hidden}.menu-item{display:flex;align-items:center;gap:13px;padding:13px 18px 13px 40px;color:#dce3ed;font-weight:900;font-size:14px;border-left:4px solid transparent;transition:.13s;border-radius:10px;margin:2px 0}.menu-item:hover{background:#242830;border-left-color:var(--yellow)}.menu-item.active,.menu-title.active{background:linear-gradient(135deg,#34302a,#2a2926);border-left-color:var(--yellow);color:#fff}.side.collapsed .menu-title{justify-content:center;padding:18px 10px}.side.collapsed .menu-item{padding:16px 10px;justify-content:center}.side.collapsed .submenu{display:none}.main{min-width:0;padding:0 34px 50px;overflow:auto;background:radial-gradient(circle at 92% -8%,rgba(16,185,129,.22),transparent 22%),radial-gradient(circle at 100% 96%,rgba(16,185,129,.12),transparent 28%),#15181d}.hero{margin:0 -34px 24px;padding:26px 34px 28px;background:radial-gradient(circle at 72% 0%,rgba(16,185,129,.20),transparent 32%),linear-gradient(120deg,#15181d 0%,#111418 62%,#24282d 100%);border-bottom:1px solid #31363d}.topbar{display:flex;align-items:center;justify-content:space-between;gap:12px}.topbar h1{margin:0;font-size:34px;letter-spacing:-1px;color:#fff}.topbar h1 .accent{color:var(--yellow)}.subtitle{color:#aeb7c3;font-size:16px;margin-top:7px}.grid{display:grid;grid-template-columns:repeat(12,1fr);gap:18px}.card{background:linear-gradient(145deg,#202329,#181b20);border:1px solid #303640;border-radius:18px;box-shadow:0 22px 55px rgba(0,0,0,.25);padding:22px;color:#eef2f7}.mini{grid-column:span 4;display:flex;align-items:center;justify-content:space-between}.mini b{font-size:28px;color:var(--yellow)}.ico{width:56px;height:56px;border-radius:16px;display:grid;place-items:center;background:linear-gradient(135deg,var(--yellow),var(--yellow2));font-size:24px;color:#17191e;box-shadow:0 12px 26px rgba(16,185,129,.18)}.span-12{grid-column:span 12}.span-8{grid-column:span 8}.span-4{grid-column:span 4}.span-6{grid-column:span 6}.span-3{grid-column:span 3}.doc-grid{display:grid;grid-template-columns:repeat(4,minmax(220px,1fr));gap:14px}.doc-card{background:linear-gradient(145deg,#24272d,#1b1f25);border:1px solid #343a43;border-radius:16px;padding:18px;min-height:158px;transition:.16s;position:relative;overflow:hidden}.doc-card:before{content:"";position:absolute;right:-34px;top:-34px;width:86px;height:86px;background:rgba(16,185,129,.17);border-radius:50%}.doc-card h3{margin:0 0 12px;font-size:17px;color:#fff}.doc-card p{margin:0 0 14px;color:#c0c8d2;font-weight:500;line-height:1.45}.doc-card:hover{transform:translateY(-2px);border-color:var(--yellow);box-shadow:0 16px 30px rgba(0,0,0,.25)}.table-wrap{overflow:auto;border:1px solid #343a43;border-radius:14px}table{width:100%;border-collapse:collapse;background:#171a20;color:#eaf3ff}th,td{text-align:left;padding:13px 14px;border-bottom:1px solid #2c323a;vertical-align:top}th{background:#111418;color:var(--yellow);font-size:13px;text-transform:uppercase}tr:hover td{background:#20242b}.form-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;align-items:end}.detail-box{background:linear-gradient(135deg,#202329,#171a20);border:1px solid #343a43;border-radius:16px;padding:15px}.detail-box small{display:block;color:#aeb7c3;margin-bottom:4px}.period-row{display:flex;gap:12px;align-items:end;flex-wrap:wrap}.mobile-head{display:none}.side-user{margin:26px 14px 14px;padding-top:20px;border-top:1px solid rgba(255,255,255,.08);display:flex;align-items:center;gap:11px;color:#e5e7eb}.avatar{width:44px;height:44px;border-radius:50%;display:grid;place-items:center;background:var(--yellow);color:#15181d;font-weight:1000}
@media(max-width:1000px){.app,.app.side-collapsed{grid-template-columns:1fr}.side{position:fixed;left:-335px;width:315px}.side.open{left:0}.side.collapsed{left:-335px}.mobile-head{display:flex;position:sticky;top:0;z-index:20;background:#17191e;color:white;padding:12px 14px;align-items:center;justify-content:space-between;border-bottom:1px solid #343a43}.main{padding:0 14px 30px}.hero{margin:0 -14px 18px;padding:20px 16px}.doc-grid{grid-template-columns:1fr}.mini,.span-8,.span-4,.span-3{grid-column:span 12}.form-grid{grid-template-columns:1fr}.topbar{align-items:flex-start;flex-direction:column}.topbar h1{font-size:24px}.subtitle{font-size:13px}.card{border-radius:16px;padding:17px}.login-card{padding:32px 28px 0}.login-card .btn-green{width:100%}}@media(min-width:1001px) and (max-width:1350px){.doc-grid{grid-template-columns:repeat(2,1fr)}}

/* === RETOQUE PRO ADMIN / FORMULARIOS === */
.login-card{border-radius:26px;background:linear-gradient(180deg,rgba(24,27,32,.98),rgba(16,18,23,.96));backdrop-filter:blur(8px)}
.login-input{border:1px solid rgba(255,255,255,.10);border-radius:16px;background:rgba(255,255,255,.055);box-shadow:inset 0 0 0 1px rgba(255,255,255,.02)}
.login-input:focus-within{background:rgba(16,185,129,.10);border-color:rgba(16,185,129,.72);box-shadow:0 0 0 4px rgba(16,185,129,.12)}
.login-input input{color:#fff;background:transparent}.login-input input:-webkit-autofill{-webkit-box-shadow:0 0 0 1000px #202329 inset!important;-webkit-text-fill-color:#fff!important}
.form-grid{grid-template-columns:repeat(12,1fr);align-items:end}.form-grid .field{grid-column:span 3}.form-grid .field:nth-child(4n+1){grid-column:span 3}.form-grid button,.form-grid .btn,.form-grid .btn-blue,.form-grid .btn-green{grid-column:span 3;justify-content:center;height:54px}.field label{display:block;margin-bottom:8px;color:#eaf0f7;font-size:13px;letter-spacing:.3px}.field input,.field select,.field textarea,.input,select,textarea{background:#0f1319;border:1px solid #3b414b;color:#fff;border-radius:14px;min-height:48px;padding:12px 14px;font-weight:800}.field input:focus,.field select:focus,.field textarea:focus,.input:focus,select:focus,textarea:focus{border-color:var(--yellow);box-shadow:0 0 0 4px rgba(16,185,129,.12);outline:none}.card form{gap:18px}.alert-card{background:linear-gradient(145deg,#202329,#16191e 65%,rgba(16,185,129,.06));}.alert-item{display:grid;grid-template-columns:48px 1fr auto;gap:12px;align-items:center;padding:13px 0;border-top:1px solid #323740}.alert-item:first-of-type{border-top:0}.bell{width:40px;height:40px;border-radius:14px;display:grid;place-items:center;background:linear-gradient(135deg,var(--yellow),var(--yellow2));box-shadow:0 12px 22px rgba(16,185,129,.18)}.alert-item span,.muted,.empty-note{color:#b8c0cb}.mini-btn{padding:9px 13px;border-radius:12px}.admin-hero{border-radius:0 0 24px 24px;margin-bottom:20px}.side .brand img{background:linear-gradient(145deg,#f7f7f7,#d8d8d8);mix-blend-mode:normal}.side .brand{background:radial-gradient(circle at 50% 28%,rgba(16,185,129,.10),transparent 44%)}
@media(max-width:1000px){.form-grid{grid-template-columns:1fr}.form-grid .field,.form-grid button,.form-grid .btn,.form-grid .btn-blue,.form-grid .btn-green{grid-column:span 1;width:100%}.alert-item{grid-template-columns:42px 1fr}.alert-item a{grid-column:1 / -1;justify-content:center}.side.open{box-shadow:0 0 0 999px rgba(0,0,0,.55),12px 0 35px rgba(0,0,0,.34)}}

.status-pill{display:inline-flex;padding:7px 10px;border-radius:999px;background:#242a32;border:1px solid #3b414b;color:#10b981;font-weight:1000;white-space:nowrap}.actions{display:flex;gap:8px;flex-wrap:wrap}.modal{position:fixed;inset:0;background:rgba(0,0,0,.70);z-index:80;display:grid;place-items:center;padding:18px}.modal-card{width:min(520px,96vw);background:#1d2128;border:1px solid #3a414b;border-radius:18px;padding:22px;box-shadow:var(--shadow)}.profile-row{display:flex;align-items:center;gap:18px;flex-wrap:wrap}.profile-img{width:92px;height:92px;border-radius:50%;object-fit:cover;background:#fff;padding:4px;border:3px solid var(--yellow)}.profile-form{flex:1;min-width:240px}.sub-mini{padding-left:58px!important;font-size:13px!important;opacity:.92}.bars{display:grid;gap:12px}.bar-row{display:grid;grid-template-columns:190px 1fr 46px;gap:12px;align-items:center}.bar-row span{height:18px;background:#111418;border-radius:999px;overflow:hidden;border:1px solid #343a43}.bar-row i{display:block;height:100%;background:linear-gradient(90deg,var(--yellow2),var(--yellow));border-radius:999px}.bar-row em{font-style:normal;color:#10b981;font-weight:1000}input[type=file]{max-width:100%;white-space:normal;overflow:hidden}.field{min-width:0}.card{min-width:0}@media(max-width:1000px){body{overflow-x:hidden}.app{overflow-x:hidden}.main{width:100%;overflow-x:hidden}.card{padding:15px}.form-grid{display:grid!important;grid-template-columns:1fr!important}.form-grid .field,.form-grid button,.form-grid a{grid-column:1!important;width:100%;min-width:0}.field input,.field select,.field textarea,input[type=file]{width:100%;max-width:100%;font-size:14px}.table-wrap{max-width:100%;overflow-x:auto;-webkit-overflow-scrolling:touch}.table-wrap table{min-width:760px}.bar-row{grid-template-columns:1fr}.profile-row{align-items:flex-start}.mobile-head{height:54px}.hero{overflow:hidden}.detail-box{grid-column:span 12!important}}.btn-warn{background:linear-gradient(135deg,#ffce4a,#059669);border:0;color:#171a20;border-radius:14px;padding:10px 14px;font-weight:1000;cursor:pointer;box-shadow:0 10px 22px rgba(22,163,74,.18)}.btn-danger{background:linear-gradient(135deg,#48131f,#7f1d1d);border:1px solid #ef4444;color:#fee2e2;border-radius:14px;padding:10px 14px;font-weight:1000;cursor:pointer}.st-aprobado{background:#113327!important;border-color:#2dd4bf!important;color:#9fffe8!important}.st-rechazado{background:#3f1520!important;border-color:#ef4444!important;color:#fecaca!important}.st-firmado{background:#182844!important;border-color:#60a5fa!important;color:#bfdbfe!important}.st-aceptado{background:#302a12!important;border-color:#10b981!important;color:#fde68a!important}.row-approved{background:linear-gradient(90deg,rgba(45,212,191,.08),transparent)}.row-rejected{background:linear-gradient(90deg,rgba(239,68,68,.10),transparent)}.nested{margin:4px 0}.nested>.menu-item{width:100%;border:0}.submenu .menu-group{margin:6px 0 6px 22px}.submenu .menu-group .menu-title{padding:13px 14px;font-size:14px;border-radius:10px}.submenu .submenu .menu-item{padding-left:48px}.nested.closed .submenu{display:none}.menu-group.closed .submenu{display:none}.menu-group .chev{margin-left:auto}.menu-group.closed .chev{transform:rotate(-90deg)}
/* Marcador de ubicación limpio: no pinta blanco, solo línea y brillo lateral */
button.menu-item{font:inherit;text-align:left;background:transparent;border-top:0;border-right:0;border-bottom:0;appearance:none;-webkit-appearance:none;width:100%;}
.menu-item.active,.menu-item.parent-active{background:linear-gradient(90deg,rgba(16,185,129,.16),rgba(16,185,129,.035) 42%,transparent)!important;border-left-color:var(--yellow)!important;color:#fff!important;box-shadow:inset 4px 0 0 var(--yellow);}
.menu-item.sub-mini.active{background:rgba(16,185,129,.12)!important;color:#fff!important;border-left-color:var(--yellow)!important;box-shadow:inset 4px 0 0 var(--yellow);}
.menu-group.force-open>.menu-title{background:linear-gradient(135deg,var(--yellow2),var(--yellow));color:#181a1f;}
.nested.force-open>.menu-item.parent-active{background:linear-gradient(90deg,rgba(16,185,129,.14),rgba(16,185,129,.04),transparent)!important;color:#fff!important;}
.nested>.menu-item:focus,.nested>.menu-item:active{background:linear-gradient(90deg,rgba(16,185,129,.16),rgba(16,185,129,.035),transparent)!important;color:#fff!important;outline:none!important;}

.module-tabs{display:grid;grid-template-columns:repeat(3,minmax(220px,1fr));gap:16px}.module-tile{padding:22px;border:1px solid #3a414b;border-radius:18px;background:linear-gradient(145deg,#24272d,#191d23)}.module-tile h2{margin:0 0 8px}.badge-green{display:inline-flex;background:#49a916;color:#fff;border-radius:7px;padding:7px 12px;font-weight:1000}.badge-orange{display:inline-flex;background:#ff8f2d;color:#fff;border-radius:7px;padding:7px 12px;font-weight:1000}.adapta-note{background:#fff;color:#142033;border-radius:14px;padding:14px;border-left:5px solid #ff8f2d}.adapta-table table{background:#fff;color:#102033}.adapta-table th{background:#f4f5f7;color:#102033;text-transform:none;font-size:14px}.adapta-table td{border-bottom:1px solid #e4e7eb}.adapta-table tr:nth-child(even) td{background:#ededed}.adapta-table tr:hover td{background:#ffe9d6}@media(max-width:1000px){.module-tabs{grid-template-columns:1fr}}

/* === AJUSTE PRO: MENÚ UNIFORME + DASHBOARDS POR GESTIÓN === */
.menu-title,.menu-item{min-height:54px;box-sizing:border-box}
.submenu>.menu-item{padding:14px 18px 14px 36px;margin:4px 0}
.menu-group.nested .menu-title,.menu-group.nested>.menu-item{font-size:14px}
.gestion-cards{grid-template-columns:repeat(2,minmax(280px,1fr))!important;display:grid;grid-template-columns:repeat(3,minmax(260px,1fr));gap:18px;margin-bottom:18px}
.gestion-card{min-height:190px;display:flex;gap:18px;align-items:flex-start;position:relative;overflow:hidden}
.gestion-card:after{content:"";position:absolute;inset:auto -35px -35px auto;width:120px;height:120px;border-radius:999px;background:rgba(16,185,129,.10)}
.gestion-card.green{border-color:rgba(57,196,99,.40)}.gestion-card.purple{border-color:rgba(160,98,255,.42)}
.gestion-icon{width:64px;height:64px;border-radius:18px;display:grid;place-items:center;font-size:30px;background:linear-gradient(135deg,var(--yellow),var(--yellow2));color:#17191e;box-shadow:0 16px 30px rgba(0,0,0,.26)}
.gestion-card.green .gestion-icon{background:linear-gradient(135deg,#27b862,#6ee78f)}.gestion-card.purple .gestion-icon{background:linear-gradient(135deg,#7f43c7,#b77cff);color:#fff}
.dashboard-panel{grid-column:span 4}.dashboard-panel .mini-grid{display:grid;grid-template-columns:repeat(2,1fr);gap:12px;margin-top:12px}.dash-metric{background:#171b21;border:1px solid #343a43;border-radius:14px;padding:14px;display:flex;justify-content:space-between;align-items:center}.dash-metric b{font-size:24px;color:var(--yellow)}
@media(max-width:1100px){.gestion-cards{grid-template-columns:1fr}.dashboard-panel{grid-column:span 12}.span-6{grid-column:span 12}}



/* === PORTAL HR PRO 2026: DASHBOARD ADMIN IGUAL A REFERENCIA === */
.app{grid-template-columns:270px minmax(0,1fr);background:radial-gradient(circle at 70% -10%,rgba(16,185,129,.08),transparent 26%),#0f141a!important}.app.side-collapsed{grid-template-columns:82px minmax(0,1fr)}.main{padding:20px 24px 32px!important;background:linear-gradient(180deg,#0f141a,#121820)!important}.side{width:270px!important;padding:8px 6px!important;background:linear-gradient(180deg,#111720,#171c23)!important;border-right:1px solid rgba(255,255,255,.08)}.side.collapsed{width:82px!important}.app.side-collapsed .side{width:82px!important}.side-top{display:flex!important;height:48px!important;background:#111720!important}.side-top .label{font-size:13px!important}.brand{display:none}.side-user{display:none}.menu-group{margin:8px 0!important}.menu-title,.menu-item{min-height:46px!important;border-radius:10px!important;font-size:14px!important;font-weight:950!important;padding:12px 16px!important;gap:12px!important}.submenu>.menu-item{padding:11px 18px 11px 36px!important;min-height:38px!important;font-size:13px!important;margin:2px 0!important}.submenu .menu-group{margin:6px 0 6px 10px!important}.submenu .menu-group .menu-title{min-height:40px!important;font-size:13px!important;padding:11px 14px!important}.menu-item.active,.menu-item.parent-active{border-left:4px solid var(--yellow)!important;background:linear-gradient(90deg,rgba(16,185,129,.16),rgba(16,185,129,.05))!important;box-shadow:inset 4px 0 0 var(--yellow)!important}.menu-group.force-open>.menu-title{background:linear-gradient(135deg,#16a34a,#10b981)!important;color:#111820!important;box-shadow:0 12px 28px rgba(22,163,74,.25)!important}.menu-group.nested.force-open>.menu-title{background:rgba(16,185,129,.08)!important;color:#f8fafc!important;border-left:4px solid var(--yellow)!important;box-shadow:inset 4px 0 0 var(--yellow)!important}.menu-group.nested.force-open>.menu-title .chev{color:#fff}.admin-shell{max-width:1560px;margin:0 auto}.admin-header{display:flex;align-items:flex-start;justify-content:space-between;gap:20px;margin:0 0 18px}.admin-title h1{margin:0 0 4px;font-size:27px;line-height:1.15}.admin-title .role{font-size:17px;font-weight:1000;color:var(--yellow);margin-bottom:18px}.admin-title p{margin:0;color:#e5edf8;font-weight:750}.hambox{width:34px;height:34px;display:grid;place-items:center;border-radius:10px;background:#151c24;border:1px solid rgba(255,255,255,.05);margin-right:14px}.admin-title-row{display:flex;align-items:flex-start}.top-actions{display:flex;align-items:center;gap:13px}.top-icon{width:39px;height:39px;border-radius:10px;background:#151c24;border:1px solid rgba(255,255,255,.05);display:grid;place-items:center;position:relative;font-size:19px}.top-icon i{position:absolute;right:-4px;top:-8px;background:#ff4d5c;color:white;border-radius:999px;font-size:11px;min-width:20px;height:20px;display:grid;place-items:center;font-style:normal}.admin-chip{display:flex;align-items:center;gap:10px;font-weight:950}.admin-chip .a{width:38px;height:38px;border-radius:50%;background:linear-gradient(135deg,#16a34a,#10b981);color:#111;display:grid;place-items:center}.gestion-cards{grid-template-columns:repeat(3,minmax(240px,1fr))!important;gap:16px!important;background:rgba(255,255,255,.025);border:1px solid rgba(255,255,255,.05);border-radius:18px;padding:18px}.gestion-card{min-height:160px!important;padding:26px!important;align-items:center!important}.gestion-card h2{font-size:20px;margin:0 0 12px}.gestion-card p{min-height:52px;line-height:1.5}.gestion-card .btn-warn,.gestion-card .btn-green,.gestion-card .btn-blue{margin-top:8px;min-width:150px;justify-content:center;border-radius:8px;padding:10px 15px}.gestion-card.green .btn-green{background:#163024!important;color:#43d96d!important;border:1px solid rgba(67,217,109,.4)!important}.gestion-card.purple .btn-blue{background:#241a33!important;color:#bc85ff!important;border:1px solid rgba(188,133,255,.45)!important}.gestion-icon{width:64px!important;height:64px!important;border-radius:14px!important}.dashboards-admin{display:grid;grid-template-columns:repeat(3,minmax(240px,1fr));gap:16px;margin-top:16px}.dashboard-panel{grid-column:auto!important;padding:22px!important}.dashboard-panel h2{font-size:16px;margin-bottom:18px}.dashboard-panel .mini-grid{grid-template-columns:repeat(2,1fr)!important;gap:14px!important}.dash-metric{min-height:86px!important;align-items:flex-start!important;position:relative;display:block!important}.dash-metric span{display:block;font-size:12px;color:#f2f6fb;margin-bottom:10px}.dash-metric b{font-size:20px!important;color:#fff!important}.dash-metric .mi{position:absolute;right:13px;bottom:13px;width:38px;height:38px;border-radius:9px;display:grid;place-items:center;background:linear-gradient(135deg,#16a34a,#10b981);color:#101418}.dashboard-panel.green .dash-metric .mi{background:linear-gradient(135deg,#27b862,#6ee78f);color:#fff}.dashboard-panel.purple .dash-metric .mi{background:linear-gradient(135deg,#7f43c7,#b77cff);color:#fff}.dashboard-panel .full-link{margin-top:16px;width:100%;justify-content:space-between;border-radius:8px;padding:12px 16px}.dashboard-panel.green .full-link{background:#142c23!important;color:#44d96d!important;border:1px solid rgba(68,217,109,.38)!important}.dashboard-panel.purple .full-link{background:#241831!important;color:#bd86ff!important;border:1px solid rgba(189,134,255,.42)!important}.admin-footer{display:flex;justify-content:space-between;color:#aab4c1;font-size:12px;margin:24px 4px 0}.test-panel-hidden{margin-top:16px}.admin-section-title{font-size:18px;margin:0 0 14px}.card{background:linear-gradient(145deg,#181e26,#14191f)!important;border-color:#313946!important}@media(max-width:1200px){.dashboards-admin,.gestion-cards{grid-template-columns:1fr!important}.app{grid-template-columns:280px 1fr}.admin-header{flex-direction:column}.top-actions{align-self:flex-end}}@media(max-width:1000px){.main{padding-top:74px!important}.app{grid-template-columns:1fr!important}.side{padding-top:60px!important}.dashboards-admin,.gestion-cards{grid-template-columns:1fr!important}}

/* === MEJORA VISUAL VACACIONES: solicitud y mis solicitudes === */
.vac-request-card{padding:28px!important;border-radius:22px!important;overflow:hidden;position:relative}.vac-request-card:before{content:"";position:absolute;right:-80px;top:-90px;width:230px;height:230px;border-radius:50%;background:rgba(16,185,129,.10);pointer-events:none}.vac-head{display:flex;justify-content:space-between;gap:18px;align-items:flex-start;margin-bottom:20px}.vac-head h2{margin:0 0 8px;font-size:26px}.vac-help{color:#c2ccd8;font-weight:800;line-height:1.45;margin:0}.period-list{display:grid;grid-template-columns:repeat(3,minmax(230px,1fr));gap:14px;margin-top:10px}.period-card{position:relative;display:grid!important;grid-template-columns:auto 1fr;gap:12px;align-items:flex-start;margin:0!important;padding:18px!important;border:1px solid #34404d;border-radius:18px;background:linear-gradient(145deg,#171c23,#11161c);cursor:pointer;min-height:104px;box-shadow:0 12px 26px rgba(0,0,0,.16);transition:.15s}.period-card:hover{transform:translateY(-1px);border-color:rgba(16,185,129,.7);box-shadow:0 16px 35px rgba(0,0,0,.28)}.period-card input{width:18px;height:18px;accent-color:#10b981;margin-top:4px}.period-main{display:grid;gap:8px}.period-years{font-size:18px;color:#fff;font-weight:1000;letter-spacing:.2px}.period-meta{display:flex;flex-wrap:wrap;gap:8px}.period-badge{background:#202732;border:1px solid #3a4654;border-radius:999px;padding:6px 10px;color:#d9e5f2;font-size:12px;font-weight:950}.period-badge strong{color:var(--yellow)}.period-card input:disabled~.period-main{opacity:.45}.vac-form-row{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-top:18px}.vac-form-row .field{min-width:0}.vac-submit-row{display:grid;grid-template-columns:1fr 1fr;gap:16px;align-items:end;margin-top:16px}.check-card{height:100%;display:flex!important;align-items:center;gap:10px;padding:15px 16px!important;border:1px solid #34404d;border-radius:16px;background:#11161c;color:#eaf3ff;font-weight:950}.check-card input{accent-color:#10b981;width:18px;height:18px}.vac-submit-row .btn-green{height:52px;justify-content:center;font-size:15px}.sol-cards{display:grid;gap:14px}.sol-card{display:grid;grid-template-columns:1.12fr 1.55fr .55fr 1.65fr 1.2fr 1fr;gap:14px;align-items:center;background:linear-gradient(145deg,#171c23,#11161d);border:1px solid #33404d;border-radius:18px;padding:18px 20px;box-shadow:0 12px 30px rgba(0,0,0,.18)}.sol-card.head{background:#0b1015;color:var(--yellow);font-size:13px;font-weight:1000;text-transform:uppercase;box-shadow:none;border-radius:18px 18px 8px 8px}.sol-card:not(.head){border-left:5px solid var(--yellow)}.sol-card b{color:#fff;font-size:16px}.sol-card .dias b{display:inline-grid;place-items:center;min-width:46px;height:38px;border-radius:12px;background:rgba(16,185,129,.12);color:var(--yellow);border:1px solid rgba(16,185,129,.28)}.sol-card .coment{color:#dbe4ee;font-weight:850}.sol-empty{padding:22px;border:1px dashed #3b4552;border-radius:16px;color:#b8c0cb}.local-note{margin-top:10px;color:#b8c0cb;font-size:13px}@media(max-width:1100px){.period-list{grid-template-columns:1fr}.vac-form-row,.vac-submit-row{grid-template-columns:1fr}}@media(max-width:900px){.sol-card,.sol-card.head{grid-template-columns:1fr}.sol-card.head{display:none}.sol-card{gap:9px}.sol-card>div:before{content:attr(data-label);display:block;color:var(--yellow);font-size:11px;text-transform:uppercase;margin-bottom:3px}.vac-head{display:block}}


/* =========================================================
   OVERRIDE UI PORTAL HR PRO - SOLO AMARILLO / SIN VERDE
   Pantallas Gestión Contratación + Crear Plantilla
   ========================================================= */
:root{
  --brand:#0b2f4a!important;
  --brand2:#08263d!important;
  --accent:#16a34a!important;
  --accent2:#10b981!important;
  --bg:#eef4f8!important;
  --paper:#ffffff!important;
  --ink:#061a33!important;
  --soft:#5b6878!important;
  --border:#cdd9e6!important;
}
body{background:var(--bg)!important;color:var(--ink)!important;font-family:Inter,Segoe UI,Arial,sans-serif!important;font-weight:750!important}
.app{background:var(--bg)!important}.main{background:linear-gradient(135deg,#f7fbff 0%,#edf4f8 45%,#e9f1f6 100%)!important;color:var(--ink)!important;padding-top:32px!important}.hero{background:transparent!important;border-bottom:0!important;padding-top:10px!important}.topbar h1,.c-title{color:var(--ink)!important;text-shadow:none!important;letter-spacing:-.7px!important;font-weight:1000!important}.subtitle,.muted,.muted2{color:#687589!important}
.side{background:linear-gradient(180deg,#083454 0%,#061d33 100%)!important;border-right:1px solid rgba(255,255,255,.12)!important;color:#eaf3ff!important}.side-top{background:#062840!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.menu-title,.menu-item{color:#eaf3ff!important}.menu-group.force-open .menu-title,.menu-title.active,.menu-item.active,.menu-group.nested.force-open>.menu-title{background:linear-gradient(135deg,var(--accent2),var(--accent))!important;color:#061a33!important;border-left-color:var(--accent)!important;box-shadow:0 14px 32px rgba(22,163,74,.22)!important}.menu-group.force-open .menu-title *, .menu-title.active *, .menu-item.active *, .menu-group.nested.force-open>.menu-title *{color:#061a33!important}.menu-title:hover,.menu-item:hover{background:rgba(16,185,129,.14)!important;border-left-color:var(--accent2)!important}.side .btn-green,.side .green,.gestion-card.green .btn-green{background:linear-gradient(135deg,var(--accent2),var(--accent))!important;color:#061a33!important;border-color:rgba(22,163,74,.45)!important}
.card,.c-card,.filter-card,.create-card{background:rgba(255,255,255,.96)!important;color:var(--ink)!important;border:1px solid var(--border)!important;border-radius:22px!important;box-shadow:0 18px 42px rgba(9,46,75,.10)!important}.c-card h2,.c-card h3,.card h2,.card h3{color:var(--ink)!important}.input,input,select,textarea{background:#fff!important;color:var(--ink)!important;border:1.5px solid #c5d2df!important;border-radius:14px!important;font-weight:900!important;box-shadow:none!important}input:focus,select:focus,textarea:focus{border-color:var(--accent)!important;box-shadow:0 0 0 4px rgba(22,163,74,.18)!important;outline:none!important}option{background:#fff!important;color:var(--ink)!important}.c-btn,.btn,.btn-blue,.btn-green,.crear-btn{background:linear-gradient(135deg,var(--accent2),var(--accent))!important;color:#061a33!important;border:0!important;border-radius:16px!important;font-weight:1000!important;box-shadow:0 12px 28px rgba(22,163,74,.28)!important}.c-btn.gray,.btn.gray,.gray{background:#6b7788!important;color:#fff!important;box-shadow:none!important}.btn-red{background:#ffe8e8!important;color:#b42318!important;border:1px solid #ffc8c8!important}.state-pill.ok,.status-pill,.pill.ok{background:#fff7d1!important;color:#8a5a00!important;border:1px solid #ffd66b!important}.state-pill.bad{background:#ffe8e8!important;color:#b42318!important;border:1px solid #ffc8c8!important}
.table-wrap{background:#fff!important;border:1px solid var(--border)!important;border-radius:20px!important;box-shadow:0 16px 36px rgba(9,46,75,.08)!important}table,.c-table,.tpl-table{background:#fff!important;color:var(--ink)!important;border-collapse:separate!important;border-spacing:0!important}.c-table th,.tpl-table th,th{background:#e9f1f7!important;color:#15304a!important;border-bottom:1px solid #cfdae6!important;text-transform:uppercase!important;letter-spacing:.04em!important;font-weight:1000!important}.c-table td,.tpl-table td,td{background:#fff!important;color:#13263e!important;border-bottom:1px solid #e1e8f0!important;font-weight:850!important}.c-table tr:nth-child(even) td,.tpl-table tr:nth-child(even) td{background:#f7fafc!important}.plantilla-table td:nth-child(n+4),.plantilla-table td:nth-child(n+4) *{color:#31445a!important;opacity:1!important}.plantilla-table td:nth-child(3){color:#061a33!important;font-weight:1000!important}.plantilla-filter{display:grid!important;grid-template-columns:180px 1fr 210px 1fr!important;gap:18px 22px!important;align-items:center!important}.plantilla-filter b,.c-form b{color:var(--ink)!important;font-weight:1000!important}.plantilla-top{display:flex!important;align-items:center!important;justify-content:space-between!important;margin-bottom:28px!important}.plantilla-top .crear-btn{padding:18px 30px!important;border-radius:999px!important;font-size:17px!important}
/* Modal Crear Plantilla */
.modal-prize{position:fixed;inset:0;background:rgba(3,14,26,.58);z-index:9999;display:none;align-items:center;justify-content:center;padding:24px}.modal-prize:target{display:flex}.modal-box{width:min(920px,96vw);max-height:92vh;overflow:auto;background:#fff;border-radius:18px;border:1px solid #d8e2ec;box-shadow:0 28px 80px rgba(0,0,0,.32);color:var(--ink)}.modal-head{display:flex;align-items:center;justify-content:space-between;padding:24px 28px;border-bottom:1px solid #e3e9f0}.modal-head h2{margin:0;color:var(--ink);font-size:28px;font-weight:1000}.modal-close{font-size:28px;color:#7a8593;font-weight:1000}.modal-body{padding:32px 48px 42px}.modal-form{display:grid;grid-template-columns:210px 1fr;gap:12px 18px;align-items:start}.modal-form label{font-size:20px;font-weight:850;color:#2d3642;text-align:right;padding-top:12px}.modal-form label.req{color:#f04438}.modal-form input,.modal-form select,.modal-form textarea{min-height:46px;font-size:18px;border-radius:9px!important}.modal-form textarea{min-height:74px}.modal-help{grid-column:2;color:#f04438;font-size:15px;font-weight:750;margin-top:-8px}.modal-actions{grid-column:2;display:flex;gap:12px;justify-content:flex-end;margin-top:24px}.modal-actions .c-btn{padding:12px 22px!important;border-radius:10px!important}.select-soft{background:#e3e7ec!important}.file-row{padding:10px 12px!important;background:#fff!important;border:1.5px solid #c5d2df!important;border-radius:9px!important;font-weight:900}.only-yellow *{--green:var(--accent)!important}
@media(max-width:900px){.plantilla-filter{grid-template-columns:1fr!important}.modal-form{grid-template-columns:1fr}.modal-form label{text-align:left}.modal-help,.modal-actions{grid-column:1}.plantilla-top{align-items:flex-start!important;gap:14px;flex-direction:column}}


/* =========================================================
   CORRECCIÓN FINAL: INTERFAZ AMARILLO + NEGRO, SIN VERDE
   Crear Plantilla + Esquemas WORD visibles
   ========================================================= */
:root{
  --bg:#0f141a!important;
  --panel:#151b22!important;
  --panel2:#10161d!important;
  --ink:#f8fafc!important;
  --muted:#aeb8c6!important;
  --border:#303946!important;
  --accent:#16a34a!important;
  --accent2:#10b981!important;
  --danger:#ff4057!important;
}
body{background:#0f141a!important;color:var(--ink)!important;font-family:Inter,Segoe UI,Arial,sans-serif!important;font-weight:850!important}.app{background:#0f141a!important}.main{background:radial-gradient(circle at 95% -8%,rgba(16,185,129,.12),transparent 24%),linear-gradient(180deg,#0f141a,#111821)!important;color:var(--ink)!important}.side{background:linear-gradient(180deg,#09131d,#101823 55%,#0b1119)!important;border-right:1px solid rgba(16,185,129,.16)!important;box-shadow:14px 0 38px rgba(0,0,0,.35)!important}.side-top{background:#08121c!important;border-bottom:1px solid rgba(16,185,129,.18)!important;color:#fff!important}.menu-title,.menu-item{background:transparent!important;color:#e9f1fb!important}.menu-title:hover,.menu-item:hover{background:rgba(16,185,129,.09)!important;color:#fff!important}.menu-group.force-open>.menu-title,.menu-title.active,.menu-item.active,.menu-item.parent-active{background:linear-gradient(135deg,#16a34a,#10b981)!important;color:#0b1119!important;border-left:4px solid #fff0a6!important;box-shadow:0 12px 28px rgba(22,163,74,.24)!important}.menu-group.nested.force-open>.menu-title,.nested.force-open>.menu-item.parent-active{background:linear-gradient(90deg,rgba(16,185,129,.23),rgba(16,185,129,.08))!important;color:#fff!important;border-left:4px solid var(--accent2)!important}.menu-group.nested.force-open>.menu-title .chev{color:#fff!important}.c-title,.plantilla-top .c-title,h1,h2,h3{color:#f8fafc!important;text-shadow:none!important}.plantilla-top{margin-bottom:22px!important}.c-card,.filter-card,.table-wrap,.card{background:linear-gradient(145deg,#171f28,#111821)!important;border:1px solid var(--border)!important;border-radius:20px!important;box-shadow:0 18px 42px rgba(0,0,0,.25)!important;color:#f8fafc!important}.plantilla-filter b,.c-form b,label{color:#f8fafc!important}.input,input,select,textarea{background:#0b1119!important;color:#f8fafc!important;border:1.5px solid #34404d!important;border-radius:12px!important;font-weight:900!important}input::placeholder,textarea::placeholder{color:#7f8b99!important}input:focus,select:focus,textarea:focus{border-color:var(--accent2)!important;box-shadow:0 0 0 4px rgba(16,185,129,.18)!important;outline:none!important}option{background:#0b1119!important;color:#f8fafc!important}.c-btn,.btn,.btn-green,.btn-blue,.crear-btn,.modal-actions .c-btn{background:linear-gradient(135deg,#16a34a,#10b981)!important;color:#111820!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 28px rgba(22,163,74,.25)!important}.c-btn.gray,.btn.gray,.gray{background:#2b3441!important;color:#f8fafc!important;border:1px solid #3e4a58!important;box-shadow:none!important}.btn-red{background:#32151c!important;color:#ffc4ca!important;border:1px solid #7f1d1d!important}.state-pill.ok,.status-pill,.pill.ok{background:rgba(16,185,129,.13)!important;color:#10b981!important;border:1px solid rgba(16,185,129,.45)!important}.state-pill.bad{background:#32151c!important;color:#ffc4ca!important;border:1px solid #7f1d1d!important}.badge-green,.gestion-card.green .gestion-icon,.gestion-card.green .btn-green,.dashboard-panel.green .dash-metric .mi,.dashboard-panel.green .full-link{background:linear-gradient(135deg,#16a34a,#10b981)!important;color:#111820!important;border-color:rgba(16,185,129,.35)!important}.gestion-card.green,.gestion-card.purple{border-color:rgba(16,185,129,.38)!important}.gestion-card.purple .gestion-icon,.dashboard-panel.purple .dash-metric .mi{background:linear-gradient(135deg,#16a34a,#10b981)!important;color:#111820!important}.dashboard-panel.green .full-link,.dashboard-panel.purple .full-link{background:rgba(16,185,129,.10)!important;color:#10b981!important;border:1px solid rgba(16,185,129,.35)!important}.c-table,.tpl-table,table{background:#111821!important;color:#f8fafc!important}.c-table th,.tpl-table th,th{background:#0b1119!important;color:#10b981!important;border-bottom:1px solid #364250!important;text-transform:uppercase!important;letter-spacing:.04em!important;font-weight:1000!important}.c-table td,.tpl-table td,td{background:#151d26!important;color:#eaf2fb!important;border-bottom:1px solid #2b3541!important;font-weight:850!important}.c-table tr:nth-child(even) td,.tpl-table tr:nth-child(even) td{background:#111821!important}.c-table tr:hover td,.tpl-table tr:hover td{background:#1d2732!important}.plantilla-table td:nth-child(n),.plantilla-table td:nth-child(n) *{color:#eaf2fb!important;opacity:1!important}.plantilla-table td:nth-child(3){color:#fff!important;font-weight:1000!important}.modal-prize{background:rgba(0,0,0,.72)!important}.modal-box{background:#10161d!important;color:#f8fafc!important;border:1px solid #34404d!important;box-shadow:0 32px 90px rgba(0,0,0,.65)!important}.modal-head{border-bottom:1px solid #2f3946!important;background:#0b1119!important}.modal-head h2{color:#fff!important}.modal-close{color:#cbd5e1!important}.modal-body{background:#10161d!important}.modal-form label{color:#eaf2fb!important;font-weight:900!important}.modal-form label.req{color:#10b981!important}.modal-help{color:#16a34a!important;font-weight:900!important}.select-soft{background:#0b1119!important}.file-row{background:#0b1119!important;border:1.5px solid #34404d!important;color:#f8fafc!important}.modal-form select#modal_esquema option{background:#0b1119!important;color:#f8fafc!important}.modal-form select#modal_esquema{border-color:#10b981!important}.only-yellow *{--green:#10b981!important;--success:#10b981!important}


/* ===== AJUSTES PRO FICHA / PLANTILLAS 2026 ===== */
.main{overflow-x:hidden!important}.plantilla-top{display:flex!important;align-items:center!important;justify-content:space-between!important;gap:16px!important}.plantilla-top .crear-btn{position:relative!important;right:auto!important;top:auto!important;white-space:nowrap!important;min-width:190px!important;justify-content:center!important}.plantilla-filter{grid-template-columns:180px minmax(240px,1fr) 180px minmax(240px,1fr)!important}.plantilla-table{min-width:1220px!important}.plantilla-table th,.plantilla-table td{white-space:normal!important;line-height:1.28!important}.plantilla-table td,.plantilla-table td *{color:#f8fafc!important;opacity:1!important;text-shadow:none!important}.plantilla-table th{color:#10b981!important}.table-wrap{max-width:100%!important;overflow:auto!important}.ficha-search{display:flex;gap:12px;margin:0 0 16px;max-width:720px}.ficha-search input{flex:1}.ficha-profile{display:grid;grid-template-columns:150px minmax(320px,1.4fr) minmax(260px,1fr) minmax(260px,1fr);gap:18px;align-items:stretch;margin-bottom:18px}.avatar-panel,.profile-main,.profile-col{background:linear-gradient(145deg,#171f28,#111821);border:1px solid #34404d;border-radius:18px;padding:18px;color:#f8fafc;box-shadow:0 16px 34px rgba(0,0,0,.22)}.avatar-panel{display:grid;place-items:center}.profile-main h2{margin:0 0 10px;color:#fff;font-size:22px}.profile-main p,.profile-col p{margin:8px 0;color:#eaf2fb}.created-box{background:#0b1119;border:1px solid #34404d;border-radius:12px;padding:10px;margin-top:12px;color:#cbd5e1}.status-dot,.status-pill{display:inline-flex;align-items:center;gap:6px;padding:7px 12px;border-radius:999px;font-weight:1000}.status-dot.ok,.status-pill.ok{background:rgba(34,197,94,.13);color:#86efac;border:1px solid rgba(34,197,94,.45)}.status-dot.bad,.status-pill.bad{background:rgba(244,63,94,.13);color:#fecdd3;border:1px solid rgba(244,63,94,.45)}.ficha-tabs{display:flex!important;padding:0!important;margin:0 0 16px!important;overflow:hidden}.ficha-tabs .tab{flex:1;background:transparent;border:0;border-bottom:3px solid transparent;color:#cbd5e1!important;cursor:pointer;font-size:16px}.ficha-tabs .tab.active{color:#10b981!important;border-bottom-color:#10b981!important}.ficha-tab-content{display:none}.ficha-tab-content.active{display:block}.laboral-grid{display:grid;grid-template-columns:1fr 1fr;gap:14px 28px;padding:24px!important}.laboral-grid label{display:grid;grid-template-columns:220px 1fr;align-items:center;gap:12px;color:#f8fafc!important}.laboral-grid input{background:#0b1119!important;color:#eaf2fb!important;border:1px solid #34404d!important}.periodos-box{margin-top:14px;background:#111821;border:1px solid #34404d;border-radius:14px;padding:14px;color:#eaf2fb}.mini-chip{display:inline-flex;margin:5px;padding:8px 12px;border-radius:999px;background:rgba(16,185,129,.12);color:#10b981;border:1px solid rgba(16,185,129,.35);font-weight:900}.period-row td{background:#0b1119!important;color:#10b981!important;font-size:16px}.check-green{display:inline-grid;place-items:center;min-width:54px;padding:7px 13px;background:#4ea60f;color:white;border-radius:8px;font-weight:1000}.check-gray{display:inline-grid;place-items:center;min-width:54px;padding:7px 13px;background:#334155;color:#cbd5e1;border-radius:8px;font-weight:1000}@media(max-width:1100px){.ficha-profile{grid-template-columns:1fr}.laboral-grid{grid-template-columns:1fr}.laboral-grid label{grid-template-columns:1fr}.plantilla-top{flex-direction:column;align-items:flex-start!important}.plantilla-filter{grid-template-columns:1fr!important}}


/* ===== CORRECCIÓN VISIBILIDAD FINAL PANEL / LIENZO ===== */
.side,.side *{color:#f8fafc!important;text-shadow:none!important;opacity:1!important;filter:none!important;}
.side .menu-title,.side .menu-title *,.side .menu-item,.side .menu-item *{color:#f1f5f9!important;font-weight:1000!important;letter-spacing:.01em!important;}
.side .menu-group.force-open>.menu-title,.side .menu-group.force-open>.menu-title *{color:#08111b!important;}
.side .menu-group.nested.force-open>.menu-title,.side .menu-group.nested.force-open>.menu-title *,.side .menu-item.parent-active,.side .menu-item.parent-active *{color:#fff!important;}
.side .submenu .menu-item,.side .submenu .menu-item *{color:#e5edf7!important;font-weight:950!important;}
.side .menu-item.active,.side .menu-item.active *{color:#08111b!important;}
.main,.main *{text-shadow:none!important;}

/* === ULTRAMEJORA 2026-05: sidebar legible, scroll estable y textos visibles === */
.app{grid-template-columns:350px 1fr!important}
.app.side-collapsed{grid-template-columns:92px 1fr!important}
.side{width:350px!important;background:linear-gradient(180deg,#07111d 0%,#0b1522 58%,#070d16 100%)!important;border-right:1px solid rgba(255,193,7,.35)!important;scrollbar-width:thin;scrollbar-color:#ffc107 #07111d}
.side::-webkit-scrollbar{width:10px}.side::-webkit-scrollbar-thumb{background:#ffc107;border-radius:999px}.side::-webkit-scrollbar-track{background:#07111d}
.side-top{background:#06111d!important;border-bottom:1px solid rgba(255,193,7,.35)!important}
.menu-group{margin:10px 14px!important;overflow:visible!important}
.menu-title{min-height:58px!important;padding:14px 16px!important;gap:13px!important;font-size:15.5px!important;line-height:1.25!important;white-space:normal!important;overflow:visible!important;background:linear-gradient(135deg,#172230,#101925)!important;border:1px solid rgba(255,255,255,.10)!important;box-shadow:0 10px 24px rgba(0,0,0,.20)!important}
.menu-title .label,.menu-item .label{display:block!important;white-space:normal!important;line-height:1.25!important;overflow:visible!important;text-overflow:clip!important;word-break:normal!important}
.menu-item{min-height:46px!important;padding:12px 16px 12px 44px!important;color:#f4f7fb!important;font-size:14.5px!important;line-height:1.25!important;opacity:1!important;overflow:visible!important}
.sub-mini{padding-left:58px!important;font-size:13.5px!important;color:#f4f7fb!important;opacity:1!important}
.menu-item:hover{background:rgba(255,193,7,.12)!important}.menu-item.active,.menu-title.active{background:linear-gradient(135deg,#ffc107,#ffb020)!important;color:#07111d!important}.menu-item.active *,.menu-title.active *{color:#07111d!important}
.side.collapsed{width:92px!important}.side.collapsed .menu-title{padding:16px 10px!important}.side.collapsed .menu-item{padding:14px 10px!important}
@media(max-width:1000px){.side{width:330px!important;left:-360px}.side.open{left:0!important}.app,.app.side-collapsed{grid-template-columns:1fr!important}}
/* FIX PRO 2026: panel lateral sin letras montadas y con scroll limpio */
.side{box-sizing:border-box!important;overflow-y:auto!important;overflow-x:hidden!important;padding-bottom:120px!important;scrollbar-gutter:stable!important;}
.side *{box-sizing:border-box!important;}
.side .menu-group{overflow:visible!important;margin:7px 10px!important;}
.side .submenu{overflow:visible!important;max-height:none!important;padding:6px 0!important;}
.side .menu-group.closed>.submenu{max-height:0!important;padding:0!important;overflow:hidden!important;}
.side .menu-title,.side .menu-item{position:relative!important;width:100%!important;min-height:44px!important;height:auto!important;line-height:1.28!important;display:flex!important;align-items:center!important;gap:10px!important;white-space:normal!important;word-break:normal!important;overflow:visible!important;}
.side .menu-title .label,.side .menu-item .label{min-width:0!important;display:block!important;white-space:normal!important;line-height:1.28!important;overflow:visible!important;color:inherit!important;}
.side .submenu>.menu-item{padding-left:38px!important;margin:3px 0!important;}
.side .submenu .submenu>.menu-item,.side .sub-mini{padding-left:48px!important;font-size:13.5px!important;}
.side-user{position:static!important;display:flex!important;margin:18px 14px 28px!important;background:rgba(15,23,42,.92)!important;border:1px solid rgba(148,163,184,.25)!important;border-radius:14px!important;padding:12px!important;box-shadow:none!important;}
.side.collapsed .side-user{display:none!important;}
.side.collapsed{overflow:hidden!important;}
.side.collapsed .submenu{display:none!important;}


.contract-detail-wrap,.contract-detail-wrap *{opacity:1!important;filter:none!important;}
.contract-detail-wrap{background:#f4f7fb!important;color:#0f172a!important;}
.contract-detail-wrap .template-head,.contract-detail-wrap .tpl-tabs,.contract-detail-wrap .preview-tools{background:#ffffff!important;color:#0f172a!important;}
.contract-detail-wrap h1,.contract-detail-wrap h2,.contract-detail-wrap h3,.contract-detail-wrap b{color:#0f172a!important;}
.contract-detail-wrap .tpl-line,.contract-detail-wrap .tpl-line *{color:#334155!important;}
.contract-detail-wrap .tpl-line b{color:#0f172a!important;}
.contract-detail-wrap .c-btn.gray{background:#263243!important;color:#f8fafc!important;}
.contract-detail-wrap .c-btn.green,.contract-detail-wrap .c-btn{color:#111827!important;}
.edit-overlay .modal-page,.edit-overlay .modal-page *{opacity:1!important;filter:none!important;text-shadow:none!important;}
.edit-overlay .modal-page{font-size:14px!important;}
.edit-overlay .modal-head h1,.edit-overlay label,.edit-overlay .actual-file{color:#111827!important;}
.edit-overlay input,.edit-overlay select,.edit-overlay textarea{color:#111827!important;background:#fff!important;}

.menu-item.doc-loaded{background:linear-gradient(135deg,#0f5132,#16a34a)!important;color:#fff!important;border-left:4px solid #86efac!important}.menu-item.doc-loaded .label,.menu-item.doc-loaded span{color:#fff!important}


/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}


/* Ajuste final: sin módulo contratación */
.gestion-cards,.dashboards-admin{grid-template-columns:repeat(2,minmax(280px,1fr))!important}
@media(max-width:1000px){.gestion-cards,.dashboards-admin{grid-template-columns:1fr!important}}

/* === FIX FINAL: PORTAL HR PRO VERDE/BLANCO, SIN AMARILLO, LETRAS NÍTIDAS === */
:root{--yellow:#10b981!important;--yellow2:#16a34a!important;--dark:#eef2f5!important;--panel:#ffffff!important;--panel2:#ffffff!important;--line:#dbe5ee!important;--txt:#0f172a!important;--mut:#64748b!important}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:800!important;text-rendering:geometricPrecision!important;-webkit-font-smoothing:antialiased!important}.main{background:#eef2f5!important;color:#0f172a!important}.hero,.card,.mini,.metric,.stat,.module-tile,.dashboard-panel,.gestion-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2,.admin-title h1{color:#0f172a!important;font-weight:1000!important}.subtitle,.muted,.muted2{color:#64748b!important}.accent,.mini b{color:#15945a!important}.btn,.btn-blue,.btn-green,.btn-warn,.btn-yellow,.c-btn,.full-link{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus,input:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}th{background:#f3f7fb!important;color:#334155!important}td{color:#1f2937!important}.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important}.side-top{background:#124f34!important}.brand:before{content:'PORTAL HR PRO'!important;color:#fff!important;font-size:22px!important;font-weight:1000!important}.brand p{font-size:0!important}.brand p:before{content:'Actualización de Datos';font-size:13px!important;color:#d1fae5!important;font-weight:900!important}.menu-title,.menu-item{color:#d7e1ea!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important;border-radius:26px!important;border:1px solid #dbe5ee!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important;text-align:center!important}.login-title b{color:#64748b!important;font-size:17px!important;display:block;text-align:center!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}



/* === PORTAL HR PRO FINAL: estilo verde/blanco limpio, sin amarillo y letras suaves === */
:root{--yellow:#16a34a!important;--yellow2:#10b981!important;--txt:#0f172a!important;--mut:#64748b!important;--line:#dbe5ee!important;--portal-green:#15945a!important;--portal-dark:#091827!important;}
html,body{font-family:'Inter','Segoe UI',Arial,sans-serif!important;font-weight:600!important;background:#eef2f5!important;color:#0f172a!important;-webkit-font-smoothing:antialiased!important;text-rendering:geometricPrecision!important;}
body *{letter-spacing:-.01em;}
.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px 48px!important;}
.hero,.card,.mini,.metric,.stat,.module-tile,.dashboard-panel,.gestion-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;border-radius:24px!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;}
.hero h1,.topbar h1,.card h2,.module-tile h2,.admin-title h1,h1,h2,h3{color:#0f172a!important;font-weight:850!important;}
.subtitle,.muted,.muted2,p,small{color:#64748b!important;font-weight:500!important;}
.accent,.mini b{color:#15945a!important;}
.btn,.btn-blue,.btn-green,.btn-warn,.btn-yellow,.c-btn,.full-link,button.btn-green{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:800!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important;}
.btn-red{border:0!important;border-radius:14px!important;}
.input,select,textarea,input:not([type=checkbox]):not([type=radio]),input[type=file],.form-grid input,.form-grid select,.form-grid textarea{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important;box-shadow:none!important;font-weight:500!important;}
.input:focus,select:focus,textarea:focus,input:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important;outline:none!important;}
label{color:#1f2937!important;font-weight:800!important;}
table{background:#fff!important;color:#243042!important;} th{background:#f3f7fb!important;color:#334155!important;font-weight:800!important;} td{color:#1f2937!important;border-bottom:1px solid #e5edf4!important;background:#fff!important;}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important;}
.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34 0%,#0e322d 45%,#091827 100%)!important;border-right:0!important;box-shadow:12px 0 35px rgba(15,23,42,.22)!important;}
.side-top{background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.10)!important;height:72px!important;}
.brand img{display:none!important}.brand{padding:24px 16px 18px!important;text-align:left!important;border-bottom:1px solid rgba(255,255,255,.10)!important;margin:0 14px 12px!important}.brand:before{content:'PORTAL HR PRO'!important;display:block;color:#fff!important;font-size:22px!important;font-weight:900!important}.brand p{font-size:0!important;margin:2px 0 0!important}.brand p:before{content:'Actualización de Datos';font-size:13px!important;color:#d1fae5!important;font-weight:700!important}
.menu-title,.menu-item{color:#d7e1ea!important;font-weight:750!important;border-radius:16px!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#1bb861,#11924f)!important;color:#fff!important;border-left-color:#d9f99d!important;box-shadow:0 12px 25px rgba(0,0,0,.18)!important}.side-user{background:rgba(255,255,255,.10)!important;border-radius:16px!important;padding:14px!important;border:1px solid rgba(255,255,255,.10)!important}.side-user small{color:#d1fae5!important}.avatar{background:#20c775!important;color:#fff!important;}
.login-body{min-height:100vh!important;display:grid!important;place-items:center!important;position:relative!important;overflow:hidden!important;background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important;}
.login-body:after{content:''!important;position:absolute!important;left:-4%!important;right:-4%!important;bottom:-6%!important;height:31%!important;background:linear-gradient(135deg,#22c55e,#064e3b)!important;clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%)!important;z-index:0!important;display:block!important;}
.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important;border-radius:26px!important;border:1px solid #dbe5ee!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important;}
.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;filter:none!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important;}
.login-card:after{display:none!important}.login-logo{display:none!important}.login-title{text-align:center!important;margin:0 0 28px!important}.login-title h1{color:#1f2937!important;font-size:42px!important;font-weight:850!important;letter-spacing:.5px!important;text-align:center!important;margin:0 0 8px!important}.login-title b{color:#64748b!important;font-size:17px!important;font-weight:500!important;display:block!important;text-align:center!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important;box-shadow:none!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important;padding:16px 10px!important;font-weight:500!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:800!important}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important;text-align:center!important}.login-links a{color:#64748b!important;font-weight:700!important}
body::-webkit-scrollbar,.side::-webkit-scrollbar{width:8px!important}body::-webkit-scrollbar-thumb,.side::-webkit-scrollbar-thumb{background:#94a3b8!important;border-radius:999px!important}body::-webkit-scrollbar-track,.side::-webkit-scrollbar-track{background:#e5e7eb!important}


/* ================= PORTAL HR PRO UI PRO 2026 =================
   Ajuste visual solicitado: login claro con campo Empresa, menú verde sobrio,
   tipografías equilibradas, iconos suaves y encabezados sin superposición. */
:root{--hr-green:#149459;--hr-green2:#1bbf70;--hr-deep:#082033;--hr-side:#0d3a31;--hr-side2:#071827;--hr-bg:#eef4f2;--hr-line:#dbe5ee;--hr-text:#152033;--hr-muted:#64748b}
body{background:var(--hr-bg)!important;color:var(--hr-text)!important;font-weight:650!important;letter-spacing:-.01em!important}.app{background:var(--hr-bg)!important}.main{background:linear-gradient(180deg,#f6faf8 0%,#eef4f2 100%)!important;color:var(--hr-text)!important;padding:24px 28px 46px!important;position:relative!important;z-index:1!important;overflow-x:hidden!important}.hero{background:rgba(255,255,255,.96)!important;border:1px solid var(--hr-line)!important;border-radius:26px!important;margin:0 0 20px!important;padding:22px 24px!important;box-shadow:0 14px 34px rgba(15,23,42,.08)!important;position:relative!important;top:auto!important;z-index:1!important}.topbar{gap:14px!important}.topbar h1{color:var(--hr-text)!important;font-size:30px!important;line-height:1.12!important;font-weight:850!important;letter-spacing:-.7px!important}.subtitle,.muted,.muted2,.empty-note{color:var(--hr-muted)!important;font-weight:600!important}.accent,.mini b,.bar-row em{color:var(--hr-green)!important}.card,.mini,.metric,.stat,.module-tile,.dashboard-panel,.gestion-card,.detail-box,.doc-card{background:rgba(255,255,255,.97)!important;border:1px solid var(--hr-line)!important;color:var(--hr-text)!important;border-radius:22px!important;box-shadow:0 12px 30px rgba(15,23,42,.075)!important}.card h2,.card h3,.module-tile h2,.doc-card h3{color:var(--hr-text)!important;font-weight:850!important}.btn,.btn-blue,.btn-green,.btn-warn,.btn-yellow,.c-btn,.full-link{background:linear-gradient(135deg,#18aa62,#0e8b4e)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:850!important;box-shadow:0 10px 20px rgba(20,148,89,.20)!important}.btn:hover,.btn-blue:hover,.btn-green:hover{filter:brightness(1.02)!important;transform:translateY(-1px)!important}.btn-red{border:0!important;border-radius:14px!important}.field label{color:#273244!important;font-weight:750!important;font-size:13px!important}.input,select,textarea,input[type=file],.field input,.field select,.field textarea{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important;min-height:48px!important;font-weight:650!important}.input:focus,select:focus,textarea:focus,input:focus{border-color:var(--hr-green)!important;box-shadow:0 0 0 4px rgba(20,148,89,.13)!important}.table-wrap{background:#fff!important;border:1px solid var(--hr-line)!important;border-radius:18px!important}table{background:#fff!important;color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important;color:#243042!important}tr:hover td{background:#f8fbfd!important}
.side{background:linear-gradient(180deg,#0e4a32 0%,#0b342d 45%,#071827 100%)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:14px 0 34px rgba(15,23,42,.22)!important;width:312px!important;z-index:30!important}.app{grid-template-columns:312px 1fr!important}.app.side-collapsed{grid-template-columns:88px 1fr!important}.side.collapsed{width:88px!important}.side-top{height:64px!important;background:rgba(7,24,39,.18)!important;border-bottom:1px solid rgba(255,255,255,.10)!important;position:sticky!important;top:0!important;z-index:8!important}.side-top b{font-size:16px!important;letter-spacing:.4px!important;color:#fff!important}.toggle{width:38px!important;height:38px!important;border-radius:13px!important;background:rgba(255,255,255,.10)!important;color:#eafff3!important;display:grid!important;place-items:center!important}.brand{padding:20px 18px 14px!important;text-align:left!important;border-bottom:1px solid rgba(255,255,255,.08)!important;margin:0 12px 10px!important;background:transparent!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO'!important;display:block!important;color:#fff!important;font-size:18px!important;font-weight:900!important;letter-spacing:.5px!important}.brand p{font-size:0!important;margin:3px 0 0!important}.brand p:before{content:'Actualización de Datos'!important;font-size:12px!important;color:#d1fae5!important;font-weight:700!important}.menu-group{margin:7px 10px!important;overflow:visible!important}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 0!important;padding:13px 15px!important;font-size:14px!important;font-weight:760!important;line-height:1.22!important}.menu-title span:first-child,.menu-item span:first-child{font-size:17px!important;min-width:22px!important;text-align:center!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#18b85d,#0f8f50)!important;color:#fff!important;border-left:0!important;box-shadow:0 12px 26px rgba(0,0,0,.18)!important}.submenu{padding:3px 0 8px!important}.sub-mini{padding-left:42px!important;font-size:13px!important}.chev{color:#d8fce5!important}.side-user{margin:22px 14px 16px!important;padding:14px!important;background:rgba(255,255,255,.10)!important;border:1px solid rgba(255,255,255,.12)!important;border-radius:18px!important;color:#fff!important}.side-user small{color:#bbf7d0!important}.avatar{background:#20c775!important;color:#fff!important}.side.collapsed .brand:before{content:'HR'!important;text-align:center!important}.side.collapsed .brand{padding:16px 6px!important;text-align:center!important;margin:0 8px 10px!important}.side.collapsed .menu-title,.side.collapsed .menu-item{justify-content:center!important;padding:14px 8px!important}.side.collapsed .side-user{display:none!important}
.login-body{min-height:100vh!important;display:grid!important;place-items:center!important;padding:28px 16px!important;position:relative!important;overflow:hidden!important;background:radial-gradient(circle at 8% 0%,rgba(34,197,94,.20) 0 18%,transparent 18.3%),radial-gradient(circle at 92% 2%,rgba(20,184,166,.17) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:''!important;position:absolute!important;left:-4%!important;right:-4%!important;bottom:-8%!important;height:33%!important;background:linear-gradient(135deg,#22c55e 0%,#059669 45%,#064e3b 100%)!important;clip-path:polygon(0 38%,25% 26%,50% 18%,75% 28%,100% 36%,100% 100%,0 100%)!important;z-index:0!important;display:block!important}.login-card{width:min(92vw,535px)!important;padding:82px 50px 32px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.97)!important;position:relative!important;z-index:2!important;border-radius:28px!important;border:1px solid #dbe5ee!important;box-shadow:0 24px 58px rgba(15,23,42,.14)!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-64px!important;transform:translateX(-50%)!important;width:128px!important;height:128px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:var(--hr-green)!important;display:grid!important;place-items:center!important;font-size:48px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title{text-align:center!important;margin:0 0 24px!important}.login-title h1{color:#1f2937!important;font-size:39px!important;font-weight:850!important;letter-spacing:.4px!important;text-align:center!important;margin:0 0 8px!important}.login-title b{color:#64748b!important;font-size:16px!important;font-weight:600!important;display:block!important}.login-title:after{content:''!important;display:block!important;margin:18px auto 0!important;width:62px!important;height:4px!important;border-radius:999px!important;background:#10b981!important}.login-input{background:#edf5ff!important;border:1px solid #cfe0f2!important;border-radius:17px!important;padding:0 14px!important;margin-bottom:16px!important;color:var(--hr-green)!important;box-shadow:none!important;display:flex!important;align-items:center!important;gap:10px!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important;padding:15px 8px!important;font-weight:650!important;width:100%!important;outline:none!important}.login-input select option{background:#fff!important;color:#0f172a!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:9px 0 7px!important;font-weight:800!important}.login-card .btn-green{width:100%!important;margin:8px 0 24px!important;font-size:19px!important;border-radius:17px!important;padding:16px!important}.login-links{margin-top:0!important;padding-bottom:0!important;text-align:center!important}.login-links a{color:#64748b!important;font-weight:750!important}.flash{background:#dcfce7!important;color:#065f46!important;border:1px solid #a7f3d0!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border:1px solid #fecaca!important}
.mobile-head{height:58px!important;display:none;align-items:center!important;justify-content:space-between!important;background:linear-gradient(135deg,#0e4a32,#071827)!important;color:#fff!important;border-bottom:1px solid rgba(255,255,255,.10)!important;z-index:60!important}.mobile-head a{color:#d1fae5!important;font-weight:750!important}.mobile-head b{font-size:14px!important;letter-spacing:.4px!important}@media(max-width:1000px){.mobile-head{display:flex!important;position:sticky!important;top:0!important}.app,.app.side-collapsed{grid-template-columns:1fr!important}.side{position:fixed!important;left:-330px!important;top:0!important;width:315px!important;height:100dvh!important;z-index:80!important;overflow-y:auto!important;padding-bottom:28px!important}.side.open{left:0!important}.side.collapsed{left:-330px!important}.main{padding:14px 14px 32px!important}.hero{margin:0 0 16px!important;padding:18px 16px!important;border-radius:20px!important;z-index:1!important}.topbar{align-items:flex-start!important;flex-direction:column!important}.topbar h1{font-size:24px!important;line-height:1.15!important}.subtitle{font-size:13px!important}.card{border-radius:18px!important;padding:16px!important}.login-card{width:min(94vw,480px)!important;padding:76px 26px 28px!important}.login-title h1{font-size:32px!important}.login-card:before{width:112px!important;height:112px!important;top:-56px!important;font-size:42px!important}.form-grid{grid-template-columns:1fr!important}.form-grid .field,.form-grid button,.form-grid a{grid-column:1!important;width:100%!important}.side-user{margin-bottom:22px!important}}@media(min-width:1001px){.content,.main{min-height:100vh!important}}
/* ============================================================ */


/* === AJUSTE FINAL SUAVE - PORTAL HR PRO ===
   Tipografía menos pesada, menú limpio tipo referencia y encabezado sin superposición. */
:root{--hr-green:#16a34a;--hr-green-dark:#065f46;--hr-navy:#0f172a;--hr-bg:#f4f8f6;--hr-line:#e2e8f0;--hr-muted:#64748b}
html,body{font-family:"Segoe UI",Inter,Arial,sans-serif!important;background:var(--hr-bg)!important;color:#172033!important;font-weight:500!important;-webkit-font-smoothing:antialiased!important;text-rendering:optimizeLegibility!important}
.main{background:var(--hr-bg)!important;padding:26px 32px 36px!important;color:#172033!important}
h1,h2,h3,.hero h1,.topbar h1,.card h2,.module-tile h2,.admin-title h1{font-weight:760!important;letter-spacing:-.035em!important;color:#111827!important;line-height:1.12!important}
p,.subtitle,.muted,.muted2,label,td,th,.card,.mini,.metric,.stat{font-weight:500!important;color:#334155!important}.subtitle,.muted,.muted2{color:#667085!important}.accent,.mini b{font-weight:700!important;color:#15945a!important}
.card,.mini,.metric,.stat,.module-tile,.dashboard-panel,.gestion-card,.hero{background:#fff!important;border:1px solid #dce8ef!important;border-radius:22px!important;box-shadow:0 12px 32px rgba(15,23,42,.075)!important}.hero{position:relative!important;z-index:1!important;overflow:hidden!important;padding:22px 24px!important;margin-bottom:18px!important}.hero:before,.hero:after{pointer-events:none!important}
.btn,.btn-blue,.btn-green,.btn-warn,.btn-yellow,.c-btn,.full-link{font-weight:720!important;border-radius:14px!important;box-shadow:0 10px 22px rgba(22,163,74,.16)!important;background:linear-gradient(135deg,#18b96d,#0f9a5d)!important;color:#fff!important}.btn:hover,.btn-green:hover{filter:brightness(.98)!important;transform:translateY(-1px)}
.input,select,textarea,input[type=file]{font-weight:520!important;background:#fff!important;color:#172033!important;border:1px solid #d8e4ec!important;border-radius:14px!important;box-shadow:none!important}.input:focus,select:focus,textarea:focus,input:focus{border-color:#17a665!important;box-shadow:0 0 0 4px rgba(22,163,74,.12)!important;outline:none!important}
/* Sidebar más suave: no letras gruesas ni elementos gigantes */
.app{background:var(--hr-bg)!important}.side{background:linear-gradient(180deg,#0b4a32 0%,#073323 52%,#071827 100%)!important;border-right:1px solid rgba(255,255,255,.08)!important;box-shadow:10px 0 26px rgba(15,23,42,.14)!important;color:#e6f6ef!important;overflow-y:auto!important;overflow-x:hidden!important;scrollbar-gutter:stable!important}.side::-webkit-scrollbar{width:8px}.side::-webkit-scrollbar-thumb{background:#facc15;border-radius:99px}.brand{min-height:64px!important;padding:14px 16px!important;border-bottom:1px solid rgba(255,255,255,.08)!important;display:flex!important;align-items:center!important;gap:10px!important}.brand img{width:34px!important;height:34px!important;border-radius:10px!important;background:rgba(255,255,255,.12)!important;padding:3px!important}.brand p{font-size:15px!important;font-weight:780!important;letter-spacing:.3px!important;color:#fff!important;margin:0!important}.menu-title,.menu-item{min-height:44px!important;margin:7px 10px!important;padding:11px 14px!important;border-radius:14px!important;font-size:14px!important;font-weight:650!important;letter-spacing:.01em!important;color:#dbe7ef!important;background:transparent!important;box-shadow:none!important;border:0!important;line-height:1.25!important}.menu-title .label,.menu-item .label{font-weight:650!important;color:inherit!important}.menu-title span:first-child,.menu-item span:first-child{font-size:17px!important;min-width:22px!important;text-align:center!important;filter:none!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.08)!important;color:#fff!important}.menu-title.active,.menu-item.active,.parent-active{background:linear-gradient(135deg,#19b86b,#0f9b59)!important;color:#04140d!important;font-weight:730!important;box-shadow:0 10px 22px rgba(5,150,105,.24)!important}.submenu{padding-left:12px!important}.sub-mini{font-size:13.5px!important;margin-left:20px!important;padding-left:15px!important}.chev{margin-left:auto!important;font-weight:700!important}.side-user{margin:18px 14px!important;padding:14px!important;border-radius:18px!important;background:rgba(255,255,255,.10)!important;border:1px solid rgba(255,255,255,.12)!important;box-shadow:none!important}.side-user b{font-size:15px!important;font-weight:750!important;color:#fff!important}.side-user small{font-size:13px!important;font-weight:550!important;color:#bbf7d0!important}.avatar{width:42px!important;height:42px!important;border-radius:50%!important;background:#18b96d!important;color:#06351f!important;font-size:20px!important;display:grid!important;place-items:center!important}
/* Header interno: fijo visual, sin montarse con el scroll */
.mobile-head{height:58px!important;z-index:120!important}.topbar,.admin-head,.header,.navbar{position:relative!important;z-index:5!important}.main > .hero:first-child,.main > .topbar:first-child{margin-top:0!important}
/* Dashboard: tarjetas más limpias */
.metric,.stat{background:#151a22!important;color:#fff!important;border:0!important}.metric b,.stat b,.metric strong,.stat strong{color:#fff!important;font-weight:720!important}.metric small,.stat small{color:#d8e1ea!important;font-weight:500!important}
/* Login tipo imagen 3: claro, sobrio y con entrada administrador */
.login-body{background:radial-gradient(circle at 8% 0%,rgba(34,197,94,.18) 0 18%,transparent 18.3%),radial-gradient(circle at 92% 2%,rgba(20,184,166,.15) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-card{width:min(92vw,520px)!important;padding:78px 48px 30px!important;background:rgba(255,255,255,.97)!important;border-radius:28px!important;border:1px solid #dce8ef!important;box-shadow:0 22px 54px rgba(15,23,42,.13)!important}.login-title h1{font-size:38px!important;font-weight:780!important;color:#182233!important;letter-spacing:-.03em!important}.login-title b{font-size:16px!important;font-weight:500!important;color:#667085!important}.login-card .field label{font-size:14px!important;font-weight:700!important;color:#1f2937!important}.login-input{background:#edf5ff!important;border:1px solid #cfe0f2!important;border-radius:17px!important;padding:0 14px!important;box-shadow:none!important}.login-input input,.login-input select{font-weight:520!important;color:#172033!important}.login-card .btn-green{font-size:18px!important;font-weight:750!important}.login-links a{color:#667085!important;font-weight:650!important;text-decoration:none!important}.login-links a:hover{color:#0f9b59!important}
@media(max-width:1000px){.main{padding:14px!important}.hero h1,.topbar h1{font-size:24px!important}.side{top:0!important;height:100dvh!important}.menu-title,.menu-item{font-size:14px!important}.login-title h1{font-size:31px!important}.login-card{padding:74px 26px 28px!important}}


/* =========================================================
   AJUSTE PRO SIDEBAR: encabezado limpio, iconos suaves,
   pestañas principales/secundarias claras y visibles
   ========================================================= */
body{font-weight:520!important;letter-spacing:-.01em!important}.side{background:linear-gradient(180deg,#0f3f2b 0%,#0b2a27 45%,#071423 100%)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:10px 0 28px rgba(15,23,42,.22)!important;width:320px!important}.app{grid-template-columns:320px 1fr!important}.app.side-collapsed{grid-template-columns:88px 1fr!important}.side.collapsed{width:88px!important}.side-head-pro{height:92px;margin:10px 10px 8px;padding:12px 10px;display:grid;grid-template-columns:42px 1fr 42px;align-items:center;gap:9px;background:linear-gradient(135deg,rgba(255,255,255,.06),rgba(255,255,255,.02));border:1px solid rgba(255,255,255,.08);border-radius:18px;box-shadow:0 12px 26px rgba(0,0,0,.13);position:sticky;top:10px;z-index:8}.side-brand-pro{display:flex;align-items:center;justify-content:center;gap:9px;min-width:0}.side-brand-icon{width:28px;height:28px;border-radius:8px;display:grid;place-items:center;background:rgba(255,255,255,.14);color:#a7f3d0;font-size:17px}.side-brand-text{display:flex;flex-direction:column;line-height:1.05;min-width:0}.side-brand-text b{font-size:15px;font-weight:780;color:#fff;letter-spacing:.01em;white-space:nowrap}.side-brand-text small{font-size:13px;font-weight:650;color:#cbd5e1;margin-top:4px}.toggle{width:42px;height:42px;border-radius:13px;background:rgba(255,255,255,.10)!important;color:#eaf8f1!important;font-size:22px;display:grid;place-items:center;transition:.16s}.toggle:hover{background:rgba(255,255,255,.17)!important;transform:translateY(-1px)}.side.collapsed .side-head-pro{grid-template-columns:1fr;height:76px;padding:10px}.side.collapsed .side-brand-pro,.side.collapsed .side-toggle-right{display:none}.side.collapsed .side-toggle-left{margin:auto}.brand{display:none!important}nav{padding:2px 12px 18px!important}.menu-group{margin:8px 0 10px!important;border-radius:18px!important;overflow:visible!important}.menu-title,.menu-item{font-family:Inter,Segoe UI,Arial,sans-serif!important;font-weight:680!important;text-shadow:none!important;letter-spacing:0!important}.menu-title{min-height:56px;padding:13px 16px!important;border-radius:18px!important;background:rgba(255,255,255,.035)!important;border:1px solid rgba(255,255,255,.08)!important;color:#e7eef7!important;box-shadow:none!important}.menu-title i,.menu-item i{width:28px;min-width:28px;height:28px;border-radius:9px;display:grid;place-items:center;font-size:18px;color:#d5dee9;background:rgba(255,255,255,.045)}.menu-title .label{line-height:1.25}.menu-title .chev{font-size:15px;color:#d3dde8!important;margin-left:auto}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.075)!important;color:#fff!important}.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#18b66b,#10a160)!important;color:#052116!important;border:1px solid rgba(134,239,172,.45)!important;box-shadow:0 12px 24px rgba(16,185,129,.22)!important}.menu-title.active i,.menu-group.force-open>.menu-title.active i{background:rgba(255,255,255,.22)!important;color:#06331e!important}.menu-title.active .chev{color:#052116!important}.menu-group.nested{margin-left:12px!important;padding-left:12px!important;border-left:1px solid rgba(203,213,225,.16)!important}.menu-group.nested>.menu-title{min-height:50px;border-radius:15px!important;background:rgba(255,255,255,.025)!important}.menu-group.nested.force-open>.menu-title{background:rgba(16,185,129,.11)!important;color:#f8fafc!important;border-left:3px solid #34d399!important}.submenu{padding:8px 0 7px!important;overflow:hidden!important}.menu-item{min-height:46px;margin:4px 0 4px 14px!important;padding:11px 13px!important;border-radius:14px!important;border-left:0!important;color:#dce5ee!important;background:transparent!important}.menu-item i{font-size:16px;background:transparent!important;color:#cbd5e1!important}.menu-item.active,.menu-item.parent-active{background:rgba(16,185,129,.16)!important;color:#ffffff!important;box-shadow:inset 3px 0 0 #34d399!important}.menu-item.active i,.menu-item.parent-active i{color:#86efac!important}.menu-item.sub-mini{font-size:13px!important;min-height:40px;margin-left:32px!important;color:#cbd5e1!important}.side-user{margin:22px 10px 18px!important;padding:16px!important;border:1px solid rgba(255,255,255,.10)!important;border-radius:18px!important;background:rgba(255,255,255,.075)!important;box-shadow:none!important}.avatar{background:#21c777!important;color:#072318!important}.side-user b{font-size:16px;font-weight:760!important;color:#fff!important}.side-user small{font-size:13px;color:#bbf7d0!important;font-weight:620!important}.side.collapsed nav{padding:4px 8px!important}.side.collapsed .menu-title{justify-content:center!important;padding:13px 8px!important}.side.collapsed .menu-title i{margin:0!important}.side.collapsed .menu-group.nested{margin-left:0!important;padding-left:0!important;border-left:0!important}.side.collapsed .label,.side.collapsed .chev,.side.collapsed .side-user{display:none!important}.side.collapsed .submenu{display:none!important}.main{background:linear-gradient(135deg,#f8fbfd 0%,#eef6f3 100%)!important;color:#142033!important}.hero,.card,.c-card,.filter-card,.table-wrap{color:#142033!important}h1,h2,h3,.c-title,.topbar h1{font-weight:780!important;color:#122033!important;text-shadow:none!important}.subtitle,.muted{font-weight:460!important;color:#667085!important}.metric b,.stat b{font-weight:680!important}.metric small,.stat small{font-weight:500!important}
@media(max-width:1000px){.app,.app.side-collapsed{grid-template-columns:1fr!important}.side{width:315px!important}.side-head-pro{top:8px;margin:8px}.side.open{left:0!important}}



/* =========================================================
   CORRECCIÓN FINAL: login tipo referencia, letras visibles,
   sidebar sin amarillo, iconos suaves y jerarquía clara
   ========================================================= */
:root{--accent:#16a34a!important;--accent2:#059669!important;--ink:#142033!important;--muted:#667085!important;--sidebar-a:#0f3e2b!important;--sidebar-b:#0a2524!important;--sidebar-c:#071421!important}
*{text-shadow:none!important}
body{font-family:"Inter","Segoe UI",Arial,sans-serif!important;font-weight:520!important;color:var(--ink)!important;background:#f3f7f8!important}
/* Login */
.login-body{min-height:100vh!important;display:grid!important;place-items:center!important;padding:28px!important;overflow:hidden!important;background:radial-gradient(circle at 8% -2%,rgba(34,197,94,.18) 0 20%,transparent 20.3%),radial-gradient(circle at 94% 0%,rgba(45,212,191,.16) 0 22%,transparent 22.3%),linear-gradient(180deg,#fbfdff 0%,#effaf5 100%)!important}
.login-body:before{content:""!important;position:absolute!important;left:0!important;right:0!important;bottom:0!important;height:34%!important;background:linear-gradient(135deg,#22c55e 0%,#047857 100%)!important;clip-path:polygon(0 42%,22% 29%,48% 22%,73% 32%,100% 25%,100% 100%,0 100%)!important;opacity:.96!important;z-index:0!important}
.login-body:after{content:"Sistema seguro y confiable\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;white-space:pre!important;position:absolute!important;bottom:20px!important;left:0!important;right:0!important;text-align:center!important;color:#ffffff!important;font-weight:650!important;line-height:1.8!important;z-index:1!important}
.login-card{width:min(92vw,520px)!important;background:rgba(255,255,255,.97)!important;border:1px solid #dce8ef!important;border-radius:30px!important;padding:84px 48px 34px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important;overflow:visible!important;position:relative!important;z-index:2!important;text-align:left!important}
.login-card:before{content:"👥"!important;position:absolute!important;top:-68px!important;left:50%!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dce8ef!important;display:grid!important;place-items:center!important;color:#14935d!important;font-size:48px!important;box-shadow:0 14px 35px rgba(15,23,42,.10)!important;opacity:1!important}
.login-card:after{display:none!important}.login-logo{display:none!important}.login-title{text-align:center!important;margin:0 0 30px!important}.login-title h1{font-size:38px!important;line-height:1.05!important;color:#182233!important;font-weight:780!important;letter-spacing:-.035em!important;margin:0 0 12px!important;text-transform:uppercase!important}.login-title b{font-size:16px!important;color:#667085!important;font-weight:500!important}.login-title:after{content:"";display:block;width:54px;height:4px;border-radius:999px;background:#16c784;margin:22px auto 0}.login-card .field{margin-bottom:18px!important}.login-card .field label{display:block!important;font-size:14px!important;font-weight:720!important;color:#1f2937!important;margin:0 0 8px!important}.login-input{height:58px!important;background:#eaf3ff!important;border:1px solid #cfe0f2!important;border-radius:17px!important;display:flex!important;align-items:center!important;gap:12px!important;padding:0 14px!important;box-shadow:none!important;margin:0!important;transition:.16s!important}.login-input:focus-within{border-color:#18b96d!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}.login-input i{width:30px!important;height:30px!important;border-radius:10px!important;display:grid!important;place-items:center!important;background:rgba(255,255,255,.55)!important;color:#15965a!important;font-size:16px!important;flex:0 0 auto!important}.login-input input,.login-input select{height:100%!important;border:0!important;background:transparent!important;color:#172033!important;font-size:16px!important;font-weight:560!important;outline:none!important;width:100%!important;padding:0 4px!important;box-shadow:none!important}.login-input input::placeholder{color:#7b8794!important;opacity:1!important}.login-input input:focus,.login-input select:focus{background:transparent!important;color:#111827!important;box-shadow:none!important}.login-card .btn-green{width:100%!important;height:62px!important;margin:14px 0 20px!important;border-radius:17px!important;background:linear-gradient(135deg,#1fac67,#0f8f55)!important;color:#fff!important;font-size:18px!important;font-weight:760!important;display:flex!important;justify-content:center!important;align-items:center!important}.login-links{margin:0!important;padding:0!important;text-align:center!important}.login-links a{color:#667085!important;font-weight:650!important;font-size:14px!important}.login-links a:hover{color:#0f9b59!important}
/* Sidebar */
.app{grid-template-columns:320px minmax(0,1fr)!important;background:#f3f7f8!important}.app.side-collapsed{grid-template-columns:88px minmax(0,1fr)!important}.side{width:320px!important;background:linear-gradient(180deg,var(--sidebar-a) 0%,var(--sidebar-b) 48%,var(--sidebar-c) 100%)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:10px 0 28px rgba(15,23,42,.20)!important;overflow-y:auto!important;overflow-x:hidden!important;scrollbar-width:thin!important;scrollbar-color:rgba(148,163,184,.55) transparent!important;padding-bottom:36px!important}.side::-webkit-scrollbar{width:8px!important}.side::-webkit-scrollbar-track{background:transparent!important}.side::-webkit-scrollbar-thumb{background:rgba(148,163,184,.45)!important;border-radius:999px!important}.side::-webkit-scrollbar-thumb:hover{background:rgba(203,213,225,.65)!important}.side-head-pro{height:88px!important;margin:10px 10px 8px!important;padding:12px 10px!important;display:grid!important;grid-template-columns:42px 1fr 42px!important;align-items:center!important;gap:9px!important;background:rgba(255,255,255,.055)!important;border:1px solid rgba(255,255,255,.09)!important;border-radius:18px!important;box-shadow:none!important;position:sticky!important;top:10px!important;z-index:8!important}.side-brand-pro{display:flex!important;align-items:center!important;justify-content:center!important;gap:9px!important;min-width:0!important}.side-brand-icon{width:30px!important;height:30px!important;border-radius:9px!important;display:grid!important;place-items:center!important;background:rgba(255,255,255,.13)!important;color:#9ee9c4!important;font-size:16px!important}.side-brand-text b{font-size:15px!important;font-weight:760!important;color:#fff!important;letter-spacing:.01em!important;white-space:nowrap!important}.side-brand-text small{font-size:13px!important;font-weight:600!important;color:#cbd5e1!important;margin-top:4px!important}.toggle{width:42px!important;height:42px!important;border-radius:13px!important;background:rgba(255,255,255,.10)!important;color:#eaf8f1!important;border:0!important;font-size:22px!important;display:grid!important;place-items:center!important}.toggle:hover{background:rgba(255,255,255,.17)!important}nav{padding:4px 12px 20px!important}.menu-group{margin:8px 0 10px!important;border-radius:18px!important;overflow:visible!important}.submenu{padding:7px 0 8px 0!important;max-height:none!important;overflow:visible!important}.menu-group.closed>.submenu{max-height:0!important;padding:0!important;overflow:hidden!important}.menu-title,.menu-item{font-family:"Inter","Segoe UI",Arial,sans-serif!important;font-weight:650!important;letter-spacing:0!important;text-shadow:none!important;white-space:normal!important;word-break:normal!important;line-height:1.26!important;color:#dbe7ef!important}.menu-title{min-height:54px!important;padding:12px 15px!important;border-radius:17px!important;background:rgba(255,255,255,.04)!important;border:1px solid rgba(255,255,255,.08)!important;box-shadow:none!important;display:flex!important;align-items:center!important;gap:12px!important}.menu-title i,.menu-item i{width:28px!important;min-width:28px!important;height:28px!important;border-radius:9px!important;display:grid!important;place-items:center!important;font-size:16px!important;color:#cbd5e1!important;background:rgba(255,255,255,.05)!important;filter:none!important}.menu-title .label,.menu-item .label{display:block!important;color:inherit!important;font-weight:inherit!important;line-height:1.26!important;overflow:visible!important;text-overflow:clip!important;white-space:normal!important}.chev{margin-left:auto!important;color:#cbd5e1!important;font-size:14px!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.085)!important;color:#fff!important}.menu-group.force-open>.menu-title,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17b96b,#0f9b59)!important;color:#052116!important;border:1px solid rgba(134,239,172,.45)!important;box-shadow:0 12px 22px rgba(16,185,129,.20)!important}.menu-group.force-open>.menu-title i,.menu-title.active i{background:rgba(255,255,255,.25)!important;color:#06331e!important}.menu-group.force-open>.menu-title .chev,.menu-title.active .chev{color:#052116!important}.menu-group.nested{margin-left:12px!important;padding-left:12px!important;border-left:1px solid rgba(203,213,225,.15)!important}.menu-group.nested>.menu-title{min-height:48px!important;border-radius:15px!important;background:rgba(255,255,255,.025)!important;color:#e6eef7!important}.menu-group.nested.force-open>.menu-title{background:rgba(16,185,129,.13)!important;color:#f8fafc!important;border-left:3px solid #34d399!important;box-shadow:none!important}.menu-item{min-height:44px!important;margin:4px 0 4px 14px!important;padding:10px 12px!important;border-radius:14px!important;border-left:0!important;background:transparent!important;display:flex!important;align-items:center!important;gap:12px!important}.menu-item i{background:transparent!important;color:#cbd5e1!important}.menu-item.active,.menu-item.parent-active{background:rgba(16,185,129,.17)!important;color:#ffffff!important;box-shadow:inset 3px 0 0 #34d399!important}.menu-item.active i,.menu-item.parent-active i{color:#86efac!important}.menu-item.active .label,.menu-item.parent-active .label{color:#ffffff!important}.menu-item.sub-mini{min-height:40px!important;font-size:13.2px!important;margin-left:32px!important;padding-left:12px!important;color:#cbd5e1!important}.side-user{margin:22px 10px 18px!important;padding:16px!important;border-radius:18px!important;background:rgba(255,255,255,.075)!important;border:1px solid rgba(255,255,255,.10)!important;box-shadow:none!important}.avatar{width:42px!important;height:42px!important;background:#21c777!important;color:#072318!important}.side-user b{font-size:16px!important;font-weight:740!important;color:#fff!important}.side-user small{font-size:13px!important;color:#bbf7d0!important;font-weight:600!important}.main{background:linear-gradient(135deg,#f8fbfd 0%,#eef6f3 100%)!important;color:#142033!important}.main h1,.main h2,.main h3,.topbar h1,.card h2{font-weight:760!important;color:#122033!important}.subtitle,.muted{font-weight:450!important;color:#667085!important}
/* collapsed */
.side.collapsed{width:88px!important;overflow:hidden!important}.side.collapsed .side-head-pro{grid-template-columns:1fr!important;height:76px!important;padding:10px!important}.side.collapsed .side-brand-pro,.side.collapsed .side-toggle-right,.side.collapsed .label,.side.collapsed .chev,.side.collapsed .side-user{display:none!important}.side.collapsed nav{padding:4px 8px!important}.side.collapsed .menu-title,.side.collapsed .menu-item{justify-content:center!important;padding:12px 8px!important;margin:6px 0!important}.side.collapsed .submenu{display:none!important}.side.collapsed .menu-group.nested{margin-left:0!important;padding-left:0!important;border-left:0!important}
@media(max-width:1000px){.app,.app.side-collapsed{grid-template-columns:1fr!important}.side{width:315px!important;left:-340px!important}.side.open{left:0!important}.main{padding:14px!important}.login-card{width:min(94vw,500px)!important;padding:78px 26px 30px!important}.login-title h1{font-size:31px!important}.login-body:after{font-size:12px!important}}



/* =========================================================
   AJUSTE FINAL SOLICITADO: login centrado, icono nuevo,
   letras visibles, footer, campanita única y sidebar sin superposición
   ========================================================= */
.login-body{
  min-height:100svh!important;
  min-height:100dvh!important;
  display:flex!important;
  align-items:center!important;
  justify-content:center!important;
  padding:92px 18px 118px!important;
  overflow-x:hidden!important;
  overflow-y:auto!important;
}
.login-card{
  width:min(92vw,540px)!important;
  margin:0 auto!important;
  padding:90px 48px 34px!important;
  border-radius:30px!important;
}
.login-logo{position:absolute!important;top:-72px!important;left:50%!important;transform:translateX(-50%)!important;margin:0!important;z-index:4!important}
.login-logo img{display:none!important}
.login-avatar-svg{
  width:142px!important;height:142px!important;border-radius:50%!important;background:#fff!important;
  border:1px solid #dbe7ef!important;box-shadow:0 16px 42px rgba(15,23,42,.10)!important;
  display:grid!important;place-items:center!important;position:relative!important;
}
.login-avatar-svg:before{
  content:""!important;width:78px!important;height:78px!important;display:block!important;
  background:no-repeat center/contain url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 96 96'%3E%3Cdefs%3E%3ClinearGradient id='g' x1='0' y1='0' x2='1' y2='1'%3E%3Cstop stop-color='%2327c77b'/%3E%3Cstop offset='1' stop-color='%230b8f5a'/%3E%3C/linearGradient%3E%3C/defs%3E%3Ccircle cx='35' cy='30' r='13' fill='url(%23g)'/%3E%3Ccircle cx='61' cy='28' r='15' fill='url(%23g)'/%3E%3Cpath d='M13 72c1-17 13-28 29-28 5 0 9 1 13 3-7 5-12 13-13 25H13z' fill='url(%23g)'/%3E%3Cpath d='M38 72c1-19 14-31 32-31 15 0 25 12 25 31H38z' fill='url(%23g)'/%3E%3C/svg%3E")!important;
}
.login-card:before{display:none!important;content:none!important}.login-card:after{display:none!important;content:none!important}
.login-title{text-align:center!important;margin:8px 0 26px!important}.login-title h1{color:#182233!important;font-size:38px!important;font-weight:800!important;letter-spacing:-.035em!important}.login-title b{color:#667085!important;font-size:17px!important;font-weight:520!important;line-height:1.35!important}.login-title:after{content:"";display:block;width:66px;height:4px;border-radius:999px;background:#16b978;margin:22px auto 0}.login-card .field{margin-bottom:16px!important}.login-card .field label{display:block!important;color:#1f2937!important;font-weight:720!important;font-size:15px!important;margin:0 0 8px!important}.login-input{height:64px!important;background:#eaf3ff!important;border:1px solid #cfe0f2!important;border-radius:18px!important;padding:0 12px!important;gap:12px!important;display:flex!important;align-items:center!important;box-shadow:none!important}.login-input i{width:34px!important;height:34px!important;border-radius:11px!important;background:rgba(255,255,255,.72)!important;color:#11945d!important;font-size:17px!important;display:grid!important;place-items:center!important;flex:0 0 auto!important}.login-input input,.login-input select{height:48px!important;min-height:48px!important;background:#fff!important;border:1px solid #d5e3ef!important;border-radius:16px!important;color:#172033!important;-webkit-text-fill-color:#172033!important;font-size:17px!important;font-weight:520!important;padding:0 14px!important;outline:none!important;box-shadow:none!important}.login-input select{appearance:auto!important}.login-input input::placeholder{color:#6b7787!important;opacity:1!important;-webkit-text-fill-color:#6b7787!important}.login-input input:focus,.login-input select:focus{background:#fff!important;color:#111827!important;-webkit-text-fill-color:#111827!important;border-color:#18b96d!important;box-shadow:0 0 0 3px rgba(22,185,120,.12)!important}.login-card .btn-green{height:64px!important;border-radius:18px!important;font-size:20px!important;font-weight:760!important;margin:18px 0 22px!important}.login-links a{font-size:15px!important;color:#667085!important;font-weight:650!important}.login-body:after{content:"🛡️  Sistema seguro y confiable\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;white-space:pre!important;bottom:20px!important;color:#ffffff!important;font-size:15px!important;line-height:1.9!important;font-weight:650!important;z-index:1!important}
/* Sidebar: evita que el encabezado se monte sobre las pestañas al hacer scroll */
.side-head-pro{position:relative!important;top:auto!important;margin:12px 10px 10px!important;height:auto!important;min-height:78px!important;z-index:2!important;background:rgba(255,255,255,.055)!important}.side{scrollbar-color:rgba(148,163,184,.55) transparent!important}.side::-webkit-scrollbar-thumb{background:rgba(148,163,184,.50)!important}.side::-webkit-scrollbar-track{background:transparent!important}.side-brand-text b{font-weight:760!important}.side-brand-text small{font-weight:570!important}.menu-title,.menu-item{font-weight:610!important;color:#e6edf5!important}.menu-title .label,.menu-item .label{color:inherit!important;line-height:1.28!important}.menu-title i,.menu-item i{color:#d3dde8!important;background:rgba(255,255,255,.06)!important}.menu-group.force-open>.menu-title,.menu-title.active{color:#052116!important}.menu-group.force-open>.menu-title .label,.menu-title.active .label{color:#052116!important}.menu-item.active,.menu-item.parent-active{background:rgba(22,185,120,.18)!important;color:#ffffff!important}.menu-item.active .label,.menu-item.parent-active .label{color:#ffffff!important}.menu-group.nested.force-open>.menu-title{color:#f8fafc!important;background:rgba(22,185,120,.13)!important}.menu-group.nested.force-open>.menu-title .label{color:#f8fafc!important}.menu-item.sub-mini{color:#d8e2ec!important}.menu-item.sub-mini.active{color:#fff!important;background:rgba(22,185,120,.20)!important}.side.collapsed .side-head-pro{min-height:64px!important}
/* Topbar: solo campanita */
.top-actions{display:flex!important;align-items:center!important;gap:12px!important}.top-actions .top-icon:not(.notif-bell){display:none!important}.top-icon.notif-bell{width:44px!important;height:44px!important;border-radius:12px!important;border:0!important;background:#101923!important;color:#fff!important;display:grid!important;place-items:center!important;position:relative!important;cursor:pointer!important;box-shadow:0 8px 18px rgba(15,23,42,.10)!important}.top-icon.notif-bell span{font-size:21px!important;line-height:1!important}.top-icon.notif-bell i{position:absolute!important;right:-7px!important;top:-8px!important;background:#ff4d5c!important;color:white!important;border-radius:999px!important;font-size:11px!important;min-width:22px!important;height:22px!important;display:grid!important;place-items:center!important;font-style:normal!important;font-weight:800!important}
@media(max-width:1000px){
  .login-body{padding:82px 12px 105px!important;align-items:center!important;justify-content:center!important}.login-card{width:calc(100vw - 24px)!important;max-width:430px!important;padding:84px 26px 30px!important;border-radius:26px!important}.login-avatar-svg{width:132px!important;height:132px!important}.login-logo{top:-68px!important}.login-title h1{font-size:32px!important;line-height:1.05!important}.login-title b{font-size:18px!important;display:block!important;max-width:300px!important;margin:auto!important}.login-input{height:64px!important}.login-input input,.login-input select{font-size:18px!important}.login-card .btn-green{height:64px!important;font-size:21px!important}.login-body:after{font-size:12px!important;bottom:14px!important}.side{padding-top:0!important}.side-head-pro{margin-top:10px!important;position:relative!important;top:auto!important}.main{padding-top:14px!important}
}
@media(max-height:760px){.login-body{align-items:flex-start!important;padding-top:86px!important}.login-card{margin-top:0!important}}


/* =========================================================
   AJUSTE FINAL SOLICITADO: login igual referencia + texto visible
   + encabezado del panel fijo al deslizar
   ========================================================= */
.login-body{
  min-height:100vh!important;min-height:100dvh!important;
  display:flex!important;align-items:center!important;justify-content:center!important;
  padding:92px 18px 118px!important;position:relative!important;overflow:hidden!important;
  background:
    radial-gradient(circle at 7% -4%,rgba(34,197,94,.20) 0 21%,transparent 21.4%),
    radial-gradient(circle at 95% 0%,rgba(20,184,166,.16) 0 22%,transparent 22.4%),
    linear-gradient(180deg,#fbfdff 0%,#effaf5 100%)!important;
}
.login-body:before{
  content:""!important;display:block!important;position:absolute!important;left:0!important;right:0!important;bottom:0!important;height:35%!important;
  background:linear-gradient(135deg,#22c55e 0%,#047857 100%)!important;
  clip-path:polygon(0 42%,22% 29%,48% 22%,73% 32%,100% 25%,100% 100%,0 100%)!important;
  opacity:.96!important;z-index:0!important;
}
.login-body:after{
  content:"🛡️  Sistema seguro y confiable\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;
  white-space:pre!important;position:absolute!important;left:0!important;right:0!important;bottom:22px!important;text-align:center!important;
  color:#fff!important;font-size:15px!important;line-height:1.85!important;font-weight:650!important;z-index:1!important;
}
.login-card{
  width:min(92vw,540px)!important;margin:0 auto!important;padding:92px 48px 36px!important;
  background:rgba(255,255,255,.97)!important;border:1px solid #dce8ef!important;border-radius:30px!important;
  box-shadow:0 22px 55px rgba(15,23,42,.13)!important;overflow:visible!important;position:relative!important;z-index:2!important;
}
.login-logo{position:absolute!important;top:-72px!important;left:50%!important;transform:translateX(-50%)!important;margin:0!important;z-index:4!important;display:block!important}
.login-logo img{display:none!important}.login-card:before,.login-card:after{display:none!important;content:none!important}
.login-avatar-svg{width:142px!important;height:142px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe7ef!important;box-shadow:0 16px 42px rgba(15,23,42,.10)!important;display:grid!important;place-items:center!important;position:relative!important}
.login-avatar-svg:before{content:""!important;width:78px!important;height:78px!important;display:block!important;background:no-repeat center/contain url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 96 96'%3E%3Cdefs%3E%3ClinearGradient id='g' x1='0' y1='0' x2='1' y2='1'%3E%3Cstop stop-color='%2327c77b'/%3E%3Cstop offset='1' stop-color='%230b8f5a'/%3E%3C/linearGradient%3E%3C/defs%3E%3Ccircle cx='35' cy='30' r='13' fill='url(%23g)'/%3E%3Ccircle cx='61' cy='28' r='15' fill='url(%23g)'/%3E%3Cpath d='M13 72c1-17 13-28 29-28 5 0 9 1 13 3-7 5-12 13-13 25H13z' fill='url(%23g)'/%3E%3Cpath d='M38 72c1-19 14-31 32-31 15 0 25 12 25 31H38z' fill='url(%23g)'/%3E%3C/svg%3E")!important}
.login-title{text-align:center!important;margin:8px 0 26px!important}.login-title h1{color:#182233!important;font-size:38px!important;line-height:1.05!important;font-weight:800!important;letter-spacing:-.035em!important;margin:0 0 12px!important;text-transform:uppercase!important}.login-title b{display:block!important;color:#667085!important;font-size:17px!important;font-weight:520!important;line-height:1.35!important}.login-title:after{content:""!important;display:block!important;width:66px!important;height:4px!important;border-radius:999px!important;background:#16b978!important;margin:22px auto 0!important}
.login-card .field{margin-bottom:16px!important}.login-card .field label{display:block!important;color:#1f2937!important;font-weight:720!important;font-size:15px!important;margin:0 0 8px!important}
.login-input{height:64px!important;background:#eaf3ff!important;border:1px solid #cfe0f2!important;border-radius:18px!important;padding:0 12px!important;gap:12px!important;display:flex!important;align-items:center!important;box-shadow:none!important;margin:0!important;color:#172033!important}
.login-input i{width:34px!important;height:34px!important;border-radius:11px!important;background:rgba(255,255,255,.72)!important;color:#11945d!important;font-size:17px!important;display:grid!important;place-items:center!important;flex:0 0 auto!important}
.login-input input,.login-input select{
  height:48px!important;min-height:48px!important;background:#fff!important;border:1px solid #d5e3ef!important;border-radius:16px!important;
  color:#172033!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;
  font-size:17px!important;font-weight:520!important;padding:0 14px!important;outline:none!important;box-shadow:none!important;width:100%!important;
}
.login-input select{appearance:auto!important;-webkit-text-fill-color:#172033!important;color:#172033!important}.login-input select option{background:#fff!important;color:#172033!important;-webkit-text-fill-color:#172033!important}
.login-input input::placeholder{color:#6b7787!important;opacity:1!important;-webkit-text-fill-color:#6b7787!important}.login-input input:focus,.login-input select:focus{background:#fff!important;color:#111827!important;-webkit-text-fill-color:#111827!important;border-color:#18b96d!important;box-shadow:0 0 0 3px rgba(22,185,120,.12)!important}
.login-input input:-webkit-autofill{-webkit-box-shadow:0 0 0 1000px #fff inset!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important}
.login-card .btn-green{width:100%!important;height:64px!important;border-radius:18px!important;background:linear-gradient(135deg,#1fac67,#0f8f55)!important;color:#fff!important;font-size:20px!important;font-weight:760!important;display:flex!important;align-items:center!important;justify-content:center!important;margin:18px 0 22px!important;border:0!important;box-shadow:0 15px 30px rgba(22,163,74,.26)!important}
.login-links{margin:0!important;padding:0!important;text-align:center!important}.login-links a{font-size:15px!important;color:#667085!important;font-weight:650!important}.login-links a:hover{color:#0f9b59!important}
/* Encabezado fijo del panel: ya no se superpone al scroll */
.side{overflow-y:auto!important;scrollbar-color:rgba(148,163,184,.55) transparent!important}.side::-webkit-scrollbar-thumb{background:rgba(148,163,184,.50)!important;border-radius:999px!important}.side::-webkit-scrollbar-track{background:transparent!important}
.side-head-pro{position:sticky!important;top:0!important;z-index:999!important;margin:0 10px 10px!important;min-height:78px!important;height:auto!important;background:linear-gradient(135deg,rgba(255,255,255,.08),rgba(255,255,255,.035))!important;border:1px solid rgba(255,255,255,.10)!important;border-radius:0 0 18px 18px!important;box-shadow:0 14px 28px rgba(0,0,0,.22)!important;backdrop-filter:blur(10px)!important}
.side-brand-text b{font-weight:760!important;color:#fff!important}.side-brand-text small{font-weight:570!important;color:#cbd5e1!important}.side-brand-icon{color:#d9f7e8!important;background:rgba(255,255,255,.12)!important}.toggle{color:#eaf8f1!important;background:rgba(255,255,255,.10)!important}
nav{padding-top:4px!important}.menu-title,.menu-item{font-weight:610!important;color:#e6edf5!important;text-shadow:none!important}.menu-title .label,.menu-item .label{color:inherit!important;line-height:1.28!important}.menu-title i,.menu-item i{color:#d3dde8!important;background:rgba(255,255,255,.06)!important}.menu-group.force-open>.menu-title,.menu-title.active{color:#052116!important}.menu-group.force-open>.menu-title .label,.menu-title.active .label{color:#052116!important}.menu-item.active,.menu-item.parent-active{background:rgba(22,185,120,.18)!important;color:#ffffff!important}.menu-item.active .label,.menu-item.parent-active .label{color:#ffffff!important}.menu-group.nested.force-open>.menu-title{color:#f8fafc!important;background:rgba(22,185,120,.13)!important}.menu-group.nested.force-open>.menu-title .label{color:#f8fafc!important}.menu-item.sub-mini{color:#d8e2ec!important}.menu-item.sub-mini.active{color:#fff!important;background:rgba(22,185,120,.20)!important}
@media(max-width:1000px){
  .login-body{padding:88px 12px 108px!important;align-items:center!important;justify-content:center!important;overflow-y:auto!important}
  .login-card{width:calc(100vw - 24px)!important;max-width:430px!important;padding:86px 26px 30px!important;border-radius:26px!important}.login-avatar-svg{width:132px!important;height:132px!important}.login-logo{top:-68px!important}.login-title h1{font-size:32px!important;line-height:1.05!important}.login-title b{font-size:18px!important;max-width:310px!important;margin:auto!important}.login-input{height:64px!important}.login-input input,.login-input select{font-size:18px!important}.login-card .btn-green{height:64px!important;font-size:21px!important}.login-body:after{font-size:12px!important;bottom:14px!important}.side-head-pro{position:sticky!important;top:0!important;margin:0 10px 10px!important}.main{padding-top:14px!important}
}
@media(max-height:760px){.login-body{align-items:flex-start!important;padding-top:86px!important}.login-card{margin-top:0!important}}



/* ===== AJUSTE FINAL LOGIN REFERENCIA IMAGEN 02 ===== */
.login-body{
  min-height:100vh!important;display:flex!important;align-items:center!important;justify-content:center!important;
  padding:36px 18px 118px!important;overflow:hidden!important;
  background:radial-gradient(circle at 8% -4%,rgba(34,197,94,.17) 0 20%,transparent 20.4%),radial-gradient(circle at 94% -2%,rgba(45,212,191,.17) 0 22%,transparent 22.4%),linear-gradient(180deg,#fbfdff 0%,#effaf5 100%)!important;
}
.login-body:before{content:""!important;position:absolute!important;left:0!important;right:0!important;bottom:0!important;height:30%!important;background:linear-gradient(135deg,#22c55e 0%,#065f46 100%)!important;clip-path:polygon(0 36%,22% 26%,48% 17%,73% 28%,100% 22%,100% 100%,0 100%)!important;z-index:0!important;opacity:.98!important}
.login-body:after{content:"🛡️  Sistema seguro y confiable\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;white-space:pre!important;position:absolute!important;left:0!important;right:0!important;bottom:18px!important;text-align:center!important;color:#fff!important;font-size:16px!important;line-height:1.9!important;font-weight:650!important;z-index:1!important;background:none!important;height:auto!important;clip-path:none!important}
.login-card{width:min(90vw,500px)!important;max-width:500px!important;margin:0 auto!important;padding:84px 46px 30px!important;border-radius:28px!important;background:rgba(255,255,255,.97)!important;border:1px solid #dce8ef!important;box-shadow:0 18px 48px rgba(15,23,42,.13)!important;overflow:visible!important;position:relative!important;z-index:2!important;text-align:left!important}
.login-logo{display:block!important;position:absolute!important;top:-62px!important;left:50%!important;transform:translateX(-50%)!important;margin:0!important;z-index:5!important}
.login-logo img{display:none!important}.login-card:before,.login-card:after{display:none!important;content:none!important}
.login-avatar-svg{width:124px!important;height:124px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe7ef!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;display:grid!important;place-items:center!important}
.login-avatar-svg:before{content:""!important;width:68px!important;height:68px!important;display:block!important;background:no-repeat center/contain url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 96 96'%3E%3Cdefs%3E%3ClinearGradient id='g' x1='0' y1='0' x2='1' y2='1'%3E%3Cstop stop-color='%2327c77b'/%3E%3Cstop offset='1' stop-color='%230b8f5a'/%3E%3C/linearGradient%3E%3C/defs%3E%3Ccircle cx='35' cy='30' r='13' fill='url(%23g)'/%3E%3Ccircle cx='61' cy='28' r='15' fill='url(%23g)'/%3E%3Cpath d='M13 72c1-17 13-28 29-28 5 0 9 1 13 3-7 5-12 13-13 25H13z' fill='url(%23g)'/%3E%3Cpath d='M38 72c1-19 14-31 32-31 15 0 25 12 25 31H38z' fill='url(%23g)'/%3E%3C/svg%3E")!important}
.login-title{text-align:center!important;margin:0 0 24px!important}.login-title h1{font-size:36px!important;line-height:1.05!important;font-weight:800!important;color:#182233!important;letter-spacing:-.035em!important;margin:0 0 10px!important;text-transform:uppercase!important}.login-title b{display:block!important;color:#667085!important;font-size:16px!important;font-weight:520!important;line-height:1.32!important}.login-title:after{content:""!important;display:block!important;width:60px!important;height:4px!important;border-radius:999px!important;background:#16b978!important;margin:18px auto 0!important}
.login-card .field{margin-bottom:13px!important}.login-card .field label{display:block!important;color:#1f2937!important;font-size:14px!important;font-weight:720!important;margin:0 0 7px!important}
.login-input{height:54px!important;background:#eaf3ff!important;border:1px solid #cfe0f2!important;border-radius:17px!important;padding:0 10px!important;gap:10px!important;display:flex!important;align-items:center!important;margin:0!important;box-shadow:none!important;color:#172033!important}.login-input:focus-within{border-color:#18b96d!important;box-shadow:0 0 0 3px rgba(22,185,120,.13)!important}
.login-input i{width:31px!important;height:31px!important;border-radius:10px!important;background:rgba(255,255,255,.72)!important;color:#11945d!important;font-size:16px!important;display:grid!important;place-items:center!important;flex:0 0 auto!important}
.login-input input,.login-input select{height:42px!important;min-height:42px!important;background:#fff!important;border:1px solid #d5e3ef!important;border-radius:15px!important;color:#172033!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;font-size:16px!important;font-weight:520!important;padding:0 13px!important;outline:none!important;box-shadow:none!important;width:100%!important}.login-input select{appearance:auto!important;color:#172033!important;-webkit-text-fill-color:#172033!important}.login-input select option{color:#172033!important;background:#fff!important}.login-input input::placeholder{color:#6b7787!important;opacity:1!important;-webkit-text-fill-color:#6b7787!important}.login-input input:focus,.login-input select:focus{background:#fff!important;color:#111827!important;-webkit-text-fill-color:#111827!important;border-color:#18b96d!important;box-shadow:none!important}.login-input input:-webkit-autofill{-webkit-box-shadow:0 0 0 1000px #fff inset!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important}
.login-card .btn-green{width:100%!important;height:58px!important;border-radius:17px!important;background:linear-gradient(135deg,#1fac67,#0f8f55)!important;color:#fff!important;font-size:19px!important;font-weight:760!important;display:flex!important;align-items:center!important;justify-content:center!important;margin:16px 0 20px!important;border:0!important;box-shadow:0 14px 28px rgba(22,163,74,.24)!important}.login-links{margin:0!important;padding:0!important;text-align:center!important}.login-links a{font-size:14px!important;color:#667085!important;font-weight:650!important;text-decoration:none!important}
@media(max-width:700px){.login-body{padding:82px 12px 108px!important;align-items:center!important;overflow-y:auto!important}.login-card{width:calc(100vw - 26px)!important;max-width:420px!important;padding:82px 26px 30px!important;border-radius:26px!important}.login-logo{top:-62px!important}.login-avatar-svg{width:124px!important;height:124px!important}.login-title h1{font-size:31px!important}.login-title b{font-size:17px!important;max-width:310px!important;margin:auto!important}.login-input{height:58px!important}.login-input input,.login-input select{height:46px!important;min-height:46px!important;font-size:17px!important}.login-card .btn-green{height:60px!important;font-size:20px!important}.login-body:after{font-size:12px!important;bottom:14px!important}}
.side-head-pro{position:sticky!important;top:0!important;z-index:999!important}


/* ===== CORRECCION FINAL SOLICITADA: HEADER FIJO, CHIP ACTIVO Y LOGIN SIN TAPAR ===== */
.side-head-pro{
  position:sticky!important;top:0!important;z-index:9999!important;margin:10px 10px 14px!important;
  background:linear-gradient(135deg,#163f31 0%,#0f2e27 58%,#0a241f 100%)!important;
  border:1px solid rgba(255,255,255,.10)!important;border-radius:18px!important;box-shadow:0 12px 28px rgba(0,0,0,.25)!important;
  overflow:hidden!important;isolation:isolate!important;backdrop-filter:none!important;
}
.side-head-pro:before{content:""!important;position:absolute!important;inset:0!important;background:linear-gradient(90deg,rgba(255,255,255,.08),rgba(255,255,255,0))!important;z-index:-1!important}.side-head-pro *{position:relative!important;z-index:2!important}.side{scrollbar-color:#7aa99a #0a241f!important}.side::-webkit-scrollbar{width:8px!important}.side::-webkit-scrollbar-track{background:#0a241f!important}.side::-webkit-scrollbar-thumb{background:#6ea58f!important;border-radius:999px!important;border:2px solid #0a241f!important}
.admin-chip{position:relative!important}.admin-chip-btn{border:0!important;background:transparent!important;display:flex!important;align-items:center!important;gap:10px!important;font-weight:850!important;color:#111827!important;cursor:pointer!important;padding:3px 4px!important;border-radius:999px!important}.admin-chip-btn:hover{background:#eef7f1!important}.admin-menu{display:none!important;position:absolute!important;right:0!important;top:48px!important;width:180px!important;background:#ffffff!important;border:1px solid #d9e5ee!important;border-radius:16px!important;box-shadow:0 18px 38px rgba(15,23,42,.16)!important;padding:8px!important;z-index:99999!important}.admin-dropdown.open .admin-menu,.admin-dropdown:focus-within .admin-menu{display:block!important}.admin-menu a{display:flex!important;align-items:center!important;gap:10px!important;padding:11px 12px!important;border-radius:12px!important;color:#172033!important;font-weight:760!important}.admin-menu a:hover{background:#eef8f2!important;color:#0f8f55!important}.admin-menu i{color:#149459!important}.top-actions{overflow:visible!important}.admin-header{overflow:visible!important}.admin-chip .a{color:#092316!important;font-weight:850!important}
.login-card{padding-top:110px!important}.login-logo{top:-62px!important;z-index:10!important}.login-title{margin-top:0!important;position:relative!important;z-index:3!important}.login-title h1{position:relative!important;z-index:3!important}.login-avatar-svg{width:124px!important;height:124px!important}.login-avatar-svg:before{width:68px!important;height:68px!important}.login-input input,.login-input select{color:#172033!important;-webkit-text-fill-color:#172033!important}.login-input input:focus,.login-input select:focus{color:#111827!important;-webkit-text-fill-color:#111827!important;background:#fff!important}
@media(max-width:700px){.login-card{padding-top:98px!important}.login-logo{top:-60px!important}.login-title h1{font-size:31px!important}.admin-menu{right:0!important;top:46px!important}.side-head-pro{margin:8px 10px 12px!important}}



/* ===== CORRECCION URGENTE DEFINITIVA ===== */
/* 1) Encabezado lateral fijo, sólido y sin mostrar contenido por detrás */
.side{
  background:linear-gradient(180deg,#104d38 0%,#073126 45%,#061923 100%)!important;
  padding-top:12px!important;
  overflow-y:auto!important;
  scrollbar-color:#7fb8a2 #061923!important;
}
.side::-webkit-scrollbar{width:8px!important}.side::-webkit-scrollbar-track{background:#061923!important}.side::-webkit-scrollbar-thumb{background:#7fb8a2!important;border-radius:999px!important;border:2px solid #061923!important}
.side-head-pro{
  position:sticky!important;top:0!important;z-index:2147483640!important;
  margin:0 12px 16px!important;min-height:82px!important;height:82px!important;
  background:#123f33!important;border:1px solid rgba(255,255,255,.16)!important;border-radius:20px!important;
  box-shadow:0 16px 34px rgba(0,0,0,.32)!important;overflow:hidden!important;isolation:isolate!important;backdrop-filter:none!important;
  display:flex!important;align-items:center!important;justify-content:space-between!important;padding:0 12px!important;
}
.side-head-pro:before{content:""!important;position:absolute!important;inset:0!important;background:linear-gradient(90deg,#123f33 0%,#196b48 52%,#123f33 100%)!important;z-index:-1!important;opacity:1!important}.side-head-pro:after{content:""!important;position:absolute!important;left:0!important;right:0!important;bottom:-1px!important;height:14px!important;background:#073126!important;z-index:1!important;opacity:1!important}.side-head-pro *{position:relative!important;z-index:3!important}.side-brand-pro{display:flex!important;align-items:center!important;gap:10px!important;min-width:0!important}.side-brand-text b{display:block!important;color:#fff!important;font-size:16px!important;font-weight:800!important;line-height:1.05!important;white-space:nowrap!important}.side-brand-text small{display:block!important;color:#d6e9e2!important;font-size:14px!important;font-weight:650!important;line-height:1.15!important}.side-brand-icon{width:38px!important;height:38px!important;border-radius:12px!important;background:rgba(255,255,255,.14)!important;color:#dff8ed!important;display:grid!important;place-items:center!important;font-size:20px!important;flex:0 0 auto!important}.side-head-pro .toggle{width:48px!important;height:48px!important;border-radius:14px!important;background:rgba(255,255,255,.12)!important;color:#fff!important;border:0!important;display:grid!important;place-items:center!important;font-size:24px!important;flex:0 0 auto!important}
nav{position:relative!important;z-index:1!important;padding-top:0!important}.menu-title,.menu-item{letter-spacing:.01em!important;text-shadow:none!important}.menu-title .label,.menu-item .label{color:inherit!important;opacity:1!important;visibility:visible!important}
/* 2) Menú del usuario administrador: Inicio / Mi cuenta / Salir */
.admin-chip{position:relative!important;z-index:999999!important}.admin-chip-btn{border:0!important;background:transparent!important;display:flex!important;align-items:center!important;gap:10px!important;color:#111827!important;font-size:17px!important;font-weight:850!important;cursor:pointer!important;padding:6px 8px!important;border-radius:999px!important}.admin-chip-btn:hover{background:#eef8f2!important}.admin-chip .a{width:46px!important;height:46px!important;border-radius:50%!important;background:#18b76f!important;color:#092316!important;display:grid!important;place-items:center!important;font-size:18px!important;font-weight:900!important}.admin-menu{display:none!important;position:absolute!important;right:0!important;top:56px!important;min-width:190px!important;background:#fff!important;border:1px solid #dbe7ef!important;border-radius:16px!important;box-shadow:0 20px 45px rgba(15,23,42,.18)!important;padding:8px!important;z-index:999999!important}.admin-dropdown.open .admin-menu,.admin-dropdown:focus-within .admin-menu{display:block!important}.admin-menu a{display:flex!important;align-items:center!important;gap:10px!important;padding:12px 13px!important;border-radius:12px!important;color:#172033!important;font-weight:760!important;text-decoration:none!important}.admin-menu a:hover{background:#eef8f2!important;color:#0f8f55!important}.admin-menu i{color:#149459!important;font-size:17px!important}.top-actions,.admin-header{overflow:visible!important}
/* 3) Login: tamaño igual a referencia y logo sin tapar el título */
.login-body{min-height:100vh!important;display:flex!important;align-items:center!important;justify-content:center!important;padding:34px 18px 118px!important;overflow:hidden!important;background:radial-gradient(circle at 8% -4%,rgba(34,197,94,.17) 0 20%,transparent 20.4%),radial-gradient(circle at 94% -2%,rgba(45,212,191,.17) 0 22%,transparent 22.4%),linear-gradient(180deg,#fbfdff 0%,#effaf5 100%)!important}.login-body:before{content:""!important;position:absolute!important;left:0!important;right:0!important;bottom:0!important;height:30%!important;background:linear-gradient(135deg,#22c55e 0%,#065f46 100%)!important;clip-path:polygon(0 36%,22% 26%,48% 17%,73% 28%,100% 22%,100% 100%,0 100%)!important;z-index:0!important;opacity:.98!important}.login-body:after{content:"🛡️  Sistema seguro y confiable\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;white-space:pre!important;position:absolute!important;left:0!important;right:0!important;bottom:18px!important;text-align:center!important;color:#fff!important;font-size:16px!important;line-height:1.9!important;font-weight:650!important;z-index:1!important;background:none!important;height:auto!important;clip-path:none!important}.login-card{width:min(90vw,500px)!important;max-width:500px!important;margin:0 auto!important;padding:96px 46px 30px!important;border-radius:28px!important;background:rgba(255,255,255,.97)!important;border:1px solid #dce8ef!important;box-shadow:0 18px 48px rgba(15,23,42,.13)!important;overflow:visible!important;position:relative!important;z-index:2!important;text-align:left!important}.login-logo{display:block!important;position:absolute!important;top:-58px!important;left:50%!important;transform:translateX(-50%)!important;margin:0!important;z-index:5!important}.login-logo img{display:none!important}.login-card:before,.login-card:after{display:none!important;content:none!important}.login-avatar-svg{width:116px!important;height:116px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe7ef!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;display:grid!important;place-items:center!important}.login-avatar-svg:before{content:""!important;width:64px!important;height:64px!important;display:block!important;background:no-repeat center/contain url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 96 96'%3E%3Cdefs%3E%3ClinearGradient id='g' x1='0' y1='0' x2='1' y2='1'%3E%3Cstop stop-color='%2327c77b'/%3E%3Cstop offset='1' stop-color='%230b8f5a'/%3E%3C/linearGradient%3E%3C/defs%3E%3Ccircle cx='35' cy='30' r='13' fill='url(%23g)'/%3E%3Ccircle cx='61' cy='28' r='15' fill='url(%23g)'/%3E%3Cpath d='M13 72c1-17 13-28 29-28 5 0 9 1 13 3-7 5-12 13-13 25H13z' fill='url(%23g)'/%3E%3Cpath d='M38 72c1-19 14-31 32-31 15 0 25 12 25 31H38z' fill='url(%23g)'/%3E%3C/svg%3E")!important}.login-title{text-align:center!important;margin:0 0 24px!important;position:relative!important;z-index:6!important}.login-title h1{font-size:36px!important;line-height:1.05!important;font-weight:800!important;color:#182233!important;letter-spacing:-.035em!important;margin:0 0 10px!important;text-transform:uppercase!important;position:relative!important;z-index:6!important}.login-title b{display:block!important;color:#667085!important;font-size:16px!important;font-weight:520!important;line-height:1.32!important}.login-title:after{content:""!important;display:block!important;width:60px!important;height:4px!important;border-radius:999px!important;background:#16b978!important;margin:18px auto 0!important}.login-card .field{margin-bottom:13px!important}.login-card .field label{display:block!important;color:#1f2937!important;font-size:14px!important;font-weight:720!important;margin:0 0 7px!important}.login-input{height:54px!important;background:#eaf3ff!important;border:1px solid #cfe0f2!important;border-radius:17px!important;padding:0 10px!important;gap:10px!important;display:flex!important;align-items:center!important;margin:0!important;box-shadow:none!important;color:#172033!important}.login-input:focus-within{border-color:#18b96d!important;box-shadow:0 0 0 3px rgba(22,185,120,.13)!important}.login-input i{width:31px!important;height:31px!important;border-radius:10px!important;background:rgba(255,255,255,.72)!important;color:#11945d!important;font-size:16px!important;display:grid!important;place-items:center!important;flex:0 0 auto!important}.login-input input,.login-input select{height:42px!important;min-height:42px!important;background:#fff!important;border:1px solid #d5e3ef!important;border-radius:15px!important;color:#172033!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;font-size:16px!important;font-weight:520!important;padding:0 13px!important;outline:none!important;box-shadow:none!important;width:100%!important}.login-input select{appearance:auto!important;color:#172033!important;-webkit-text-fill-color:#172033!important}.login-input select option{color:#172033!important;background:#fff!important}.login-input input::placeholder{color:#6b7787!important;opacity:1!important;-webkit-text-fill-color:#6b7787!important}.login-input input:focus,.login-input select:focus{background:#fff!important;color:#111827!important;-webkit-text-fill-color:#111827!important;border-color:#18b96d!important;box-shadow:none!important}.login-input input:-webkit-autofill{-webkit-box-shadow:0 0 0 1000px #fff inset!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important}.login-card .btn-green{width:100%!important;height:58px!important;border-radius:17px!important;background:linear-gradient(135deg,#1fac67,#0f8f55)!important;color:#fff!important;font-size:19px!important;font-weight:760!important;display:flex!important;align-items:center!important;justify-content:center!important;margin:16px 0 20px!important;border:0!important;box-shadow:0 14px 28px rgba(22,163,74,.24)!important}.login-links{margin:0!important;padding:0!important;text-align:center!important}.login-links a{font-size:14px!important;color:#667085!important;font-weight:650!important;text-decoration:none!important}
@media(max-width:700px){.login-body{padding:82px 12px 108px!important;align-items:center!important;overflow-y:auto!important}.login-card{width:calc(100vw - 26px)!important;max-width:420px!important;padding:96px 26px 30px!important;border-radius:26px!important}.login-logo{top:-56px!important}.login-avatar-svg{width:112px!important;height:112px!important}.login-avatar-svg:before{width:62px!important;height:62px!important}.login-title h1{font-size:31px!important}.login-title b{font-size:17px!important;max-width:310px!important;margin:auto!important}.login-input{height:58px!important}.login-input input,.login-input select{height:46px!important;min-height:46px!important;font-size:17px!important}.login-card .btn-green{height:60px!important;font-size:20px!important}.login-body:after{font-size:12px!important;bottom:14px!important}.side-head-pro{height:82px!important;margin:0 10px 14px!important}}



/* =========================================================
   CORRECCIÓN DEFINITIVA - HEADER SIN TRANSPARENCIA + LOGIN
   ========================================================= */
/* 1. Login igual a referencia, título visible y sin icono encima */
.login-body{
  min-height:100vh!important;
  display:flex!important;
  align-items:center!important;
  justify-content:center!important;
  padding:36px 18px 118px!important;
  overflow:hidden!important;
  background:
    radial-gradient(circle at 8% -5%,rgba(34,197,94,.18) 0 20%,transparent 20.5%),
    radial-gradient(circle at 94% -4%,rgba(45,212,191,.18) 0 22%,transparent 22.5%),
    linear-gradient(180deg,#fbfdff 0%,#effaf5 100%)!important;
}
.login-body:before{
  content:""!important;
  position:absolute!important;left:0!important;right:0!important;bottom:0!important;
  height:30%!important;
  background:linear-gradient(135deg,#22c55e 0%,#065f46 100%)!important;
  clip-path:polygon(0 36%,23% 26%,48% 17%,73% 28%,100% 22%,100% 100%,0 100%)!important;
  z-index:0!important;opacity:.98!important;
}
.login-body:after{
  content:"🛡️  Sistema seguro y confiable\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;
  white-space:pre!important;position:absolute!important;left:0!important;right:0!important;bottom:18px!important;
  text-align:center!important;color:#fff!important;font-size:16px!important;line-height:1.9!important;font-weight:650!important;
  z-index:1!important;background:none!important;height:auto!important;clip-path:none!important;
}
.login-card{
  width:min(90vw,500px)!important;max-width:500px!important;
  margin:0 auto!important;padding:120px 46px 30px!important;
  border-radius:28px!important;background:rgba(255,255,255,.975)!important;border:1px solid #dce8ef!important;
  box-shadow:0 18px 48px rgba(15,23,42,.13)!important;overflow:visible!important;position:relative!important;z-index:2!important;
  text-align:left!important;
}
.login-logo{display:block!important;position:absolute!important;top:-58px!important;left:50%!important;transform:translateX(-50%)!important;margin:0!important;z-index:4!important;}
.login-logo img{display:none!important}.login-card:before,.login-card:after{display:none!important;content:none!important}
.login-avatar-svg{width:116px!important;height:116px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe7ef!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;display:grid!important;place-items:center!important;}
.login-avatar-svg:before{width:64px!important;height:64px!important;}
.login-title{position:relative!important;z-index:10!important;text-align:center!important;margin:0 0 24px!important;}
.login-title h1{
  position:relative!important;z-index:10!important;margin:0 0 10px!important;
  font-size:36px!important;line-height:1.05!important;font-weight:800!important;letter-spacing:-.035em!important;
  color:#182233!important;text-transform:uppercase!important;white-space:nowrap!important;
}
.login-title b{display:block!important;color:#667085!important;font-size:16px!important;font-weight:520!important;line-height:1.32!important;}
.login-title:after{content:""!important;display:block!important;width:60px!important;height:4px!important;border-radius:999px!important;background:#16b978!important;margin:18px auto 0!important;}
.login-input{height:54px!important;background:#eaf3ff!important;border:1px solid #cfe0f2!important;border-radius:17px!important;padding:0 10px!important;gap:10px!important;display:flex!important;align-items:center!important;margin:0!important;box-shadow:none!important;color:#172033!important;}
.login-input input,.login-input select{height:42px!important;min-height:42px!important;background:#fff!important;border:1px solid #d5e3ef!important;border-radius:15px!important;color:#172033!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;font-size:16px!important;font-weight:520!important;padding:0 13px!important;outline:none!important;box-shadow:none!important;width:100%!important;}
.login-input input:focus,.login-input select:focus{background:#fff!important;color:#111827!important;-webkit-text-fill-color:#111827!important;border-color:#18b96d!important;box-shadow:none!important;}
.login-input input:-webkit-autofill{-webkit-box-shadow:0 0 0 1000px #fff inset!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;}
.login-card .btn-green{width:100%!important;height:58px!important;border-radius:17px!important;font-size:19px!important;font-weight:760!important;margin:16px 0 20px!important;}
.login-links{margin:0!important;padding:0!important;text-align:center!important}.login-links a{font-size:14px!important;color:#667085!important;font-weight:650!important;text-decoration:none!important;}

/* 2. Encabezado lateral: fondo opaco que tapa cualquier pestaña al hacer scroll */
.side{background:#073126!important;overflow-y:auto!important;overflow-x:hidden!important;}
.side-head-pro{
  position:sticky!important;top:0!important;z-index:2147483646!important;
  margin:0 12px 18px!important;height:82px!important;min-height:82px!important;
  padding:0 12px!important;display:flex!important;align-items:center!important;justify-content:space-between!important;
  background:transparent!important;border:1px solid rgba(255,255,255,.16)!important;border-radius:20px!important;
  box-shadow:0 16px 34px rgba(0,0,0,.34)!important;overflow:visible!important;isolation:isolate!important;backdrop-filter:none!important;
}
.side-head-pro:before{
  content:""!important;position:absolute!important;left:-22px!important;right:-22px!important;top:-18px!important;bottom:-18px!important;
  background:#073126!important;z-index:-3!important;border-radius:0!important;opacity:1!important;
}
.side-head-pro:after{
  content:""!important;position:absolute!important;inset:0!important;
  background:linear-gradient(90deg,#123f33 0%,#196b48 52%,#123f33 100%)!important;
  z-index:-2!important;border-radius:20px!important;opacity:1!important;
}
.side-head-pro *{position:relative!important;z-index:3!important;}
.side-brand-text b{color:#fff!important;font-size:16px!important;font-weight:800!important;line-height:1.05!important;white-space:nowrap!important;}
.side-brand-text small{color:#d6e9e2!important;font-size:14px!important;font-weight:650!important;line-height:1.15!important;display:block!important;}
.side-head-pro .toggle{width:48px!important;height:48px!important;border-radius:14px!important;background:rgba(255,255,255,.12)!important;color:#fff!important;border:0!important;display:grid!important;place-items:center!important;font-size:24px!important;}
nav{position:relative!important;z-index:1!important;padding-top:4px!important;}

/* 3. Texto del menú bien visible, especialmente en celular */
.menu-title,.menu-item{color:#eaf7f1!important;opacity:1!important;text-shadow:none!important;font-weight:760!important;}
.menu-title .label,.menu-item .label{color:inherit!important;opacity:1!important;visibility:visible!important;display:inline!important;}
.menu-item.active,.menu-item.parent-active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:rgba(22,185,120,.30)!important;color:#ffffff!important;box-shadow:inset 4px 0 0 #34d399!important;border:1px solid rgba(134,239,172,.22)!important;}
.menu-item.active .label,.menu-item.parent-active .label,.menu-title.active .label,.menu-group.force-open>.menu-title.active .label{color:#ffffff!important;}
.menu-title i,.menu-item i{color:#d7eee5!important;background:rgba(255,255,255,.08)!important;}
.mobile-head{background:#0f4a35!important;color:#fff!important;z-index:2147483645!important;}
.mobile-head b,.mobile-head span,.mobile-head a{color:#fff!important;opacity:1!important;visibility:visible!important;}

@media(max-width:1000px){
  .side{position:fixed!important;top:0!important;left:-340px!important;width:315px!important;height:100dvh!important;background:#073126!important;z-index:2147483644!important;}
  .side.open{left:0!important;}
  .side-head-pro{top:0!important;margin:0 10px 18px!important;height:82px!important;min-height:82px!important;}
  .side-head-pro:before{left:-14px!important;right:-14px!important;top:-14px!important;bottom:-16px!important;background:#073126!important;}
  .menu-title,.menu-item{font-size:13.5px!important;color:#f8fffc!important;font-weight:760!important;}
  .menu-title .label,.menu-item .label{color:#f8fffc!important;}
  .login-body{padding:86px 12px 108px!important;align-items:center!important;overflow-y:auto!important;}
  .login-card{width:calc(100vw - 26px)!important;max-width:420px!important;padding:112px 26px 30px!important;border-radius:26px!important;}
  .login-logo{top:-56px!important;}
  .login-avatar-svg{width:112px!important;height:112px!important;}
  .login-avatar-svg:before{width:62px!important;height:62px!important;}
  .login-title h1{font-size:31px!important;}
  .login-title b{font-size:17px!important;max-width:310px!important;margin:auto!important;}
  .login-body:after{font-size:12px!important;bottom:14px!important;}
}


/* =========================================================
   AJUSTE FINAL EXACTO LOGIN REFERENCIA 1443
   - fondo suave verde/menta
   - tarjeta compacta centrada
   - icono arriba sin tapar letras
   - textos y campos visibles
   ========================================================= */
.login-body{
  min-height:100vh!important;
  display:flex!important;
  align-items:center!important;
  justify-content:center!important;
  padding:38px 18px 118px!important;
  overflow:hidden!important;
  position:relative!important;
  background:
    radial-gradient(circle at 7% -6%,rgba(34,197,94,.18) 0 20%,transparent 20.4%),
    radial-gradient(circle at 94% -6%,rgba(45,212,191,.17) 0 22%,transparent 22.4%),
    linear-gradient(180deg,#fbfdff 0%,#effaf5 100%)!important;
}
.login-body:before{
  content:""!important;
  position:absolute!important;
  left:0!important;right:0!important;bottom:0!important;
  height:31%!important;
  background:linear-gradient(135deg,#22c55e 0%,#059669 45%,#064e3b 100%)!important;
  clip-path:polygon(0 39%,22% 29%,48% 20%,73% 30%,100% 24%,100% 100%,0 100%)!important;
  z-index:0!important;
  opacity:.98!important;
}
.login-body:after{
  content:"🛡️  Sistema seguro y confiable\A© 2025 Prize Pro. Todos los derechos reservados."!important;
  white-space:pre!important;
  position:absolute!important;
  left:0!important;right:0!important;bottom:18px!important;
  text-align:center!important;
  color:#fff!important;
  font-size:16px!important;
  line-height:1.9!important;
  font-weight:650!important;
  z-index:1!important;
  background:none!important;
  height:auto!important;
  clip-path:none!important;
}
.login-card{
  width:min(90vw,500px)!important;
  max-width:500px!important;
  min-height:auto!important;
  margin:0 auto!important;
  padding:82px 46px 30px!important;
  border-radius:28px!important;
  background:rgba(255,255,255,.975)!important;
  border:1px solid #dce8ef!important;
  box-shadow:0 20px 54px rgba(15,23,42,.12)!important;
  overflow:visible!important;
  position:relative!important;
  z-index:2!important;
  text-align:left!important;
}
.login-logo{
  display:block!important;
  position:absolute!important;
  top:-62px!important;
  left:50%!important;
  transform:translateX(-50%)!important;
  margin:0!important;
  z-index:4!important;
}
.login-logo img{display:none!important;}
.login-card:before,.login-card:after{display:none!important;content:none!important;}
.login-avatar-svg{
  width:124px!important;
  height:124px!important;
  border-radius:50%!important;
  background:#fff!important;
  border:1px solid #dbe7ef!important;
  box-shadow:0 16px 40px rgba(15,23,42,.10)!important;
  display:grid!important;
  place-items:center!important;
}
.login-avatar-svg:before{
  content:""!important;
  display:block!important;
  width:68px!important;
  height:68px!important;
  background:center/contain no-repeat url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 96 96'%3E%3Cdefs%3E%3ClinearGradient id='g' x1='0' y1='0' x2='1' y2='1'%3E%3Cstop stop-color='%2327c77b'/%3E%3Cstop offset='1' stop-color='%230b8f5a'/%3E%3C/linearGradient%3E%3C/defs%3E%3Ccircle cx='35' cy='30' r='13' fill='url(%23g)'/%3E%3Ccircle cx='61' cy='28' r='15' fill='url(%23g)'/%3E%3Cpath d='M13 72c1-17 13-28 29-28 5 0 9 1 13 3-7 5-12 13-13 25H13z' fill='url(%23g)'/%3E%3Cpath d='M38 72c1-19 14-31 32-31 15 0 25 12 25 31H38z' fill='url(%23g)'/%3E%3C/svg%3E")!important;
}
.login-title{
  position:relative!important;
  z-index:6!important;
  text-align:center!important;
  margin:0 0 23px!important;
}
.login-title h1{
  position:relative!important;
  z-index:6!important;
  margin:0 0 11px!important;
  color:#182233!important;
  font-size:38px!important;
  line-height:1.06!important;
  font-weight:800!important;
  letter-spacing:-.035em!important;
  text-transform:uppercase!important;
  white-space:nowrap!important;
}
.login-title b{
  display:block!important;
  color:#667085!important;
  font-size:16px!important;
  font-weight:520!important;
  line-height:1.35!important;
}
.login-title:after{
  content:""!important;
  display:block!important;
  width:60px!important;
  height:4px!important;
  border-radius:999px!important;
  background:#16b978!important;
  margin:18px auto 0!important;
}
.login-card .field{margin-bottom:13px!important;}
.login-card .field label{
  display:block!important;
  color:#1f2937!important;
  font-size:14px!important;
  font-weight:720!important;
  margin:0 0 7px!important;
}
.login-input{
  height:54px!important;
  background:#eaf3ff!important;
  border:1px solid #cfe0f2!important;
  border-radius:17px!important;
  padding:0 10px!important;
  gap:10px!important;
  display:flex!important;
  align-items:center!important;
  margin:0!important;
  box-shadow:none!important;
  color:#172033!important;
}
.login-input:focus-within{border-color:#18b96d!important;box-shadow:0 0 0 3px rgba(22,185,120,.13)!important;}
.login-input i{
  width:31px!important;height:31px!important;border-radius:10px!important;
  background:rgba(255,255,255,.72)!important;color:#11945d!important;font-size:16px!important;
  display:grid!important;place-items:center!important;flex:0 0 auto!important;
}
.login-input input,.login-input select{
  height:42px!important;min-height:42px!important;background:#fff!important;border:1px solid #d5e3ef!important;border-radius:15px!important;
  color:#172033!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;
  font-size:16px!important;font-weight:520!important;padding:0 13px!important;outline:none!important;box-shadow:none!important;width:100%!important;
}
.login-input select{appearance:auto!important;color:#172033!important;-webkit-text-fill-color:#172033!important;}
.login-input select option{background:#fff!important;color:#172033!important;}
.login-input input::placeholder{color:#6b7787!important;opacity:1!important;-webkit-text-fill-color:#6b7787!important;}
.login-input input:focus,.login-input select:focus{background:#fff!important;color:#111827!important;-webkit-text-fill-color:#111827!important;border-color:#18b96d!important;box-shadow:none!important;}
.login-input input:-webkit-autofill{-webkit-box-shadow:0 0 0 1000px #fff inset!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;}
.login-card .btn-green{
  width:100%!important;height:58px!important;border-radius:17px!important;background:linear-gradient(135deg,#1fac67,#0f8f55)!important;
  color:#fff!important;font-size:19px!important;font-weight:760!important;display:flex!important;align-items:center!important;justify-content:center!important;
  margin:16px 0 20px!important;border:0!important;box-shadow:0 14px 28px rgba(22,163,74,.24)!important;
}
.login-links{margin:0!important;padding:0!important;text-align:center!important;}
.login-links a{font-size:14px!important;color:#667085!important;font-weight:650!important;text-decoration:none!important;}
@media(max-width:700px){
  .login-body{padding:74px 12px 104px!important;align-items:center!important;overflow-y:auto!important;}
  .login-card{width:calc(100vw - 26px)!important;max-width:420px!important;padding:90px 26px 30px!important;border-radius:26px!important;}
  .login-logo{top:-56px!important;}
  .login-avatar-svg{width:112px!important;height:112px!important;}
  .login-avatar-svg:before{width:62px!important;height:62px!important;}
  .login-title h1{font-size:31px!important;}
  .login-title b{font-size:17px!important;max-width:310px!important;margin:auto!important;}
  .login-input{height:58px!important;}
  .login-input input,.login-input select{height:46px!important;min-height:46px!important;font-size:17px!important;}
  .login-card .btn-green{height:60px!important;font-size:20px!important;}
  .login-body:after{font-size:12px!important;bottom:14px!important;}
}


/* === IA HR PRO - Widget global === */
.iahr-fab{position:fixed;right:22px;bottom:22px;z-index:9998;background:linear-gradient(135deg,#16a34a,#22c55e);color:#fff;border-radius:999px;padding:13px 17px;box-shadow:0 18px 40px rgba(0,0,0,.35);display:flex;align-items:center;gap:8px;font-weight:900;cursor:pointer;border:1px solid rgba(255,255,255,.22)}
.iahr-fab i{font-size:18px}.iahr-panel{position:fixed;right:22px;bottom:82px;width:min(430px,calc(100vw - 28px));max-height:70vh;z-index:9999;background:#111827;border:1px solid rgba(148,163,184,.35);border-radius:20px;box-shadow:0 25px 70px rgba(0,0,0,.48);display:none;overflow:hidden}.iahr-panel.open{display:block}.iahr-head{display:flex;justify-content:space-between;align-items:center;padding:15px 16px;background:linear-gradient(135deg,#064e3b,#111827);border-bottom:1px solid rgba(255,255,255,.08)}.iahr-head b{display:block;color:#fff}.iahr-head small{display:block;color:#bbf7d0;font-size:11px;margin-top:3px}.iahr-head button{background:rgba(255,255,255,.12);color:#fff;border:0;border-radius:10px;width:32px;height:32px;font-size:20px;cursor:pointer}.iahr-body{padding:14px}.iahr-ex{font-size:12px;color:#cbd5e1;background:#0f172a;border:1px solid rgba(148,163,184,.25);border-radius:12px;padding:10px;margin-bottom:10px}.iahr-body textarea{width:100%;height:92px;border-radius:14px;border:1px solid #334155;background:#f8fafc;color:#0f172a;padding:12px;font-weight:700;resize:vertical}.iahr-send{width:100%;margin-top:10px;justify-content:center}.iahr-respuesta{margin-top:12px;max-height:310px;overflow:auto}.ia-result{border-radius:14px;padding:13px;border:1px solid rgba(148,163,184,.25);background:#0f172a;color:#e5e7eb}.ia-result h3{margin:0 0 8px;color:#fff}.ia-result h4{margin:10px 0 6px}.ia-result p{margin:7px 0;line-height:1.45}.ia-result ul{margin:8px 0 0 18px;padding:0}.ia-result.ok{border-color:rgba(34,197,94,.45);background:linear-gradient(135deg,rgba(6,78,59,.75),#0f172a)}.ia-result.warn{border-color:rgba(245,158,11,.5);background:linear-gradient(135deg,rgba(120,53,15,.58),#0f172a)}.ia-result.bad{border-color:rgba(239,68,68,.55);background:linear-gradient(135deg,rgba(127,29,29,.62),#0f172a)}.ia-result.legal{border-color:rgba(59,130,246,.55);background:linear-gradient(135deg,rgba(30,64,175,.58),#0f172a)}.ia-result a{color:#86efac;font-weight:900}.ia-result small,.ia-result .muted{color:#d1d5db}.iahr-respuesta,.iahr-respuesta *{color:#e5e7eb}.ia-result p,.ia-result li{color:#f1f5f9}.ia-result.legal{background:linear-gradient(135deg,rgba(30,64,175,.72),#0b1220)}.ia-admin-form{display:grid;grid-template-columns:repeat(2,minmax(180px,1fr));gap:12px}.ia-admin-form textarea{grid-column:1/-1;min-height:110px}.ia-admin-form .full{grid-column:1/-1}.ia-table td{vertical-align:top}@media(max-width:700px){.iahr-fab{right:14px;bottom:14px}.iahr-panel{right:14px;bottom:72px}}


/* === AJUSTE PRO: IA HR en verde/blanco con lectura clara === */
.iahr-panel{
  background:#ffffff!important;
  border:1px solid #d6efe3!important;
  box-shadow:0 26px 80px rgba(15,118,80,.24)!important;
}
.iahr-head{
  background:linear-gradient(135deg,#0f8f5a,#16a34a)!important;
  border-bottom:0!important;
}
.iahr-head b,.iahr-head small{color:#ffffff!important;}
.iahr-head button{background:rgba(255,255,255,.22)!important;color:#ffffff!important;}
.iahr-body{background:#ffffff!important;}
.iahr-ex{
  color:#0f5132!important;
  background:#ecfdf5!important;
  border:1px solid #bbf7d0!important;
}
.iahr-body textarea{
  background:#ffffff!important;
  color:#0f172a!important;
  border:2px solid #bbf7d0!important;
  box-shadow:0 0 0 4px rgba(34,197,94,.08)!important;
}
.iahr-body textarea:focus{
  outline:none!important;
  border-color:#16a34a!important;
  box-shadow:0 0 0 4px rgba(22,163,74,.16)!important;
}
.iahr-respuesta{background:#ffffff!important;}
.ia-result,.ia-result.ok,.ia-result.warn,.ia-result.bad,.ia-result.legal{
  background:#ffffff!important;
  color:#0f172a!important;
  border:1.5px solid #bbf7d0!important;
  box-shadow:0 12px 30px rgba(15,118,80,.10)!important;
}
.ia-result h3,.ia-result h4{
  color:#064e3b!important;
}
.ia-result p,.ia-result li,.ia-result small,.ia-result .muted,.iahr-respuesta,.iahr-respuesta *{
  color:#1f2937!important;
}
.ia-result b,.ia-result strong{
  color:#065f46!important;
}
.ia-result.legal{
  border-left:7px solid #16a34a!important;
}
.ia-result.ok{
  border-left:7px solid #22c55e!important;
}
.ia-result.warn{
  border-left:7px solid #f59e0b!important;
}
.ia-result.bad{
  border-left:7px solid #ef4444!important;
}
.ia-result a{
  color:#047857!important;
  font-weight:900!important;
}
.iahr-send,.iahr-body .btn-green{
  background:linear-gradient(135deg,#16a34a,#0f9f61)!important;
  color:#ffffff!important;
  border:0!important;
  box-shadow:0 14px 34px rgba(16,185,129,.22)!important;
}
.iahr-send:hover{filter:brightness(.97)!important;transform:translateY(-1px);}
.iahr-fab{
  background:linear-gradient(135deg,#16a34a,#0f9f61)!important;
  color:#ffffff!important;
  box-shadow:0 18px 42px rgba(16,185,129,.35)!important;
}


/* === AJUSTE PRO: entradas separadas trabajador / administrador === */
.login-links{
  display:flex!important;
  justify-content:center!important;
  align-items:center!important;
  margin-top:18px!important;
}
.login-links a.login-switch{
  display:inline-flex!important;
  align-items:center!important;
  justify-content:center!important;
  gap:8px!important;
  padding:12px 18px!important;
  border-radius:18px!important;
  font-weight:900!important;
  text-decoration:none!important;
  font-size:16px!important;
  transition:.18s ease!important;
}
.login-links a.login-switch.admin{
  background:#ecfdf5!important;
  color:#047857!important;
  border:1px solid #bbf7d0!important;
}
.login-links a.login-switch.worker{
  background:#eff6ff!important;
  color:#1d4ed8!important;
  border:1px solid #bfdbfe!important;
}
.login-links a.login-switch:hover{
  transform:translateY(-1px)!important;
  box-shadow:0 12px 28px rgba(15,118,80,.15)!important;
}
.login-mode-pill{
  width:100%!important;
  display:flex!important;
  align-items:center!important;
  justify-content:center!important;
  gap:9px!important;
  padding:12px 16px!important;
  margin:-2px 0 10px!important;
  border-radius:18px!important;
  font-weight:900!important;
  letter-spacing:.2px!important;
}
.login-mode-pill.admin-pill{
  background:linear-gradient(135deg,#ecfdf5,#ffffff)!important;
  color:#047857!important;
  border:1px solid #bbf7d0!important;
}

</style>
<script>
function side(){return document.querySelector('.side')}
function appShell(){return document.querySelector('.app')}
function saveSideScroll(){const s=side(); if(s){localStorage.setItem('sideScroll',s.scrollTop||0)}}
function restoreSideScroll(){const s=side(); if(s){s.scrollTop=parseInt(localStorage.getItem('sideScroll')||'0')}}
function toggleSide(){const s=side(), a=appShell(); if(!s)return; if(window.innerWidth<1000){s.classList.toggle('open')}else{const c=!s.classList.contains('collapsed'); s.classList.toggle('collapsed',c); if(a)a.classList.toggle('side-collapsed',c); localStorage.setItem('sideCollapsed',c?'1':'0')}}
function toggleGroup(id){const g=document.getElementById(id); if(!g)return; g.classList.toggle('closed'); localStorage.setItem('group_'+id,g.classList.contains('closed')?'1':'0')}
function initSide(){const s=side(), a=appShell(); if(!s)return; const c=localStorage.getItem('sideCollapsed')==='1' && window.innerWidth>=1000; s.classList.toggle('collapsed',c); if(a)a.classList.toggle('side-collapsed',c); document.querySelectorAll('.menu-group[data-group]').forEach(g=>{const id=g.id; const saved=localStorage.getItem('group_'+id); if(saved==='1' && !g.classList.contains('force-open')) g.classList.add('closed')}); if(!location.hash){setTimeout(restoreSideScroll,60)}; document.querySelectorAll('.menu-item').forEach(a=>a.addEventListener('click',()=>{saveSideScroll(); if(window.innerWidth<1000){const s=side(); if(s)s.classList.remove('open')}}));}
function filterCards(){const q=(document.getElementById('cardSearch')?.value||'').toLowerCase();document.querySelectorAll('.doc-card').forEach(c=>{c.style.display=c.innerText.toLowerCase().includes(q)?'block':'none'})}
window.addEventListener('DOMContentLoaded',()=>{initSide(); if(location.hash){document.querySelectorAll('.menu-item').forEach(x=>{if(x.getAttribute('href')&&x.getAttribute('href').endsWith(location.hash)) x.classList.add('active')}); setTimeout(()=>{document.querySelector(location.hash)?.scrollIntoView({block:'start'});},120)}});window.addEventListener('beforeunload',saveSideScroll)


function iahrToggle(){const p=document.getElementById('iahrPanel'); if(p)p.classList.toggle('open')}
async function iahrAsk(ev){
  if(ev && ev.preventDefault) ev.preventDefault();
  const q=(document.getElementById('iahrPregunta')?.value||'').trim();
  const modulo=document.getElementById('iahrModulo')?.value||'';
  const box=document.getElementById('iahrRespuesta');
  if(!box)return false;
  if(!q){box.innerHTML='<div class="ia-result warn"><h3>IA HR</h3><p>Escribe una pregunta para poder ayudarte.</p></div>';return false;}
  box.innerHTML='<div class="ia-result warn"><p>Analizando...</p></div>';
  try{
    const r=await fetch('/ia_hr/api',{method:'POST',credentials:'same-origin',headers:{'Content-Type':'application/json','Accept':'application/json'},body:JSON.stringify({pregunta:q,modulo:modulo})});
    const txt=await r.text();
    let data={};
    try{data=JSON.parse(txt)}catch(parseErr){throw new Error('Respuesta no válida del servidor: '+txt.slice(0,180));}
    if(!r.ok || data.ok===false){throw new Error(data.error || ('HTTP '+r.status));}
    box.innerHTML=data.html||'<div class="ia-result bad">No se obtuvo respuesta.</div>';
  }catch(e){box.innerHTML='<div class="ia-result bad"><h3>Error IA</h3><p>No se pudo consultar la IA.</p><small>'+String(e.message||e)+'</small></div>'; console.error('IA HR error:',e);}
  return false;
}
document.addEventListener('DOMContentLoaded',function(){
  document.querySelectorAll('.iahr-send').forEach(function(btn){
    btn.addEventListener('click',iahrAsk);
  });
  const txt=document.getElementById('iahrPregunta');
  if(txt){txt.addEventListener('keydown',function(e){if(e.key==='Enter' && !e.shiftKey){e.preventDefault(); iahrAsk(e);}});}
});

document.addEventListener('click',function(e){
  document.querySelectorAll('.admin-dropdown.open').forEach(function(d){
    if(!d.contains(e.target)){d.classList.remove('open')}
  });
});

</script></head><body>{{ body|safe }}
<script>
window.addEventListener('DOMContentLoaded',()=>{
  const tipo=document.querySelector("select[name='tipo_plantilla']");
  const esquema=document.getElementById('modal_esquema');
  if(tipo && esquema){
    const wordOptions=[
      'Trabajador Contrato Laboral',
      'Trabajador Datos Laborales',
      'Esquema Trabajador Datos Laborales GR',
      'Trabajador Declaración Jurada Datos Personales',
      'Trabajador Declaración Jurada Parentesco'
    ];
    const pdfOptions=['Trabajador Contrato Laboral'];
    function cargarEsquemas(){
      const lista=(tipo.value||'').toUpperCase()==='WORD'?wordOptions:pdfOptions;
      const actual=esquema.value;
      esquema.innerHTML=lista.map(v=>`<option>${v}</option>`).join('');
      if(lista.includes(actual)) esquema.value=actual;
    }
    tipo.addEventListener('change', cargarEsquemas);
    cargarEsquemas();
  }
});
</script>
</body></html>
'''


def render_page(content, title="PORTAL HR PRO", active="Inicio"):
    user_label = session.get('admin_nombre') or session.get('nombre') or 'Usuario HR'
    primer_nombre = user_label.split()[0] if user_label else 'Usuario'
    body = f'''
    <div class="mobile-head"><button class="toggle" onclick="toggleSide()">☰</button><b>PORTAL HR PRO</b><a href="/logout">Salir</a></div>
    <div class="app"><aside class="side"><div class="side-head-pro"><button class="toggle side-toggle-left" title="Expandir / contraer panel" onclick="toggleSide()"><i class="bi bi-list"></i></button><div class="side-brand-pro"><div class="side-brand-icon"><i class="bi bi-kanban"></i></div><div class="side-brand-text"><b>PORTAL HR PRO</b><small>Dashboard</small></div></div><button class="toggle side-toggle-right" title="Expandir / contraer panel" onclick="toggleSide()"><i class="bi bi-list"></i></button></div>
      {sidebar(active)}<div class="side-user"><div class="avatar"><i class="bi bi-person-fill"></i></div><div><b>{primer_nombre}</b><br><small>{'Administrador' if session.get('admin_id') else 'Trabajador'}</small></div></div></aside><main class="main">{flashes()}{content}</main>{ia_widget_hr(active)}</div>'''
    return render_template_string(BASE, body=body, title=title)


def flashes():
    out = ""
    for cat, msg in list(getattr(request, 'flashes', []) or []):
        out += f"<div class='flash {'err' if cat=='error' else ''}'>{msg}</div>"
    # Flask get_flashed_messages unavailable without import? import below lazily
    from flask import get_flashed_messages
    out = "".join([f"<div class='flash {'err' if c=='error' else ''}'>{m}</div>" for c, m in get_flashed_messages(with_categories=True)])
    return out


def item(tipo, label, icon, active):
    active_type = str(active).split(':', 1)[0]
    cls = "menu-item active" if active_type == tipo else "menu-item"
    return f"<a class='{cls}' onclick='saveSideScroll()' href='{url_for('panel_tipo', tipo=tipo)}'><span>{icon}</span><span class='label'>{label}</span></a>"


def sidebar(active):
    active_txt = str(active or '')
    active_type = active_txt.split(':', 1)[0]
    active_sub = active_txt.split(':', 1)[1] if ':' in active_txt else ''
    pago_parts=[]
    for k,l,i in TIPOS_PAGO:
        if k=='Normal':
            sub_open = ' force-open' if active_type == k else ''
            base_cls = 'menu-item parent-active' if active_type == k else 'menu-item'
            cls_mensual = 'menu-item sub-mini active' if active_type == k and active_sub == 'Mensual' else 'menu-item sub-mini'
            cls_semanal = 'menu-item sub-mini active' if active_type == k and active_sub == 'Semanal' else 'menu-item sub-mini'
            pago_parts.append(f"<div id='grp_normal' data-group='normal' class='menu-group nested{sub_open}'><button type='button' class='{base_cls}' onclick=\"toggleGroup('grp_normal')\"><span>{i}</span><span class='label'>{l}</span><span class='chev'>∨</span></button><div class='submenu'>")
            pago_parts.append(f"<a class='{cls_mensual}' onclick='saveSideScroll()' href='{url_for('panel_tipo', tipo=k, sub='Mensual')}'><i class='bi bi-calendar3'></i><span class='label'>Normal mensual</span></a>")
            pago_parts.append(f"<a class='{cls_semanal}' onclick='saveSideScroll()' href='{url_for('panel_tipo', tipo=k, sub='Semanal')}'><i class='bi bi-calendar-week'></i><span class='label'>Normal semanal</span></a></div></div>")
        else:
            pago_parts.append(item(k,l,i,active))
    pago = ''.join(pago_parts)
    emp = ''.join(item(k,l,i,active) for k,l,i in TIPOS_EMPRESA)
    per = ''.join(item(k,l,i,active) for k,l,i in TIPOS_PERSONALES)
    def gclass(keys):
        return 'menu-group force-open' if active_type in keys else 'menu-group'
    pago_cls = gclass([k for k,_,_ in TIPOS_PAGO])
    emp_cls = gclass([k for k,_,_ in TIPOS_EMPRESA])
    per_cls = gclass([k for k,_,_ in TIPOS_PERSONALES])
    admin = ""
    if session.get('admin_id'):
        admin_keys = ['Admin','Trabajadores','Usuarios','Modulo documentos','Subir documentos','Gestion Vacacional','Modo prueba'] + [k for k,_,_ in TIPOS_PAGO] + [k for k,_,_ in TIPOS_EMPRESA] + [k for k,_,_ in TIPOS_PERSONALES]
        admin_cls = 'menu-group force-open' if active_type in admin_keys else 'menu-group'
        cls_dash = 'menu-item active' if active == 'Admin' else 'menu-item'
        cls_trab = 'menu-item active' if active == 'Trabajadores' else 'menu-item'
        cls_docs = 'menu-item active' if active == 'Subir documentos' else 'menu-item'
        cls_users = 'menu-item active' if active == 'Usuarios' else 'menu-item'
        cls_moddocs = 'menu-item active' if active == 'Modulo documentos' else 'menu-item'
        cls_vac = 'menu-item active' if active == 'Gestion Vacacional' else 'menu-item'
        cls_con = 'menu-item active' if active == 'Gestion Contratacion' else 'menu-item'
        cls_test = 'menu-item active' if active == 'Modo prueba' else 'menu-item'
        docs_mod_keys = [k for k,_,_ in TIPOS_PAGO] + [k for k,_,_ in TIPOS_EMPRESA] + [k for k,_,_ in TIPOS_PERSONALES] + ['Modulo documentos','Subir documentos']
        docs_mod_cls = 'menu-group nested force-open' if active_type in docs_mod_keys else 'menu-group nested'
        vac_cls = 'menu-group nested force-open' if active == 'Gestion Vacacional' else 'menu-group nested'
        con_cls = 'menu-group nested force-open' if active_type == 'Gestion Contratacion' else 'menu-group nested'
        docs_head = 'menu-title' + (' active' if active_type in docs_mod_keys else '')
        vac_head = 'menu-title' + (' active' if active == 'Gestion Vacacional' else '')
        con_head = 'menu-title' + (' active' if active_type == 'Gestion Contratacion' else '')
        docs_count_con = 0
        try:
            with db() as _conx:
                docs_count_con = _conx.execute('SELECT COUNT(*) FROM contratacion_docs').fetchone()[0]
        except Exception:
            docs_count_con = 0
        archivos_cls = 'menu-item sub-mini doc-loaded ' + ('active' if active_sub == 'documentaria' else '') if docs_count_con else 'menu-item sub-mini ' + ('active' if active_sub == 'documentaria' else '')
        admin = f"""
        <div id='grp_admin' data-group='admin' class='{admin_cls}'>
          <button type='button' class='menu-title' onclick="toggleGroup('grp_admin')"><i class='bi bi-grid-1x2'></i><span class='label'>Administrador</span><span class='chev'>∨</span></button>
          <div class='submenu'>
            <a class='{cls_dash}' onclick='saveSideScroll()' href='/admin'><i class='bi bi-speedometer2'></i><span class='label'>Dashboard</span></a>
            <div id='grp_modulo_documentos' data-group='modulo_documentos' class='{docs_mod_cls}'>
              <button type='button' class='{docs_head}' onclick="toggleGroup('grp_modulo_documentos')"><i class='bi bi-folder2-open'></i><span class='label'>1. Gestión Documental</span><span class='chev'>∨</span></button>
              <div class='submenu'>
                <a class='{cls_moddocs}' onclick='saveSideScroll()' href='/admin/modulo/documentos'><i class='bi bi-speedometer2'></i><span class='label'>Dashboard</span></a>
                <div id='grp_pago' data-group='pago' class='{pago_cls}'>
                  <button type='button' class='menu-title' onclick="toggleGroup('grp_pago')"><i class='bi bi-file-earmark-text'></i><span class='label'>Documentos de pago</span><span class='chev'>∨</span></button>
                  <div class='submenu'>{pago}</div>
                </div>
                <div id='grp_empresa' data-group='empresa' class='{emp_cls}'>
                  <button type='button' class='menu-title' onclick="toggleGroup('grp_empresa')"><i class='bi bi-building'></i><span class='label'>Documentos de la empresa</span><span class='chev'>∨</span></button>
                  <div class='submenu'>{emp}</div>
                </div>
                <div id='grp_personal' data-group='personal' class='{per_cls}'>
                  <button type='button' class='menu-title' onclick="toggleGroup('grp_personal')"><i class='bi bi-person-lines-fill'></i><span class='label'>Documentos personales</span><span class='chev'>∨</span></button>
                  <div class='submenu'>{per}</div>
                </div>
                <a class='{cls_docs}' onclick='saveSideScroll()' href='/admin/documentos'><i class='bi bi-cloud-arrow-up'></i><span class='label'>Subir / gestionar documentos</span></a>
                <a class='menu-item' onclick='saveSideScroll()' href='/admin/plantilla_gestion/documental'><i class='bi bi-file-earmark-arrow-down'></i><span class='label'>Plantilla Documental</span></a>
              </div>
            </div>
            <div id='grp_vacacional' data-group='vacacional' class='{vac_cls}'>
              <button type='button' class='{vac_head}' onclick="toggleGroup('grp_vacacional')"><i class='bi bi-calendar2-week'></i><span class='label'>2. Gestión Vacacional</span><span class='chev'>∨</span></button>
              <div class='submenu'>
                <a class='{cls_vac}' onclick='saveSideScroll()' href='/admin/vacaciones'><i class='bi bi-bar-chart-line'></i><span class='label'>Dashboard vacacional</span></a>
                <a class='menu-item' onclick='saveSideScroll()' href='/admin/vacaciones#cargar-saldos'><i class='bi bi-calendar-check'></i><span class='label'>Saldos Vacacionales</span></a>
                <a class='menu-item' onclick='saveSideScroll()' href='/admin/vacaciones#solicitudes'><i class='bi bi-file-earmark-check'></i><span class='label'>Solicitudes de Vacaciones</span></a>
                <a class='menu-item' onclick='saveSideScroll()' href='/admin/vacaciones#aprobaciones'><i class='bi bi-check2-circle'></i><span class='label'>Aprobaciones</span></a>
                <a class='menu-item' onclick='saveSideScroll()' href='/admin/vacaciones#reportes'><i class='bi bi-clipboard-data'></i><span class='label'>Reportes</span></a>
                <a class='menu-item' onclick='saveSideScroll()' href='/admin/plantilla_gestion/vacacional'><i class='bi bi-file-earmark-arrow-down'></i><span class='label'>Plantilla Vacacional</span></a>
              </div>
            </div>
            <div id='grp_trabajadores_admin' data-group='trabajadores_admin' class='menu-group nested {'force-open' if active in ['Trabajadores','Usuarios'] or active_sub == 'anuncios' else ''}'>
              <button type='button' class='menu-title {'active' if active in ['Trabajadores','Usuarios'] or active_sub == 'anuncios' else ''}' onclick="toggleGroup('grp_trabajadores_admin')"><i class='bi bi-people'></i><span class='label'>Trabajadores / Usuarios y claves</span><span class='chev'>∨</span></button>
              <div class='submenu'>
                <a class='{cls_trab}' onclick='saveSideScroll()' href='/admin/trabajadores'><i class='bi bi-people'></i><span class='label'>Trabajadores</span></a>
                <a class='{cls_users}' onclick='saveSideScroll()' href='/admin/usuarios'><i class='bi bi-lock'></i><span class='label'>Usuarios y claves</span></a>
              </div>
            </div>
            <a class='menu-item {'active' if active == 'IA HR' else ''}' onclick='saveSideScroll()' href='/admin/ia_hr'><i class='bi bi-robot'></i><span class='label'>Base IA HR</span></a>
            <a class='{cls_test}' onclick='saveSideScroll()' href='/admin/modo_prueba'><i class='bi bi-magic'></i><span class='label'>Modo prueba y limpieza</span></a>
          </div>
        </div>"""
    user_docs_keys = [k for k,_,_ in TIPOS_PAGO] + [k for k,_,_ in TIPOS_EMPRESA] + [k for k,_,_ in TIPOS_PERSONALES]
    user_docs_cls = 'menu-group force-open' if active_type in user_docs_keys else 'menu-group'
    documentos_generales = '' if session.get('admin_id') else f"""
      <div id='grp_user_documental' data-group='user_documental' class='{user_docs_cls}'>
        <button type='button' class='menu-title' onclick="toggleGroup('grp_user_documental')"><i class='bi bi-folder2-open'></i><span class='label'>Gestión Documental</span><span class='chev'>∨</span></button>
        <div class='submenu'><a class='menu-item' onclick='saveSideScroll()' href='/panel'><i class='bi bi-speedometer2'></i><span class='label'>Dashboard documental</span></a>
          <div id='grp_pago' data-group='pago' class='{pago_cls}'>
            <button type='button' class='menu-title' onclick="toggleGroup('grp_pago')"><i class='bi bi-file-earmark-text'></i><span class='label'>Documentos de pago</span><span class='chev'>∨</span></button>
            <div class='submenu'>{pago}</div>
          </div>
          <div id='grp_empresa' data-group='empresa' class='{emp_cls}'>
            <button type='button' class='menu-title' onclick="toggleGroup('grp_empresa')"><i class='bi bi-building'></i><span class='label'>Documentos de la empresa</span><span class='chev'>∨</span></button>
            <div class='submenu'>{emp}</div>
          </div>
          <div id='grp_personal' data-group='personal' class='{per_cls}'>
            <button type='button' class='menu-title' onclick="toggleGroup('grp_personal')"><i class='bi bi-person-lines-fill'></i><span class='label'>Documentos personales</span><span class='chev'>∨</span></button>
            <div class='submenu'>{per}</div>
          </div>
        </div>
      </div>"""
    user_gestiones = '' if session.get('admin_id') else f"""
      <div id='grp_user_vacacional' data-group='user_vacacional' class='menu-group {'force-open' if active == 'Gestion Vacacional' else ''}'>
        <button type='button' class='menu-title {'active' if active == 'Gestion Vacacional' else ''}' onclick="toggleGroup('grp_user_vacacional')"><i class='bi bi-calendar2-week'></i><span class='label'>Gestión Vacacional</span><span class='chev'>∨</span></button>
        <div class='submenu'><a class='menu-item {'active' if active == 'Gestion Vacacional' else ''}' onclick='saveSideScroll()' href='/vacaciones/mi_solicitud'><i class='bi bi-bar-chart-line'></i><span class='label'>Dashboard vacacional</span></a><a class='menu-item {'active' if active == 'Gestion Vacacional' else ''}' onclick='saveSideScroll()' href='/vacaciones/mi_solicitud#solicitar'><i class='bi bi-calendar-check'></i><span class='label'>Saldo y solicitud</span></a><a class='menu-item {'active' if active == 'Gestion Vacacional' else ''}' onclick='saveSideScroll()' href='/vacaciones/aprobaciones_jefe'><i class='bi bi-check2-circle'></i><span class='label'>Aprobaciones jefe</span></a></div>
      </div>
"""
    return f"""
    <nav>
      {documentos_generales}
      {user_gestiones}
      {admin}
      <div id='grp_cuenta' data-group='cuenta' class='menu-group {'force-open' if active == 'Inicio' else ''}'>
        <button type='button' class='menu-title {'active' if active == 'Inicio' else ''}' onclick="toggleGroup('grp_cuenta')"><i class='bi bi-person-circle'></i><span class='label'>Mi cuenta</span><span class='chev'>∨</span></button>
        <div class='submenu'><a class='menu-item {'active' if active == 'Inicio' else ''}' onclick='saveSideScroll()' href='/panel'><i class='bi bi-house'></i><span class='label'>Inicio</span></a><a class='menu-item' href='/logout'><i class='bi bi-box-arrow-right'></i><span class='label'>Salir</span></a></div>
      </div>
    </nav>"""


def empresas_disponibles_login():
    """Empresas visibles en el login, tomadas de la columna EMPRESA de trabajadores activos."""
    empresas = []
    try:
        with db() as con:
            rows = con.execute("SELECT DISTINCT empresa FROM trabajadores WHERE activo=1 AND COALESCE(TRIM(empresa),'')<>'' ORDER BY empresa").fetchall()
        for r in rows:
            for raw in clean(r['empresa']).replace('|','/').replace(';','/').replace(',','/').split('/'):
                e = clean(raw)
                if not e:
                    continue
                if e.upper() == 'PORTAL HR PRO SUPERFRUITS':
                    e = 'AQUANQA'
                if e not in empresas:
                    empresas.append(e)
    except Exception:
        pass
    return empresas or ['AQUANQA']

def login_template(admin=False, error=""):
    action = url_for('admin_login') if admin else url_for('login')
    title = "PORTAL HR ADMIN" if admin else "PORTAL HR"
    sub = "Panel de administración" if admin else "Actualización de datos de trabajadores"
    if admin:
        # Entrada Administrador: solo credenciales de administrador.
        # La empresa se guarda como valor interno para mantener compatibilidad con session['empresa_login'].
        fields = f"""
          <input type='hidden' name='empresa' value='PORTAL HR PRO'>
          <div class='login-mode-pill admin-pill'><i class='bi bi-shield-lock'></i> Entrada administrador</div>
          <div class='field'><label>Usuario</label><div class='login-input'><i class='bi bi-person'></i><input name='usuario' placeholder='admin' required></div></div>
          <div class='field'><label>Clave</label><div class='login-input'><i class='bi bi-lock'></i><input name='clave' type='password' placeholder='••••••••' required></div></div>
        """
    else:
        opts = ''.join([f"<option value='{e}'>{e}</option>" for e in empresas_disponibles_login()])
        fields = f"""
          <div class='field'><label>Empresa</label><div class='login-input'><i class='bi bi-building'></i><select name='empresa' required><option value=''>Seleccione empresa</option>{opts}</select></div></div>
          <div class='field'><label>DNI</label><div class='login-input'><i class='bi bi-person-badge'></i><input name='dni' maxlength='8' placeholder='Ingrese su DNI' required></div></div>
          <div class='field'><label>Clave</label><div class='login-input'><i class='bi bi-lock'></i><input name='correo' type='password' placeholder='Fecha nacimiento: ddmmaaaa' required></div></div>
        """
    body = f"""
    <div class='login-body'><form class='login-card' method='post' action='{action}'><div class='login-inner'>
      <div class='login-logo'><div class='login-avatar-svg' aria-label='Portal HR Pro'></div></div><div class='login-title'><h1>{title}</h1><b>{sub}</b></div>
      {f"<div class='flash err'>{error}</div>" if error else ""}{fields}<button class='btn-green'>Ingresar</button>
    </div><div class='login-links'>{'<a class="login-switch worker" href="/"><i class="bi bi-person"></i> Entrada trabajador</a>' if admin else '<a class="login-switch admin" href="/admin/login"><i class="bi bi-shield-lock"></i> Entrada administrador</a>'}</div></form></div>"""
    return render_template_string(BASE, body=body, title=title)


# =============================
# ROUTES ESTÁTICAS / LOGO
# =============================
@app.route('/_logo/<path:filename>')
def logo_file(filename):
    p = BASE_DIR / filename
    if p.exists():
        return send_file(p)
    abort(404)

@app.route('/logo_svg')
def logo_svg():
    # Logo genérico PORTAL HR PRO: sin texto fijo de empresa.
    svg = """<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 360 140'><defs><linearGradient id='g' x1='0' x2='1'><stop offset='0' stop-color='#22c55e'/><stop offset='1' stop-color='#059669'/></linearGradient></defs><rect width='360' height='140' rx='26' fill='white'/><circle cx='63' cy='68' r='34' fill='url(#g)'/><text x='105' y='78' font-family='Segoe UI,Arial' font-size='58' font-style='italic' font-weight='900' fill='#111827'>PORTAL HR PRO</text><text x='112' y='112' font-family='Arial' font-size='24' font-weight='900' fill='#2b668d'>RRHH</text></svg>"""
    return app.response_class(svg, mimetype='image/svg+xml')

# =============================
# LOGIN USUARIO
# =============================
@app.route('/', methods=['GET','POST'])
def login():
    # Pantalla principal tipo imagen de referencia: acceso administrador limpio.
    if request.method == 'POST':
        u = clean(request.form.get('usuario'))
        c = clean(request.form.get('clave'))
        empresa_login = clean(request.form.get('empresa')) or 'PORTAL HR PRO'
        if u == ADMIN_USER and c == ADMIN_PASS:
            session.clear(); session['admin_id']='admin'; session['admin_nombre']='Administrador'; session['empresa_login']=empresa_login
            return redirect(url_for('admin'))
        return login_template(True, 'Usuario o clave incorrecta.')
    return login_template(True)

@app.route('/logout')
def logout():
    session.clear(); return redirect(url_for('login'))


@app.route('/seleccionar_empresa', methods=['GET','POST'])
@worker_required
def seleccionar_empresa():
    dni=session['dni']; t=get_trabajador(dni)
    emp_real = clean(t['empresa']) if t and 'empresa' in t.keys() else ''
    # Solo mostrar la(s) empresa(s) que vienen desde la columna EMPRESA del trabajador.
    # Se elimina cualquier opción fija como PORTAL HR PRO SUPERFRUITS.
    empresas=[]
    for raw in emp_real.replace('|','/').replace(';','/').replace(',','/').split('/'):
        e=clean(raw)
        if not e: continue
        # corrección defensiva para bases antiguas/demostración
        if e.upper() == 'PORTAL HR PRO SUPERFRUITS':
            e = 'AQUANQA'
        if e not in empresas:
            empresas.append(e)
    if not empresas:
        empresas=['AQUANQA']
    if request.method=='POST':
        emp=clean(request.form.get('empresa'))
        if emp not in empresas: emp=empresas[0]
        session['empresa']=emp
        flash(f'Empresa seleccionada: {emp}', 'ok')
        return redirect(url_for('panel'))
    opts=''.join([f"<option value='{e}'>{e}</option>" for e in empresas])
    content=f"""<div class='login-body'><form class='login-card' method='post'><div class='login-inner'>
      <div class='login-logo'><img src='{logo_url()}'></div><div class='login-title'><h1>Elegir empresa</h1><b>{t['nombre'] if t else dni}</b></div>
      <div class='field'><label>Empresa asignada</label><div class='login-input'>🏢<select name='empresa' style='width:100%;background:transparent;color:#fff;border:0;padding:15px;font-weight:900'>{opts}</select></div></div>
      <button class='btn-green'>Ingresar al portal</button></div></form></div>"""
    return render_template_string(BASE, body=content, title='Elegir empresa')

@app.route('/panel')
@worker_required
def panel():
    dni = session['dni']; sincronizar_documentos_carpeta(dni); t = get_trabajador(dni)
    if not t:
        flash('No se encontró tu trabajador activo. Vuelve a iniciar sesión o contacta a RRHH.', 'err')
        return redirect(url_for('logout'))
    docs = listar_documentos(dni=dni, limit=999)
    ultimo = docs[0]['tipo'] if docs else 'Sin documento'
    cards = ''.join(doc_card(k,l,i) for k,l,i in (TIPOS_PAGO+TIPOS_EMPRESA+TIPOS_PERSONALES))
    dashboard_gestiones = """
      <div class='card span-12'><h2>Dashboards de gestión</h2><div class='module-tabs'>
        <a class='module-tile' href='/panel'><h2>📁 Gestión Documental</h2><p class='muted'>Documentos, pagos, empresa y personales.</p></a>
        <a class='module-tile' href='/vacaciones/mi_solicitud'><h2>🏖️ Gestión Vacacional</h2><p class='muted'>Saldo, solicitud y seguimiento de aprobaciones.</p></a>
      </div></div>
    """
    content = f"""
    <div class='hero'><div class='topbar'><div><h1>Portal Documental <span class='accent'>PORTAL HR PRO</span></h1><div class='subtitle'>{t['nombre']} · DNI {t['dni']} · {session.get('empresa') or t['empresa']}</div></div><div style='display:flex;gap:10px;align-items:center'><span class='btn'>● Activo</span><a class='btn-blue' href='/panel'>Ver todo</a></div></div></div>
    <section class='grid'><div class='card mini'><div><span>Documentos</span><br><b>{len(docs)}</b></div><div class='ico'>🗂️</div></div><div class='card mini'><div><span>Último tipo</span><br><b>{ultimo}</b></div><div class='ico'>📄</div></div><div class='card mini'><div><span>Estado</span><br><b>Activo</b></div><div class='ico'>✅</div></div>{dashboard_gestiones}<div class='card span-12 profile-card'><div><h2>Mi perfil y foto</h2><p class='muted'>Actualiza tu foto para que el portal quede como panel profesional.</p></div><div class='profile-row'><img class='profile-img' src='{url_for('foto_trabajador', dni=dni) if t['foto_ruta'] else logo_url()}'><form method='post' action='/mi_foto' enctype='multipart/form-data' class='form-grid profile-form'><div class='field'><label>Foto personal</label><input type='file' name='foto' accept='.png,.jpg,.jpeg,.webp' required></div><button class='btn-green'>Cargar foto</button></form></div></div>
    <div class='card span-12'><div style='display:flex;justify-content:space-between;gap:12px;align-items:center;flex-wrap:wrap'><h2>Accesos por pestaña</h2><input id='cardSearch' onkeyup='filterCards()' class='input' style='max-width:310px' placeholder='Buscar pestaña...'></div><div class='doc-grid'>{cards}</div></div>
    <div class='card span-12'><h2>🔔 Notificaciones</h2>{notificaciones_trabajador(dni)}</div><div class='card span-12'><h2>Últimos documentos</h2>{tabla_docs(docs)}</div></section>"""
    return render_page(content, active='Inicio')


def notificaciones_trabajador(dni):
    with db() as con:
        rows = con.execute("SELECT evento,fecha,detalle FROM eventos_documento WHERE dni=? ORDER BY id DESC LIMIT 10", (normalizar_dni(dni),)).fetchall()
    if not rows:
        return "<p class='muted'>Sin notificaciones por ahora.</p>"
    return "".join([f"<div class='alert-item'><div class='bell'>🔔</div><div><b>{r['evento']}</b><br><span>{r['fecha']} · {r['detalle'] or ''}</span></div></div>" for r in rows])

def doc_card(k,l,i):
    return f"<div class='doc-card'><h3>{i} {l}</h3><p>Consulta, filtra por periodo y revisa el detalle del documento.</p><a class='btn-blue' href='{url_for('panel_tipo', tipo=k)}'>Abrir</a></div>"

@app.route('/documentos/<tipo>')
@portal_required
def panel_tipo(tipo):
    if tipo not in ALL_TIPOS: abort(404)
    asegurar_carpetas_documentales(tipo)
    is_admin = bool(session.get('admin_id'))
    dni = session.get('dni', '')
    if is_admin:
        sincronizar_documentos_carpeta()
    else:
        sincronizar_documentos_carpeta(dni)
    label, icon, categoria = ALL_TIPOS[tipo]
    periodo = clean(request.args.get('periodo'))
    sub = clean(request.args.get('sub'))
    pers = periodos_disponibles(dni=None if is_admin else dni, tipo=tipo) if is_admin else periodos_desde_ingreso(dni, tipo)
    docs = listar_documentos(dni=None if is_admin else dni, tipo=tipo, periodo=periodo or None, limit=999)
    if tipo == 'Normal' and sub:
        docs = [d for d in docs if sub.lower() in clean(d['detalle']).lower()]
    opts = "<option value=''>Todos los periodos</option>" + ''.join([f"<option {'selected' if p==periodo else ''}>{p}</option>" for p in pers])
    detalle = detalle_tipo(tipo, docs)
    upload_extra = ""
    if tipo in ['Otros','Contrato Personal'] and not is_admin:
        upload_extra = f"""
        <div class='card span-12'><h2>Adjuntar nuevo documento personal</h2><form method='post' action='/subir_personal' enctype='multipart/form-data' class='form-grid'>
        <input type='hidden' name='tipo' value='{tipo}'><div class='field'><label>Periodo</label><input name='periodo' value='{datetime.now(APP_TZ).strftime('%Y-%m')}'></div><div class='field'><label>Detalle</label><input name='detalle' placeholder='Ej: Certificado, solicitud, evidencia'></div><div class='field'><label>Archivo</label><input type='file' name='archivo' accept='.pdf,.png,.jpg,.jpeg,.webp,.doc,.docx,.xls,.xlsx' required></div><div class='field'><label>Observación</label><textarea name='observacion' rows='2' placeholder='Comentario u observación'></textarea></div><button class='btn-green'>Subir documento</button></form></div>"""
    content = f"""
    <div class='hero'><div class='topbar'><div><h1>{icon} {label}</h1><div class='subtitle'>Consulta filtrada por pestaña y periodo seleccionado.</div></div><a class='btn-blue' href='{url_for('admin') if is_admin else url_for('panel')}'>Volver</a></div></div>
    <section class='grid'><div class='card mini'><div><span>Total</span><br><b>{len(docs)}</b></div><div class='ico'>{icon}</div></div><div class='card mini'><div><span>Periodo</span><br><b>{periodo or 'Todos'}</b></div><div class='ico'>📅</div></div><div class='card mini'><div><span>Filtro</span><br><b>{tipo}{' - '+sub if sub else ''}</b></div><div class='ico'>🔎</div></div>
    <div class='card span-12'><form method='get' class='period-row'><div class='field'><label>Elegir periodo</label><select name='periodo'>{opts}</select></div><button class='btn-blue'>Aplicar filtro</button><a class='btn' href='{url_for('panel_tipo', tipo=tipo)}'>Limpiar</a></form></div>
    <div class='card span-12'><h2>Detalle de {label}</h2>{detalle}</div>{upload_extra}<div class='card span-12'><h2>Listado</h2>{tabla_docs(docs)}</div></section>"""
    return render_page(content, active=(f'{tipo}:{sub}' if tipo == 'Normal' and sub else tipo))

@app.route('/subir_personal', methods=['POST'])
@worker_required
def subir_personal():
    try:
        guardar_documento(request.files.get('archivo'), session['dni'], clean(request.form.get('tipo')) or 'Otros', request.form.get('periodo'), request.form.get('detalle'), request.form.get('observacion'), session.get('dni'))
        flash('Documento personal subido correctamente.', 'ok')
    except Exception as e:
        flash(f'No se pudo subir: {e}', 'error')
    return redirect(url_for('panel_tipo', tipo=clean(request.form.get('tipo')) or 'Otros'))


def detalle_tipo(tipo, docs):
    ult = docs[0] if docs else None
    label, icon, cat = ALL_TIPOS.get(tipo, (tipo,'📄',''))
    texto = {
        'Utilidad':'Boletas de participación de utilidades por periodo anual.',
        'Vacaciones':'Documentos relacionados a pago o liquidación de vacaciones.',
        'Normal':'Boletas de pago normal mensual, quincenal o semanal.',
        'Constancia Gratificación':'Constancias asociadas a gratificación.',
        'CTS':'Boletas o constancias de Compensación por Tiempo de Servicios.',
        'Liquidación':'Documentos de liquidación de beneficios sociales.',
        'Gratificación':'Boletas de gratificación ordinaria o extraordinaria.',
        'Otros':'Espacio para documentos personales adjuntos por el usuario o administrador.',
    }.get(tipo, 'Documento disponible para consulta y descarga.')
    return f"""
    <div class='grid'><div class='detail-box span-4'><small>Tipo</small><b>{icon} {label}</b></div><div class='detail-box span-4'><small>Último periodo</small><b>{ult['periodo'] if ult else 'Sin periodo'}</b></div><div class='detail-box span-4'><small>Última carga</small><b>{ult['fecha_subida'] if ult else 'Sin carga'}</b></div><div class='detail-box span-12'><small>Descripción</small>{texto}</div></div>"""


def tabla_docs(rows):
    headers = "<tr><th>Tipo</th><th>DNI</th><th>Trabajador</th><th>Periodo</th><th>Detalle</th><th>Observación</th><th>Estado</th><th>Cargado por</th><th>Fecha</th><th>Archivo</th><th>Acciones</th></tr>"
    if not rows:
        return f"<div class='table-wrap'><table>{headers}<tr><td colspan='11'>No hay documentos en esta pestaña.</td></tr></table></div>"
    body = ''
    is_admin = bool(session.get('admin_id'))
    dni_sess = session.get('dni')
    with db() as con:
        nombres = {r['dni']: r['nombre'] for r in con.execute("SELECT dni,nombre FROM trabajadores").fetchall()}
    for r in rows:
        rid = r['id']; estado = r['estado'] if 'estado' in r.keys() and r['estado'] else 'Pendiente'
        row_cls = 'row-approved' if estado == 'Aprobado' else ('row-rejected' if estado == 'Rechazado' else '')
        pill_cls = 'status-pill ' + ('st-aprobado' if estado == 'Aprobado' else 'st-rechazado' if estado == 'Rechazado' else 'st-firmado' if estado == 'Firmado' else 'st-aceptado' if estado == 'Aceptado' else '')
        ver = f"<a class='btn-blue' target='_blank' href='{url_for('ver_doc', doc_id=rid)}'>Ver/Descargar</a>"
        acciones = []
        # Trabajador: puede aceptar/rechazar/firmar, y solo eliminar documentos personales propios.
        if r['categoria'] in ['pago','empresa','personal'] and dni_sess and (r['dni'] == dni_sess or r['categoria']=='empresa'):
            if estado not in ['Aceptado','Firmado','Aprobado','Rechazado']:
                acciones.append(f"<a class='btn-green mini-btn' href='{url_for('flujo_doc', doc_id=rid, accion='aceptar')}'>Aceptar</a>")
                acciones.append(f"<button class='btn-danger mini-btn' onclick=\"showReject({rid})\">Rechazar</button>")
            if estado in ['Aceptado','Firmado']:
                acciones.append(f"<a class='btn-blue mini-btn' href='{url_for('flujo_doc', doc_id=rid, accion='firmar')}'>Firmar</a>")
        # Administrador: aprueba o rechaza. Eliminar queda disponible solo para admin o personal propio del trabajador.
        if is_admin and r['categoria'] in ['pago','personal','empresa']:
            if estado != 'Aprobado':
                acciones.append(f"<a class='btn-warn mini-btn' href='{url_for('flujo_doc', doc_id=rid, accion='aprobar')}'>Aprobar</a>")
            if estado != 'Rechazado':
                acciones.append(f"<button class='btn-danger mini-btn' onclick=\"showReject({rid})\">Rechazar</button>")
        if is_admin or (dni_sess and r['dni'] == dni_sess and r['categoria'] == 'personal'):
            acciones.append(f"<a class='btn-red mini-btn' onclick=\"return confirm('¿Eliminar este documento?')\" href='{url_for('eliminar_doc', doc_id=rid)}'>Eliminar</a>")
        dni_val = r['dni'] or 'EMPRESA'
        trabajador = nombres.get(r['dni'], 'Documento general' if r['categoria']=='empresa' else '-')
        cargado_por = r['uploaded_by'] or 'sistema'
        if cargado_por == 'auto': cargado_por = 'Carpeta automática'
        body += f"<tr class='{row_cls}'><td>{r['tipo']}</td><td>{dni_val}</td><td>{trabajador}</td><td>{r['periodo'] or ''}</td><td>{r['detalle'] or '-'}</td><td>{r['observacion'] or '-'}</td><td><span class='{pill_cls}'>{estado}</span></td><td><b>{cargado_por}</b></td><td>{r['fecha_subida']}</td><td>{ver}</td><td><div class='actions'>{''.join(acciones) or '-'}</div></td></tr>"
    modal = """<div id='rejectBox' class='modal hidden'><form method='post' id='rejectForm' class='modal-card'><h2>Rechazar documento</h2><label>Comentario obligatorio</label><textarea name='comentario' required rows='4' placeholder='Indique el motivo del rechazo'></textarea><div class='actions'><button class='btn-red'>Rechazar</button><button type='button' class='btn' onclick='hideReject()'>Cancelar</button></div></form></div><script>function showReject(id){let m=document.getElementById('rejectBox'),f=document.getElementById('rejectForm');f.action='/documento/'+id+'/rechazar';m.classList.remove('hidden')}function hideReject(){document.getElementById('rejectBox').classList.add('hidden')}</script>"""
    return f"<div class='table-wrap'><table><thead>{headers}</thead><tbody>{body}</tbody></table></div>{modal}"

@app.route('/documento/<int:doc_id>/eliminar')
def eliminar_doc(doc_id):
    with db() as con:
        r = con.execute("SELECT * FROM documentos WHERE id=?", (doc_id,)).fetchone()
        if not r: abort(404)
        dni_sess = session.get('dni')
        permitido = bool(session.get('admin_id')) or (dni_sess and r['dni'] == dni_sess and r['categoria'] == 'personal')
        if not permitido: abort(403)
        try:
            path = Path(r['ruta_archivo'])
            if path.exists(): path.unlink()
        except Exception:
            pass
        con.execute("DELETE FROM documentos WHERE id=?", (doc_id,))
        con.commit()
    flash('Documento eliminado correctamente.', 'ok')
    return redirect(request.referrer or url_for('panel'))

@app.route('/documento/<int:doc_id>/<accion>')
def flujo_doc(doc_id, accion):
    if accion not in ['aceptar','firmar','aprobar']: abort(404)
    with db() as con:
        r = con.execute("SELECT * FROM documentos WHERE id=?", (doc_id,)).fetchone()
        if not r: abort(404)
        dni_sess = session.get('dni')
        if accion in ['aceptar','firmar'] and (not dni_sess or (r['categoria']!='empresa' and r['dni'] != dni_sess)):
            abort(403)
        if accion == 'aprobar' and not session.get('admin_id'):
            abort(403)
        if accion == 'aceptar':
            con.execute("UPDATE documentos SET estado='Aceptado', fecha_aceptacion=? WHERE id=?", (now_txt(), doc_id))
            flash('Documento aceptado. Ahora puede firmarlo.', 'ok')
        elif accion == 'firmar':
            con.execute("UPDATE documentos SET estado='Firmado', fecha_firma=? WHERE id=?", (now_txt(), doc_id))
            flash('Documento firmado correctamente.', 'ok')
        elif accion == 'aprobar':
            con.execute("UPDATE documentos SET estado='Aprobado', fecha_aprobacion=? WHERE id=?", (now_txt(), doc_id))
            con.execute("INSERT INTO eventos_documento(documento_id,dni,evento,fecha,detalle) VALUES(?,?,?,?,?)", (doc_id, r['dni'] or '', 'Aprobado', now_txt(), 'Aprobado por administrador'))
            flash('Documento aprobado por administrador.', 'ok')
        con.commit()
    return redirect(request.referrer or url_for('panel'))

@app.route('/documento/<int:doc_id>/rechazar', methods=['POST'])
def rechazar_doc(doc_id):
    comentario = clean(request.form.get('comentario'))
    with db() as con:
        r = con.execute("SELECT * FROM documentos WHERE id=?", (doc_id,)).fetchone()
        if not r: abort(404)
        dni_sess = session.get('dni')
        is_admin = bool(session.get('admin_id'))
        if not is_admin and (not dni_sess or (r['categoria']!='empresa' and r['dni'] != dni_sess)): abort(403)
        con.execute("UPDATE documentos SET estado='Rechazado', comentario_rechazo=?, observacion=? WHERE id=?", (comentario, comentario, doc_id))
        con.execute("INSERT INTO eventos_documento(documento_id,dni,evento,fecha,detalle) VALUES(?,?,?,?,?)", (doc_id, r['dni'] or dni_sess or '', 'Rechazado', now_txt(), comentario))
        con.commit()
    flash('Documento rechazado. Se registró notificación para el trabajador.', 'ok')
    return redirect(request.referrer or url_for('panel'))

@app.route('/ver/<int:doc_id>')
def ver_doc(doc_id):
    with db() as con:
        r = con.execute("SELECT * FROM documentos WHERE id=?", (doc_id,)).fetchone()
    if not r: abort(404)
    if not session.get('admin_id'):
        dni = session.get('dni')
        if not dni or (r['categoria'] != 'empresa' and r['dni'] != dni): abort(403)
    path = Path(r['ruta_archivo'])
    if not path.exists(): abort(404)
    if session.get('dni'):
        try: registrar_evento_documento(doc_id, session.get('dni'), 'Abierto', 'Trabajador abrió/descargó el documento')
        except Exception: pass
    return send_file(path, as_attachment=False, download_name=r['archivo_nombre'])



def alertas_admin(limit=8):
    with db() as con:
        return con.execute("""
            SELECT d.*, t.nombre AS trabajador
            FROM documentos d
            LEFT JOIN trabajadores t ON t.dni=d.dni
            ORDER BY d.id DESC
            LIMIT ?
        """, (limit,)).fetchall()

# =============================
# ADMIN
# =============================
@app.route('/admin/login', methods=['GET','POST'])
def admin_login():
    if request.method == 'POST':
        u, c = clean(request.form.get('usuario')), clean(request.form.get('clave'))
        with db() as con:
            user = con.execute("SELECT * FROM usuarios_admin WHERE usuario=? AND activo=1", (u,)).fetchone()
        if not user or not check_password_hash(user['clave_hash'], c):
            return login_template(True, 'Usuario o clave incorrecta.')
        empresa_login = clean(request.form.get('empresa')) or 'PORTAL HR PRO'
        session.clear(); session['admin_id']=user['id']; session['admin_user']=user['usuario']; session['admin_nombre']=user['nombre']; session['empresa_login']=empresa_login
        return redirect(url_for('admin'))
    return login_template(True)

@app.route('/admin')
@admin_required
def admin():
    sincronizar_documentos_carpeta()
    desde = clean(request.args.get('desde'))
    hasta = clean(request.args.get('hasta'))
    with db() as con:
        trabajadores = con.execute("SELECT COUNT(*) FROM trabajadores").fetchone()[0]
        docs = con.execute("SELECT COUNT(*) FROM documentos").fetchone()[0]
        emp = con.execute("SELECT COUNT(*) FROM documentos WHERE categoria='empresa'").fetchone()[0]
        leidos = con.execute("SELECT COUNT(*) FROM documentos WHERE fecha_lectura IS NOT NULL AND fecha_lectura<>''").fetchone()[0]
        aprobados = con.execute("SELECT COUNT(*) FROM documentos WHERE estado='Aprobado'").fetchone()[0]
        rechazados = con.execute("SELECT COUNT(*) FROM documentos WHERE estado='Rechazado'").fetchone()[0]
        ult = con.execute("SELECT * FROM documentos ORDER BY id DESC LIMIT 12").fetchall()
    alerts = alertas_admin(8)
    with db() as con:
        chart_rows = con.execute("SELECT tipo, COUNT(*) c FROM documentos GROUP BY tipo ORDER BY c DESC LIMIT 8").fetchall()
        fechas_docs = con.execute("SELECT fecha_subida FROM documentos").fetchall()
    hoy_dt = datetime.now(APP_TZ).date()
    doc_dia = doc_semana = doc_mes = doc_rango = 0
    desde_dt = parse_fecha_any(desde) if desde else None
    hasta_dt = parse_fecha_any(hasta) if hasta else None
    for rr in fechas_docs:
        try:
            dd = datetime.strptime((rr['fecha_subida'] or '')[:10], '%d/%m/%Y').date()
            if dd == hoy_dt: doc_dia += 1
            if (hoy_dt - dd).days <= 7: doc_semana += 1
            if dd.year == hoy_dt.year and dd.month == hoy_dt.month: doc_mes += 1
            if (not desde_dt or dd >= desde_dt) and (not hasta_dt or dd <= hasta_dt): doc_rango += 1
        except Exception:
            pass
    maxc = max([x['c'] for x in chart_rows] or [1])
    chart_html = ''.join([f"<div class='bar-row'><b>{x['tipo']}</b><span><i style='width:{max(6, int(x['c']*100/maxc))}%'></i></span><em>{x['c']}</em></div>" for x in chart_rows]) or "<p class='muted'>Sin información para graficar.</p>"
    alert_items = ''.join([f"<div class='alert-item'><div class='bell'>🔔</div><div><b>{(a['trabajador'] or a['dni'] or 'Documento empresa')}</b><br><span>{a['tipo']} · {a['periodo'] or 'Sin periodo'} · {a['fecha_subida']} · Cargado por: {a['uploaded_by'] or 'sistema'}</span></div><a class='btn-blue mini-btn' target='_blank' href='{url_for('ver_doc', doc_id=a['id'])}'>Ver</a></div>" for a in alerts]) or "<div class='empty-note'>Aún no hay documentos cargados.</div>"
    with db() as con:
        ind_rows = con.execute("""
            SELECT tipo,
                   COUNT(*) total,
                   SUM(CASE WHEN estado='Aprobado' THEN 1 ELSE 0 END) aprobados,
                   SUM(CASE WHEN estado='Rechazado' THEN 1 ELSE 0 END) rechazados,
                   SUM(CASE WHEN fecha_lectura IS NOT NULL AND fecha_lectura<>'' THEN 1 ELSE 0 END) leidos
            FROM documentos
            GROUP BY tipo
            ORDER BY tipo
        """).fetchall()
    ind_html = ''.join([f"<tr><td>{r['tipo']}</td><td>{r['total']}</td><td><span class='status-pill st-aprobado'>{r['aprobados']}</span></td><td><span class='status-pill st-rechazado'>{r['rechazados']}</span></td><td><span class='status-pill'>{r['leidos']}</span></td></tr>" for r in ind_rows]) or "<tr><td colspan='5'>Sin documentos.</td></tr>"
    modo_txt = 'ACTIVO' if modo_prueba_activo() else 'INACTIVO'
    with db() as con:
        vac_saldos = con.execute("SELECT COUNT(*) FROM vacaciones_saldos").fetchone()[0]
        vac_solicitudes = con.execute("SELECT COUNT(*) FROM vacaciones_solicitudes").fetchone()[0]
        vac_pendientes = con.execute("SELECT COUNT(*) FROM vacaciones_solicitudes WHERE estado LIKE 'Pendiente%'").fetchone()[0]
        vac_aprobadas = con.execute("SELECT COUNT(*) FROM vacaciones_solicitudes WHERE estado LIKE 'Aprobado%'").fetchone()[0]
    content = f"""
    <div class='admin-shell'>
      <div class='admin-header'>
        <div class='admin-title-row'>
          <button class='hambox' onclick='toggleSide()'>☰</button>
          <div class='admin-title'>
            <h1>Panel de Administración</h1>
            <div class='role'>Administrador</div>
            <p>Bienvenido al panel de administración. Seleccione una gestión para comenzar.</p>
          </div>
        </div>
        <div class='top-actions'>
          <button class='top-icon notif-bell' title='Notificaciones' onclick="alert('No tienes notificaciones pendientes')"><span>🔔</span><i>0</i></button>
          <div class='admin-chip admin-dropdown'>
            <button type='button' class='admin-chip-btn' onclick="this.parentElement.classList.toggle('open')">
              <span class='a'>A</span><span>Administrador⌄</span>
            </button>
            <div class='admin-menu'>
              <a href='/admin'><i class='bi bi-house'></i> Inicio</a>
              <a href='/admin/usuarios'><i class='bi bi-person-circle'></i> Mi cuenta</a>
              <a href='/logout'><i class='bi bi-box-arrow-right'></i> Salir</a>
            </div>
          </div>
        </div>
      </div>

      <div class='gestion-cards'>
        <div class='card gestion-card'>
          <div class='gestion-icon'>📁</div>
          <div><h2>Gestión Documental</h2><p class='muted'>Administre y controle todos los documentos de la organización.</p><a class='btn-warn' href='/admin/modulo/documentos'>Ir al Dashboard <span>→</span></a></div>
        </div>
        <div class='card gestion-card green'>
          <div class='gestion-icon'>☂️</div>
          <div><h2>Gestión Vacacional</h2><p class='muted'>Administre saldos y solicitudes de vacaciones de los trabajadores.</p><a class='btn-green' href='/admin/vacaciones'>Ir al Dashboard <span>→</span></a></div>
        </div>
      </div>

      <div class='dashboards-admin'>
        <div class='card dashboard-panel'>
          <h2>📁 Dashboard - Gestión Documental</h2>
          <div class='mini-grid'>
            <div class='dash-metric'><span>Trabajadores</span><b>{trabajadores}</b><em class='mi'>👥</em></div>
            <div class='dash-metric'><span>Documentos</span><b>{docs}</b><em class='mi'>📄</em></div>
            <div class='dash-metric'><span>Recibidos / Abiertos</span><b>{leidos}</b><em class='mi'>👁️</em></div>
            <div class='dash-metric'><span>Aprobados</span><b>{aprobados}</b><em class='mi'>✓</em></div>
            <div class='dash-metric'><span>Rechazados</span><b>{rechazados}</b><em class='mi'>−</em></div>
            <div class='dash-metric'><span>Empresas</span><b>{emp}</b><em class='mi'>🏢</em></div>
          </div>
          <a class='btn-warn full-link' href='/admin/modulo/documentos'>Ver Dashboard Completo <span>→</span></a>
        </div>
        <div class='card dashboard-panel green'>
          <h2>☂️ Dashboard - Gestión Vacacional</h2>
          <div class='mini-grid'>
            <div class='dash-metric'><span>Saldos Registrados</span><b>{vac_saldos}</b><em class='mi'>🗓️</em></div>
            <div class='dash-metric'><span>Solicitudes</span><b>{vac_solicitudes}</b><em class='mi'>📄</em></div>
            <div class='dash-metric'><span>Pendientes</span><b>{vac_pendientes}</b><em class='mi'>⏱️</em></div>
            <div class='dash-metric'><span>Aprobadas</span><b>{vac_aprobadas}</b><em class='mi'>✓</em></div>
            <div class='dash-metric'><span>Rechazadas</span><b>0</b><em class='mi'>−</em></div>
            <div class='dash-metric'><span>En Proceso</span><b>{vac_pendientes}</b><em class='mi'>…</em></div>
          </div>
          <a class='btn-green full-link' href='/admin/vacaciones'>Ver Dashboard Completo <span>→</span></a>
        </div>
      </div>

      <div class='admin-footer'><span>© 2026 PORTAL HR PRO</span><span>Versión 1.0.0</span></div>
    </div>
    """
    return render_page(content, active='Admin')


def normalizar_header_excel(valor):
    """Normaliza encabezados de Excel para que el usuario pueda cargar plantillas amplias sin errores."""
    txt = clean(valor).upper()
    txt = re.sub(r"[ÁÀÂÄ]", "A", txt)
    txt = re.sub(r"[ÉÈÊË]", "E", txt)
    txt = re.sub(r"[ÍÌÎÏ]", "I", txt)
    txt = re.sub(r"[ÓÒÔÖ]", "O", txt)
    txt = re.sub(r"[ÚÙÛÜ]", "U", txt)
    txt = txt.replace("Ñ", "N")
    txt = re.sub(r"[^A-Z0-9]+", "_", txt).strip("_")
    aliases = {
        'TRABAJADOR':'NOMBRE','APELLIDOS_Y_NOMBRES':'NOMBRE','NOMBRES_Y_APELLIDOS':'NOMBRE','NOMBRE_TRABAJADOR':'NOMBRE','NOMBRE_COMPLETO':'NOMBRE',
        'FECHA_NAC':'FECHA_NACIMIENTO','F_NACIMIENTO':'FECHA_NACIMIENTO','NACIMIENTO':'FECHA_NACIMIENTO',
        'FECHA_ING':'FECHA_INGRESO','F_INGRESO':'FECHA_INGRESO','INGRESO':'FECHA_INGRESO',
        'JEFE_INMEDIATO':'JEFE_DNI','JEFE':'JEFE_DNI','DNI_JEFE':'JEFE_DNI','JEFE_INMEDIATO_DNI':'JEFE_DNI',
        'JEFE_NOMBRE':'JEFE_NOMBRE','NOMBRE_JEFE':'JEFE_NOMBRE',
        'EMAIL':'CORREO','MAIL':'CORREO','CORREO_ELECTRONICO':'CORREO',
        'TELEFONO':'CELULAR','NRO_TELEFONO_MOVIL':'CELULAR','NUMERO_CELULAR':'CELULAR',
        'TELEFONO_EMERGENCIA':'TELEFONO_EMERGENCIA','NUMERO_TELEFONICO_DE_EMERGENCIA':'TELEFONO_EMERGENCIA',
        'CONTACTO_DE_EMERGENCIA':'CONTACTO_EMERGENCIA','CONTACTO_EMERGENCIA':'CONTACTO_EMERGENCIA',
        'TIPO_CONTRATO':'TIPO_CONTRATO','FECHA_FIN_CONTRATO':'FECHA_FIN_CONTRATO','REMUNERACION_BASICA':'REMUNERACION_BASICA',
        'DIRECCION_SIMPLE':'DIRECCION','DIRECCION_ACTUAL':'DIRECCION',
        'NIVEL_EDUCACION':'NIVEL_EDUCATIVO','CARNET_CONADIS':'CARNET_CONADIS','CONADIS':'CARNET_CONADIS'
    }
    return aliases.get(txt, txt)

def valor_fila(row, idx_map, *names, default=''):
    for name in names:
        k = normalizar_header_excel(name)
        i = idx_map.get(k, -1)
        if i >= 0 and i < len(row) and row[i] is not None:
            return row[i]
    return default

def construir_plantilla_trabajadores_xlsx(path):
    """Plantilla maestra para la pestaña Trabajadores: sirve para gestión documental, vacaciones y contratos."""
    wb = Workbook()
    ws = wb.active
    ws.title = 'TRABAJADORES'
    headers = [
        'EMPRESA','DNI','TRABAJADOR','CARGO','AREA','GERENCIA','SEDE','JEFE INMEDIATO','JEFE NOMBRE','PLANILLA','CORREO','CELULAR',
        'CONTACTO EMERGENCIA','TELEFONO EMERGENCIA','FECHA NACIMIENTO','FECHA INGRESO','TIPO CONTRATO','FECHA FIN CONTRATO','REMUNERACION BASICA',
        'DIRECCION','DEPARTAMENTO','PROVINCIA','DISTRITO','NIVEL EDUCATIVO','PROCEDENCIA','INDUMENTARIA','CARNET CONADIS','ACTIVO','OBSERVACION'
    ]
    ws.append(headers)
    ws.append(['AQUANQA','74324033','APELLIDOS Y NOMBRES','ANALISTA DE RRHH','GESTION DEL TALENTO HUMANO','ADMINISTRACION','TRUJILLO','43043999','JEFE DIRECTO','MENSUAL','correo@empresa.com','999999999','FAMILIAR DIRECTO','988888888','01/01/1990','01/05/2024','INTERMITENTE','31/12/2026','1200','AV. EJEMPLO 123','LA LIBERTAD','TRUJILLO','TRUJILLO','TECNICO','COSTA','SI','NO','SI','FILA DE EJEMPLO, BORRAR ANTES DE CARGAR'])
    ws.append(['AQUANCA II','00123456','OTRO TRABAJADOR','OPERARIO','OPERACIONES','PLANTA','CHAO','43043999','JEFE DIRECTO','SEMANAL','correo2@empresa.com','977777777','MADRE/PADRE','966666666','15/03/1995','10/02/2025','TEMPORAL','10/08/2026','1130','DIRECCION REFERENCIAL','LA LIBERTAD','VIRU','CHAO','SECUNDARIA COMPLETA','SIERRA','NO','SI','SI','FILA DE EJEMPLO, BORRAR ANTES DE CARGAR'])
    ws.freeze_panes = 'A2'
    dark = PatternFill('solid', fgColor='111827')
    green = PatternFill('solid', fgColor='22C55E')
    thin = Side(style='thin', color='D1D5DB')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = dark
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border
    for row in ws.iter_rows(min_row=2, max_row=500, max_col=len(headers)):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
    for col in range(1, len(headers)+1):
        letter = ws.cell(1, col).column_letter
        ws.column_dimensions[letter].width = 22
    for c in ['B','H','L','N']:
        ws.column_dimensions[c].width = 16
    for c in ['C','T','AC']:
        ws.column_dimensions[c].width = 34
    # Resalta campos mínimos obligatorios
    for col in [1,2,3,11,15]:
        ws.cell(1, col).fill = green
    # Validaciones
    def add_list(col, values):
        dv = DataValidation(type='list', formula1='"' + ','.join(values) + '"', allow_blank=True)
        ws.add_data_validation(dv); dv.add(f'{col}2:{col}5000')
    add_list('A', ['AQUANQA','AQUANCA II'])
    add_list('J', ['MENSUAL','SEMANAL','QUINCENAL','OTROS'])
    add_list('Q', ['INDETERMINADO','INTERMITENTE','TEMPORAL','SUPLENCIA','PRACTICANTE','OTROS'])
    add_list('X', ['PRIMARIA','SECUNDARIA COMPLETA','TECNICO','UNIVERSITARIO','BACHILLER','TITULADO','OTROS'])
    add_list('Z', ['SI','NO'])
    add_list('AA', ['SI','NO'])
    add_list('AB', ['SI','NO'])
    # Hoja instrucciones
    ins = wb.create_sheet('INSTRUCCIONES')
    ins.append(['USO DE LA PLANTILLA'])
    ins.append(['1. Cargar este Excel desde Admin > Trabajadores > Importar Excel.'])
    ins.append(['2. Campos obligatorios marcados en verde: EMPRESA, DNI, TRABAJADOR, CORREO y FECHA NACIMIENTO.'])
    ins.append(['3. DNI debe tener 8 dígitos. La clave del trabajador se genera con su fecha de nacimiento: ddmmaaaa.'])
    ins.append(['4. Las columnas extra quedan preparadas para gestión documental, vacacional y contratos.'])
    ins.append(['5. Puede borrar las filas de ejemplo antes de importar.'])
    ins.column_dimensions['A'].width = 120
    ins['A1'].font = Font(bold=True, size=14, color='FFFFFF')
    ins['A1'].fill = dark
    for row in ins.iter_rows(min_row=1, max_row=6, max_col=1):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical='top')
            cell.border = border
    wb.save(path)
    return path

@app.route('/admin/trabajadores', methods=['GET','POST'])
@admin_required
def admin_trabajadores():
    if request.method == 'POST':
        if 'excel' in request.files and request.files['excel'].filename:
            f = request.files['excel']; path = UPLOAD_DIR / f"base_{now_file()}_{secure_filename(f.filename)}"; f.save(path)
            wb = load_workbook(path, data_only=True); ws = wb.active
            headers = [normalizar_header_excel(c.value) for c in ws[1]]
            idx_map = {h:i for i,h in enumerate(headers) if h}
            n=0; omitidos=0
            extras_cols = ['celular','contacto_emergencia','telefono_emergencia','tipo_contrato','fecha_fin_contrato','remuneracion_basica','direccion','departamento','provincia','distrito','nivel_educativo','procedencia','indumentaria','carnet_conadis','observacion']
            with db() as con:
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if not any(row):
                        continue
                    dni = normalizar_dni(valor_fila(row, idx_map, 'DNI'))
                    nombre = clean(valor_fila(row, idx_map, 'NOMBRE','TRABAJADOR'))
                    correo = clean(valor_fila(row, idx_map, 'CORREO','EMAIL')).lower()
                    if not dni or not nombre:
                        omitidos += 1
                        continue
                    cargo = clean(valor_fila(row, idx_map, 'CARGO','PUESTO'))
                    area = clean(valor_fila(row, idx_map, 'AREA'))
                    jefe_dni = normalizar_dni(valor_fila(row, idx_map, 'JEFE_DNI','JEFE INMEDIATO'))
                    jefe_nombre = clean(valor_fila(row, idx_map, 'JEFE_NOMBRE'))
                    empresa = clean(valor_fila(row, idx_map, 'EMPRESA', default='AQUANQA')) or 'AQUANQA'
                    fecha_nac = excel_cell_fecha(valor_fila(row, idx_map, 'FECHA_NACIMIENTO'))
                    planilla = clean(valor_fila(row, idx_map, 'PLANILLA'))
                    fecha_ing = excel_cell_fecha(valor_fila(row, idx_map, 'FECHA_INGRESO'))
                    clave = generar_clave_trabajador(dni, fecha_nac)
                    con.execute("INSERT OR REPLACE INTO trabajadores(dni,nombre,correo,cargo,area,jefe_dni,jefe_nombre,empresa,planilla,fecha_nacimiento,fecha_ingreso,usuario_portal,clave_portal,activo,fecha_registro) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,1,?)", (dni,nombre,correo,cargo,area,jefe_dni,jefe_nombre,empresa,planilla,fecha_nac,fecha_ing,dni,clave,now_txt()))
                    extra_values = {
                        'celular': clean(valor_fila(row, idx_map, 'CELULAR')),
                        'contacto_emergencia': clean(valor_fila(row, idx_map, 'CONTACTO_EMERGENCIA')),
                        'telefono_emergencia': clean(valor_fila(row, idx_map, 'TELEFONO_EMERGENCIA')),
                        'tipo_contrato': clean(valor_fila(row, idx_map, 'TIPO_CONTRATO')),
                        'fecha_fin_contrato': excel_cell_fecha(valor_fila(row, idx_map, 'FECHA_FIN_CONTRATO')),
                        'remuneracion_basica': clean(valor_fila(row, idx_map, 'REMUNERACION_BASICA')),
                        'direccion': clean(valor_fila(row, idx_map, 'DIRECCION')),
                        'departamento': clean(valor_fila(row, idx_map, 'DEPARTAMENTO')),
                        'provincia': clean(valor_fila(row, idx_map, 'PROVINCIA')),
                        'distrito': clean(valor_fila(row, idx_map, 'DISTRITO')),
                        'nivel_educativo': clean(valor_fila(row, idx_map, 'NIVEL_EDUCATIVO')),
                        'procedencia': clean(valor_fila(row, idx_map, 'PROCEDENCIA')),
                        'indumentaria': clean(valor_fila(row, idx_map, 'INDUMENTARIA')),
                        'carnet_conadis': clean(valor_fila(row, idx_map, 'CARNET_CONADIS')),
                        'observacion': clean(valor_fila(row, idx_map, 'OBSERVACION'))
                    }
                    set_sql = ','.join([f"{c}=?" for c in extras_cols])
                    con.execute(f"UPDATE trabajadores SET {set_sql} WHERE dni=?", [extra_values[c] for c in extras_cols] + [dni])
                    n+=1
                con.commit()
            respaldar_exceles_locales()
            flash(f'Base cargada correctamente: {n} trabajadores. Omitidos por falta de DNI/NOMBRE: {omitidos}. Respaldo actualizado en REGISTROS_EXCEL_LOCAL.', 'ok')
        else:
            dni=normalizar_dni(request.form.get('dni'))
            with db() as con:
                fecha_nac=clean(request.form.get('fecha_nacimiento')); clave=generar_clave_trabajador(dni, fecha_nac); con.execute("INSERT OR REPLACE INTO trabajadores(dni,nombre,correo,cargo,area,jefe_dni,jefe_nombre,empresa,planilla,fecha_nacimiento,fecha_ingreso,usuario_portal,clave_portal,activo,fecha_registro) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,1,?)", (dni,clean(request.form.get('nombre')),clean(request.form.get('correo')).lower(),clean(request.form.get('cargo')),clean(request.form.get('area')),normalizar_dni(request.form.get('jefe_dni')),clean(request.form.get('jefe_nombre')),clean(request.form.get('empresa')) or 'AQUANQA',clean(request.form.get('planilla')),fecha_nac,fecha_sin_hora(request.form.get('fecha_ingreso')),dni,clave,now_txt()))
                con.commit()
            respaldar_exceles_locales()
            flash('Trabajador guardado y respaldo Excel actualizado.', 'ok')
        return redirect(url_for('admin_trabajadores'))
    with db() as con:
        rows = con.execute("SELECT * FROM trabajadores ORDER BY nombre LIMIT 300").fetchall()
    table = ''.join([f"<tr><td>{r['dni']}</td><td>{r['nombre']}</td><td>{r['correo']}</td><td>{r['cargo'] or ''}</td><td>{r['empresa'] or ''}</td><td>{r['jefe_dni'] if 'jefe_dni' in r.keys() and r['jefe_dni'] else ''}</td><td>{r['planilla'] if 'planilla' in r.keys() and r['planilla'] else ''}</td></tr>" for r in rows])
    content = f"""
    <div class='topbar'><div><h1>Trabajadores</h1><div class='subtitle'>Carga manual o masiva por Excel.</div><div class='local-note'>Respaldo local automático: REGISTROS_EXCEL_LOCAL / 01_TRABAJADORES_LOCAL.xlsx</div></div></div><section class='grid'>
    <div class='card span-12'><h2>Nuevo trabajador</h2><form method='post' class='form-grid'><div class='field'><label>DNI</label><input name='dni' required></div><div class='field'><label>Trabajador</label><input name='nombre' required></div><div class='field'><label>Correo</label><input name='correo' type='email' required></div><div class='field'><label>Cargo</label><input name='cargo'></div><div class='field'><label>Área</label><input name='area'></div><div class='field'><label>Empresa</label><select name='empresa'><option>AQUANQA</option><option>AQUANCA II</option></select></div><div class='field'><label>Jefe inmediato DNI</label><input name='jefe_dni' placeholder='DNI del jefe'></div><div class='field'><label>Jefe nombre</label><input name='jefe_nombre' placeholder='Opcional'></div><div class='field'><label>Planilla</label><input name='planilla'></div><div class='field'><label>Fecha nacimiento</label><input name='fecha_nacimiento' placeholder='dd/mm/aaaa'></div><div class='field'><label>Fecha de ingreso</label><input name='fecha_ingreso' placeholder='dd/mm/aaaa'></div><button class='btn-green'>Guardar + crear usuario</button></form></div>
    <div class='card span-12'><h2>Carga Excel</h2><p class='muted'>Plantilla oficial maestra para Gestión Documental y Vacacional. Acepta columnas de datos laborales, jefe inmediato, emergencia, ubicación, CONADIS e indumentaria. Crea usuario masivo con DNI y clave automática.</p><form method='post' enctype='multipart/form-data' class='form-grid'><div class='field'><label>Excel plantilla masiva</label><input type='file' name='excel' accept='.xlsx' required></div><button class='btn-blue'>Importar Excel</button><a class='btn-green' href='/admin/plantilla_trabajadores'>Plantilla Trabajadores</a> <a class='btn-blue' href='/admin/plantilla_gestion/documental'>Plantilla Documental</a></form></div>
    <div class='card span-12'><h2>Listado</h2><div class='table-wrap'><table><tr><th>DNI</th><th>Nombre</th><th>Correo</th><th>Cargo</th><th>Empresa</th><th>Jefe DNI</th><th>Planilla</th></tr>{table}</table></div></div></section>"""
    return render_page(content, active='Trabajadores')



def aplicar_formato_plantilla(ws, headers, color='111827'):
    dark = PatternFill('solid', fgColor=color)
    green = PatternFill('solid', fgColor='22C55E')
    thin = Side(style='thin', color='D1D5DB')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws.freeze_panes = 'A2'
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = dark
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border
    for row in ws.iter_rows(min_row=2, max_row=500, max_col=len(headers)):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
    for col in range(1, len(headers)+1):
        ws.column_dimensions[ws.cell(1,col).column_letter].width = 22
    for col in range(1, min(6, len(headers))+1):
        ws.cell(1, col).fill = green

def agregar_validacion_lista(ws, col, values, end=5000):
    dv = DataValidation(type='list', formula1='"' + ','.join(values) + '"', allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(f'{col}2:{col}{end}')

def construir_plantilla_gestion_documental_xlsx(path):
    wb = Workbook(); ws = wb.active; ws.title = 'GESTION_DOCUMENTAL'
    headers = ['EMPRESA','DNI','TRABAJADOR','CATEGORIA DOCUMENTO','TIPO DOCUMENTO','PERIODO','DETALLE','RUTA ARCHIVO O NOMBRE PDF','ESTADO','REQUIERE ACEPTACION','REQUIERE FIRMA','FECHA LIMITE','OBSERVACION']
    ws.append(headers)
    ws.append(['AQUANQA','74324033','APELLIDOS Y NOMBRES','DOCUMENTOS DE PAGO','BOLETA NORMAL','2026-05','SEMANA 20','74324033_BOLETA.pdf','PENDIENTE','SI','NO','31/05/2026','Ejemplo, borrar antes de cargar'])
    aplicar_formato_plantilla(ws, headers, '1F2937')
    agregar_validacion_lista(ws,'A',['AQUANQA','AQUANCA II'])
    agregar_validacion_lista(ws,'D',['DOCUMENTOS DE PAGO','DOCUMENTOS DE LA EMPRESA','DOCUMENTOS PERSONALES'])
    agregar_validacion_lista(ws,'E',['BOLETA NORMAL','CTS','GRATIFICACION','UTILIDAD','CONTRATO DE TRABAJO','REGLAMENTO INTERNO','POLITICA','OTROS'])
    agregar_validacion_lista(ws,'I',['PENDIENTE','ACEPTADO','FIRMADO','APROBADO','RECHAZADO','ARCHIVADO'])
    agregar_validacion_lista(ws,'J',['SI','NO']); agregar_validacion_lista(ws,'K',['SI','NO'])
    ins=wb.create_sheet('INSTRUCCIONES'); ins.append(['Plantilla para cargar/controlar documentos por trabajador, periodo y estado. Usar en Gestión Documental.'])
    ins.column_dimensions['A'].width=120
    wb.save(path); return path

def construir_plantilla_gestion_vacacional_xlsx(path):
    wb = Workbook(); ws = wb.active; ws.title = 'GESTION_VACACIONAL'
    headers = ['EMPRESA','DNI','TRABAJADOR','AREA','JEFE INMEDIATO','JEFE NOMBRE','FECHA INGRESO','I_PERIODO','F_PERIODO','DIAS GANADOS','DIAS GOZADOS','SALDO','ADELANTO VACACIONES','OBSERVACION']
    ws.append(headers)
    ws.append(['AQUANQA','74324033','APELLIDOS Y NOMBRES','GESTION DEL TALENTO HUMANO','43043999','JEFE DIRECTO','01/05/2024','2025','2026',30,0,30,'NO','Ejemplo, borrar antes de cargar'])
    aplicar_formato_plantilla(ws, headers, '166534')
    agregar_validacion_lista(ws,'A',['AQUANQA','AQUANCA II']); agregar_validacion_lista(ws,'M',['SI','NO'])
    for col in ('H','I'):
        for row in range(2, 5001): ws[f'{col}{row}'].number_format='@'
    ins=wb.create_sheet('INSTRUCCIONES'); ins.append(['Plantilla para saldos vacacionales. I_PERIODO y F_PERIODO deben ser años: 2025 / 2026.'])
    ins.column_dimensions['A'].width=120
    wb.save(path); return path

def construir_plantilla_gestion_contratacion_xlsx(path):
    wb = Workbook(); ws = wb.active; ws.title = 'GESTION_CONTRATACION'
    headers = ['EMPRESA','DNI','TRABAJADOR','AREA','CARGO','PLANILLA','TIPO TRABAJADOR','TIPO CONTRATO','FECHA INICIO CONTRATO','FECHA FIN CONTRATO','REMUNERACION BASICA','MONEDA','SEDE','ZONA','DIRECCION','DEPARTAMENTO','PROVINCIA','DISTRITO','MODALIDAD FIRMA','ESTADO FIRMA','PROVEEDOR FIRMA','REQUIERE RECONOCIMIENTO FACIAL','REQUIERE FIRMA DIGITAL','OBSERVACION']
    ws.append(headers)
    ws.append(['AQUANQA','74324033','APELLIDOS Y NOMBRES','RRHH','ANALISTA','MENSUAL','EMPLEADO','INTERMITENTE','01/05/2024','31/12/2026','1200','SOLES','TRUJILLO','OFICINA','AV. EJEMPLO 123','LA LIBERTAD','TRUJILLO','TRUJILLO','FACIAL + FIRMA DIGITAL','PENDIENTE','INTERNO','SI','SI','Ejemplo, borrar antes de cargar'])
    aplicar_formato_plantilla(ws, headers, '4C1D95')
    agregar_validacion_lista(ws,'A',['AQUANQA','AQUANCA II']); agregar_validacion_lista(ws,'H',['INDETERMINADO','INTERMITENTE','TEMPORAL','SUPLENCIA','PRACTICANTE','OTROS'])
    agregar_validacion_lista(ws,'S',['RECONOCIMIENTO FACIAL','FIRMA DIGITAL','FACIAL + FIRMA DIGITAL','CARGA MANUAL RRHH'])
    agregar_validacion_lista(ws,'T',['PENDIENTE','ENVIADO','VALIDADO FACIAL','FIRMADO DIGITAL','OBSERVADO','ANULADO'])
    agregar_validacion_lista(ws,'V',['SI','NO']); agregar_validacion_lista(ws,'W',['SI','NO'])
    campos=wb.create_sheet('CAMPOS_WORD')
    campos.append(['CAMPO PARA WORD','EJEMPLO DE USO'])
    for campo,_ in CAMPOS_ESQUEMA_TRABAJADOR_CONTRATO_LABORAL:
        campos.append([campo, '{{'+campo+'}}'])
    campos.column_dimensions['A'].width=34; campos.column_dimensions['B'].width=38
    aplicar_formato_plantilla(campos, ['CAMPO PARA WORD','EJEMPLO DE USO'], '4C1D95')
    ins=wb.create_sheet('INSTRUCCIONES'); ins.append(['Plantilla para contratación: sirve para contratos, renovaciones, estados de firma y campos de correspondencia Word.'])
    ins.column_dimensions['A'].width=120
    wb.save(path); return path

@app.route('/admin/plantilla_gestion/<gestion>')
@admin_required
def plantilla_gestion(gestion):
    gestion = clean(gestion).lower()
    if gestion == 'documental':
        path = PERSIST_DIR / 'PLANTILLA_GESTION_DOCUMENTAL.xlsx'; construir_plantilla_gestion_documental_xlsx(path); name='PLANTILLA_GESTION_DOCUMENTAL.xlsx'
    elif gestion == 'vacacional':
        path = PERSIST_DIR / 'PLANTILLA_GESTION_VACACIONAL.xlsx'; construir_plantilla_gestion_vacacional_xlsx(path); name='PLANTILLA_GESTION_VACACIONAL.xlsx'
    elif gestion in ['contratacion','contratación']:
        path = PERSIST_DIR / 'PLANTILLA_GESTION_CONTRATACION.xlsx'; construir_plantilla_gestion_contratacion_xlsx(path); name='PLANTILLA_GESTION_CONTRATACION.xlsx'
    else:
        abort(404)
    return send_file(path, as_attachment=True, download_name=name)

@app.route('/admin/plantilla_trabajadores')
@admin_required
def plantilla_trabajadores():
    path = PERSIST_DIR / 'PLANTILLA_GENERAL_TRABAJADORES_GESTION_RRHH.xlsx'
    construir_plantilla_trabajadores_xlsx(path)
    return send_file(path, as_attachment=True, download_name='PLANTILLA_GENERAL_TRABAJADORES_GESTION_RRHH.xlsx')

@app.route('/foto/<dni>')
def foto_trabajador(dni):
    t = get_trabajador(dni)
    if not t or not t['foto_ruta']: abort(404)
    path = Path(t['foto_ruta'])
    if not path.exists(): abort(404)
    return send_file(path, as_attachment=False)

@app.route('/mi_foto', methods=['POST'])
@worker_required
def mi_foto():
    f = request.files.get('foto')
    if not f or not f.filename:
        flash('Seleccione una foto.', 'error'); return redirect(url_for('panel'))
    ext = Path(secure_filename(f.filename)).suffix.lower()
    if ext not in ['.png','.jpg','.jpeg','.webp']:
        flash('Formato de foto no permitido.', 'error'); return redirect(url_for('panel'))
    folder = UPLOAD_DIR / 'fotos'; folder.mkdir(parents=True, exist_ok=True)
    path = folder / f"{session['dni']}_{now_file()}{ext}"; f.save(path)
    with db() as con:
        con.execute('UPDATE trabajadores SET foto_ruta=? WHERE dni=?', (str(path), session['dni'])); con.commit()
    flash('Foto actualizada correctamente.', 'ok')
    return redirect(url_for('panel'))


@app.route('/admin/usuarios')
@admin_required
def admin_usuarios():
    with db() as con:
        rows = con.execute("SELECT t.dni,t.nombre,t.correo,t.empresa,t.cargo,t.fecha_nacimiento,t.usuario_portal,t.clave_portal,t.activo,COALESCE(l.intentos,0) intentos,COALESCE(l.bloqueado,0) bloqueado FROM trabajadores t LEFT JOIN login_intentos l ON l.dni=t.dni ORDER BY t.nombre LIMIT 10000").fetchall()
    trs=[]
    for r in rows:
        dni = r['dni']
        clave = r['clave_portal'] or generar_clave_trabajador(r['dni'], r['fecha_nacimiento'])
        estado = '🔒 Bloqueado' if int(r['bloqueado'] or 0)==1 else '✅ Activo'
        desbloq = f" <a class='btn-green mini-btn' href='/admin/usuario/{dni}/desbloquear'>Desbloquear</a>" if int(r['bloqueado'] or 0)==1 or int(r['intentos'] or 0)>0 else ''
        trs.append(f"<tr><td>{dni}</td><td>{r['nombre']}</td><td>{r['usuario_portal'] or dni}</td><td><b>{clave}</b></td><td>{estado}<br><small>Intentos: {int(r['intentos'] or 0)}</small></td><td>{r['empresa'] or ''}</td><td><a class='btn-blue mini-btn' href='/admin/usuario/{dni}/reset'>Regenerar</a>{desbloq} <a class='btn-red mini-btn' onclick='return confirm(\"¿Eliminar trabajador/usuario?\")' href='/admin/usuario/{dni}/eliminar'>Eliminar</a></td></tr>")
    table=''.join(trs)
    content=f"""
    <div class='hero'><div class='topbar'><div><h1>Usuarios y contraseñas</h1><div class='subtitle'>Control para más de 10 mil trabajadores. Usuario = DNI; clave = fecha de nacimiento sin / (ddmmaaaa).</div></div><a class='btn-green' href='/admin/plantilla_trabajadores'>Plantilla masiva</a></div></div>
    <section class='grid'><div class='card span-12'><h2>Listado de accesos</h2><p class='muted'>El trabajador ingresa con usuario = DNI y clave = fecha nacimiento sin / (ddmmaaaa).</p><div class='table-wrap'><table><tr><th>DNI</th><th>Trabajador</th><th>Usuario</th><th>Clave</th><th>Estado login</th><th>Empresa</th><th>Opciones</th></tr>{table}</table></div></div></section>"""
    return render_page(content, active='Usuarios')

@app.route('/admin/usuario/<dni>/reset')
@admin_required
def admin_usuario_reset(dni):
    t=get_trabajador(dni)
    if not t: abort(404)
    clave=generar_clave_trabajador(dni, t['fecha_nacimiento'] if 'fecha_nacimiento' in t.keys() else '')
    with db() as con:
        con.execute("UPDATE trabajadores SET usuario_portal=?, clave_portal=? WHERE dni=?", (normalizar_dni(dni), clave, normalizar_dni(dni))); con.execute('DELETE FROM login_intentos WHERE dni=?',(normalizar_dni(dni),)); con.commit()
    flash('Usuario regenerado correctamente.', 'ok')
    return redirect(url_for('admin_usuarios'))


@app.route('/admin/usuario/<dni>/desbloquear')
@admin_required
def admin_usuario_desbloquear(dni):
    reset_intentos_login(dni)
    flash('Usuario desbloqueado. Ya puede ingresar con DNI y fecha de nacimiento sin /.', 'ok')
    return redirect(url_for('admin_usuarios'))

@app.route('/admin/usuario/<dni>/eliminar')
@admin_required
def admin_usuario_eliminar(dni):
    with db() as con:
        con.execute("DELETE FROM trabajadores WHERE dni=?", (normalizar_dni(dni),)); con.commit()
    flash('Trabajador/usuario eliminado.', 'ok')
    return redirect(url_for('admin_usuarios'))

@app.route('/admin/documentos', methods=['GET','POST'])
@admin_required
def admin_documentos():
    if request.method == 'POST':
        tipo = clean(request.form.get('tipo'))
        dni = normalizar_dni(request.form.get('dni'))
        periodo = request.form.get('periodo')
        detalle = request.form.get('detalle')
        obs = request.form.get('observacion')
        per_norm = clean(request.form.get('periodicidad_normal'))
        if tipo == 'Normal' and per_norm and per_norm.lower() not in clean(detalle).lower():
            detalle = (clean(detalle) + ' - ' + per_norm).strip(' -')
        files = request.files.getlist('archivos')
        ok=0
        try:
            for f in files:
                if f and f.filename:
                    guardar_documento(f, dni, tipo, periodo, detalle, obs, marca_carga(session.get('admin_user','admin'))); ok += 1
            flash(f'Carga completada: {ok} archivo(s).', 'ok')
        except Exception as e:
            flash(f'Error en carga: {e}', 'error')
        return redirect(url_for('admin_documentos', tipo=tipo))
    tipo = clean(request.args.get('tipo')) or 'Utilidad'
    buscar = clean(request.args.get('buscar'))
    periodo = clean(request.args.get('periodo'))
    sub = clean(request.args.get('sub'))
    tipo_options = ''.join([f"<option value='{k}' {'selected' if k==tipo else ''}>{l}</option>" for k,l,i in TIPOS_PAGO+TIPOS_EMPRESA+TIPOS_PERSONALES])
    pers = periodos_disponibles(tipo=tipo)
    periodo_options = "<option value=''>Todos</option>" + ''.join([f"<option {'selected' if p==periodo else ''}>{p}</option>" for p in pers])
    rows = listar_documentos(tipo=tipo if tipo else None, periodo=periodo or None, buscar=buscar, limit=500)
    content = f"""
    <div class='hero'><div class='topbar'><div><h1>Subir y gestionar documentos</h1><div class='subtitle'>Administrador: pago, empresa y documentos personales.</div></div><a class='btn-warn' href='/admin/sincronizar'>Actualizar / detectar PDFs</a><a class='btn-blue' href='/admin/crear_carpetas'>Crear carpetas + detectar</a></div></div><section class='grid'>
    <div class='card span-12'><h2>📁 Carpeta local automática</h2><p class='muted'>Ruta actual: <b>{DOCUMENTOS_BASE_DIR}</b><br>Coloca PDFs en DOCUMENTOS DE PAGO / BOLETAS NORMAL / SEMANAL o MENSUAL y presiona <b>Actualizar / detectar PDFs</b>. Solo se cargarán trabajadores activos. El DNI se detecta por nombre/ruta y también leyendo el contenido del PDF.</p><div class='actions'><a class='btn-warn' href='/admin/sincronizar'>Actualizar / detectar PDFs</a><a class='btn-blue' href='/admin/crear_carpetas'>Crear estructura</a></div></div><div class='card span-12'><h2>Carga de documentos</h2><form method='post' enctype='multipart/form-data' class='form-grid'><div class='field'><label>Tipo</label><select name='tipo'>{tipo_options}</select></div><div class='field'><label>DNI trabajador</label><input name='dni' placeholder='Vacío si es documento de empresa'></div><div class='field'><label>Periodo</label><input name='periodo' value='{datetime.now(APP_TZ).strftime('%Y-%m')}' list='periodos'></div><div class='field'><label>Detalle</label><input name='detalle' placeholder='Ej: Boleta semanal / Política actualizada'></div><div class='field'><label>Boleta Normal</label><select name='periodicidad_normal'><option value=''>No aplica</option><option>Mensual</option><option>Semanal</option></select></div><div class='field'><label>Archivos</label><input type='file' name='archivos' accept='.pdf,.png,.jpg,.jpeg,.webp,.doc,.docx,.xls,.xlsx' multiple required></div><div class='field'><label>Observación</label><textarea name='observacion' rows='2'></textarea></div><button class='btn-green'>Subir</button></form></div>
    <div class='card span-12'><h2>Filtros</h2><form method='get' class='form-grid'><div class='field'><label>Tipo</label><select name='tipo'>{tipo_options}</select></div><div class='field'><label>Periodo</label><select name='periodo'>{periodo_options}</select></div><div class='field'><label>Buscar</label><input name='buscar' value='{buscar}' placeholder='DNI, detalle, observación'></div><button class='btn-blue'>Filtrar</button><a class='btn' href='/admin/documentos'>Limpiar</a></form></div>
    <div class='card span-12'><h2>Listado</h2>{tabla_docs(rows)}</div></section>"""
    return render_page(content, active='Subir documentos')



def dias_entre_texto(fi, ff):
    a=parse_fecha_any(fi); b=parse_fecha_any(ff)
    if not a or not b: return 0
    return max((b-a).days+1, 0)


def _ids_periodos_texto(valor):
    return [int(x) for x in str(valor or '').replace(',', '|').split('|') if str(x).strip().isdigit()]


def dias_reservados_periodos(con, dni, excluir_id=None):
    """Devuelve días ya comprometidos por periodo para no usar el mismo saldo 2 veces.
    Considera solicitudes pendientes y aprobadas; ignora rechazadas/anuladas.
    """
    params = [dni]
    extra = ''
    if excluir_id:
        extra = ' AND id<>?'
        params.append(excluir_id)
    filas = con.execute(f"""
        SELECT id, dias, estado, periodo_ids
        FROM vacaciones_solicitudes
        WHERE dni=?
          AND COALESCE(estado,'') NOT LIKE 'Rechazado%'
          AND COALESCE(estado,'') NOT LIKE 'Anulado%'
          {extra}
    """, params).fetchall()
    usados = {}
    for row in filas:
        ids = _ids_periodos_texto(row['periodo_ids'] if 'periodo_ids' in row.keys() else '')
        if not ids:
            continue
        dias = float(row['dias'] or 0)
        por_periodo = dias / max(len(ids), 1)
        for pid in ids:
            usados[pid] = usados.get(pid, 0) + por_periodo
    return usados


def saldo_disponible_real(row, usados_por_periodo):
    return max(float(row['saldo'] or 0) - float(usados_por_periodo.get(int(row['id']), 0)), 0)

def obtener_jefe_dni_trabajador(con, trabajador_dni, periodo_ids=None):
    """Obtiene el DNI del jefe inmediato de forma robusta.
    Prioridad: periodo seleccionado en saldos -> cualquier saldo del trabajador -> ficha trabajadores.
    Esto evita que una solicitud quede sin aprobador cuando el jefe viene de la plantilla de saldos.
    """
    trabajador_dni = normalizar_dni(trabajador_dni)
    periodo_ids = [int(x) for x in (periodo_ids or []) if str(x).isdigit()]
    if periodo_ids:
        marks = ','.join(['?'] * len(periodo_ids))
        row = con.execute(f"""
            SELECT jefe_dni FROM vacaciones_saldos
            WHERE dni=? AND id IN ({marks}) AND COALESCE(jefe_dni,'')<>''
            ORDER BY periodo_inicio, periodo_fin LIMIT 1
        """, [trabajador_dni] + periodo_ids).fetchone()
        if row:
            jd = normalizar_dni(row['jefe_dni'])
            if jd: return jd
    row = con.execute("""
        SELECT jefe_dni FROM vacaciones_saldos
        WHERE dni=? AND COALESCE(jefe_dni,'')<>''
        ORDER BY periodo_inicio, periodo_fin LIMIT 1
    """, (trabajador_dni,)).fetchone()
    if row:
        jd = normalizar_dni(row['jefe_dni'])
        if jd: return jd
    row = con.execute("SELECT jefe_dni FROM trabajadores WHERE dni=?", (trabajador_dni,)).fetchone()
    if row and 'jefe_dni' in row.keys():
        jd = normalizar_dni(row['jefe_dni'])
        if jd: return jd
    return ''

def sql_solicitudes_jefe(extra_where=''):
    return f"""
        SELECT vs.* FROM vacaciones_solicitudes vs
        LEFT JOIN trabajadores tr ON tr.dni = vs.dni
        WHERE (
            normalizar_dni_sql(COALESCE(vs.jefe_dni,'')) = ?
            OR normalizar_dni_sql(COALESCE(tr.jefe_dni,'')) = ?
            OR EXISTS (
                SELECT 1 FROM vacaciones_saldos s
                WHERE s.dni = vs.dni
                  AND normalizar_dni_sql(COALESCE(s.jefe_dni,'')) = ?
            )
        )
        {extra_where}
    """

@app.route('/admin/modulo/documentos')
@admin_required
def admin_modulo_documentos():
    with db() as con:
        total=con.execute('SELECT COUNT(*) c FROM documentos').fetchone()['c']
        pendientes=con.execute("SELECT COUNT(*) c FROM documentos WHERE COALESCE(estado,'Pendiente') IN ('Pendiente','Aceptado','Firmado')").fetchone()['c']
        aprob=con.execute("SELECT COUNT(*) c FROM documentos WHERE estado='Aprobado'").fetchone()['c']
    content=f"""
    <div class='hero'><div class='topbar'><div><h1>Gestión <span class='accent'>Documental</span></h1><div class='subtitle'>Concentra todo lo ya implementado: cargas, PDFs, carpetas locales, aceptación/firma/aprobación y trazabilidad.</div></div><a class='btn-green' href='/admin/documentos'>Entrar a documentos</a></div></div>
    <section class='grid'><div class='card mini'><div><h3>Total documentos</h3><b>{total}</b></div><div class='ico'>🗃️</div></div><div class='card mini'><div><h3>Pendientes</h3><b>{pendientes}</b></div><div class='ico'>⏳</div></div><div class='card mini'><div><h3>Aprobados</h3><b>{aprob}</b></div><div class='ico'>✅</div></div>
    <div class='card span-12'><div class='module-tabs'><a class='module-tile' href='/admin/documentos'><h2>📤 Subir documentos</h2><p class='muted'>Pago, empresa y personales.</p></a><a class='module-tile' href='/admin/sincronizar'><h2>🔎 Detectar PDFs</h2><p class='muted'>Lee carpetas locales y detecta DNI.</p></a><a class='module-tile' href='/admin/crear_carpetas'><h2>📁 Crear carpetas</h2><p class='muted'>Estructura automática DOCUMENTOS_PRIZE_AUTO.</p></a></div></div></section>"""
    return render_page(content, active='Modulo documentos')

@app.route('/admin/vacaciones', methods=['GET','POST'])
@admin_required
def admin_vacaciones():
    if request.method=='POST':
        f=request.files.get('excel')
        ok=0
        if f and f.filename:
            path=UPLOAD_DIR/'vacaciones'; path.mkdir(parents=True, exist_ok=True)
            x=path/(now_file()+'_'+secure_filename(f.filename)); f.save(x)
            wb=load_workbook(x, data_only=True); ws=wb.active
            headers=[str(c.value or '').strip().upper() for c in ws[1]]
            def val(row, names):
                # Normaliza cabeceras: acepta tildes, espacios y dos puntos (ej. DIAS GANADO:)
                hdr_norm=[re.sub(r'[^A-Z0-9ÑÁÉÍÓÚ]+',' ', h).strip().replace('Á','A').replace('É','E').replace('Í','I').replace('Ó','O').replace('Ú','U') for h in headers]
                for n in names:
                    n2=re.sub(r'[^A-Z0-9ÑÁÉÍÓÚ]+',' ', str(n).upper()).strip().replace('Á','A').replace('É','E').replace('Í','I').replace('Ó','O').replace('Ú','U')
                    if n in headers: return row[headers.index(n)].value
                    if n2 in hdr_norm: return row[hdr_norm.index(n2)].value
                return ''
            def num(v, default=0):
                try:
                    if v is None or str(v).strip()=='': return default
                    return float(str(v).replace(',', '.').strip())
                except Exception:
                    return default
            with db() as con:
                for row in ws.iter_rows(min_row=2):
                    dni=normalizar_dni(val(row,['DNI','DOCUMENTO','CODIGO','CÓDIGO']))
                    if not dni: continue
                    trabajador=clean(val(row,['TRABAJADOR','NOMBRE','APELLIDOS Y NOMBRES']))
                    gan=num(val(row,['DIAS GANADOS','DÍAS GANADOS','DIAS GANADO','DÍAS GANADO','GANADOS']), 0)
                    saldo=num(val(row,['SALDO','SALDO VACACIONAL']), gan)
                    trabajador_db = con.execute('SELECT * FROM trabajadores WHERE dni=?', (dni,)).fetchone()
                    if trabajador_db:
                        trabajador = trabajador or trabajador_db['nombre']
                        fecha_ing = trabajador_db['fecha_ingreso'] if 'fecha_ingreso' in trabajador_db.keys() else ''
                    else:
                        fecha_ing = ''
                    p_ini=periodo_year_value(val(row,['I_PERIODO','PERIODO INICIO','INICIO PERIODO','FECHA INICIO PERIODO']))
                    p_fin=periodo_year_value(val(row,['F_PERIODO','PERIODO FIN','FIN PERIODO','FECHA FIN PERIODO']))
                    periodo=clean(val(row,['PERIODO','PERÍODO'])) or periodo_anual_texto(p_ini, p_fin)
                    jefe_raw=clean(val(row,['JEFE DNI','DNI JEFE','JEFE INMEDIATO','JEFE']))
                    jefe_dni=normalizar_dni(jefe_raw)
                    jefe_nombre=''
                    if jefe_dni:
                        jr=con.execute('SELECT nombre FROM trabajadores WHERE dni=?', (jefe_dni,)).fetchone()
                        jefe_nombre = jr['nombre'] if jr else jefe_raw
                    else:
                        jr=con.execute('SELECT dni,nombre FROM trabajadores WHERE UPPER(nombre)=UPPER(?)', (jefe_raw,)).fetchone()
                        jefe_dni = jr['dni'] if jr else ''
                        jefe_nombre = jr['nombre'] if jr else jefe_raw
                    existente = con.execute("SELECT id FROM vacaciones_saldos WHERE dni=? AND COALESCE(periodo_inicio,'')=? AND COALESCE(periodo_fin,'')=? ORDER BY id LIMIT 1", (dni, p_ini, p_fin)).fetchone()
                    data_saldo = (trabajador,clean(val(row,['EMPRESA'])),clean(val(row,['AREA','ÁREA'])),jefe_nombre,jefe_dni,fecha_ing,p_ini,p_fin,gan,0,saldo,periodo,now_txt(),marca_carga(session.get('admin_user','admin')))
                    if existente:
                        con.execute('''UPDATE vacaciones_saldos
                                          SET trabajador=?, empresa=?, area=?, jefe=?, jefe_dni=?, fecha_ingreso=?,
                                              periodo_inicio=?, periodo_fin=?, dias_ganados=?, dias_gozados=?, saldo=?,
                                              periodo=?, fecha_carga=?, uploaded_by=?
                                        WHERE id=?''', data_saldo + (existente['id'],))
                    else:
                        con.execute('''INSERT INTO vacaciones_saldos
                          (dni,trabajador,empresa,area,jefe,jefe_dni,fecha_ingreso,periodo_inicio,periodo_fin,dias_ganados,dias_gozados,saldo,periodo,fecha_carga,uploaded_by)
                          VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''', (dni,) + data_saldo)
                    if jefe_dni:
                        con.execute("UPDATE trabajadores SET jefe_dni=?, jefe_nombre=COALESCE(NULLIF(?, ''), jefe_nombre) WHERE dni=?", (jefe_dni, jefe_nombre, dni))
                        con.execute("UPDATE vacaciones_solicitudes SET jefe_dni=? WHERE dni=? AND (estado IN ('Pendiente jefe','Pendiente Jefe','Pendiente') OR COALESCE(jefe_dni,'')='')", (jefe_dni, dni))
                    ok+=1
                sincronizar_jefes_vacaciones(con)
                con.commit()
                respaldar_exceles_locales()
        flash(f'Saldos vacacionales cargados/actualizados: {ok}. Respaldo Excel local actualizado.','ok')
        return redirect(url_for('admin_vacaciones'))
    q_sol=clean(request.args.get('q_sol'))
    q_sal=clean(request.args.get('q_sal'))
    with db() as con:
        params=[]; where=''
        if q_sal:
            where='WHERE dni LIKE ? OR UPPER(trabajador) LIKE UPPER(?)'
            params=[q_sal+'%', '%'+q_sal+'%']
        saldos=con.execute(f'SELECT * FROM vacaciones_saldos {where} ORDER BY trabajador, periodo_inicio LIMIT 500', params).fetchall()
        params=[]; where=''
        if q_sol:
            where='WHERE dni LIKE ? OR UPPER(trabajador) LIKE UPPER(?)'
            params=[q_sol+'%', '%'+q_sol+'%']
        solicitudes=con.execute(f'SELECT * FROM vacaciones_solicitudes {where} ORDER BY id DESC LIMIT 500', params).fetchall()
    sal=''.join([f"<tr><td>{r['dni']}</td><td>{r['trabajador']}</td><td>{r['empresa'] or ''}</td><td>{r['area'] or ''}</td><td>{r['jefe_dni'] if 'jefe_dni' in r.keys() else ''}</td><td>{r['jefe'] or ''}</td><td>{r['periodo_inicio'] or ''}</td><td>{r['periodo_fin'] or ''}</td><td>{r['dias_ganados']}</td><td><b>{r['saldo']}</b></td></tr>" for r in saldos])
    sol=''.join([f"<tr><td>{r['id']}</td><td>{r['dni']}</td><td>{r['trabajador']}</td><td>{r['jefe_dni'] if 'jefe_dni' in r.keys() else ''}</td><td>{r['fecha_inicio']} al {r['fecha_fin']}</td><td>{r['dias']}</td><td>{r['periodo_detalle'] if 'periodo_detalle' in r.keys() and r['periodo_detalle'] else ''}</td><td><span class='status-pill'>{r['estado']}</span></td><td class='actions'><a class='btn-green mini-btn' href='/admin/vacaciones/{r['id']}/jefe/aprobar'>Apr. jefe</a><a class='btn-green mini-btn' href='/admin/vacaciones/{r['id']}/gh/aprobar'>Apr. GTH</a><a class='btn-red mini-btn' href='/admin/vacaciones/{r['id']}/rechazar'>Rechazar</a></td></tr>" for r in solicitudes])
    content=f"""
    <div class='hero'><div class='topbar'><div><h1>Gestión <span class='accent'>Vacacional</span></h1><div class='subtitle'>Administrador carga saldos; usuario solicita goce; flujo: jefe inmediato → Gestión del Talento Humano.</div></div><a class='btn-green' href='/admin/vacaciones/plantilla'>Descargar plantilla</a><a class='btn-blue' href='/admin/vacaciones/sincronizar_jefes'>Sincronizar jefes</a></div></div>
    <section class='grid'><div id='aprobaciones' class='card mini'><div><h3>Pendientes de aprobación</h3><b>{len([r for r in solicitudes if 'Pendiente' in (r['estado'] or '')])}</b></div><div class='ico'>✅</div></div><div class='card mini'><div><h3>Saldos registrados</h3><b>{len(saldos)}</b></div><div class='ico'>🗓️</div></div><div class='card mini'><div><h3>Solicitudes totales</h3><b>{len(solicitudes)}</b></div><div class='ico'>📄</div></div><div id='cargar-saldos' class='card span-12'><h2>🏖️ Saldos Vacacionales</h2><form method='post' enctype='multipart/form-data' class='form-grid'><div class='field'><label>Excel saldos</label><input type='file' name='excel' accept='.xlsx' required></div><button class='btn-green'>Subir saldos</button></form><p class='muted'>Columnas: EMPRESA, DNI, TRABAJADOR, AREA, JEFE INMEDIATO (DNI), I_PERIODO, F_PERIODO, DIAS GANADOS, SALDO. No usar FECHA INGRESO ni PERIODO ni DÍAS GOZADOS.</p></div>
    <div id='solicitudes' class='card span-12'><h2>📄 Solicitudes de vacaciones</h2><form method='get' class='form-grid'><div class='field'><label>Buscar por DNI o apellidos</label><input name='q_sol' value='{q_sol}' placeholder='Ej: 473 o QUINTANA'></div><button class='btn-green'>Filtrar solicitudes</button><a class='btn' href='/admin/vacaciones#solicitudes'>Limpiar</a></form><div class='table-wrap'><table><tr><th>ID</th><th>DNI</th><th>Trabajador</th><th>DNI jefe</th><th>Rango</th><th>Días</th><th>Periodo usado</th><th>Estado</th><th>Acciones</th></tr>{sol or '<tr><td colspan=9>No hay solicitudes.</td></tr>'}</table></div></div>
    <div id='reportes' class='card span-12'><h2>📑 Reporte de saldos cargados</h2><form method='get' class='form-grid'><div class='field'><label>Buscar por DNI o apellidos</label><input name='q_sal' value='{q_sal}' placeholder='Ej: 473 o QUINTANA'></div><button class='btn-green'>Filtrar saldos</button><a class='btn' href='/admin/vacaciones#reportes'>Limpiar</a></form><div class='table-wrap'><table><tr><th>DNI</th><th>Trabajador</th><th>Empresa</th><th>Área</th><th>DNI jefe</th><th>Jefe</th><th>I_Periodo</th><th>F_Periodo</th><th>Ganados</th><th>Saldo</th></tr>{sal or '<tr><td colspan=10>No hay saldos cargados.</td></tr>'}</table></div></div></section>"""
    return render_page(content, active='Gestion Vacacional')

@app.route('/admin/vacaciones/plantilla')
@admin_required
def admin_vacaciones_plantilla():
    path=PERSIST_DIR/'PLANTILLA_SALDOS_VACACIONALES.xlsx'
    wb=Workbook(); ws=wb.active; ws.title='SALDOS'
    headers=['EMPRESA','DNI','TRABAJADOR','AREA','JEFE INMEDIATO','I_PERIODO','F_PERIODO','DIAS GANADOS','SALDO']
    ws.append(headers); ws.append(['AQUANQA','74324033','APELLIDOS Y NOMBRES','RRHH','43043999','2025','2026',30,30])
    # Formato simple: I_PERIODO y F_PERIODO son SOLO AÑOS, no fechas completas.
    for i,h in enumerate(headers,1):
        ws.column_dimensions[chr(64+i)].width=24
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='166534')
        cell.alignment = Alignment(horizontal='center')
    for col in ('F','G'):
        for row in range(2, 502):
            ws[f'{col}{row}'].number_format = '@'
    ws.freeze_panes='A2'
    wb.save(path); return send_file(path, as_attachment=True, download_name='PLANTILLA_SALDOS_VACACIONALES.xlsx')

@app.route('/admin/vacaciones/<int:sid>/<rol>/<accion>')
@admin_required
def admin_vacaciones_accion(sid, rol, accion):
    if accion=='aprobar' and rol=='jefe': estado='Pendiente GTH'; col='fecha_jefe'
    elif accion=='aprobar' and rol=='gh': estado='Aprobado GTH'; col='fecha_gh'
    else: estado='Rechazado'; col='fecha_gh'
    with db() as con:
        con.execute(f'UPDATE vacaciones_solicitudes SET estado=?, {col}=? WHERE id=?', (estado, now_txt(), sid)); con.commit(); respaldar_exceles_locales()
    flash('Solicitud actualizada.', 'ok'); return redirect(url_for('admin_vacaciones'))


@app.route('/admin/vacaciones/sincronizar_jefes')
@admin_required
def admin_vacaciones_sincronizar_jefes():
    with db() as con:
        sincronizar_jefes_vacaciones(con)
        con.commit()
    respaldar_exceles_locales()
    flash('Jefes inmediatos sincronizados: las solicitudes pendientes fueron migradas al DNI del jefe cargado en saldos/trabajadores.', 'ok')
    return redirect(url_for('admin_vacaciones'))

@app.route('/vacaciones/mi_solicitud', methods=['GET','POST'])
@worker_required
def trabajador_vacaciones():
    dni=session['dni']; t=get_trabajador(dni)
    with db() as con:
        saldos_usuario=con.execute('SELECT * FROM vacaciones_saldos WHERE dni=? ORDER BY periodo_inicio, periodo_fin',(dni,)).fetchall()
        saldo=saldos_usuario[0] if saldos_usuario else None
    if request.method=='POST':
        fi_raw=clean(request.form.get('fecha_inicio')); ff_raw=clean(request.form.get('fecha_fin'))
        fi_date=parse_fecha_any(fi_raw); ff_date=parse_fecha_any(ff_raw)
        hoy=hoy_lima()
        hoy_iso=hoy.isoformat()
        if not fi_raw or not ff_raw or not fi_date or not ff_date:
            flash('No se registró la solicitud: debe seleccionar fechas válidas.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        # BLOQUEO DEFINITIVO EN SERVIDOR: aunque manipulen el HTML o escriban la fecha manualmente.
        if fi_date < hoy:
            flash(f'No se registró la solicitud: la fecha de inicio {fi_date.strftime("%d/%m/%Y")} es anterior a hoy {hoy.strftime("%d/%m/%Y")}.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        if ff_date < hoy:
            flash(f'No se registró la solicitud: la fecha fin {ff_date.strftime("%d/%m/%Y")} es anterior a hoy {hoy.strftime("%d/%m/%Y")}.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        if ff_date < fi_date:
            flash('No se registró la solicitud: la fecha fin no puede ser menor que la fecha inicio.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        # Bloqueo adicional: no permitir solicitudes demasiado lejanas.
        if (fi_date - hoy).days > 90:
            flash('No se registró la solicitud: la fecha de inicio supera el límite permitido de 90 días.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        fi=fi_date.isoformat(); ff=ff_date.isoformat(); dias=dias_entre_texto(fi,ff)
        adelanto = '1' if request.form.get('adelanto') else ''
        periodo_ids = [int(x) for x in request.form.getlist('periodos') if str(x).isdigit()]
        if not periodo_ids:
            flash('Seleccione con check el periodo que usará para gozar vacaciones.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        with db() as con:
            marks=','.join(['?']*len(periodo_ids))
            saldos_sel=con.execute(f'SELECT * FROM vacaciones_saldos WHERE dni=? AND id IN ({marks}) ORDER BY periodo_inicio, periodo_fin', [dni]+periodo_ids).fetchall()
            usados = dias_reservados_periodos(con, dni)
        saldo_disponible = sum(saldo_disponible_real(r, usados) for r in saldos_sel)
        if dias <= 0:
            flash('Rango de fechas inválido. Revisa inicio y fin.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        if not saldos_sel or saldo_disponible <= 0:
            flash('No se registró la solicitud: el/los periodo(s) seleccionado(s) no tienen saldo disponible o ya fueron usados en otra solicitud.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        if dias > saldo_disponible:
            flash(f'No se registró la solicitud: solicita {dias} día(s), pero el saldo real disponible es {saldo_disponible}.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        estado = 'Pendiente jefe'
        motivo_base = clean(request.form.get('motivo'))
        periodo_detalle = ' | '.join([f"{r['periodo_inicio']}-{r['periodo_fin']} (saldo disponible {saldo_disponible_real(r, usados)})" for r in saldos_sel])
        periodo_ids_txt = '|'.join(str(x) for x in periodo_ids)
        if adelanto:
            motivo_base = (motivo_base + ' | ' if motivo_base else '') + 'Solicitud marcada como comentario especial; validada dentro del saldo disponible.'
        try:
            
            with db() as con:
                jefe_dni = obtener_jefe_dni_trabajador(con, dni, periodo_ids)
            if not jefe_dni:
                flash('No se registró la solicitud: este trabajador no tiene DNI de jefe inmediato. Cargue la plantilla de trabajadores o saldos con JEFE INMEDIATO = DNI del jefe.', 'err')
                return redirect(url_for('trabajador_vacaciones'))
            # Segundo candado justo antes de grabar: evita registrar fechas pasadas aunque el formulario haya sido alterado.
            if parse_fecha_any(fi) < hoy_lima() or parse_fecha_any(ff) < hoy_lima():
                flash('No se registró la solicitud: las fechas no pueden ser anteriores a hoy.', 'err')
                return redirect(url_for('trabajador_vacaciones'))
            con.execute('INSERT INTO vacaciones_solicitudes(dni,trabajador,jefe_dni,fecha_inicio,fecha_fin,dias,motivo,estado,fecha_solicitud,periodo_detalle,periodo_ids) VALUES(?,?,?,?,?,?,?,?,?,?,?)',(dni,t['nombre'] if t else '',jefe_dni,fi,ff,dias,motivo_base,estado,now_txt(),periodo_detalle,periodo_ids_txt)); con.commit(); respaldar_exceles_locales()
        except Exception as e:
            flash(f'Error interno al registrar vacaciones: {str(e)}', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        flash('Solicitud registrada. Pasará por jefe inmediato y Gestión del Talento Humano.','ok')
        return redirect(url_for('trabajador_vacaciones'))
    with db() as con:
        saldos_usuario=con.execute('SELECT * FROM vacaciones_saldos WHERE dni=? ORDER BY periodo_inicio, periodo_fin',(dni,)).fetchall()
        saldo=saldos_usuario[0] if saldos_usuario else None
        solicitudes=con.execute('SELECT * FROM vacaciones_solicitudes WHERE dni=? ORDER BY id DESC',(dni,)).fetchall()
        por_aprobar=con.execute(sql_solicitudes_jefe("AND vs.estado='Pendiente jefe' ORDER BY vs.id DESC"),(dni,dni,dni)).fetchall()
    sol=''.join([f"<div class='sol-card'><div data-label='Fecha'><b>{r['fecha_solicitud']}</b></div><div data-label='Rango'><b>{r['fecha_inicio']} al {r['fecha_fin']}</b></div><div data-label='Días' class='dias'><b>{r['dias']}</b></div><div data-label='Periodo usado'><b>{r['periodo_detalle'] if 'periodo_detalle' in r.keys() and r['periodo_detalle'] else '-'}</b></div><div data-label='Estado'><span class='status-pill'>{r['estado']}</span></div><div data-label='Comentario' class='coment'>{r['motivo'] or '-'}</div></div>" for r in solicitudes])
    sol_aprobar=''.join([f"<tr><td>{r['fecha_solicitud']}</td><td>{r['dni']}</td><td>{r['trabajador']}</td><td>{r['fecha_inicio']} al {r['fecha_fin']}</td><td>{r['dias']}</td><td><span class='status-pill'>{r['estado']}</span></td><td class='actions'><a class='btn-green mini-btn' href='/vacaciones/aprobar_jefe/{r['id']}'>Aprobar</a><a class='btn-red mini-btn' href='/vacaciones/rechazar_jefe/{r['id']}'>Rechazar</a></td></tr>" for r in por_aprobar])
    with db() as con:
        usados_periodos = dias_reservados_periodos(con, dni)
    saldo_val = sum(saldo_disponible_real(r, usados_periodos) for r in saldos_usuario)
    periodos_html = ''.join([f"<label class='period-card {'disabled' if saldo_disponible_real(r, usados_periodos) <= 0 else ''}'><input type='checkbox' name='periodos' value='{r['id']}' {'disabled' if saldo_disponible_real(r, usados_periodos) <= 0 else ''}><span class='period-main'><span class='period-years'>{r['periodo_inicio'] or ''} - {r['periodo_fin'] or ''}</span><span class='period-meta'><span class='period-badge'>Ganados: <strong>{r['dias_ganados']}</strong></span><span class='period-badge'>Saldo real: <strong>{saldo_disponible_real(r, usados_periodos)}</strong></span><span class='period-badge'>Usado/Pendiente: <strong>{usados_periodos.get(int(r['id']), 0)}</strong></span></span></span></label>" for r in saldos_usuario]) or '<p class=\'muted\'>No tiene periodos vacacionales cargados.</p>'
    content=f"""
    <div class='hero'><div class='topbar'><div><h1>Gestión <span class='accent'>Vacacional</span></h1><div class='subtitle'>Consulta tu saldo, valida días disponibles y registra solicitudes.</div></div></div></div>
    <section class='grid'><div class='card mini'><div><h3>Saldo disponible</h3><b>{saldo_val}</b></div><div class='ico'>🏖️</div></div><div class='card mini'><div><h3>Días ganados</h3><b>{sum(float(r['dias_ganados'] or 0) for r in saldos_usuario)}</b></div><div class='ico'>📈</div></div><div class='card mini'><div><h3>Periodos</h3><b>{len(saldos_usuario)}</b></div><div class='ico'>📅</div></div><div class='card mini'><div><h3>Fecha ingreso</h3><b>{fecha_sin_hora(t['fecha_ingreso'] if t and 'fecha_ingreso' in t.keys() else '') or '-'}</b></div><div class='ico'>🗓️</div></div>
    <div class='card span-12' style='{"display:block" if sol_aprobar else "display:none"}'><h2>✅ Solicitudes pendientes por aprobar como jefe inmediato</h2><p class='muted'>Te aparecen aquí solo los trabajadores que tienen tu DNI como jefe inmediato en la plantilla de saldos.</p><div class='table-wrap'><table><tr><th>Fecha</th><th>DNI</th><th>Trabajador</th><th>Rango</th><th>Días</th><th>Estado</th><th>Acciones</th></tr>{sol_aprobar or '<tr><td colspan=7>No tienes solicitudes pendientes por aprobar.</td></tr>'}</table></div></div>
    <div id='solicitar' class='card span-12 vac-request-card'><div class='vac-head'><div><h2>🗓️ Nueva solicitud</h2><p class='vac-help'>Marca el periodo que vas a utilizar. Puedes seleccionar más de uno cuando el descanso consuma saldos acumulados.</p></div></div><form method='post' id='formSolicitudVacaciones'><div class='field'><label>Periodos disponibles</label><div class='period-list'>{periodos_html}</div></div><div class='vac-form-row'><div class='field'><label>Inicio</label><input type='date' id='fecha_inicio_vac' name='fecha_inicio' min='{hoy_lima().isoformat()}' required></div><div class='field'><label>Fin</label><input type='date' id='fecha_fin_vac' name='fecha_fin' min='{hoy_lima().isoformat()}' required></div><div class='field'><label>Motivo / comentario</label><input name='motivo' placeholder='Goce vacacional'></div></div><div class='vac-submit-row'><label class='check-card'><input type='checkbox' name='adelanto' value='1'> Requiere revisión especial</label><button class='btn-green'>Registrar solicitud</button></div></form><script>
(function(){{
  const hoyISO = '{hoy_lima().isoformat()}';
  const f = document.getElementById('formSolicitudVacaciones');
  const ini = document.getElementById('fecha_inicio_vac');
  const fin = document.getElementById('fecha_fin_vac');

  function valorISO(campo){{
    if(!campo || !campo.value) return '';
    const v = campo.value.trim();
    // Navegadores modernos devuelven YYYY-MM-DD.
    if(/^\d{{4}}-\d{{2}}-\d{{2}}$/.test(v)) return v;
    // Soporte adicional si el navegador deja escribir DD/MM/AAAA.
    const m = v.match(/^(\d{{1,2}})[\/\-](\d{{1,2}})[\/\-](\d{{4}})$/);
    if(m){{
      const d = m[1].padStart(2,'0');
      const mo = m[2].padStart(2,'0');
      return `${{m[3]}}-${{mo}}-${{d}}`;
    }}
    return '';
  }}

  [ini, fin].forEach(x => {{
    if(x){{
      x.setAttribute('min', hoyISO);
      x.setAttribute('autocomplete', 'off');
      // No mostrar alertas mientras la fecha está incompleta; solo validar al registrar.
    }}
  }});

  if(f){{ f.addEventListener('submit', function(e){{
    const iniISO = valorISO(ini);
    const finISO = valorISO(fin);
    if(!iniISO || !finISO){{
      e.preventDefault(); alert('Seleccione fecha de inicio y fin válidas.'); return false;
    }}
    if(iniISO < hoyISO || finISO < hoyISO){{
      e.preventDefault(); alert('No se registró: la fecha de inicio y fin no pueden ser anteriores a hoy (' + hoyISO + ').'); return false;
    }}
    if(finISO < iniISO){{
      e.preventDefault(); alert('La fecha fin no puede ser menor que inicio.'); return false;
    }}
  }});}}
}})();
</script></div>
    <div class='card span-12'><h2>Mis solicitudes</h2><div class='sol-cards'><div class='sol-card head'><div>Fecha</div><div>Rango</div><div>Días</div><div>Periodo usado</div><div>Estado</div><div>Comentario</div></div>{sol or "<div class='sol-empty'>No hay solicitudes registradas.</div>"}</div></div></section>"""
    return render_page(content, active='Gestion Vacacional')


@app.route('/vacaciones/aprobar_jefe/<int:sid>')
@worker_required
def vacaciones_aprobar_jefe(sid):
    dni=session['dni']
    with db() as con:
        r=con.execute("SELECT * FROM vacaciones_solicitudes WHERE id=?", (sid,)).fetchone()
        jefe_solicitud = normalizar_dni(r['jefe_dni'] if r and 'jefe_dni' in r.keys() else '') if r else ''
        jefe_trabajador = obtener_jefe_dni_trabajador(con, r['dni']) if r else ''
        if not r or (jefe_solicitud != dni and jefe_trabajador != dni):
            flash('No autorizado: esta solicitud no corresponde a tu aprobación como jefe inmediato.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        con.execute("UPDATE vacaciones_solicitudes SET estado='Pendiente GTH', fecha_jefe=?, comentario_jefe=? WHERE id=?", (now_txt(), 'Aprobado por jefe inmediato', sid))
        con.commit()
    flash('Solicitud aprobada por jefe inmediato. Ahora queda pendiente de GTH.', 'ok')
    return redirect(url_for('trabajador_vacaciones'))

@app.route('/vacaciones/rechazar_jefe/<int:sid>')
@worker_required
def vacaciones_rechazar_jefe(sid):
    dni=session['dni']
    with db() as con:
        r=con.execute("SELECT * FROM vacaciones_solicitudes WHERE id=?", (sid,)).fetchone()
        jefe_solicitud = normalizar_dni(r['jefe_dni'] if r and 'jefe_dni' in r.keys() else '') if r else ''
        jefe_trabajador = obtener_jefe_dni_trabajador(con, r['dni']) if r else ''
        if not r or (jefe_solicitud != dni and jefe_trabajador != dni):
            flash('No autorizado: esta solicitud no corresponde a tu aprobación como jefe inmediato.', 'err')
            return redirect(url_for('trabajador_vacaciones'))
        con.execute("UPDATE vacaciones_solicitudes SET estado='Rechazado por jefe', fecha_jefe=?, comentario_jefe=? WHERE id=?", (now_txt(), 'Rechazado por jefe inmediato', sid))
        con.commit()
    flash('Solicitud rechazada por jefe inmediato.', 'ok')
    return redirect(url_for('trabajador_vacaciones'))


@app.route('/vacaciones/aprobaciones_jefe')
@worker_required
def vacaciones_aprobaciones_jefe():
    dni=session['dni']; t=get_trabajador(dni)
    with db() as con:
        rows=con.execute(sql_solicitudes_jefe("ORDER BY CASE WHEN vs.estado='Pendiente jefe' THEN 0 ELSE 1 END, vs.id DESC"), (dni,dni,dni)).fetchall()
    pendientes=sum(1 for r in rows if (r['estado'] or '') == 'Pendiente jefe')
    aprobadas=sum(1 for r in rows if 'GTH' in (r['estado'] or '') or 'Aprobado' in (r['estado'] or ''))
    rechazadas=sum(1 for r in rows if 'Rechazado' in (r['estado'] or ''))
    cards=[]
    for r in rows:
        acciones = ""
        if (r['estado'] or '') == 'Pendiente jefe':
            acciones = f"<div class='actions'><a class='btn-green mini-btn' href='/vacaciones/aprobar_jefe/{r['id']}'>Aprobar</a><a class='btn-red mini-btn' href='/vacaciones/rechazar_jefe/{r['id']}'>Rechazar</a></div>"
        cards.append(f"""
        <div class='sol-card'>
          <div data-label='Fecha'><b>{r['fecha_solicitud']}</b></div>
          <div data-label='Trabajador'><b>{r['dni']} - {r['trabajador']}</b></div>
          <div data-label='Rango'><b>{r['fecha_inicio']} al {r['fecha_fin']}</b></div>
          <div data-label='Días' class='dias'><b>{r['dias']}</b></div>
          <div data-label='Periodo usado'><b>{r['periodo_detalle'] if 'periodo_detalle' in r.keys() and r['periodo_detalle'] else '-'}</b></div>
          <div data-label='Estado'><span class='status-pill'>{r['estado']}</span></div>
          <div data-label='Acción'>{acciones or '<span class="muted">Sin acción pendiente</span>'}</div>
        </div>""")
    content=f"""
    <div class='hero'><div class='topbar'><div><h1>Aprobaciones <span class='accent'>Jefe inmediato</span></h1><div class='subtitle'>Este panel aparece para cualquier trabajador que figure como JEFE INMEDIATO por DNI en la plantilla de saldos.</div></div></div></div>
    <section class='grid'>
      <div class='card mini'><div><h3>Jefe</h3><b>{t['nombre'] if t else dni}</b></div><div class='ico'>👤</div></div>
      <div class='card mini'><div><h3>Pendientes</h3><b>{pendientes}</b></div><div class='ico'>⏳</div></div>
      <div class='card mini'><div><h3>Aprobadas / GTH</h3><b>{aprobadas}</b></div><div class='ico'>✅</div></div>
      <div class='card mini'><div><h3>Rechazadas</h3><b>{rechazadas}</b></div><div class='ico'>🚫</div></div>
      <div class='card span-12'><h2>Solicitudes asignadas a tu DNI</h2><p class='muted'>Si no aparece una solicitud, revise que en la plantilla de saldos el campo JEFE INMEDIATO tenga exactamente tu DNI.</p><div class='sol-cards'><div class='sol-card head'><div>Fecha</div><div>Trabajador</div><div>Rango</div><div>Días</div><div>Periodo usado</div><div>Estado</div><div>Acción</div></div>{''.join(cards) or "<div class='sol-empty'>No tienes solicitudes asignadas para aprobar.</div>"}</div></div>
    </section>"""
    return render_page(content, active='Gestion Vacacional')

@app.route('/contratacion/mis_documentos')
@worker_required
def trabajador_contratacion():
    dni=session['dni']; t=get_trabajador(dni)
    with db() as con:
        docs=con.execute('SELECT * FROM contratacion_docs WHERE dni=? ORDER BY id DESC', (dni,)).fetchall()
        firma_rows=con.execute('SELECT documento_id,firma_token,estado FROM firma_solicitudes WHERE dni=? ORDER BY id DESC', (dni,)).fetchall()
    tokens_por_doc={}
    for fr in firma_rows:
        if fr['documento_id'] and fr['documento_id'] not in tokens_por_doc and fr['firma_token']:
            tokens_por_doc[fr['documento_id']]=fr['firma_token']
    filas=[]
    for r in docs:
        token=tokens_por_doc.get(r['id'])
        accion=f"<a class='btn-blue mini-btn' target='_blank' href='/contratacion/ver/{r['id']}'>Ver</a>"
        if token and 'FIRMADO' not in (r['estado'] or '').upper():
            accion += f" <a class='btn-green mini-btn' target='_blank' href='/firma/{token}'>Firmar ahora</a>"
        filas.append(f"<tr><td>{r['tipo_doc']}</td><td>{r['etapa']}</td><td><span class='status-pill'>{r['estado']}</span></td><td>{r['fecha_registro']}</td><td>{accion}</td></tr>")
    rows=''.join(filas)
    content=f"""
    <div class='hero'><div class='topbar'><div><h1>Gestión de <span class='accent'>Contrato</span></h1><div class='subtitle'>Visualiza y descarga tus contratos, anexos y documentos de incorporación o renovación.</div></div></div></div>
    <section class='grid'><div class='card mini'><div><h3>Trabajador</h3><b>{t['nombre'] if t else dni}</b></div><div class='ico'>👤</div></div><div class='card mini'><div><h3>Empresa</h3><b>{session.get('empresa') or (t['empresa'] if t else '')}</b></div><div class='ico'>🏢</div></div><div class='card mini'><div><h3>Documentos</h3><b>{len(docs)}</b></div><div class='ico'>🧾</div></div>
    <div id='mis-contratos' class='card span-12'><h2>Mis documentos contractuales</h2><div class='table-wrap'><table><tr><th>Documento</th><th>Etapa</th><th>Estado</th><th>Fecha</th><th>Acción</th></tr>{rows or '<tr><td colspan=5>No hay documentos de contratación cargados.</td></tr>'}</table></div></div></section>"""
    return render_page(content, active=f'Gestion Contratacion:{sec}')

@app.route('/contratacion/ver/<int:cid>')
@worker_required
def ver_contratacion(cid):
    with db() as con:
        r=con.execute('SELECT * FROM contratacion_docs WHERE id=?', (cid,)).fetchone()
    if not r or r['dni'] != session.get('dni'): abort(404)
    path=Path(r['ruta_archivo'])
    if not path.exists(): abort(404)
    return send_file(path, as_attachment=False, download_name=r['archivo_nombre'])



@app.route('/admin/contratacion/plantilla/<int:pid>')
@admin_required
def contratacion_plantilla_detalle(pid):
    tab=request.args.get('tab','contenido')
    with db() as con:
        pl=con.execute('SELECT * FROM contratacion_plantillas WHERE id=?',(pid,)).fetchone()
        campos=con.execute('SELECT * FROM contratacion_plantilla_campos WHERE plantilla_id=? ORDER BY id',(pid,)).fetchall()
        condiciones=con.execute('SELECT * FROM contratacion_plantilla_condiciones WHERE plantilla_id=? ORDER BY id',(pid,)).fetchall()
    if not pl: abort(404)
    estado = 'Activo' if pl['activo'] else 'Inactivo'
    condicion_habilitada = (pl['condicion'] or '').upper() == 'CONDICIONES'
    edit_url=url_for('contratacion_plantilla_editar',pid=pid)
    tabs=f"""
    <div class='tpl-tabs'>
      <a class='tpl-tab {'active' if tab=='contenido' else ''}' href='{url_for('contratacion_plantilla_detalle',pid=pid,tab='contenido')}'>Contenido</a>
      <a class='tpl-tab {'active' if tab=='campos' else ''}' href='{url_for('contratacion_plantilla_detalle',pid=pid,tab='campos')}'>Campos</a>
      <a class='tpl-tab {'active' if tab=='condiciones' else ''}' href='{url_for('contratacion_plantilla_detalle',pid=pid,tab='condiciones')}'>Condiciones</a>
    </div>"""
    if tab=='campos':
        rows=''.join([f"<tr><td><a class='icon-btn small' href='{url_for('contratacion_campo_editar',pid=pid,campo_id=c['id'])}'>Editar</a></td><td><span class='state-pill {'ok' if c['activo'] else 'bad'}'>{'ACTIVE' if c['activo'] else 'INACTIVE'}</span></td><td>{html.escape(c['descripcion'] or '')}</td><td>{html.escape(c['tipo_campo'] or '')}</td><td>{html.escape(c['nombre_campo'] or '')}</td><td><code>{{{{{html.escape(c['campo_origen'] or '')}}}}}</code></td><td>{html.escape(c['tipo_dato'] or '')}</td><td>{html.escape(c['requerido'] or '')}</td></tr>" for c in campos])
        body=f"""<div class='tpl-toolbar schema-toolbar'><span>Campos de correspondencia para usar en Word como <b>«CampoOrigen»</b> o <b>{{{{CampoOrigen}}}}</b>.</span><span><a class='c-btn' href='{url_for('contratacion_campo_editar',pid=pid)}'>+ Crear Campo</a> <a class='c-btn gray' href='{url_for('contratacion_campos_esquema',pid=pid)}'>⌕ Campos de Esquema</a> <a class='c-btn gray' href='{url_for('contratacion_campos_esquema_excel',pid=pid)}'>⬇ Descargar campos</a></span></div><div class='tpl-table-wrap'><table class='tpl-table'><tr><th>Editar</th><th>Estado</th><th>Descripción</th><th>Tipo Campo</th><th>Nombre Campo</th><th>Campo Origen</th><th>Tipo de Dato</th><th>Requerido</th></tr>{rows or '<tr><td colspan="8">Sin campos registrados.</td></tr>'}</table></div>"""
    elif tab=='condiciones':
        crear_btn = f"<a class='c-btn' href='{url_for('contratacion_condicion_editar',pid=pid)}'>⊕ Crear Condición</a>" if condicion_habilitada else "<span class='c-btn gray disabled'>Condiciones deshabilitadas</span>"
        info = "" if condicion_habilitada else "<div class='notice'>Para crear o editar condiciones primero entra al lápiz de <b>Editar Plantilla</b> y cambia el campo <b>Condición</b> a <b>CONDICIONES</b>.</div>"
        rows=''.join([f"""
          <tr>
            <td class='actions'><a class='icon-btn small' href='{url_for('contratacion_condicion_editar',pid=pid,cid=c['id'])}'>✎</a>
              <form method='post' action='{url_for('contratacion_condicion_eliminar',pid=pid,cid=c['id'])}' onsubmit="return confirm('¿Eliminar condición?')"><button class='trash'>🗑</button></form></td>
            <td><span class='state-pill {'ok' if c['activo'] else 'bad'}'>{'ACTIVE' if c['activo'] else 'INACTIVE'}</span></td>
            <td>{html.escape(c['nombre_campo'] or '')}</td><td>{html.escape(c['condicion'] or '=')}</td><td>{html.escape(c['valor'] or '')}</td>
          </tr>""" for c in condiciones])
        body=f"""<div class='cond-head'><div>{crear_btn}</div></div>{info}<div class='tpl-table-wrap'><table class='tpl-table'><tr><th>Proceso</th><th>Estado</th><th>Label</th><th>Condición</th><th>Valor</th></tr>{rows or '<tr><td colspan="5">Sin condiciones registradas.</td></tr>'}</table></div>"""
    else:
        dni_preview = normalizar_dni(request.args.get('dni_preview'))
        trabajador_preview = None
        with db() as con:
            trabajador_preview = con.execute('SELECT * FROM trabajadores WHERE dni=?', (dni_preview,)).fetchone() if dni_preview else con.execute('SELECT * FROM trabajadores ORDER BY nombre LIMIT 1').fetchone()
        valores_preview = mapa_campos_trabajador(trabajador_preview)
        cumple_cond, detalle_cond = plantilla_cumple_condiciones(pl, condiciones, trabajador_preview)
        ruta_preview = Path(pl['ruta_archivo']) if pl['ruta_archivo'] else None
        if ruta_preview and ruta_preview.exists() and str(pl['archivo_nombre']).lower().endswith('.pdf'):
            preview=f"<iframe class='pdf-frame' src='{url_for('contratacion_plantilla_archivo',pid=pid)}'></iframe>"
        elif ruta_preview and ruta_preview.exists() and ruta_preview.suffix.lower()=='.docx':
            preview=docx_to_preview_html(ruta_preview, valores_preview)
        elif ruta_preview and ruta_preview.exists() and ruta_preview.suffix.lower()=='.doc':
            preview=f"<div class='preview-empty'><b>Archivo Word .doc cargado:</b> {html.escape(pl['archivo_nombre'] or '')}<br><br><a class='c-btn gray' href='{url_for('contratacion_plantilla_archivo',pid=pid)}'>⬇ Descargar / abrir archivo</a><p>Para vista previa dentro del sistema se recomienda guardar la plantilla como .docx.</p></div>"
        elif ruta_preview and ruta_preview.exists():
            preview=f"<div class='preview-empty'><b>Archivo cargado:</b> {html.escape(pl['archivo_nombre'] or '')}<br><br><a class='c-btn gray' href='{url_for('contratacion_plantilla_archivo',pid=pid)}'>⬇ Descargar / abrir archivo</a></div>"
        else:
            preview=f"<div class='preview-empty'><b>Sin archivo cargado.</b><br>Haz clic en el lápiz ✎ para editar y cargar la plantilla Word/PDF.</div>"
        dni_val = html.escape(dni_preview or row_get(trabajador_preview, 'dni') or '')
        gen_url = url_for('contratacion_plantilla_generar', pid=pid, dni=dni_val) if dni_val else url_for('contratacion_plantilla_generar', pid=pid)
        estado_cond = '✅ Cumple condiciones' if cumple_cond else '⚠ No cumple condiciones'
        body=f"""<form class='preview-tools' method='get' action='{url_for('contratacion_plantilla_detalle',pid=pid)}'><input type='hidden' name='tab' value='contenido'><input name='dni_preview' maxlength='8' value='{dni_val}' placeholder='DNI del trabajador para previsualizar'><button class='c-btn green'>Previsualizar conectado</button><a class='c-btn gray' href='{gen_url}'>⬇ Generar Word combinado</a></form><div class='tpl-toolbar'><b>{estado_cond}</b> &nbsp; {detalle_cond}</div>{preview}"""
    file_btn = f"<a class='c-btn gray' href='{url_for('contratacion_plantilla_archivo',pid=pid)}'>⬇ Descargar Archivo</a>"
    content=f"""
    <style>
      .contract-detail-wrap{{color:#162033!important;background:#f5f7fb;margin:-8px -10px 0;padding:14px 18px 32px;min-height:calc(100vh - 80px)}}
      .contract-detail-wrap *{{box-sizing:border-box}}
      .detail-title{{font-size:24px;font-weight:950;margin:0 0 14px;color:#101827;letter-spacing:-.4px}}
      .template-head{{background:linear-gradient(135deg,#fff,#f9fbff);border:1px solid #dce4ef;border-radius:18px;padding:24px 28px;display:grid;grid-template-columns:1.15fr .85fr;gap:28px;margin-bottom:12px;box-shadow:0 16px 35px rgba(15,23,42,.08)}}
      .template-head h1{{margin:0 0 16px;font-size:30px;line-height:1.15;color:#162033;font-weight:950;display:flex;align-items:center;gap:10px;flex-wrap:wrap}}
      .tpl-line{{margin:14px 0;color:#697487;font-size:15px;font-weight:800;display:flex;gap:10px;align-items:center;flex-wrap:wrap}}
      .tpl-line b{{color:#162033;font-weight:950;min-width:122px}}
      .icon-btn{{display:inline-flex!important;align-items:center;justify-content:center;min-width:36px;height:36px;border-radius:11px;background:#e9edf4!important;color:#111827!important;border:1px solid #d8dee8;text-decoration:none;font-size:20px;font-weight:950;box-shadow:0 6px 16px rgba(15,23,42,.08)}}
      .icon-btn.small{{min-width:32px;height:32px;font-size:17px}}
      .icon-btn:hover{{background:#dfe6f0!important;transform:translateY(-1px)}}
      .green-dot,.red-dot{{display:inline-block;width:12px;height:12px;border-radius:50%;margin-right:6px}}.green-dot{{background:#22c55e}}.red-dot{{background:#ef4444}}
      .meta-box{{margin-top:16px;background:#eef2f7;color:#647084;padding:10px 12px;font-weight:850;font-size:13px;display:flex;gap:30px;flex-wrap:wrap;border-radius:12px}}
      .tpl-side{{border-left:2px dashed #ccd6e3;padding-left:26px}}
      .tpl-tabs{{display:flex;background:#fff;border:1px solid #dce4ef;border-radius:14px;margin:12px 0 0;overflow:hidden;box-shadow:0 10px 24px rgba(15,23,42,.06)}}
      .tpl-tab{{padding:16px 22px;color:#697487!important;text-decoration:none;font-weight:900;border-bottom:3px solid transparent;min-width:190px;text-align:center}}
      .tpl-tab.active{{color:#0f172a!important;border-bottom-color:#ff963b;background:#fff7ed}}
      .preview-tools{{display:flex;gap:14px;margin:12px 0 0;padding:14px;background:#fff;border:1px solid #dce4ef;border-radius:14px 14px 0 0}}
      .preview-tools input{{max-width:450px;background:#fff!important;color:#111!important;border:1px solid #d1d5db;border-radius:10px;padding:12px 14px;font-weight:800}}
      .c-btn{{display:inline-flex;align-items:center;justify-content:center;border:0;border-radius:12px;padding:11px 16px;background:linear-gradient(135deg,#ff963b,#ff7a1a);color:#fff!important;font-weight:950;text-decoration:none;cursor:pointer;box-shadow:0 10px 22px rgba(255,122,26,.22)}}
      .c-btn.gray{{background:#e7ebf2;color:#111827!important;box-shadow:none}}.c-btn.green{{background:#22c55e;color:#fff!important}}.c-btn.disabled{{opacity:.62;cursor:not-allowed}}
      .tpl-toolbar{{color:#647084;font-weight:850;padding:13px 18px;background:#fff;border-left:1px solid #dce4ef;border-right:1px solid #dce4ef}}
      .schema-toolbar{{display:flex;justify-content:space-between;gap:12px;align-items:center;flex-wrap:wrap}}
      .pdf-frame{{width:100%;height:720px;border:1px solid #d1d5db;background:#eee;border-radius:0 0 14px 14px}}
      .preview-empty{{background:#fff!important;color:#111827!important;border:1px dashed #cbd5e1;border-radius:14px;padding:30px;margin:12px 0;font-size:16px}}.preview-empty p{{color:#6b7280!important}}
      .cond-head{{display:flex;justify-content:flex-end;background:#fff;border:1px solid #dce4ef;border-top:0;padding:14px 18px}}
      .notice{{margin:12px 0;padding:14px 16px;border-radius:14px;background:#fff7ed;border:1px solid #fed7aa;color:#7c2d12;font-weight:850}}
      .tpl-table-wrap{{overflow:auto;background:#fff;border:1px solid #dce4ef;border-radius:0 0 14px 14px;max-height:620px}}
      .tpl-table{{width:100%;border-collapse:collapse;color:#111827!important;background:#fff!important}}
      .tpl-table th{{background:#f8fafc!important;color:#111827!important;text-align:left;font-size:14px;border:1px solid #e1e7ef;padding:13px;font-weight:950}}
      .tpl-table td{{border:1px solid #e1e7ef;padding:13px;color:#111827!important;background:#fff!important;font-weight:750}}
      .tpl-table tr:nth-child(even) td{{background:#f5f6f8!important}}
      .state-pill{{display:inline-flex;align-items:center;border:1px solid #d1d5db;border-radius:20px;padding:6px 12px;background:#fff;font-weight:900;color:#111827}}
      .state-pill.ok:before{{content:'';width:10px;height:10px;border-radius:50%;background:#22c55e;margin-right:7px}}.state-pill.bad:before{{content:'';width:10px;height:10px;border-radius:50%;background:#ef4444;margin-right:7px}}
      .actions{{display:flex;gap:8px;align-items:center}}.actions form{{margin:0}}.trash{{height:32px;border:0;background:#fee2e2;border-radius:10px;cursor:pointer}}
      .word-preview{{background:#fff;color:#111827!important;border:1px solid #dce4ef;border-radius:0 0 14px 14px;padding:32px 42px;min-height:620px;line-height:1.65;font-family:Arial,serif;font-size:15.5px;box-shadow:inset 0 0 0 1px #eef2f7}}
      .word-preview p{{margin:0 0 12px;color:#111827!important;font-weight:500;text-align:justify}}
      .word-table{{width:100%;border-collapse:collapse;margin:14px 0;background:#fff;color:#111827!important}}.word-table td{{border:1px solid #cbd5e1;padding:8px 10px;background:#fff!important;color:#111827!important;font-weight:500}}
      .cond-ok{{display:inline-flex;margin:3px 6px 3px 0;padding:5px 9px;border-radius:999px;background:#dcfce7;color:#166534;font-weight:900;border:1px solid #86efac}}.cond-bad{{display:inline-flex;margin:3px 6px 3px 0;padding:5px 9px;border-radius:999px;background:#fee2e2;color:#991b1b;font-weight:900;border:1px solid #fecaca}}
      .back-top{{display:flex;justify-content:flex-end;margin:0 0 10px}}
      code{{background:#eef2ff;border:1px solid #dbeafe;border-radius:8px;padding:4px 7px;color:#1e3a8a}}

      .contract-detail-wrap h1,.contract-detail-wrap h2,.contract-detail-wrap h3,.contract-detail-wrap b,.contract-detail-wrap label,.contract-detail-wrap span,.contract-detail-wrap div{{text-shadow:none!important;opacity:1!important}}
      .contract-detail-wrap .detail-title,.contract-detail-wrap .template-head h1{{color:#0f172a!important}}
      .contract-detail-wrap .tpl-line,.contract-detail-wrap .tpl-line span{{color:#334155!important;font-weight:900!important}}
      .contract-detail-wrap .tpl-line b{{color:#0f172a!important;font-weight:1000!important}}
      .contract-detail-wrap .meta-box,.contract-detail-wrap .meta-box *{{color:#475569!important}}
      .contract-detail-wrap .tpl-toolbar{{color:#475569!important;background:#fff!important}}
      .contract-detail-wrap .preview-empty,.contract-detail-wrap .preview-empty *{{color:#0f172a!important;font-weight:950!important}}
      @media(max-width:900px){{.template-head{{grid-template-columns:1fr}}.tpl-side{{border-left:0;border-top:2px dashed #cbd5e1;padding-left:0;padding-top:18px}}.tpl-tabs{{overflow:auto}}}}
    

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
    <div class='contract-detail-wrap'>
      <div class='back-top'><a class='c-btn' href='/admin/contratacion?sec=plantillas'>← Atrás</a></div>
      <h2 class='detail-title'>Detalle Plantilla</h2>
      <div class='template-head'>
        <div>
          <h1>{html.escape(pl['nombre_plantilla'] or '')} <a class='icon-btn' title='Editar plantilla' href='{edit_url}'>✎</a></h1>
          <div class='tpl-line'><b>Descripción:</b>{html.escape(pl['descripcion'] or '')}</div>
          <div class='tpl-line'><b>Nombre Archivo:</b>{html.escape(pl['archivo_nombre'] or 'SIN ARCHIVO')} <a class='icon-btn' title='Cargar o cambiar plantilla' href='{edit_url}'>⌕</a></div>
          {file_btn}
          <div class='meta-box'><span><b>Fecha de Creación:</b> {html.escape(pl['fecha_creacion'] or '')}</span><span><b>Creado por:</b> {html.escape(pl['creado_por'] or '')}</span></div>
        </div>
        <div class='tpl-side'>
          <div class='tpl-line'><b>Versión:</b>{html.escape(pl['version'] or 'Version 01')}</div>
          <div class='tpl-line'><b>Modo selección:</b>{'Usar criterios de selección que cumplan los datos del trabajador' if condicion_habilitada else 'Sin criterios / plantilla general'}</div>
          <div class='tpl-line'><b>Esquema:</b>{html.escape(pl['esquema'] or 'Trabajador Contrato Laboral')} <a class='icon-btn' title='Ver campos de esquema' href='{url_for('contratacion_campos_esquema',pid=pid)}'>⌕</a></div>
          <div class='tpl-line'><b>Tipo Documento:</b>{html.escape(pl['tipo_documento'] or '')}</div>
          <div class='tpl-line'><b>Estado:</b><span class='{'green-dot' if pl['activo'] else 'red-dot'}'></span>{estado}</div>
        </div>
      </div>
      {tabs}{body}
    </div>
    """
    return render_page(content, active='Gestion Contratacion:plantillas')


@app.route('/admin/contratacion/plantilla/<int:pid>/campo', defaults={'campo_id': None}, methods=['GET','POST'])
@app.route('/admin/contratacion/plantilla/<int:pid>/campo/<int:campo_id>', methods=['GET','POST'])
@admin_required
def contratacion_campo_editar(pid, campo_id=None):
    with db() as con:
        pl=con.execute('SELECT * FROM contratacion_plantillas WHERE id=?',(pid,)).fetchone()
        campo=con.execute('SELECT * FROM contratacion_plantilla_campos WHERE id=? AND plantilla_id=?',(campo_id,pid)).fetchone() if campo_id else None
    if not pl: abort(404)
    if request.method == 'POST':
        nombre=clean(request.form.get('nombre_campo')) or 'Campo Manual'
        origen=clean(request.form.get('campo_origen')) or re.sub(r'[^A-Za-z0-9_]','',nombre.replace(' ',''))
        tipo_campo=clean(request.form.get('tipo_campo')) or 'Origen de Datos'
        tipo_dato=clean(request.form.get('tipo_dato')) or 'Text'
        requerido=clean(request.form.get('requerido')) or 'SI'
        activo=1 if request.form.get('activo','1')=='1' else 0
        descripcion=clean(request.form.get('descripcion'))
        valor_default=clean(request.form.get('valor_default'))
        opciones=clean(request.form.get('opciones'))
        with db() as con:
            if campo_id:
                con.execute("UPDATE contratacion_plantilla_campos SET descripcion=?, tipo_campo=?, nombre_campo=?, campo_origen=?, tipo_dato=?, requerido=?, activo=?, valor_default=?, opciones=? WHERE id=? AND plantilla_id=?", (descripcion,tipo_campo,nombre,origen,tipo_dato,requerido,activo,valor_default,opciones,campo_id,pid))
            else:
                con.execute("INSERT INTO contratacion_plantilla_campos(plantilla_id,descripcion,tipo_campo,nombre_campo,campo_origen,tipo_dato,requerido,activo,valor_default,opciones) VALUES(?,?,?,?,?,?,?,?,?,?)", (pid,descripcion,tipo_campo,nombre,origen,tipo_dato,requerido,activo,valor_default,opciones))
            con.commit()
        flash('Campo guardado y enlazado a la plantilla Word.', 'ok')
        return redirect(url_for('contratacion_plantilla_detalle',pid=pid,tab='campos'))
    def cv(k, default=''):
        try: return campo[k] if campo and k in campo.keys() else default
        except Exception: return default
    campo_options=''.join([f"<option value='{html.escape(n)}' data-origen='{html.escape(o)}' data-tipo='{html.escape(td)}'>{html.escape(n)} / {{{{{html.escape(o)}}}}}</option>" for n,o,td in CONTRATACION_CAMPOS_CORRESPONDENCIA])
    content=f"""<div class='cond-overlay'><div class='cond-modal'><h2>Campo de Plantilla</h2><p class='muted2'>Configura campos automaticos desde Excel, manuales o desplegables. En Word usa {{{{CampoOrigen}}}}.</p><form method='post' class='cond-form'><label>Campo base</label><input list='campos_base' id='campoBase' oninput='autoCampoBase()'><datalist id='campos_base'>{campo_options}</datalist><label>Nombre visible</label><input name='nombre_campo' id='nombreCampo' value='{html.escape(cv('nombre_campo'))}' required><label>Campo origen Word</label><input name='campo_origen' id='campoOrigen' value='{html.escape(cv('campo_origen'))}' required><label>Tipo campo</label><select name='tipo_campo'><option {'selected' if cv('tipo_campo')=='Origen de Datos' else ''}>Origen de Datos</option><option {'selected' if cv('tipo_campo')=='Manual' else ''}>Manual</option><option {'selected' if cv('tipo_campo')=='Desplegable' else ''}>Desplegable</option></select><label>Tipo dato</label><select name='tipo_dato' id='tipoDato'><option {'selected' if cv('tipo_dato')=='Text' else ''}>Text</option><option {'selected' if cv('tipo_dato')=='Number' else ''}>Number</option><option {'selected' if cv('tipo_dato')=='DateTime' else ''}>DateTime</option></select><label>Valor manual/default</label><input name='valor_default' value='{html.escape(cv('valor_default'))}' placeholder='Ej: Básico / 1500 / 18.05.2026'><label>Opciones desplegable</label><textarea name='opciones' placeholder='Una opción por línea: Básico&#10;Intermedio&#10;Avanzado'>{html.escape(cv('opciones'))}</textarea><label>Requerido</label><select name='requerido'><option {'selected' if cv('requerido')=='SI' else ''}>SI</option><option {'selected' if cv('requerido')=='NO' else ''}>NO</option></select><label>Estado</label><select name='activo'><option value='1' {'selected' if str(cv('activo','1'))!='0' else ''}>ACTIVO</option><option value='0' {'selected' if str(cv('activo','1'))=='0' else ''}>INACTIVO</option></select><label>Descripcion</label><textarea name='descripcion'>{html.escape(cv('descripcion'))}</textarea><div class='modal-actions'><a class='c-btn gray' href='{url_for('contratacion_plantilla_detalle',pid=pid,tab='campos')}'>Cancelar</a><button class='c-btn'>Guardar campo</button></div></form></div></div><script>function autoCampoBase(){{const inp=document.getElementById('campoBase'); const op=[...document.querySelectorAll('#campos_base option')].find(o=>o.value===inp.value); if(!op)return; document.getElementById('nombreCampo').value=op.value.split(' / ')[0]; document.getElementById('campoOrigen').value=op.dataset.origen; document.getElementById('tipoDato').value=op.dataset.tipo||'Text';}}</script>"""
    return render_page(content, active='Gestion Contratacion:plantillas')
@app.route('/admin/contratacion/plantilla/<int:pid>/condicion', defaults={'cid': None}, methods=['GET','POST'])
@app.route('/admin/contratacion/plantilla/<int:pid>/condicion/<int:cid>', methods=['GET','POST'])
@admin_required
def contratacion_condicion_editar(pid, cid=None):
    with db() as con:
        pl=con.execute('SELECT * FROM contratacion_plantillas WHERE id=?',(pid,)).fetchone()
        cond=con.execute('SELECT * FROM contratacion_plantilla_condiciones WHERE id=? AND plantilla_id=?',(cid,pid)).fetchone() if cid else None
    if not pl: abort(404)
    condicion_habilitada = (pl['condicion'] or '').upper() == 'CONDICIONES'
    if request.method == 'POST':
        if not condicion_habilitada:
            flash('Primero habilita CONDICIONES en Editar Plantilla.', 'error')
            return redirect(url_for('contratacion_plantilla_detalle', pid=pid, tab='condiciones'))
        nombre_campo=clean(request.form.get('nombre_campo')) or 'Planilla'
        operador=clean(request.form.get('condicion')) or '='
        valor=clean(request.form.get('valor'))
        activo=1 if request.form.get('activo','1') == '1' else 0
        if operador not in CONDICION_OPERADORES:
            operador='='
        with db() as con:
            if cid:
                con.execute('UPDATE contratacion_plantilla_condiciones SET nombre_campo=?, condicion=?, valor=?, activo=? WHERE id=? AND plantilla_id=?', (nombre_campo,operador,valor,activo,cid,pid))
                flash('Condición actualizada correctamente.', 'ok')
            else:
                con.execute('INSERT INTO contratacion_plantilla_condiciones(plantilla_id,nombre_campo,condicion,valor,activo,fecha_registro,creado_por) VALUES(?,?,?,?,?,?,?)', (pid,nombre_campo,operador,valor,activo,now_txt(),marca_carga(session.get('admin_user','admin'))))
                flash('Condición creada correctamente.', 'ok')
            con.commit()
        return redirect(url_for('contratacion_plantilla_detalle', pid=pid, tab='condiciones'))
    campo_actual = cond['nombre_campo'] if cond else 'Planilla'
    operador_actual = cond['condicion'] if cond else '='
    valor_actual = cond['valor'] if cond else ''
    activo_actual = int(cond['activo']) if cond else 1
    campo_options=''.join([f"<option value='{html.escape(n)}' {'selected' if n==campo_actual else ''}>{html.escape(n)}</option>" for n,ori,td in CONTRATACION_CAMPOS_CORRESPONDENCIA])
    op_options=''.join([f"<option value='{html.escape(o)}' {'selected' if o==operador_actual else ''}>{html.escape(o)}</option>" for o in CONDICION_OPERADORES])
    valores=[]
    for arr in VALORES_CONDICION.values():
        valores += arr
    valores=sorted(set(valores))
    data_values=''.join([f"<option value='{html.escape(v)}'></option>" for v in valores])
    valores_js = {k:v for k,v in VALORES_CONDICION.items()}
    # Permite que el desplegable funcione por nombre visible y por CampoOrigen Word.
    for nom, ori, td in CONTRATACION_CAMPOS_CORRESPONDENCIA:
        if nom in VALORES_CONDICION and ori not in valores_js:
            valores_js[ori] = VALORES_CONDICION[nom]
        if ori in VALORES_CONDICION and nom not in valores_js:
            valores_js[nom] = VALORES_CONDICION[ori]
    tipos_js = {nom:td for nom,ori,td in CONTRATACION_CAMPOS_CORRESPONDENCIA}
    tipos_js.update({ori:td for nom,ori,td in CONTRATACION_CAMPOS_CORRESPONDENCIA})
    import json
    valores_json = json.dumps(valores_js, ensure_ascii=False)
    tipos_json = json.dumps(tipos_js, ensure_ascii=False)
    disabled_attr = '' if condicion_habilitada else 'disabled'
    content=f"""
    <style>
      .cond-overlay{{min-height:calc(100vh - 40px);display:flex;align-items:flex-start;justify-content:center;padding:22px;background:rgba(15,23,42,.62);margin:-20px -24px -40px}}
      .cond-modal{{width:min(780px,96vw);background:#fff;color:#111827;border-radius:18px;border:1px solid #dbe3ee;box-shadow:0 32px 90px rgba(0,0,0,.40);overflow:hidden}}
      .cond-head2{{display:flex;justify-content:space-between;align-items:center;padding:22px 26px;border-bottom:1px solid #e5eaf1;background:linear-gradient(135deg,#fff,#f8fafc)}}
      .cond-head2 h1{{margin:0;font-size:27px;font-weight:950;color:#0f172a}}
      .close-x{{width:38px;height:38px;border-radius:12px;background:#f1f5f9;color:#6b7280!important;border:1px solid #dbe3ee;display:flex;align-items:center;justify-content:center;text-decoration:none;font-size:26px;font-weight:950}}
      .cond-form{{padding:28px 42px 32px;display:grid;grid-template-columns:160px 1fr;gap:14px 16px;align-items:center}}
      .cond-form label{{text-align:right;font-size:18px;color:#374151;font-weight:850}}
      .cond-form select,.cond-form input{{height:48px;background:#fff!important;color:#111827!important;border:1px solid #cbd5e1;border-radius:12px;padding:10px 13px;font-size:17px;font-weight:750;outline:none}}
      .cond-form select:focus,.cond-form input:focus{{border-color:#60a5fa;box-shadow:0 0 0 4px rgba(96,165,250,.22)}}
      .cond-form select:disabled{{background:#e5e7eb!important;color:#6b7280!important;cursor:not-allowed}}
      .hint{{grid-column:2/3;background:#fff7ed;border:1px solid #fed7aa;color:#7c2d12;border-radius:12px;padding:12px 14px;font-weight:850}}
      .cond-actions{{grid-column:2/3;display:flex;gap:12px;justify-content:flex-end;margin-top:10px}}
      .c-btn{{display:inline-flex;align-items:center;justify-content:center;border:0;border-radius:12px;padding:12px 18px;background:linear-gradient(135deg,#ff963b,#ff7a1a);color:#fff!important;font-weight:950;text-decoration:none;cursor:pointer;font-size:16px}}
      .c-btn.gray{{background:#e7ebf2;color:#111827!important}}
      @media(max-width:700px){{.cond-form{{grid-template-columns:1fr;padding:22px}}.cond-form label{{text-align:left}}.hint,.cond-actions{{grid-column:1/2}}}}
    

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
    <div class='cond-overlay'><div class='cond-modal'>
      <div class='cond-head2'><h1>{'Editar' if cid else 'Crear'} Condición</h1><a class='close-x' href='{url_for('contratacion_plantilla_detalle',pid=pid,tab='condiciones')}'>×</a></div>
      <form method='post' class='cond-form'>
        <label>Nombre Campo:</label><select name='nombre_campo' id='nombre_campo' {disabled_attr}>{campo_options}</select>
        <label>Condición:</label><select name='condicion' {disabled_attr}>{op_options}</select>
        <label>Valor:</label><input name='valor' id='valor_condicion' list='valores_condicion' value='{html.escape(valor_actual)}' {disabled_attr} placeholder='Selecciona o escribe el valor'>
        <datalist id='valores_condicion'>{data_values}</datalist>
        <label>Estado:</label><select name='activo'><option value='1' {'selected' if activo_actual else ''}>Activo</option><option value='0' {'selected' if not activo_actual else ''}>Inactivo</option></select>
        {'' if condicion_habilitada else '<div class="hint">El campo condición está bloqueado porque la plantilla está en SIN CONDICIONES. Activa CONDICIONES en Editar Plantilla.</div>'}
        <div class='cond-actions'><button class='c-btn' {'disabled' if not condicion_habilitada else ''}>Guardar</button><a class='c-btn gray' href='{url_for('contratacion_plantilla_detalle',pid=pid,tab='condiciones')}'>Cerrar</a></div>
      </form>
      <script>
        const valoresPorCampo = {valores_json};
        const tiposPorCampo = {tipos_json};
        const campo = document.getElementById('nombre_campo');
        const dl = document.getElementById('valores_condicion');
        const valorInput = document.getElementById('valor_condicion');
        function cargarValores(){{
          if(!campo || !dl || !valorInput) return;
          const vals = valoresPorCampo[campo.value] || [];
          dl.innerHTML = vals.map(v => `<option value="${{v}}"></option>`).join('');
          const tipo = (tiposPorCampo[campo.value] || '').toLowerCase();
          if(tipo.includes('number') || tipo.includes('numeric')){{ valorInput.type='number'; valorInput.removeAttribute('list'); }}
          else if(tipo.includes('date')){{ valorInput.type='date'; valorInput.removeAttribute('list'); }}
          else {{ valorInput.type='text'; valorInput.setAttribute('list','valores_condicion'); }}
          valorInput.placeholder = vals.length ? 'Seleccione del desplegable o escriba manualmente' : 'Escriba valor manual';
        }}
        if(campo) campo.addEventListener('change', cargarValores);
        cargarValores();
      </script>
    </div></div>
    """
    return render_page(content, active='Gestion Contratacion:plantillas')

@app.route('/admin/contratacion/plantilla/<int:pid>/condicion/<int:cid>/eliminar', methods=['POST'])
@admin_required
def contratacion_condicion_eliminar(pid, cid):
    with db() as con:
        con.execute('DELETE FROM contratacion_plantilla_condiciones WHERE id=? AND plantilla_id=?', (cid,pid))
        con.commit()
    flash('Condición eliminada correctamente.', 'ok')
    return redirect(url_for('contratacion_plantilla_detalle', pid=pid, tab='condiciones'))



@app.route('/admin/firma/configuracion', methods=['GET','POST'])
@admin_required
def firma_configuracion():
    """Panel base para dejar preparada la integración con reconocimiento facial, RENIEC o firma digital."""
    if request.method == 'POST':
        proveedor = clean(request.form.get('proveedor')) or 'INTERNO'
        modo = clean(request.form.get('modo')) or 'RECONOCIMIENTO FACIAL / FIRMA DIGITAL'
        reniec_activo = 1 if request.form.get('reniec_activo') == '1' else 0
        firma_digital_activo = 1 if request.form.get('firma_digital_activo') == '1' else 0
        url_api = clean(request.form.get('url_api'))
        token_ref = clean(request.form.get('token_ref'))
        observacion = clean(request.form.get('observacion'))
        with db() as con:
            con.execute("INSERT INTO firma_configuracion(proveedor,modo,reniec_activo,firma_digital_activo,url_api,token_ref,observacion,fecha_registro) VALUES(?,?,?,?,?,?,?,?)", (proveedor,modo,reniec_activo,firma_digital_activo,url_api,token_ref,observacion,now_txt()))
            con.commit()
        flash('Configuración de firma guardada. Para RENIEC o firma digital real se debe colocar proveedor autorizado/API oficial.', 'ok')
        return redirect(url_for('firma_configuracion'))
    with db() as con:
        cfg = con.execute('SELECT * FROM firma_configuracion ORDER BY id DESC LIMIT 1').fetchone()
    content=f"""
    <section class='grid'>
      <div class='card span-12'><h1>Firma / Reconocimiento Facial</h1>
        <p class='muted'>Panel preparado para controlar documentos que serán firmados por celular. La conexión RENIEC o firma digital debe realizarse con proveedor autorizado y credenciales oficiales.</p>
        <form method='post' class='form-grid'>
          <div class='field'><label>Proveedor</label><select name='proveedor'><option>INTERNO</option><option>RENIEC / PROVEEDOR AUTORIZADO</option><option>FIRMA DIGITAL CERTIFICADA</option><option>BIOMETRÍA TERCERO</option></select></div>
          <div class='field'><label>Modo</label><select name='modo'><option>RECONOCIMIENTO FACIAL / FIRMA DIGITAL</option><option>SOLO RECONOCIMIENTO FACIAL</option><option>SOLO FIRMA DIGITAL</option><option>FIRMA CON OTP + EVIDENCIA</option></select></div>
          <div class='field'><label>RENIEC activo</label><select name='reniec_activo'><option value='0'>NO</option><option value='1'>SI</option></select></div>
          <div class='field'><label>Firma digital activa</label><select name='firma_digital_activo'><option value='0'>NO</option><option value='1'>SI</option></select></div>
          <div class='field span-6'><label>URL API / Endpoint</label><input name='url_api' value='{html.escape(row_get(cfg,'url_api'))}' placeholder='https://api.proveedor.com/...'></div>
          <div class='field span-6'><label>Token / referencia segura</label><input name='token_ref' value='{html.escape(row_get(cfg,'token_ref'))}' placeholder='Guardar token real como variable de entorno'></div>
          <div class='field span-12'><label>Observación / flujo</label><textarea name='observacion' rows='4'>{html.escape(row_get(cfg,'observacion'))}</textarea></div>
          <button class='btn-yellow'>Guardar configuración</button>
        </form>
      </div>
      <div class='card span-12'><h2>Flujo recomendado</h2>
        <div class='table-wrap'><table><tr><th>Paso</th><th>Control</th><th>Resultado</th></tr>
        <tr><td>1</td><td>Generar Word/PDF combinado</td><td>Documento con campos «CampoOrigen» llenos</td></tr>
        <tr><td>2</td><td>Enviar enlace al celular</td><td>Trabajador revisa documento</td></tr>
        <tr><td>3</td><td>Validación biométrica / OTP / firma digital</td><td>Evidencia y fecha de firma</td></tr>
        <tr><td>4</td><td>Archivado automático</td><td>Ficha del trabajador y renovaciones</td></tr>
        </table></div>
      </div>
    </section>
    """
    return render_page(content, active='Gestion Contratacion')

@app.route('/admin/contratacion/plantilla/<int:pid>/campos_esquema')
@admin_required
def contratacion_campos_esquema(pid):
    dni = normalizar_dni(request.args.get('dni'))
    with db() as con:
        pl = con.execute('SELECT * FROM contratacion_plantillas WHERE id=?', (pid,)).fetchone()
        trabajador = con.execute('SELECT * FROM trabajadores WHERE dni=?', (dni,)).fetchone() if dni else con.execute('SELECT * FROM trabajadores ORDER BY nombre LIMIT 1').fetchone()
    if not pl:
        abort(404)
    valores = valores_esquema_desde_trabajador(trabajador)
    rows = ''.join([f"<tr><td><code>«{html.escape(campo)}»</code><br><small><code>{{{{{html.escape(campo)}}}}}</code></small></td><td>{html.escape(valor)}</td></tr>" for campo, valor in valores])
    dni_val = html.escape(dni or row_get(trabajador, 'dni') or '')
    content = f"""
    <style>
      .schema-overlay{{min-height:calc(100vh - 40px);display:flex;justify-content:center;align-items:flex-start;padding:18px;background:rgba(15,23,42,.55);margin:-20px -24px -40px}}
      .schema-modal{{width:min(930px,96vw);background:#fff;color:#111827!important;border-radius:12px;border:1px solid #dbe1e8;box-shadow:0 24px 70px rgba(0,0,0,.35);overflow:hidden}}
      .schema-head{{display:flex;justify-content:space-between;align-items:center;padding:20px 24px;border-bottom:1px solid #e5e7eb}}
      .schema-head h1{{margin:0;font-size:28px;font-weight:950;color:#111827}}
      .close-x{{width:42px;height:42px;border-radius:10px;background:#fff;color:#8b9097!important;border:2px solid #e5e7eb;display:flex;align-items:center;justify-content:center;text-decoration:none;font-size:30px;font-weight:900}}
      .schema-tools{{display:flex;justify-content:space-between;gap:14px;align-items:center;padding:18px 24px;background:#fff;flex-wrap:wrap}}
      .schema-search{{display:flex;gap:10px;align-items:center;flex-wrap:wrap}}
      .schema-search input{{width:160px;background:#fff!important;color:#111!important;border:1px solid #cbd5e1;border-radius:10px;padding:11px 12px;font-weight:850}}
      .schema-btn{{display:inline-flex;align-items:center;gap:8px;border:0;border-radius:10px;background:#eef2f7;color:#111827!important;padding:12px 16px;text-decoration:none;font-weight:950;cursor:pointer}}
      .schema-btn.orange{{background:#ff963b;color:#fff!important}}
      .schema-note{{padding:0 24px 12px;color:#64748b;font-weight:800}}
      .schema-table-wrap{{max-height:560px;overflow:auto;border-top:1px solid #e5e7eb}}
      .schema-table{{width:100%;border-collapse:collapse;color:#111827!important;background:#fff!important}}
      .schema-table th{{position:sticky;top:0;background:#f8fafc;color:#111827;text-align:left;border:1px solid #e5e7eb;padding:15px;font-size:18px}}
      .schema-table td{{border:1px solid #e5e7eb;padding:13px 15px;font-size:17px;font-weight:760;background:#fff;color:#111827!important}}
      .schema-table tr:nth-child(even) td{{background:#f6f6f6}}
      code{{background:#eef2ff;border:1px solid #dbeafe;border-radius:8px;padding:4px 7px;color:#1e3a8a}}
      @media(max-width:700px){{.schema-tools{{align-items:stretch}}.schema-search input{{width:100%}}.schema-btn{{width:100%;justify-content:center}}}}
    

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
    <div class='schema-overlay'><div class='schema-modal'>
      <div class='schema-head'><h1>Campos de Esquema</h1><a class='close-x' href='{url_for('contratacion_plantilla_detalle',pid=pid,tab='campos')}'>×</a></div>
      <div class='schema-tools'>
        <form class='schema-search' method='get'>
          <b>DNI trabajador:</b><input name='dni' maxlength='8' value='{dni_val}' placeholder='Buscar DNI'>
          <button class='schema-btn orange'>🔎 Traer datos</button>
        </form>
        <a class='schema-btn' href='{url_for('contratacion_campos_esquema_excel',pid=pid,dni=dni_val)}'>⬇ Descargar</a>
      </div>
      <div class='schema-note'>Plantilla: <b>{html.escape(pl['nombre_plantilla'] or '')}</b> | Esquema: <b>{html.escape(pl['esquema'] or 'Trabajador Contrato Laboral')}</b>. Los valores se llenan con la base de trabajadores cuando el dato exista; si falta, se muestra el modelo.</div>
      <div class='schema-table-wrap'><table class='schema-table'><tr><th>Campo</th><th>Descripción / valor</th></tr>{rows}</table></div>
    </div></div>
    """
    return render_page(content, active='Gestion Contratacion:plantillas')


@app.route('/admin/contratacion/plantilla/<int:pid>/campos_esquema.xlsx')
@admin_required
def contratacion_campos_esquema_excel(pid):
    dni = normalizar_dni(request.args.get('dni'))
    with db() as con:
        pl = con.execute('SELECT * FROM contratacion_plantillas WHERE id=?', (pid,)).fetchone()
        trabajador = con.execute('SELECT * FROM trabajadores WHERE dni=?', (dni,)).fetchone() if dni else con.execute('SELECT * FROM trabajadores ORDER BY nombre LIMIT 1').fetchone()
    if not pl:
        abort(404)
    wb = Workbook(); ws = wb.active; ws.title = 'Campos Esquema'
    ws.append(['Campo Word', 'Campo Sistema', 'Descripción / Valor'])
    for campo, valor in valores_esquema_desde_trabajador(trabajador):
        ws.append(['«' + campo + '»', '{{' + campo + '}}', valor])
    ws.column_dimensions['A'].width = 34; ws.column_dimensions['B'].width = 34; ws.column_dimensions['C'].width = 58
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='1F2937')
        cell.alignment = Alignment(horizontal='center')
    ws.freeze_panes = 'A2'
    out = PERSIST_DIR / f"Campos_Esquema_{pid}_{now_file()}.xlsx"
    wb.save(out)
    return send_file(out, as_attachment=True, download_name=f"Campos_Esquema_{pid}.xlsx")



@app.route('/admin/contratacion/plantilla/<int:pid>/generar')
@admin_required
def contratacion_plantilla_generar(pid):
    """Genera y descarga el Word final con campos {{CampoOrigen}} reemplazados por datos del trabajador."""
    dni = normalizar_dni(request.args.get('dni'))
    try:
        out, pl, trabajador, cumple, detalle = generar_docx_desde_plantilla(pid, dni)
    except Exception as e:
        flash(f'No se pudo generar el Word combinado: {e}', 'error')
        return redirect(url_for('contratacion_plantilla_detalle', pid=pid, tab='contenido'))
    if not cumple:
        flash('El trabajador no cumple las condiciones configuradas. Se descarga en modo revisión para que puedas validar.', 'error')
    nombre_trab = re.sub(r'[^A-Za-z0-9_ -]+', '', row_get(trabajador, 'nombre', 'TRABAJADOR')).strip() or 'TRABAJADOR'
    safe_pl = re.sub(r'[^A-Za-z0-9_ -]+', '', pl['nombre_plantilla'] or 'plantilla').strip() or 'plantilla'
    return send_file(out, as_attachment=True, download_name=f"{safe_pl}_{dni or nombre_trab}.docx")

@app.route('/admin/contratacion/plantilla/<int:pid>/archivo')
@admin_required
def contratacion_plantilla_archivo(pid):
    """Descarga la plantilla. Si el archivo cargado no es Word, genera una copia Word editable con campos."""
    with db() as con:
        pl=con.execute('SELECT * FROM contratacion_plantillas WHERE id=?',(pid,)).fetchone()
        campos=con.execute('SELECT * FROM contratacion_plantilla_campos WHERE plantilla_id=? ORDER BY id',(pid,)).fetchall()
    if not pl:
        abort(404)
    ruta = Path(pl['ruta_archivo']) if pl['ruta_archivo'] else None
    nombre = pl['archivo_nombre'] or f"{pl['nombre_plantilla'] or 'plantilla'}.docx"
    # Si el archivo original ya es Word, descargarlo directamente en modo adjunto.
    if ruta and ruta.exists() and ruta.suffix.lower() in ('.doc', '.docx'):
        return send_file(ruta, as_attachment=True, download_name=nombre)
    # Si no hay Word cargado, se genera un DOCX base editable con los campos de correspondencia.
    if Document is None:
        if ruta and ruta.exists():
            return send_file(ruta, as_attachment=True, download_name=nombre)
        abort(404)
    doc = Document()
    doc.add_heading(pl['nombre_plantilla'] or 'Plantilla de contrato', level=1)
    doc.add_paragraph(f"Tipo Documento: {pl['tipo_documento'] or ''}")
    doc.add_paragraph(f"Esquema: {pl['esquema'] or 'Trabajador Contrato Laboral'}")
    doc.add_paragraph(f"Versión: {pl['version'] or 'Version 01'}")
    doc.add_paragraph('')
    doc.add_paragraph('CAMPOS DE CORRESPONDENCIA DISPONIBLES PARA WORD:')
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'Nombre Campo'
    hdr[1].text = 'Campo Origen'
    hdr[2].text = 'Uso en plantilla'
    for c in campos:
        row = table.add_row().cells
        row[0].text = str(c['nombre_campo'] or '')
        row[1].text = str(c['campo_origen'] or '')
        row[2].text = '«' + str(c['campo_origen'] or '') + '»'
    safe_name = re.sub(r'[^A-Za-z0-9_ -]+', '', pl['nombre_plantilla'] or 'plantilla').strip() or 'plantilla'
    out = PERSIST_DIR / f"{safe_name}_{pid}_{now_file()}.docx"
    doc.save(out)
    return send_file(out, as_attachment=True, download_name=f"{safe_name}.docx")

@app.route('/admin/contratacion/plantilla/<int:pid>/historial')
@admin_required
def contratacion_plantilla_historial(pid):
    """Ventana tipo lupita: historial real de cargas de plantillas de contrato."""
    with db() as con:
        pl = con.execute('SELECT * FROM contratacion_plantillas WHERE id=?', (pid,)).fetchone()
        historial = con.execute("""
            SELECT id, fecha_creacion, fecha_actualizacion, creado_por, nombre_plantilla,
                   tipo_documento, esquema, condicion, version, archivo_nombre, ruta_archivo, activo
              FROM contratacion_plantillas
             WHERE UPPER(COALESCE(tipo_documento,'')) = UPPER(COALESCE((SELECT tipo_documento FROM contratacion_plantillas WHERE id=?),''))
                OR UPPER(COALESCE(nombre_plantilla,'')) LIKE '%' || UPPER(COALESCE((SELECT nombre_plantilla FROM contratacion_plantillas WHERE id=?),'')) || '%'
             ORDER BY id DESC
             LIMIT 250
        """, (pid, pid)).fetchall()
    if not pl:
        abort(404)
    def h(v):
        return html.escape(str(v or ''))
    rows = ''.join([f"""
      <tr>
        <td>{h(r['fecha_actualizacion'] or r['fecha_creacion'])}</td>
        <td>{h(r['creado_por'] or 'Admin_AQ...')}</td>
        <td>{h(r['nombre_plantilla'])}</td>
        <td>{h(r['tipo_documento'])}</td>
        <td>{h(r['archivo_nombre'] or (r['nombre_plantilla'] + '.docx'))}</td>
        <td><span class='state {'ok' if r['activo'] else 'bad'}'>{'ACTIVO' if r['activo'] else 'INACTIVO'}</span></td>
        <td><a class='mini-download' href='{url_for('contratacion_plantilla_archivo', pid=r['id'])}'>⬇ Word</a></td>
      </tr>""" for r in historial])
    content = f"""
    <style>
      .hist-overlay{{min-height:calc(100vh - 70px);display:flex;align-items:flex-start;justify-content:center;padding:18px 10px;background:rgba(17,24,39,.55);margin:-20px -24px -40px;overflow:auto}}
      .hist-modal{{width:min(1120px,97vw);max-height:calc(100vh - 38px);background:#fff;color:#111827!important;border-radius:8px;border:1px solid #dbe1e8;box-shadow:0 22px 65px rgba(0,0,0,.35);overflow:hidden}}
      .hist-head{{display:flex;justify-content:space-between;align-items:center;padding:16px 22px;border-bottom:1px solid #dbe1e8;background:#fff}}
      .hist-head h1{{margin:0;font-size:25px;font-weight:950;color:#111827!important}}
      .close-x{{width:34px;height:34px;border-radius:8px;background:#fff;color:#697386!important;border:2px solid #e5e7eb;display:flex;align-items:center;justify-content:center;text-decoration:none;font-size:25px;font-weight:900}}
      .hist-info{{padding:12px 22px;background:#f8fafc;color:#475569;font-weight:850;border-bottom:1px solid #e5e7eb}}
      .hist-table-wrap{{height:min(650px,72vh);overflow:auto;background:#fff}}
      .hist-table{{width:100%;min-width:1050px;border-collapse:collapse;background:#fff;color:#111827!important}}
      .hist-table th{{position:sticky;top:0;background:#f8fafc;color:#111827!important;border:1px solid #e5e7eb;padding:13px;text-align:left;font-size:16px;font-weight:950}}
      .hist-table td{{border:1px solid #e5e7eb;padding:13px;color:#111827!important;background:#fff;font-size:16px;font-weight:760;vertical-align:middle}}
      .hist-table tr:nth-child(even) td{{background:#f3f4f6}}
      .state{{border:1px solid #d1d5db;border-radius:999px;padding:6px 10px;font-weight:950;background:#fff;color:#16a34a}}
      .state.bad{{color:#e11d48;background:#fff1f2}}
      .mini-download{{display:inline-flex;align-items:center;gap:6px;background:#e5e9ef;color:#111827!important;text-decoration:none;border-radius:8px;padding:8px 12px;font-weight:950;white-space:nowrap}}
    

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
    <div class='hist-overlay'><div class='hist-modal'>
      <div class='hist-head'><h1>Editar Plantilla</h1><a class='close-x' href='{url_for('contratacion_plantilla_editar', pid=pid)}'>×</a></div>
      <div class='hist-info'>Historial de cargas de contratos relacionado con: <b>{h(pl['nombre_plantilla'])}</b>. Desde esta ventana puedes descargar cada plantilla en Word.</div>
      <div class='hist-table-wrap'><table class='hist-table'>
        <tr><th>Fecha Registro</th><th>Creado por</th><th>Nombre Archivo</th><th>Tipo Documento</th><th>Identificador</th><th>Estado</th><th>Descarga</th></tr>
        {rows or '<tr><td colspan="8">No hay historial de cargas.</td></tr>'}
      </table></div>
    </div></div>
    """
    return render_page(content, active='Gestion Contratacion:plantillas')

@app.route('/admin/contratacion/plantilla/<int:pid>/editar')
@admin_required
def contratacion_plantilla_editar(pid):
    with db() as con:
        pl=con.execute('SELECT * FROM contratacion_plantillas WHERE id=?',(pid,)).fetchone()
    if not pl: abort(404)
    tipos_opts = ['CONTRATO TRABAJADOR','CONTRATO TRABAJADOR(RENOVACIÓN)','ACUERDO PREFERENCIAL','AUTODECLARACION BUENAS PRACTICAS','CARGO ENTREGA RENOVACION','ELECCIÓN DE BENEFICIOS SOCIALES','NOTA DE CARGO','CARTA DE COMPROMISO','CARGO DE ENTREGA']
    tipo_options=''.join([f"<option value='{html.escape(x)}' {'selected' if (pl['tipo_documento'] or '')==x else ''}>{html.escape(x)}</option>" for x in tipos_opts])
    modo_actual = 'criterios' if (pl['condicion'] or '').upper() == 'CONDICIONES' else 'todos'
    archivo_actual = html.escape(pl['archivo_nombre'] or 'Sin archivo cargado')
    content=f"""
    <style>
      .edit-overlay{{min-height:calc(100vh - 70px);display:flex;align-items:flex-start;justify-content:center;padding:12px 10px;background:rgba(3,10,18,.72);margin:-20px -24px -40px;overflow:auto}}
      .modal-page{{width:min(640px,94vw);max-height:calc(100vh - 28px);overflow:auto;background:#fff;color:#111827!important;border-radius:10px;border:1px solid #dbe1e8;box-shadow:0 22px 65px rgba(0,0,0,.42)}}
      .modal-page *{{box-sizing:border-box}}
      .modal-head{{position:sticky;top:0;z-index:2;display:flex;justify-content:space-between;align-items:center;padding:12px 18px;border-bottom:1px solid #dbe1e8;background:#fff}}
      .modal-head h1{{margin:0;font-size:21px;font-weight:950;color:#111827!important;letter-spacing:-.2px}}
      .close-x{{width:38px;height:38px;border-radius:8px;background:#fff;color:#8b9097!important;border:2px solid #e5e7eb;display:flex;align-items:center;justify-content:center;text-decoration:none;font-size:28px;font-weight:900}}
      .edit-grid{{display:grid;grid-template-columns:150px minmax(0,1fr);gap:9px 12px;align-items:center;padding:16px 22px}}
      .edit-grid label{{color:#1f2937!important;font-size:15px;text-align:right;font-weight:950;line-height:1.18}}
      .edit-grid input,.edit-grid select,.edit-grid textarea{{background:#fff!important;color:#111827!important;border:1px solid #b8c2cf;border-radius:8px;padding:8px 11px;width:100%;font-size:14px;font-weight:900;min-height:36px;box-shadow:none!important}}
      .edit-grid textarea{{min-height:54px;resize:vertical}}
      .locked-field{{background:#eef1f5!important;color:#6b7280!important;cursor:not-allowed;border-color:#d8dee8!important}}
      .schema-line{{display:grid;grid-template-columns:1fr 42px 42px;gap:8px}}
      .schema-line .icon-btn{{height:38px;min-width:42px;border-radius:8px;background:#e9edf4;color:#111827!important;border:1px solid #d8dee8;text-decoration:none;display:flex;align-items:center;justify-content:center;font-size:20px;font-weight:950}}
      .upload-box{{background:#eef1f5;border-radius:7px;padding:0;color:#111827;overflow:hidden}}
      .upload-row{{display:flex;align-items:center;gap:10px;padding:10px 12px;white-space:nowrap;overflow:auto}}
      .upload-row input[type=file]{{border:0!important;padding:0!important;background:transparent!important;min-height:auto;font-size:15px}}
      .warn{{background:#fff3cd;border:1px solid #ffe08a;color:#6b4a00!important;border-radius:6px;margin:0 0 8px;padding:9px;text-align:center;font-size:13px;font-weight:950}}
      .actual-file{{display:block;margin-top:8px;color:#111827!important;font-weight:900}}
      .modal-actions{{position:sticky;bottom:0;background:#fff;display:flex;justify-content:flex-end;gap:10px;padding:12px 22px 16px;border-top:1px solid #eef2f7}}
      .c-btn{{display:inline-flex;align-items:center;justify-content:center;border:0;border-radius:8px;padding:9px 15px;background:#ff963b;color:#fff!important;font-weight:950;text-decoration:none;cursor:pointer;font-size:14px}}
      .c-btn.gray{{background:#e5e9ef;color:#111827!important}}
      @media(max-width:720px){{.edit-overlay{{padding:8px}}.modal-page{{width:100%;max-height:calc(100vh - 16px)}}.edit-grid{{grid-template-columns:1fr;padding:18px}}.edit-grid label{{text-align:left;font-size:15px}}.modal-actions{{padding:14px 18px 20px}}}}
    

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
    <div class='edit-overlay'>
      <div class='modal-page'>
        <div class='modal-head'><h1>Editar Plantilla</h1><a class='close-x' href='{url_for('contratacion_plantilla_detalle',pid=pid)}'>×</a></div>
        <form method='post' action='/admin/contratacion?sec=plantillas' enctype='multipart/form-data'>
          <div class='edit-grid'>
            <input type='hidden' name='accion' value='plantilla'><input type='hidden' name='plantilla_id' value='{pid}'>
            <input type='hidden' name='tipo_documento' value='{html.escape(pl['tipo_documento'] or '')}'>
            <input type='hidden' name='esquema' value='{html.escape(pl['esquema'] or 'Trabajador Contrato Laboral')}'>
            <input type='hidden' name='nombre_plantilla' value='{html.escape(pl['nombre_plantilla'] or '')}'>

            <label>Tipo Documento</label><input class='locked-field' value='{html.escape(pl['tipo_documento'] or '')}' readonly>
            <label>Esquema</label><div class='schema-line'><input class='locked-field' value='{html.escape(pl['esquema'] or 'Trabajador Contrato Laboral')}' readonly><a class='icon-btn' title='Campos de esquema' href='{url_for('contratacion_campos_esquema',pid=pid)}'>⌕</a><a class='icon-btn' title='Historial de cargas' href='{url_for('contratacion_plantilla_historial',pid=pid)}'>🔍</a></div>
            <label>Nombre Plantilla</label><input class='locked-field' value='{html.escape(pl['nombre_plantilla'] or '')}' readonly>
            <label>Descripción</label><textarea name='descripcion'>{html.escape(pl['descripcion'] or '')}</textarea>
            <label>Modo de selección de la plantilla</label><select name='condicion'>
              <option value='SIN CONDICIONES' {'selected' if modo_actual=='todos' else ''}>Utilizar para todos los trabajadores</option>
              <option value='CONDICIONES' {'selected' if modo_actual=='criterios' else ''}>Usar criterios de selección que cumplan con los datos del trabajador</option>
            </select>
            <label>Estado</label><select name='activo'><option value='1' {'selected' if pl['activo'] else ''}>Activo</option><option value='0' {'selected' if not pl['activo'] else ''}>Inactivo</option></select>
            <label>Versión</label><input name='version' value='{html.escape(pl['version'] or 'Version 01')}'>
            <label>Plantilla contrato</label><div><div class='warn'><b>Advertencia!</b> Tipo de archivo permitido .doc y .docx para Word. También acepta .pdf.</div><div class='upload-box'><div class='upload-row'>⬆ <input type='file' name='archivo' accept='.doc,.docx,.pdf'></div></div><small class='actual-file'>Actual: {archivo_actual}</small><div style='margin-top:10px;display:flex;gap:8px;flex-wrap:wrap'><a class='c-btn gray' href='{url_for('contratacion_plantilla_historial',pid=pid)}'>🔍 Ver historial de cargas</a><a class='c-btn gray' href='{url_for('contratacion_plantilla_archivo',pid=pid)}'>⬇ Descargar Archivo Word</a></div></div>
          </div>
          <div class='modal-actions'><button class='c-btn'>Actualizar</button><a class='c-btn gray' href='{url_for('contratacion_plantilla_detalle',pid=pid)}'>Cerrar</a></div>
        </form>
      </div>
    </div>
    """
    return render_page(content, active='Gestion Contratacion:plantillas')

@app.route('/admin/contratacion/plantilla/<int:pid>/eliminar', methods=['POST'])
@admin_required
def contratacion_plantilla_eliminar(pid):
    """Elimina una plantilla de contrato y sus campos sin tocar contratos/documentos ya generados."""
    with db() as con:
        pl = con.execute('SELECT * FROM contratacion_plantillas WHERE id=?', (pid,)).fetchone()
        if not pl:
            flash('La plantilla no existe o ya fue eliminada.', 'error')
            return redirect(url_for('admin_contratacion', sec='plantillas'))
        con.execute('DELETE FROM contratacion_plantilla_campos WHERE plantilla_id=?', (pid,))
        con.execute('DELETE FROM contratacion_plantillas WHERE id=?', (pid,))
        con.commit()
    flash('Plantilla eliminada correctamente. Los documentos/contratos históricos no se borran.', 'ok')
    return redirect(url_for('admin_contratacion', sec='plantillas'))



def crear_token_firma():
    return uuid.uuid4().hex + uuid.uuid4().hex[:8]


def firma_url_token(token):
    # En Render/producción request.host_url será HTTPS. En local funciona por localhost.
    try:
        return url_for('firma_publica_token', token=token, _external=True)
    except Exception:
        return '/firma/' + str(token or '')


def guardar_selfie_firma(data_url, firma_id, dni, origen='WEB'):
    """Guarda captura PNG/JPG enviada desde getUserMedia. No realiza identificación biométrica interna;
    deja evidencia y trazabilidad para integración con proveedor facial/firma digital autorizado."""
    if not data_url or 'base64,' not in data_url:
        raise ValueError('No se recibió imagen válida desde la cámara.')
    meta, b64 = data_url.split('base64,', 1)
    ext = '.jpg' if 'jpeg' in meta.lower() or 'jpg' in meta.lower() else '.png'
    raw = base64.b64decode(b64)
    if len(raw) < 1000:
        raise ValueError('La imagen capturada parece estar vacía.')
    if len(raw) > 8 * 1024 * 1024:
        raise ValueError('La imagen supera el tamaño permitido.')
    carpeta = UPLOAD_DIR / 'firma_digital' / normalizar_dni(dni or 'SIN_DNI')
    carpeta.mkdir(parents=True, exist_ok=True)
    nombre = f"{now_file()}_firma_{firma_id}_{origen}{ext}"
    path = carpeta / secure_filename(nombre)
    path.write_bytes(raw)
    return str(path), hashlib.sha256(raw).hexdigest()


@app.route('/admin/firma/camara_demo')
@admin_required
def firma_camara_demo():
    content = """
    <section class='grid'><div class='card span-12'><h1>Prueba de cámara para firma facial</h1>
    <p class='muted'>Valida cámara en laptop o celular. No guarda datos.</p>
    <div class='camera-box'><video id='video' autoplay playsinline></video><canvas id='canvas' style='display:none'></canvas><img id='preview' style='display:none'></div>
    <div class='actions'><button class='btn-green' type='button' onclick='startCamera()'>Activar cámara</button><button class='btn-blue' type='button' onclick='capturePhoto()'>Capturar vista previa</button><a class='btn' href='/admin/contratacion?sec=firma'>Volver a firma</a></div><p id='camStatus' class='muted'></p>
    </div></section>
    <style>.camera-box{background:#0f172a;border-radius:18px;padding:16px;display:grid;place-items:center;min-height:360px}.camera-box video,.camera-box img{max-width:100%;border-radius:14px;background:#000}

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
    <script>
    let stream=null;
    async function startCamera(){const st=document.getElementById('camStatus');try{if(!navigator.mediaDevices||!navigator.mediaDevices.getUserMedia)throw new Error('Navegador sin getUserMedia');const tries=[{video:{facingMode:{ideal:'user'},width:{ideal:1280},height:{ideal:720}},audio:false},{video:{width:{ideal:640},height:{ideal:480}},audio:false},{video:true,audio:false}];let last=null;for(const cfg of tries){try{stream=await navigator.mediaDevices.getUserMedia(cfg);break;}catch(e){last=e;}}if(!stream)throw(last||new Error('Sin cámara'));const v=document.getElementById('video');v.srcObject=stream;await v.play();st.textContent='Cámara activa correctamente.';}catch(e){st.textContent='No se pudo activar la cámara: '+(e.name||e.message)+'. Revisa permisos, HTTPS/localhost o cierra otras apps que usen cámara.';}}
    function capturePhoto(){const v=document.getElementById('video'),c=document.getElementById('canvas'),img=document.getElementById('preview');if(!v.videoWidth){document.getElementById('camStatus').textContent='Primero activa la cámara.';return;}c.width=v.videoWidth;c.height=v.videoHeight;c.getContext('2d').drawImage(v,0,0);img.src=c.toDataURL('image/png');img.style.display='block';}
    document.addEventListener('DOMContentLoaded',()=>setTimeout(startCamera,600));
    </script>
    """
    return render_page(content, active='Gestion Contratacion:firma')


@app.route('/firma/<token>', methods=['GET','POST'])
def firma_publica_token(token):
    token = clean(token)
    with db() as con:
        sol = con.execute('SELECT * FROM firma_solicitudes WHERE firma_token=? ORDER BY id DESC LIMIT 1', (token,)).fetchone()
        if not sol:
            abort(404)
        doc = con.execute('SELECT * FROM contratacion_docs WHERE id=?', (sol['documento_id'],)).fetchone() if sol['documento_id'] else None
    if request.method == 'POST':
        acepta = 1 if request.form.get('acepta') == '1' else 0
        firma_texto = clean(request.form.get('firma_texto'))
        img_data = request.form.get('captura_base64')
        cam_origen = clean(request.form.get('camara_origen')) or 'CELULAR/LAPTOP'
        if not acepta:
            flash('Debe aceptar la declaración para registrar la firma.', 'err')
            return redirect(url_for('firma_publica_token', token=token))
        if not img_data:
            flash('Debe capturar una foto de evidencia desde la cámara.', 'err')
            return redirect(url_for('firma_publica_token', token=token))
        try:
            selfie_path, hash_ev = guardar_selfie_firma(img_data, sol['id'], sol['dni'], cam_origen)
            evidencia = f"SELFIE:{Path(selfie_path).name}; HASH:{hash_ev[:16]}; FIRMA:{firma_texto or sol['trabajador']}"
            with db() as con:
                con.execute('UPDATE firma_solicitudes SET estado=?, evidencia_ref=?, selfie_path=?, camara_origen=?, acepta_terminos=?, ip_registro=?, user_agent=?, hash_evidencia=?, fecha_captura=?, fecha_firma=?, validacion_estado=?, proveedor_respuesta=? WHERE id=?',
                            ('Firmado con evidencia facial', evidencia, selfie_path, cam_origen, 1, request.headers.get('X-Forwarded-For', request.remote_addr), request.headers.get('User-Agent','')[:500], hash_ev, now_txt(), now_txt(), 'EVIDENCIA CAPTURADA - PENDIENTE VALIDACIÓN PROVEEDOR', 'Captura local registrada. Integración biométrica externa configurable.', sol['id']))
                if sol['documento_id']:
                    con.execute("UPDATE contratacion_docs SET estado='FIRMADO CON EVIDENCIA FACIAL' WHERE id=?", (sol['documento_id'],))
                con.execute('INSERT INTO eventos_documento(dni,evento,fecha,detalle) VALUES(?,?,?,?)',(sol['dni'],'Contrato firmado disponible',now_txt(),'Su contrato firmado ya está visible en Gestión Contrato / Mis documentos contractuales.'))
                con.commit()
            flash('Firma registrada correctamente con evidencia de cámara.', 'ok')
            return redirect(url_for('firma_publica_token', token=token, ok='1'))
        except Exception as e:
            flash('No se pudo guardar la evidencia: ' + str(e), 'err')
            return redirect(url_for('firma_publica_token', token=token))
    ok = request.args.get('ok') == '1'
    estado = sol['estado'] or 'Pendiente'
    contrato_info = f"{html.escape(doc['tipo_doc'])} - {html.escape(doc['archivo_nombre'] or '')}" if doc else 'Contrato/documento de contratación'
    content = f"""
    <!doctype html><html lang='es'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>Firma digital / facial</title>
    <style>body{{margin:0;font-family:Arial,Helvetica,sans-serif;background:#0f172a;color:#111827}}.wrap{{max-width:920px;margin:0 auto;padding:22px}}.card{{background:#fff;border-radius:22px;padding:22px;box-shadow:0 18px 45px #0005;margin:18px 0}}h1{{margin:0;color:#0f172a}}.muted{{color:#64748b;line-height:1.5}}.badge{{display:inline-block;background:#fef3c7;color:#92400e;border-radius:999px;padding:8px 12px;font-weight:800}}video,img{{width:100%;min-height:330px;max-height:430px;background:#000;border-radius:18px;object-fit:cover}}.camera{{background:#111827;border-radius:22px;padding:14px}}button,.btn{{border:0;border-radius:12px;padding:13px 16px;font-weight:900;cursor:pointer;text-decoration:none;display:inline-block}}.primary{{background:#2563eb;color:white}}.green{{background:#16a34a;color:white}}.gray{{background:#475569;color:white}}input[type=text]{{width:100%;padding:13px;border:1px solid #cbd5e1;border-radius:12px;font-size:16px;box-sizing:border-box}}.actions{{display:flex;flex-wrap:wrap;gap:10px;margin-top:14px}}.ok{{background:#dcfce7;color:#166534;padding:12px;border-radius:12px;font-weight:900}}label{{font-weight:800}}canvas{{display:none}}@media(max-width:700px){{.wrap{{padding:12px}}.card{{padding:16px;border-radius:16px}}}}

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style></head>
    <body><div class='wrap'><div class='card'><h1>Firma de contrato con cámara</h1><p class='muted'>Trabajador: <b>{html.escape(sol['trabajador'] or '')}</b> · DNI: <b>{html.escape(sol['dni'] or '')}</b></p><p>Documento: <b>{contrato_info}</b></p><span class='badge'>Estado: {html.escape(estado)}</span></div>
    {"<div class='card ok'>✅ Firma registrada correctamente. Ya puede cerrar esta ventana.</div>" if ok else ""}
    <form method='post' class='card' onsubmit='return prepararEnvio()'><h2>1. Cámara facial y captura de evidencia</h2><p class='muted'>La cámara intenta encender automáticamente. En celular debe abrirse en HTTPS; si es local con IP usa APP_SSL=1.</p><div class='camera'><video id='video' autoplay playsinline muted></video><canvas id='canvas'></canvas><img id='preview' style='display:none'></div><input type='hidden' name='captura_base64' id='captura_base64'><input type='hidden' name='camara_origen' id='camara_origen' value='WEB'><div class='actions'><button type='button' class='primary' onclick='startCamera(event)'>🔄 Reintentar cámara</button><button type='button' class='green' onclick='capturePhoto()'>✅ Capturar foto</button><button type='button' class='gray' onclick='stopCamera()'>Detener cámara</button><label class='btn gray' style='cursor:pointer'>📁 Cámara/archivo<input id='fileCamFallback' type='file' accept='image/*' capture='user' onchange='loadFileFallback(this)' style='display:none'></label></div><p id='camStatus' class='muted'></p><h2>2. Aceptación / firma digital simple</h2><p class='muted'>Declaro que soy el titular del DNI indicado, que he revisado el documento y acepto registrar mi firma/aceptación electrónica con evidencia de cámara.</p><label><input type='checkbox' name='acepta' value='1' required> Acepto firmar/validar este documento</label><br><br><label>Nombre completo como firma</label><input type='text' name='firma_texto' value='{html.escape(sol['trabajador'] or '')}' required><div class='actions'><button class='green' type='submit'>✍️ Registrar firma</button></div></form></div>
    <script>
let stream=null,captured=false,starting=false;
function stMsg(t,ok=false,err=false){{const st=document.getElementById('camStatus'); if(st){{st.innerHTML=t; st.style.fontWeight='900'; st.style.color=err?'#b91c1c':(ok?'#15803d':'#64748b');}}}}
function secureOk(){{return window.isSecureContext || location.protocol==='https:' || ['localhost','127.0.0.1','::1'].includes(location.hostname);}}
function isMobile(){{return /Android|iPhone|iPad|iPod|Mobile/i.test(navigator.userAgent);}}
async function waitVideo(v,ms=12000){{const ini=Date.now(); while(Date.now()-ini<ms){{if(v.videoWidth&&v.videoHeight&&v.readyState>=2)return true; await new Promise(r=>setTimeout(r,150));}} return false;}}
async function startCamera(ev){{
  if(ev)ev.preventDefault(); if(starting)return false; starting=true;
  const v=document.getElementById('video'); const preview=document.getElementById('preview');
  try{{
    captured=false; document.getElementById('captura_base64').value=''; if(preview)preview.style.display='none';
    if(!secureOk()) throw new Error('CONTEXTO_NO_SEGURO');
    if(!navigator.mediaDevices||!navigator.mediaDevices.getUserMedia) throw new Error('MEDIADEVICES_NO_DISPONIBLE');
    if(stream){{stream.getTracks().forEach(t=>t.stop()); stream=null;}}
    document.getElementById('camara_origen').value=isMobile()?'CELULAR':'LAPTOP/PC';
    v.muted=true; v.autoplay=true; v.playsInline=true; v.setAttribute('playsinline',''); v.setAttribute('webkit-playsinline',''); v.style.display='block';
    stMsg('Activando cámara real... acepta el permiso del navegador.');
    const tries=[
      {{video:{{facingMode:{{ideal:'user'}},width:{{ideal:1280}},height:{{ideal:720}}}},audio:false}},
      {{video:{{facingMode:'user'}},audio:false}},
      {{video:{{width:{{ideal:640}},height:{{ideal:480}}}},audio:false}},
      {{video:true,audio:false}}
    ];
    let last=null;
    for(const cfg of tries){{try{{stream=await navigator.mediaDevices.getUserMedia(cfg); break;}}catch(e){{last=e;}}}}
    if(!stream) throw(last||new Error('SIN_CAMARA'));
    v.srcObject=stream;
    try{{await v.play();}}catch(e){{}}
    if(!await waitVideo(v)) throw new Error('VIDEO_NEGRO_O_SIN_IMAGEN');
    stMsg('✅ Cámara activa. Coloca el rostro al centro y presiona <b>Capturar foto</b>.',true);
    return true;
  }}catch(e){{
    if(stream){{stream.getTracks().forEach(t=>t.stop()); stream=null;}} if(v)v.srcObject=null;
    const n=(e&&e.name)?e.name:((e&&e.message)?e.message:'Error'); let ayuda='';
    if(n==='CONTEXTO_NO_SEGURO') ayuda=' En celular no abre cámara con HTTP/IP local. Usa HTTPS de Render o ejecuta local con APP_SSL=1 y entra por https://IP-DE-TU-PC:5000.';
    else if(n==='NotAllowedError'||n==='PermissionDeniedError') ayuda=' Dale PERMITIR a Cámara en el candado del navegador.';
    else if(n==='NotReadableError'||n==='TrackStartError') ayuda=' Cierra Zoom/Meet/Teams/Cámara de Windows u otra app que esté usando la cámara.';
    else if(n==='NotFoundError'||n==='DevicesNotFoundError') ayuda=' No se encontró cámara conectada.';
    else ayuda=' Revisa permisos de cámara del navegador y del sistema operativo.';
    stMsg('❌ No se pudo activar cámara: '+n+'.'+ayuda+' También puedes usar el botón <b>Cámara/archivo</b> como respaldo.',false,true);
    return false;
  }}finally{{starting=false;}}
}}
function capturePhoto(){{
  const v=document.getElementById('video'),c=document.getElementById('canvas'),img=document.getElementById('preview');
  if(!v||!v.srcObject||!v.videoWidth){{stMsg('Primero activa la cámara y acepta el permiso.',false,true);return false;}}
  c.width=v.videoWidth;c.height=v.videoHeight;c.getContext('2d').drawImage(v,0,0,c.width,c.height);
  const data=c.toDataURL('image/jpeg',0.88); document.getElementById('captura_base64').value=data;
  img.src=data;img.style.display='block';captured=true;stMsg('✅ Foto capturada correctamente. Ya puedes registrar la firma.',true);return true;
}}
function loadFileFallback(input){{const f=input&&input.files?input.files[0]:null; if(!f)return; const rd=new FileReader(); rd.onload=()=>{{document.getElementById('captura_base64').value=rd.result; const img=document.getElementById('preview'); img.src=rd.result; img.style.display='block'; captured=true; document.getElementById('camara_origen').value=isMobile()?'CELULAR-FALLBACK':'PC-FALLBACK'; stMsg('✅ Evidencia cargada desde cámara/archivo. Ya puedes registrar la firma.',true);}}; rd.readAsDataURL(f);}}
function stopCamera(){{if(stream){{stream.getTracks().forEach(t=>t.stop()); stream=null;}} const v=document.getElementById('video'); if(v){{v.pause(); v.srcObject=null;}} stMsg('Cámara detenida.');}}
function prepararEnvio(){{if(!captured||!document.getElementById('captura_base64').value){{alert('Primero capture la foto de evidencia.');return false;}}return true;}}
document.addEventListener('DOMContentLoaded',()=>{{stMsg(secureOk()?'Intentando encender cámara automáticamente...':'⚠️ Para celular necesitas HTTPS. Usa Render o APP_SSL=1 local.'); setTimeout(()=>{{if(!captured)startCamera();}},350);}});
</script>
    </body></html>"""
    return content


@app.route('/admin/contratacion/carga/plantilla/<tipo>')
@admin_required
def descargar_plantilla_carga_contratacion(tipo):
    """Descarga plantillas oficiales para Carga Masiva: actualización de datos y bajas."""
    wb = Workbook(); ws = wb.active
    tipo = (tipo or '').lower().strip()
    if tipo in ('baja','bajas','cese'):
        ws.title = 'BAJA_TRABAJADOR'
        headers = ['Código Trabajador', 'Motivo Cese', 'Fecha', 'Comentario']
        ejemplo = ['74324033', 'CESE / RENUNCIA / FIN DE CONTRATO', '31/05/2026', 'Comentario opcional']
        nombre = 'PLANTILLA_CARGA_MASIVA_BAJA_TRABAJADOR.xlsx'
    else:
        ws.title = 'ACTUALIZACION_DATOS'
        headers = ['Legajo', 'Número teléfono móvil', 'Correo electrónico']
        ejemplo = ['74324033', '999999999', 'correo@empresa.com']
        nombre = 'PLANTILLA_CARGA_MASIVA_ACTUALIZACION_DATOS_TRABAJADOR.xlsx'
    ws.append(headers); ws.append(ejemplo)
    fill = PatternFill('solid', fgColor='BFBFBF')
    font = Font(bold=True, color='000000')
    border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    for c in ws[1]:
        c.fill = fill; c.font = font; c.alignment = Alignment(horizontal='center'); c.border = border
    for col in range(1, len(headers)+1):
        ws.column_dimensions[chr(64+col)].width = 28
    out = PERSIST_DIR / nombre
    wb.save(out)
    return send_file(out, as_attachment=True, download_name=nombre)


def _headers_excel(ws):
    return {str(c.value or '').strip().upper(): i for i, c in enumerate(ws[1])}

def _celda(row, idx, *names):
    for n in names:
        i = idx.get(str(n).strip().upper())
        if i is not None:
            return clean(row[i].value)
    return ''


@app.route('/admin/contratacion', methods=['GET','POST'])
@admin_required
def admin_contratacion():
    """Gestión Contratos estilo Adapta: flujos, cargas, reportes, maestros, anuncios y documentaria."""
    sec = request.args.get('sec','flujo')
    if request.method=='POST':
        accion = request.form.get('accion','doc')
        if accion == 'anuncio':
            f = request.files.get('archivo')
            titulo = clean(request.form.get('titulo')) or 'Anuncio de contratación'
            carpeta = UPLOAD_DIR/'contratacion'/'anuncios'; carpeta.mkdir(parents=True, exist_ok=True)
            if f and f.filename:
                name = now_file()+'_'+secure_filename(f.filename); path = carpeta/name; f.save(path)
                flash(f'Anuncio cargado: {titulo} / {f.filename}', 'ok')
            else:
                flash('Completa el archivo del anuncio.', 'error')
            return redirect(url_for('admin_contratacion', sec='anuncios'))
        if accion == 'estado_trabajador':
            dni_estado = normalizar_dni(request.form.get('dni'))
            nuevo_estado = 1 if request.form.get('nuevo_estado') == '1' else 0
            with db() as con:
                con.execute('UPDATE trabajadores SET activo=? WHERE dni=?', (nuevo_estado, dni_estado)); con.commit()
            flash(('Trabajador reactivado.' if nuevo_estado else 'Trabajador cesado/inactivado. Sus documentos se mantienen archivados.'), 'ok')
            return redirect(url_for('admin_contratacion', sec='actualizar'))
        if accion == 'carga_actualizar_datos':
            empresa = clean(request.form.get('empresa')) or 'AQUANQA'
            f = request.files.get('archivo')
            if not f or not f.filename:
                flash('Selecciona el Excel de actualización de datos.', 'error')
                return redirect(url_for('admin_contratacion', sec='carga'))
            carpeta = UPLOAD_DIR/'contratacion'/'carga_masiva'; carpeta.mkdir(parents=True, exist_ok=True)
            path = carpeta/(now_file()+'_ACTUALIZACION_'+secure_filename(f.filename)); f.save(path)
            ok = no = err = 0
            try:
                wb = load_workbook(path, data_only=True); ws = wb.active; idx = _headers_excel(ws)
                with db() as con:
                    for row in ws.iter_rows(min_row=2):
                        legajo = normalizar_dni(_celda(row, idx, 'Legajo', 'Código Trabajador', 'Codigo Trabajador', 'DNI'))
                        celular = _celda(row, idx, 'Número teléfono móvil', 'Numero telefono movil', 'Telefono movil', 'Celular')
                        correo = _celda(row, idx, 'Correo electrónico', 'Correo electronico', 'Correo')
                        if not legajo:
                            continue
                        existe = con.execute('SELECT dni FROM trabajadores WHERE dni=?', (legajo,)).fetchone()
                        if not existe:
                            no += 1; continue
                        sets=[]; vals=[]
                        if celular:
                            sets.append('celular=?'); vals.append(celular)
                        if correo:
                            sets.append('correo=?'); vals.append(correo.lower())
                        if empresa:
                            sets.append('empresa=?'); vals.append(empresa)
                        if sets:
                            vals.append(legajo)
                            con.execute('UPDATE trabajadores SET '+','.join(sets)+' WHERE dni=?', vals)
                            ok += 1
                    con.commit()
                exportar_excels_locales()
                flash(f'Actualización de datos completada: {ok} trabajadores actualizados, {no} no encontrados.', 'ok')
            except Exception as e:
                err += 1; flash('Error procesando actualización de datos: '+str(e), 'error')
            return redirect(url_for('admin_contratacion', sec='carga'))
        if accion == 'carga_baja_trabajador':
            empresa = clean(request.form.get('empresa')) or 'AQUANQA'
            f = request.files.get('archivo')
            if not f or not f.filename:
                flash('Selecciona el Excel de baja de trabajador.', 'error')
                return redirect(url_for('admin_contratacion', sec='carga'))
            carpeta = UPLOAD_DIR/'contratacion'/'carga_masiva'; carpeta.mkdir(parents=True, exist_ok=True)
            path = carpeta/(now_file()+'_BAJA_'+secure_filename(f.filename)); f.save(path)
            ok = no = 0
            try:
                wb = load_workbook(path, data_only=True); ws = wb.active; idx = _headers_excel(ws)
                with db() as con:
                    for row in ws.iter_rows(min_row=2):
                        dni = normalizar_dni(_celda(row, idx, 'Código Trabajador', 'Codigo Trabajador', 'Legajo', 'DNI'))
                        motivo = _celda(row, idx, 'Motivo Cese', 'Motivo')
                        fecha = fecha_sin_hora(_celda(row, idx, 'Fecha', 'Fecha Cese'))
                        comentario = _celda(row, idx, 'Comentario', 'Observacion', 'Observación')
                        if not dni:
                            continue
                        existe = con.execute('SELECT dni FROM trabajadores WHERE dni=?', (dni,)).fetchone()
                        if not existe:
                            no += 1; continue
                        obs = f"BAJA MASIVA | EMPRESA: {empresa} | MOTIVO: {motivo} | FECHA: {fecha} | COMENTARIO: {comentario}".strip()
                        con.execute('UPDATE trabajadores SET activo=0, fecha_fin_contrato=COALESCE(NULLIF(?,\'\'), fecha_fin_contrato), observacion=? WHERE dni=?', (fecha, obs, dni))
                        ok += 1
                    con.commit()
                exportar_excels_locales()
                flash(f'Baja masiva completada: {ok} trabajadores inactivados, {no} no encontrados.', 'ok')
            except Exception as e:
                flash('Error procesando baja masiva: '+str(e), 'error')
            return redirect(url_for('admin_contratacion', sec='carga'))
        if accion == 'guardar_tipo_etapa':
            tid = request.form.get('tipo_id')
            codigo = clean(request.form.get('codigo'))
            descripcion = clean(request.form.get('descripcion')).upper()
            etapa = clean(request.form.get('etapa')) or 'Incorporacion'
            activo = 1 if request.form.get('activo') == '1' else 0
            obligatorio = 1 if request.form.get('obligatorio') == '1' else 0
            req_fd = 1 if request.form.get('requiere_firma_digital') == '1' else 0
            req_fe = 1 if request.form.get('requiere_firma_electronica') == '1' else 0
            req_cr = 1 if request.form.get('requiere_confirmacion_recepcion') == '1' else 0
            if not codigo or not descripcion:
                flash('Completa código y tipo de documento.', 'error')
                return redirect(url_for('admin_contratacion', sec='tipos_etapa'))
            with db() as con:
                if tid:
                    con.execute('''UPDATE contratacion_tipos SET codigo=?, descripcion=?, etapa=?, obligatorio=?, activo=?, requiere_firma_digital=?, requiere_firma_electronica=?, requiere_confirmacion_recepcion=? WHERE id=?''',
                                (codigo, descripcion, etapa, obligatorio, activo, req_fd, req_fe, req_cr, tid))
                    flash('Tipo de documento por etapa actualizado.', 'ok')
                else:
                    con.execute('''INSERT INTO contratacion_tipos(codigo,descripcion,etapa,obligatorio,activo,requiere_firma_digital,requiere_firma_electronica,requiere_confirmacion_recepcion) VALUES(?,?,?,?,?,?,?,?)''',
                                (codigo, descripcion, etapa, obligatorio, activo, req_fd, req_fe, req_cr))
                    flash('Tipo de documento por etapa creado.', 'ok')
                con.commit()
            return redirect(url_for('admin_contratacion', sec='tipos_etapa'))
        if accion == 'estado_tipo_etapa':
            tid = request.form.get('tipo_id')
            activo = 1 if request.form.get('activo') == '1' else 0
            with db() as con:
                con.execute('UPDATE contratacion_tipos SET activo=? WHERE id=?', (activo, tid)); con.commit()
            flash('Estado actualizado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='tipos_etapa'))
        if accion == 'eliminar_tipo_etapa':
            tid = request.form.get('tipo_id')
            with db() as con:
                con.execute('DELETE FROM contratacion_tipos WHERE id=?', (tid,)); con.commit()
            flash('Tipo de documento eliminado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='tipos_etapa'))

        if accion == 'guardar_tipo_empleado':
            did = request.form.get('doc_emp_id')
            codigo = clean(request.form.get('codigo'))
            descripcion = clean(request.form.get('descripcion')).upper()
            nombre_corto = clean(request.form.get('nombre_corto')).upper()
            grupo = clean(request.form.get('grupo')) or 'Contratación Trabajadores'
            activo = 1 if request.form.get('activo') == '1' else 0
            sin_cargo = 1 if request.form.get('permite_doc_sin_cargo') == '1' else 0
            validacion = 1 if request.form.get('requiere_validacion') == '1' else 0
            vencimiento = 1 if request.form.get('tiene_fecha_vencimiento') == '1' else 0
            contrato = 1 if request.form.get('es_contrato') == '1' else 0
            if not descripcion or not nombre_corto or not grupo:
                flash('Completa Descripción, Nombre Corto y Grupo.', 'error')
                return redirect(url_for('admin_contratacion', sec='tipo_empleado'))
            if not codigo:
                codigo = str(abs(hash(descripcion)) % 9000 + 100).zfill(3)
            with db() as con:
                if did:
                    con.execute('''UPDATE contratacion_tipo_empleado SET codigo=?,descripcion=?,nombre_corto=?,grupo=?,activo=?,permite_doc_sin_cargo=?,requiere_validacion=?,tiene_fecha_vencimiento=?,es_contrato=? WHERE id=?''',
                                (codigo, descripcion, nombre_corto, grupo, activo, sin_cargo, validacion, vencimiento, contrato, did))
                    flash('Tipo documento de empleado actualizado.', 'ok')
                else:
                    con.execute('''INSERT INTO contratacion_tipo_empleado(codigo,descripcion,nombre_corto,grupo,activo,permite_doc_sin_cargo,requiere_validacion,tiene_fecha_vencimiento,es_contrato) VALUES(?,?,?,?,?,?,?,?,?)''',
                                (codigo, descripcion, nombre_corto, grupo, activo, sin_cargo, validacion, vencimiento, contrato))
                    flash('Tipo documento de empleado creado.', 'ok')
                con.commit()
            return redirect(url_for('admin_contratacion', sec='tipo_empleado'))
        if accion == 'estado_tipo_empleado':
            did = request.form.get('doc_emp_id')
            activo = 1 if request.form.get('activo') == '1' else 0
            with db() as con:
                con.execute('UPDATE contratacion_tipo_empleado SET activo=? WHERE id=?', (activo, did)); con.commit()
            flash('Estado actualizado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='tipo_empleado'))
        if accion == 'eliminar_tipo_empleado':
            did = request.form.get('doc_emp_id')
            with db() as con:
                con.execute('DELETE FROM contratacion_tipo_empleado WHERE id=?', (did,)); con.commit()
            flash('Tipo documento de empleado eliminado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='tipo_empleado'))
        if accion == 'guardar_cargo':
            cid = request.form.get('cargo_id')
            codigo = clean(request.form.get('codigo'))
            nombre = clean(request.form.get('nombre')).upper()
            nombre_corto = clean(request.form.get('nombre_corto')).upper()
            resumen = clean(request.form.get('resumen'))
            funciones = clean(request.form.get('funciones'))
            activo = 1 if request.form.get('activo') == '1' else 0
            if not nombre or not nombre_corto:
                flash('Completa Nombre y Nombre Corto del cargo.', 'error')
                return redirect(url_for('admin_contratacion', sec='cargo'))
            if not codigo:
                codigo = str(abs(hash(nombre)) % 9000 + 100).zfill(3)
            with db() as con:
                if cid:
                    con.execute('''UPDATE contratacion_cargos SET codigo=?,nombre=?,nombre_corto=?,resumen=?,funciones=?,activo=? WHERE id=?''', (codigo, nombre, nombre_corto, resumen, funciones, activo, cid))
                    flash('Cargo actualizado.', 'ok')
                else:
                    con.execute('''INSERT INTO contratacion_cargos(codigo,nombre,nombre_corto,resumen,funciones,activo) VALUES(?,?,?,?,?,?)''', (codigo, nombre, nombre_corto, resumen, funciones, activo))
                    flash('Cargo creado.', 'ok')
                con.commit()
            return redirect(url_for('admin_contratacion', sec='cargo'))
        if accion == 'estado_cargo':
            cid = request.form.get('cargo_id')
            activo = 1 if request.form.get('activo') == '1' else 0
            with db() as con:
                con.execute('UPDATE contratacion_cargos SET activo=? WHERE id=?', (activo, cid)); con.commit()
            flash('Estado de cargo actualizado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='cargo'))
        if accion == 'eliminar_cargo':
            cid = request.form.get('cargo_id')
            with db() as con:
                con.execute('DELETE FROM contratacion_cargos WHERE id=?', (cid,)); con.commit()
            flash('Cargo eliminado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='cargo'))
        if accion == 'crear_observado':
            tipo_persona = clean(request.form.get('tipo_persona')) or 'Trabajador'
            trabajador_sel = clean(request.form.get('trabajador_sel'))
            numero_documento = normalizar_dni(request.form.get('numero_documento') or trabajador_sel)
            nombre = clean(request.form.get('nombre_externo'))
            if tipo_persona.lower() == 'trabajador' and numero_documento:
                tr = get_trabajador(numero_documento)
                if tr:
                    nombre = tr['nombre']
            motivo_lista = clean(request.form.get('motivo'))
            motivo_otro = clean(request.form.get('motivo_otro'))
            motivo = motivo_otro.upper() if motivo_otro else motivo_lista.upper()
            nivel = clean(request.form.get('nivel_restriccion')) or 'NIVEL 3'
            comentario = clean(request.form.get('comentario'))
            if not numero_documento or not nombre or not motivo:
                flash('Completa trabajador/DNI, nombre y motivo para registrar observado.', 'error')
                return redirect(url_for('admin_contratacion', sec='observados'))
            with db() as con:
                con.execute('''INSERT INTO trabajadores_observados(tipo_persona,tipo_documento,numero_documento,nombre,motivo,nivel_restriccion,comentario,estado,creado_por,fecha_creacion,fecha_edicion)
                               VALUES(?,?,?,?,?,?,?,?,?,?,?)''', (tipo_persona, 'DNI', numero_documento, nombre.upper(), motivo, nivel, comentario.upper(), 'Active', marca_carga(session.get('admin_user','admin')), now_txt(), ''))
                con.commit()
            flash('Trabajador observado registrado correctamente.', 'ok')
            return redirect(url_for('admin_contratacion', sec='observados'))
        if accion == 'estado_observado':
            oid = request.form.get('observado_id')
            estado = 'Inactive' if request.form.get('estado') == 'Inactive' else 'Active'
            with db() as con:
                con.execute('UPDATE trabajadores_observados SET estado=?, editor_por=?, fecha_edicion=? WHERE id=?', (estado, marca_carga(session.get('admin_user','admin')), now_txt(), oid)); con.commit()
            flash('Estado de trabajador observado actualizado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='observados'))
        if accion == 'eliminar_observado':
            oid = request.form.get('observado_id')
            with db() as con:
                con.execute('DELETE FROM trabajadores_observados WHERE id=?', (oid,)); con.commit()
            flash('Registro observado eliminado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='observados'))
        if accion == 'estado_plantilla':
            pid_estado = request.form.get('plantilla_id')
            nuevo_estado = 1 if request.form.get('activo') == '1' else 0
            with db() as con:
                con.execute('UPDATE contratacion_plantillas SET activo=?, fecha_actualizacion=? WHERE id=?', (nuevo_estado, now_txt(), pid_estado))
                con.commit()
            flash('Estado de plantilla actualizado.', 'ok')
            return redirect(url_for('admin_contratacion', sec='plantillas'))
        if accion == 'plantilla':
            pid = request.form.get('plantilla_id')
            nombre = clean(request.form.get('nombre_plantilla')) or 'PLANTILLA SIN NOMBRE'
            descripcion = clean(request.form.get('descripcion')) or nombre
            tipo_doc = clean(request.form.get('tipo_documento')) or descripcion
            esquema = clean(request.form.get('esquema')) or 'Trabajador Contrato Laboral'
            condicion = clean(request.form.get('condicion')) or 'SIN CONDICIONES'
            version = clean(request.form.get('version')) or 'Version 01'
            # REGLA PRO: si se carga o existe archivo Word/PDF => ACTIVO; si no hay archivo => INACTIVO.
            # Esto evita que una plantilla sin documento aparezca como lista para firma.
            f = request.files.get('archivo')
            ruta = None; archivo_nombre = None
            carpeta = UPLOAD_DIR/'contratacion'/'plantillas'; carpeta.mkdir(parents=True, exist_ok=True)
            if f and f.filename:
                archivo_nombre = secure_filename(f.filename)
                path = carpeta/(now_file()+'_'+archivo_nombre)
                f.save(path); ruta = str(path)
            with db() as con:
                if pid:
                    row = con.execute('SELECT ruta_archivo,archivo_nombre FROM contratacion_plantillas WHERE id=?',(pid,)).fetchone()
                    if not ruta and row:
                        ruta=row['ruta_archivo']; archivo_nombre=row['archivo_nombre']
                    activo = 1 if ruta else 0
                    con.execute('''UPDATE contratacion_plantillas SET nombre_plantilla=?,descripcion=?,tipo_documento=?,esquema=?,condicion=?,version=?,activo=?,archivo_nombre=?,ruta_archivo=?,fecha_actualizacion=? WHERE id=?''', (nombre,descripcion,tipo_doc,esquema,condicion,version,activo,archivo_nombre,ruta,now_txt(),pid))
                    detectados = sincronizar_campos_desde_word(pid, ruta) if ruta else 0
                    flash(('Plantilla actualizada correctamente. Campos Word detectados: ' + str(detectados)) if detectados else 'Plantilla actualizada correctamente.', 'ok')
                    redir = url_for('contratacion_plantilla_detalle', pid=pid, tab='contenido')
                else:
                    activo = 1 if ruta else 0
                    cur=con.execute('''INSERT INTO contratacion_plantillas(nombre_plantilla,descripcion,tipo_documento,esquema,condicion,version,activo,archivo_nombre,ruta_archivo,fecha_creacion,fecha_actualizacion,creado_por) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)''', (nombre,descripcion,tipo_doc,esquema,condicion,version,activo,archivo_nombre,ruta,now_txt(),now_txt(),marca_carga(session.get('admin_user','admin'))))
                    nuevo_id=cur.lastrowid
                    detectados = sincronizar_campos_desde_word(nuevo_id, ruta) if ruta else 0
                    for nom,origen,td in CONTRATACION_CAMPOS_CORRESPONDENCIA:
                        con.execute('INSERT INTO contratacion_plantilla_campos(plantilla_id,nombre_campo,campo_origen,tipo_dato) VALUES(?,?,?,?)',(nuevo_id,nom,origen,td))
                    if condicion == 'CONDICIONES':
                        con.execute('INSERT INTO contratacion_plantilla_condiciones(plantilla_id,nombre_campo,condicion,valor,activo,fecha_registro,creado_por) VALUES(?,?,?,?,?,?,?)',(nuevo_id,'Planilla','=','OBREROS RÉGIMEN AGRÍCOLA',1,now_txt(),marca_carga(session.get('admin_user','admin'))))
                        con.execute('INSERT INTO contratacion_plantilla_condiciones(plantilla_id,nombre_campo,condicion,valor,activo,fecha_registro,creado_por) VALUES(?,?,?,?,?,?,?)',(nuevo_id,'Tipo Contrato','=','INTERMITENTE OBRERO',1,now_txt(),marca_carga(session.get('admin_user','admin'))))
                    flash(('Plantilla creada correctamente. Campos Word detectados: ' + str(detectados)) if 'detectados' in locals() and detectados else 'Plantilla creada correctamente.', 'ok')
                    redir = url_for('contratacion_plantilla_detalle', pid=nuevo_id, tab='contenido')
                con.commit()
            return redirect(redir)
        if accion == 'generar_docs_plantillas':
            dni = normalizar_dni(request.form.get('dni_plantilla_firma'))
            trab = get_trabajador(dni)
            ids_raw = request.form.getlist('plantillas_firma')
            creados = 0
            if not dni or not trab:
                flash('Selecciona un trabajador valido para generar documentos desde plantillas.', 'error')
                return redirect(url_for('admin_contratacion', sec='firma'))
            with db() as con:
                for pid_raw in ids_raw:
                    if not str(pid_raw).isdigit():
                        continue
                    plx = con.execute('SELECT * FROM contratacion_plantillas WHERE id=? AND activo=1', (int(pid_raw),)).fetchone()
                    if not plx:
                        continue
                    nombre_doc = plx['archivo_nombre'] or (plx['nombre_plantilla'] + '.docx')
                    con.execute('INSERT INTO contratacion_docs(dni,trabajador,empresa,etapa,tipo_doc,estado,archivo_nombre,ruta_archivo,fecha_registro,uploaded_by) VALUES(?,?,?,?,?,?,?,?,?,?)',
                                (dni, trab['nombre'], trab['empresa'], 'Generado desde plantilla', plx['tipo_documento'] or plx['nombre_plantilla'], 'GENERADO - PENDIENTE FIRMA', nombre_doc, plx['ruta_archivo'], now_txt(), marca_carga(session.get('admin_user','admin'))))
                    con.execute('INSERT INTO eventos_documento(dni,evento,fecha,detalle) VALUES(?,?,?,?)', (dni, 'Documento generado desde plantilla', now_txt(), f"{plx['nombre_plantilla']} listo para firma facial/digital."))
                    creados += 1
                con.commit()
            flash(f'Documentos generados desde plantillas: {creados}. Ahora ya aparecen abajo para marcarlos y enviarlos a firma masiva.', 'ok' if creados else 'error')
            return redirect(url_for('admin_contratacion', sec='firma'))

        if accion == 'firma_masiva':
            ids_raw = request.form.get('documentos_lote') or ''
            metodo = clean(request.form.get('metodo_masivo')) or 'FACIAL + FIRMA DIGITAL'
            obs = clean(request.form.get('observacion_masiva')) or 'Envío masivo a firma facial/digital'
            ids=[]
            for x in re.split(r'[\n,; ]+', ids_raw):
                x=str(x).strip()
                if x.isdigit() and int(x) not in ids:
                    ids.append(int(x))
            creadas=0
            with db() as con:
                for doc_id in ids:
                    doc=con.execute('SELECT * FROM contratacion_docs WHERE id=?',(doc_id,)).fetchone()
                    if not doc: continue
                    token=crear_token_firma()
                    con.execute('INSERT INTO firma_solicitudes(documento_id,dni,trabajador,metodo,estado,evidencia_ref,fecha_envio,observacion,firma_token,validacion_estado) VALUES(?,?,?,?,?,?,?,?,?,?)',(doc['id'],doc['dni'],doc['trabajador'],metodo,'Pendiente de captura facial','',now_txt(),obs,token,'PENDIENTE'))
                    con.execute("UPDATE contratacion_docs SET estado='ENVIADO A FIRMA' WHERE id=?",(doc['id'],))
                    con.execute('INSERT INTO eventos_documento(dni,evento,fecha,detalle) VALUES(?,?,?,?)',(doc['dni'],'Contrato pendiente de firma',now_txt(),f"Tiene pendiente firmar: {doc['tipo_doc']}. Ingrese a Gestión Contrato / Bandeja de firmas."))
                    creadas += 1
                con.commit()
            flash(f'Solicitudes masivas creadas: {creadas}. Revisa la Bandeja de Firmas para copiar/enviar enlaces móviles.', 'ok' if creadas else 'error')
            return redirect(url_for('admin_contratacion', sec='firma'))
        if accion == 'firma_solicitud':
            doc_id = request.form.get('documento_id')
            metodo = clean(request.form.get('metodo')) or 'FACIAL + FIRMA DIGITAL'
            obs = clean(request.form.get('observacion'))
            with db() as con:
                doc = con.execute('SELECT * FROM contratacion_docs WHERE id=?',(doc_id,)).fetchone()
                if doc:
                    token = crear_token_firma()
                    con.execute('INSERT INTO firma_solicitudes(documento_id,dni,trabajador,metodo,estado,evidencia_ref,fecha_envio,observacion,firma_token,validacion_estado) VALUES(?,?,?,?,?,?,?,?,?,?)',(doc['id'],doc['dni'],doc['trabajador'],metodo,'Pendiente de captura facial','',now_txt(),obs,token,'PENDIENTE'))
                    con.execute("UPDATE contratacion_docs SET estado='ENVIADO A FIRMA' WHERE id=?",(doc['id'],))
                    con.execute('INSERT INTO eventos_documento(dni,evento,fecha,detalle) VALUES(?,?,?,?)',(doc['dni'],'Contrato pendiente de firma',now_txt(),f"Tiene pendiente firmar: {doc['tipo_doc']}. Ingrese a Gestión Contrato / Bandeja de firmas."))
                    con.commit(); flash('Solicitud de firma creada. Copia el enlace móvil/web para que el trabajador abra cámara, capture evidencia y acepte/firme.', 'ok')
                else:
                    flash('Documento no encontrado para enviar a firma.', 'error')
            return redirect(url_for('admin_contratacion', sec='firma'))
        f=request.files.get('archivo'); dni=normalizar_dni(request.form.get('dni')); trab=get_trabajador(dni); tipo=clean(request.form.get('tipo_doc')); etapa=clean(request.form.get('etapa')) or 'Incorporación'
        if f and f.filename and dni:
            folder=UPLOAD_DIR/'contratacion'/dni; folder.mkdir(parents=True, exist_ok=True)
            name=now_file()+'_'+secure_filename(f.filename); path=folder/name; f.save(path)
            with db() as con:
                con.execute('INSERT INTO contratacion_docs(dni,trabajador,empresa,etapa,tipo_doc,estado,archivo_nombre,ruta_archivo,fecha_registro,uploaded_by) VALUES(?,?,?,?,?,?,?,?,?,?)',(dni, trab['nombre'] if trab else '', trab['empresa'] if trab else '', etapa, tipo, 'FIRMADO', f.filename, str(path), now_txt(), marca_carga(session.get('admin_user','admin')))); con.commit()
            flash('Documento de contratación registrado y archivado.', 'ok')
        else:
            flash('Selecciona trabajador y archivo.', 'error')
        return redirect(url_for('admin_contratacion', sec='documentaria'))
    with db() as con:
        tipos=con.execute('SELECT * FROM contratacion_tipos ORDER BY etapa, descripcion').fetchall()
        docs=con.execute('SELECT * FROM contratacion_docs ORDER BY id DESC LIMIT 300').fetchall()
        firma_sols=con.execute('SELECT * FROM firma_solicitudes ORDER BY id DESC LIMIT 300').fetchall()
        trabajadores=con.execute('SELECT dni,nombre,empresa,cargo,area,correo,activo,fecha_registro FROM trabajadores ORDER BY nombre LIMIT 700').fetchall()
        observados=con.execute('SELECT * FROM trabajadores_observados ORDER BY id DESC LIMIT 500').fetchall()
        tipo_empleado=con.execute('SELECT * FROM contratacion_tipo_empleado ORDER BY descripcion LIMIT 1000').fetchall()
        cargos=con.execute('SELECT * FROM contratacion_cargos ORDER BY nombre LIMIT 2000').fetchall()

        # Filtros reales de Plantilla Documentos
        f_nombre = clean(request.args.get('f_nombre'))
        f_tipo = clean(request.args.get('f_tipo'))
        f_esquema = clean(request.args.get('f_esquema'))
        f_condicion = clean(request.args.get('f_condicion'))
        where_pl, params_pl = [], []
        if f_nombre:
            where_pl.append('(nombre_plantilla LIKE ? OR descripcion LIKE ? OR archivo_nombre LIKE ?)')
            params_pl += [f'%{f_nombre}%', f'%{f_nombre}%', f'%{f_nombre}%']
        if f_tipo:
            where_pl.append('tipo_documento LIKE ?')
            params_pl.append(f'%{f_tipo}%')
        if f_esquema:
            where_pl.append('esquema LIKE ?')
            params_pl.append(f'%{f_esquema}%')
        if f_condicion:
            where_pl.append('condicion LIKE ?')
            params_pl.append(f'%{f_condicion}%')
        sql_pl = 'SELECT * FROM contratacion_plantillas'
        if where_pl:
            sql_pl += ' WHERE ' + ' AND '.join(where_pl)
        sql_pl += ' ORDER BY id DESC'
        plantillas=con.execute(sql_pl, params_pl).fetchall()
    tipos_doc_opciones = []
    for r in tipos:
        if r['descripcion'] and r['descripcion'] not in tipos_doc_opciones:
            tipos_doc_opciones.append(r['descripcion'])
    for r in plantillas:
        if r['tipo_documento'] and r['tipo_documento'] not in tipos_doc_opciones:
            tipos_doc_opciones.append(r['tipo_documento'])
    opt_tipo=''.join([f"<option value='{html.escape(str(x))}'>{html.escape(str(x))}</option>" for x in tipos_doc_opciones])
    opt_trab=''.join([f"<option value='{r['dni']}'>{r['dni']} - {r['nombre']}</option>" for r in trabajadores])
    sample_trab = trabajadores[0] if trabajadores else None
    docs_rows=''.join([f"<tr><td><input type='checkbox'></td><td>🔍 📄</td><td>{r['dni']}</td><td>{r['trabajador']}</td><td>{r['tipo_doc']}</td><td><span class='c-badge cyan'>{r['estado'][:1] or 'F'}</span></td><td>{r['fecha_registro']}</td></tr>" for r in docs])
    renov_rows=''.join([f"<tr><td><input type='checkbox'></td><td>{t['dni']}</td><td>{t['nombre']}</td><td>INICIO</td><td>{fecha_sin_hora(t['fecha_registro'])}</td><td></td><td>30/06/2026</td><td><span class='c-badge green'>✓</span></td><td><span class='c-badge green'>✓</span></td><td>0</td></tr>" for t in trabajadores[:12]])
    def chk_badge(v):
        return "<span class='c-badge green'>✓</span>" if int(v or 0)==1 else "<span class='c-badge gray'>—</span>"
    def escjs(v):
        return html.escape(str(v or ''), quote=True)
    tipos_rows=''.join([f"""
        <tr>
          <td class='nowrap'><button type='button' class='icon-btn' onclick='editarTipoEtapa(\"{r['id']}\",\"{escjs(r['codigo'])}\",\"{escjs(r['descripcion'])}\",\"{escjs(r['etapa'])}\",\"{1 if r['activo'] else 0}\",\"{1 if r['obligatorio'] else 0}\",\"{1 if r['requiere_firma_digital'] else 0}\",\"{1 if r['requiere_firma_electronica'] else 0}\",\"{1 if r['requiere_confirmacion_recepcion'] else 0}\")'>✎</button><form method='post' style='display:inline' onsubmit='return confirm(\"¿Eliminar tipo de documento?\")'><input type='hidden' name='accion' value='eliminar_tipo_etapa'><input type='hidden' name='tipo_id' value='{r['id']}'><button class='icon-btn'>🗑</button></form></td>
          <td><form method='post'><input type='hidden' name='accion' value='estado_tipo_etapa'><input type='hidden' name='tipo_id' value='{r['id']}'><select name='activo' onchange='this.form.submit()' class='select-mini'><option value='1' {'selected' if r['activo'] else ''}>🟢 Active</option><option value='0' {'selected' if not r['activo'] else ''}>🔴 Inactive</option></select></form></td>
          <td>{html.escape(str(r['codigo'] or ''))}</td><td>{html.escape(str(r['descripcion'] or ''))}</td><td>{html.escape(str(r['etapa'] or ''))}</td>
          <td>{chk_badge(r['requiere_firma_digital'])}</td><td>{chk_badge(r['requiere_firma_electronica'])}</td><td>{chk_badge(r['obligatorio'])}</td><td>{chk_badge(r['requiere_confirmacion_recepcion'])}</td>
        </tr>""" for r in tipos])
    def h(v):
        return html.escape(str(v or ''))
    def estado_select(nombre_accion, id_name, rid, activo):
        return f"<form method='post'><input type='hidden' name='accion' value='{nombre_accion}'><input type='hidden' name='{id_name}' value='{rid}'><select name='activo' onchange='this.form.submit()' class='select-mini'><option value='1' {'selected' if activo else ''}>🟢 Active</option><option value='0' {'selected' if not activo else ''}>🔴 Inactive</option></select></form>"
    tipo_emp_rows=''.join([f"""
        <tr>
          <td class='nowrap'><button type='button' class='icon-btn' onclick='editarTipoEmpleado("{r['id']}","{escjs(r['codigo'])}","{escjs(r['descripcion'])}","{escjs(r['nombre_corto'])}","{escjs(r['grupo'])}","{1 if r['activo'] else 0}","{1 if r['permite_doc_sin_cargo'] else 0}","{1 if r['requiere_validacion'] else 0}","{1 if r['tiene_fecha_vencimiento'] else 0}","{1 if r['es_contrato'] else 0}")'>✎</button><form method='post' style='display:inline' onsubmit='return confirm("¿Eliminar tipo documento de empleado?")'><input type='hidden' name='accion' value='eliminar_tipo_empleado'><input type='hidden' name='doc_emp_id' value='{r['id']}'><button class='icon-btn'>🗑</button></form></td>
          <td>{estado_select('estado_tipo_empleado','doc_emp_id',r['id'],r['activo'])}</td><td>{h(r['codigo'])}</td><td>{h(r['descripcion'])}</td><td>{h(r['nombre_corto'])}</td><td>{h(r['grupo'])}</td>
        </tr>""" for r in tipo_empleado])
    cargo_rows=''.join([f"""
        <tr>
          <td class='nowrap'><button type='button' class='icon-btn' onclick='editarCargo("{r['id']}","{escjs(r['codigo'])}","{escjs(r['nombre'])}","{escjs(r['nombre_corto'])}","{escjs(r['resumen'])}","{escjs(r['funciones'])}","{1 if r['activo'] else 0}")'>✎</button><form method='post' style='display:inline' onsubmit='return confirm("¿Eliminar cargo?")'><input type='hidden' name='accion' value='eliminar_cargo'><input type='hidden' name='cargo_id' value='{r['id']}'><button class='icon-btn'>🗑</button></form></td>
          <td>{estado_select('estado_cargo','cargo_id',r['id'],r['activo'])}</td><td>{h(r['codigo'])}</td><td>{h(r['nombre'])}</td><td>{h(r['nombre_corto'])}</td>
        </tr>""" for r in cargos])
    plantillas_rows=''.join([
        f"""<tr>
          <td class='tpl-actions'>
            <a class='icon-btn action-edit' title='Abrir detalle / editar' href='{url_for('contratacion_plantilla_detalle', pid=r['id'])}'>✎</a><a class='icon-btn' title='Historial de cargas' href='{url_for('contratacion_plantilla_historial', pid=r['id'])}'>🔍</a><a class='icon-btn' title='Descargar plantilla Word' href='{url_for('contratacion_plantilla_archivo', pid=r['id'])}'>⬇</a>
            <form method='post' action='{url_for('contratacion_plantilla_eliminar', pid=r['id'])}' style='display:inline' onsubmit="return confirm('¿Eliminar esta plantilla? No se borrarán contratos/documentos históricos.');">
              <button class='icon-btn action-delete' title='Eliminar plantilla' type='submit'>🗑</button>
            </form>
          </td>
          <td><form method='post' class='state-form'><input type='hidden' name='accion' value='estado_plantilla'><input type='hidden' name='plantilla_id' value='{r['id']}'><select name='activo' class='state-select {'inactive' if not r['activo'] else ''}' onchange='this.form.submit()'><option value='1' {'selected' if r['activo'] else ''}>Activo</option><option value='0' {'selected' if not r['activo'] else ''}>Inactivo</option></select></form></td>
          <td><a class='tpl-link' href='{url_for('contratacion_plantilla_detalle', pid=r['id'])}'>{h(r['nombre_plantilla'])}</a></td>
          <td>{h(r['tipo_documento'])}</td><td>{h(r['esquema'])}</td><td>{h(r['descripcion'])}</td><td>{h(r['version'])}</td><td>{h(r['condicion'])}</td><td>{h(r['archivo_nombre'])}</td>
        </tr>""" for r in plantillas])
    if not observados:
        with db() as con:
            for t in trabajadores[:10]:
                con.execute('''INSERT INTO trabajadores_observados(tipo_persona,tipo_documento,numero_documento,nombre,motivo,nivel_restriccion,comentario,estado,creado_por,fecha_creacion,fecha_edicion)
                               VALUES(?,?,?,?,?,?,?,?,?,?,?)''', ('Trabajador','DNI',t['dni'],t['nombre'],'SINDICALISTA','NIVEL 3','SINDICALISTA','Active','Sistema',now_txt(),''))
            con.commit()
            observados=con.execute('SELECT * FROM trabajadores_observados ORDER BY id DESC LIMIT 500').fetchall()
    obs_activos=[r for r in observados if (r['estado'] or 'Active') == 'Active']
    obs_inactivos=[r for r in observados if (r['estado'] or 'Active') != 'Active']
    def estado_observado_form(r):
        return f"<form method='post' class='state-form'><input type='hidden' name='accion' value='estado_observado'><input type='hidden' name='observado_id' value='{r['id']}'><select name='estado' class='state-select {'inactive' if (r['estado'] or '')!='Active' else ''}' onchange='this.form.submit()'><option value='Active' {'selected' if (r['estado'] or '')=='Active' else ''}>🟢 Active</option><option value='Inactive' {'selected' if (r['estado'] or '')=='Inactive' else ''}>🔴 Inactive</option></select></form>"
    obs_rows=''.join([f"<tr><td><input type='checkbox'></td><td class='tpl-actions'><a class='icon-btn' title='Log'>🔗</a><a class='icon-btn' title='Editar'>✎</a><form method='post' style='display:inline' onsubmit=\"return confirm('¿Eliminar registro observado?');\"><input type='hidden' name='accion' value='eliminar_observado'><input type='hidden' name='observado_id' value='{r['id']}'><button class='icon-btn' type='submit'>🗑</button></form></td><td>{estado_observado_form(r)}</td><td>{h(r['tipo_documento'] or 'DNI')}</td><td>{h(r['numero_documento'])}</td><td>{h(r['nombre'])}</td><td>{h(r['motivo'])}</td><td>{h(r['nivel_restriccion'])}</td><td>{h(r['comentario'])}</td><td>{h(r['creado_por'])}</td><td>{h(r['fecha_creacion'])}</td><td>{h(r['editor_por'])}</td><td>{h(r['fecha_edicion'])}</td></tr>" for r in obs_activos]) or "<tr><td colspan='13'>No hay trabajadores observados activos.</td></tr>"
    obs_anulados_rows=''.join([f"<tr><td>🔗</td><td>{h(r['tipo_documento'] or 'DNI')}</td><td>{h(r['numero_documento'])}</td><td>{h(r['nombre'])}</td><td>{h(r['motivo'])}</td><td>{h(r['nivel_restriccion'])}</td><td>{h(r['comentario'])}</td><td>{h(r['creado_por'])}</td><td>{h(r['fecha_creacion'])}</td><td>{h(r['editor_por'])}</td><td>{h(r['fecha_edicion'])}</td></tr>" for r in obs_inactivos]) or "<tr><td colspan='11'>No hay registros anulados/inactivos.</td></tr>"
    motivo_options=''.join([f"<option value='{h(x)}'>{h(x)}</option>" for x in MOTIVOS_TRABAJADOR_OBSERVADO])
    nivel_options=''.join([f"<option value='{h(x)}'>{h(x)}</option>" for x in NIVELES_RESTRICCION_OBSERVADO])

    report_rows="<tr><td>✎ ⬇</td><td>RHTR01</td><td>Reporte Datos Trabajador</td><td>Reporte Datos Trabajador</td><td>TR_REPORT_1</td><td>Admin</td></tr>"
    trabajadores_estado_rows=''.join([f"<tr><td>{t['dni']}</td><td>{t['nombre']}</td><td>{t['empresa']}</td><td>{t['cargo'] or ''}</td><td><span class='state'>{'ACTIVO' if t['activo'] else 'CESADO/INACTIVO'}</span></td><td><form method='post' style='display:flex;gap:8px'><input type='hidden' name='accion' value='estado_trabajador'><input type='hidden' name='dni' value='{t['dni']}'><button class='c-btn {'gray' if t['activo'] else ''}' name='nuevo_estado' value='0'>Cesar</button><button class='c-btn green' name='nuevo_estado' value='1'>Reactivar</button></form></td></tr>" for t in trabajadores])
    # CSS local de contenido: se eliminó el segundo menú blanco tipo Adapta.
    # Todo se maneja desde el panel principal oscuro, con una sola pestaña activa por vez.
    css="""
    <style>
    .main{background:#f4f6f8!important;color:#111827!important}.c-title{font-size:24px;margin:0 0 20px;font-weight:950;color:#111827}.c-bar{display:flex;justify-content:space-between;gap:14px;align-items:center;margin-bottom:14px}.c-filter{display:grid;grid-template-columns:180px minmax(260px,420px) 160px minmax(260px,420px);gap:10px 24px;align-items:center;margin-bottom:22px}.c-filter input,.c-filter select,.c-form input,.c-form select{background:#fff!important;color:#111827!important;border:1px solid #cfd6df!important;border-radius:8px!important;padding:10px!important}.c-btn{background:#ff8d35;color:#fff;border:0;border-radius:10px;padding:10px 16px;font-weight:950;display:inline-flex;gap:8px;align-items:center;text-decoration:none}.c-btn.gray{background:#66707c}.c-card{background:#fff;border:1px solid #dde2e7;border-radius:12px;box-shadow:0 6px 18px #0000000d;margin-bottom:18px}.tabs{display:flex;border-bottom:1px solid #dde2e7}.tab{padding:14px 24px;font-weight:900;color:#7a7f87;border-bottom:3px solid transparent}.tab.active{color:#1f2937;border-bottom-color:#ff8d35}.c-table{width:100%;border-collapse:collapse;background:#fff;font-size:15px}.c-table th{font-weight:900;text-align:left;background:#f7f8fa;border:1px solid #dde2e7;padding:12px}.c-table td{border:1px solid #e1e5ea;padding:11px;vertical-align:middle}.c-table tr:nth-child(even) td{background:#eeeeee}.c-table tr.selected td{background:#bcd7fb!important}.c-badge{display:inline-grid;place-items:center;min-width:70px;border-radius:7px;padding:6px 10px;color:#fff;font-weight:950}.c-badge.green{background:#55ad11}.c-badge.cyan{background:#51c2d4}.state{border:1px solid #d1d5db;border-radius:99px;padding:7px 12px;background:white;color:#16a34a;font-weight:900}.tile-grid{display:grid;grid-template-columns:repeat(2,minmax(260px,1fr));gap:30px;max-width:1100px;margin:45px auto}.c-tile{background:#fff;border:1px solid #dde2e7;border-radius:14px;min-height:185px;padding:28px;display:flex;gap:18px;align-items:flex-start;position:relative;box-shadow:0 6px 18px #0000000d}.tile-icon{width:74px;height:74px;border-radius:50%;background:#ffe5cc;color:#ff8d35;display:grid;place-items:center;font-size:34px}.download-corner{position:absolute;right:18px;bottom:14px;font-size:24px}.toolbar{display:flex;justify-content:flex-end;gap:18px;color:#7b8088;margin:12px 0}.c-form{display:grid;grid-template-columns:180px minmax(260px,430px) 180px minmax(260px,430px);gap:12px 22px;align-items:center}.profile{display:grid;grid-template-columns:150px 1fr 1fr 1fr;gap:24px;align-items:start}.avatar-big{width:135px;height:135px;border-radius:50%;background:#f7c26d;display:grid;place-items:center;font-size:80px}.divider{border-left:2px dashed #c8c8c8;padding-left:20px}.muted2{color:#667085}.anuncio-upload{background:#fff;padding:22px;border-radius:12px;border:1px dashed #ff8d35;max-width:900px}.anuncio-upload input[type=file]{background:#fff!important;color:#111!important;border:1px solid #ddd!important;padding:10px!important;border-radius:8px!important}.video-box{margin-top:16px;background:#111827;color:#fff;padding:18px;border-radius:10px}.video-box video{width:100%;max-height:260px;background:#000;border-radius:8px}.table-wrap{overflow:auto}.plantilla-top{display:flex;justify-content:space-between;align-items:center;margin-bottom:10px}.crear-btn{border-radius:22px}.plantilla-filter{display:grid;grid-template-columns:190px minmax(280px,438px) 190px minmax(280px,438px);gap:8px 28px;align-items:center}.filter-card{margin-bottom:18px}.create-card{display:none}.create-card:target{display:block}.tpl-actions{white-space:nowrap;text-align:center}.icon-btn{display:inline-grid;place-items:center;width:34px;height:34px;border:0;background:transparent;color:#111827!important;font-size:20px;font-weight:950;text-decoration:none;cursor:pointer}.icon-btn:hover{background:#eef2f7;border-radius:8px}.action-delete{color:#111827!important}.state.inactive{color:#6b7280;background:#f3f4f6}.state-form{margin:0}.state-select{background:white!important;color:#16a34a!important;border:1px solid #d1d5db!important;border-radius:99px!important;padding:7px 12px!important;font-weight:900;min-width:110px}.state-select.inactive{color:#e11d48!important}.tpl-link{font-weight:900;color:#0f172a!important}.plantilla-table th:first-child,.plantilla-table td:first-child{min-width:92px;text-align:center}@media(max-width:1000px){.c-filter,.c-form,.profile{grid-template-columns:1fr}.tile-grid{grid-template-columns:1fr;margin:20px 0}.c-table{min-width:1000px}}

    .main{background:radial-gradient(circle at 95% -8%,rgba(16,185,129,.12),transparent 24%),linear-gradient(180deg,#0f141a,#111821)!important;color:#f8fafc!important}.c-title,h1,h2,h3{color:#f8fafc!important}.c-card,.filter-card,.table-wrap{background:linear-gradient(145deg,#171f28,#111821)!important;border:1px solid #34404d!important;color:#f8fafc!important}.c-filter input,.c-filter select,.c-form input,.c-form select,input,select,textarea{background:#0b1119!important;color:#f8fafc!important;border:1.5px solid #34404d!important}.c-table,table{background:#111821!important;color:#f8fafc!important}.c-table th,th{background:#0b1119!important;color:#10b981!important;border-color:#364250!important}.c-table td,td{background:#151d26!important;color:#eaf2fb!important;border-color:#2b3541!important}.c-table tr:nth-child(even) td{background:#111821!important}.plantilla-table td,.plantilla-table td *{color:#eaf2fb!important;opacity:1!important}.plantilla-table th{color:#10b981!important}.plantilla-top .crear-btn{position:relative!important;right:auto!important;top:auto!important;min-width:190px!important}.plantilla-top{gap:16px!important}.icon-btn{color:#10b981!important}.tpl-link{color:#fff!important}.state-select{background:#0b1119!important;color:#10b981!important;border:1px solid #10b981!important}.ficha-search{display:flex;gap:12px;margin:0 0 16px;max-width:720px}.ficha-search input{flex:1}.ficha-profile{display:grid;grid-template-columns:150px minmax(320px,1.4fr) minmax(260px,1fr) minmax(260px,1fr);gap:18px;align-items:stretch;margin-bottom:18px}.avatar-panel,.profile-main,.profile-col{background:linear-gradient(145deg,#171f28,#111821);border:1px solid #34404d;border-radius:18px;padding:18px;color:#f8fafc;box-shadow:0 16px 34px rgba(0,0,0,.22)}.avatar-panel{display:grid;place-items:center}.profile-main h2{margin:0 0 10px;color:#fff;font-size:22px}.profile-main p,.profile-col p{margin:8px 0;color:#eaf2fb}.created-box{background:#0b1119;border:1px solid #34404d;border-radius:12px;padding:10px;margin-top:12px;color:#cbd5e1}.status-dot,.status-pill{display:inline-flex;align-items:center;gap:6px;padding:7px 12px;border-radius:999px;font-weight:1000}.status-dot.ok,.status-pill.ok{background:rgba(34,197,94,.13);color:#86efac;border:1px solid rgba(34,197,94,.45)}.status-dot.bad,.status-pill.bad{background:rgba(244,63,94,.13);color:#fecdd3;border:1px solid rgba(244,63,94,.45)}.ficha-tabs{display:flex!important;padding:0!important;margin:0 0 16px!important;overflow:hidden}.ficha-tabs .tab{flex:1;background:transparent;border:0;border-bottom:3px solid transparent;color:#cbd5e1!important;cursor:pointer;font-size:16px}.ficha-tabs .tab.active{color:#10b981!important;border-bottom-color:#10b981!important}.ficha-tab-content{display:none}.ficha-tab-content.active{display:block}.laboral-grid{display:grid;grid-template-columns:1fr 1fr;gap:14px 28px;padding:24px!important}.laboral-grid label{display:grid;grid-template-columns:220px 1fr;align-items:center;gap:12px;color:#f8fafc!important}.periodos-box{margin-top:14px;background:#111821;border:1px solid #34404d;border-radius:14px;padding:14px;color:#eaf2fb}.mini-chip{display:inline-flex;margin:5px;padding:8px 12px;border-radius:999px;background:rgba(16,185,129,.12);color:#10b981;border:1px solid rgba(16,185,129,.35);font-weight:900}.period-row td{background:#0b1119!important;color:#10b981!important;font-size:16px}.check-green{display:inline-grid;place-items:center;min-width:54px;padding:7px 13px;background:#4ea60f;color:white;border-radius:8px;font-weight:1000}.check-gray{display:inline-grid;place-items:center;min-width:54px;padding:7px 13px;background:#334155;color:#cbd5e1;border-radius:8px;font-weight:1000}@media(max-width:1100px){.ficha-profile{grid-template-columns:1fr}.laboral-grid{grid-template-columns:1fr}.laboral-grid label{grid-template-columns:1fr}.plantilla-filter{grid-template-columns:1fr!important}}

    

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>"""
    def wrap(inner):
        return css + inner
    if sec=='flujo':
        rows=''.join([f"<tr class='{ 'selected' if i==0 else ''}'><td><input type='checkbox' {'checked' if i==0 else ''}></td><td>🔍 📄</td><td>{232105-i}</td><td><span class='c-badge green'>APROBADO</span></td><td>{'Eliminar Contrato' if i%2==0 else 'Eliminar Alta Trabajador'}</td><td>{now_txt()}</td><td>{now_txt()}</td><td>{(trabajadores[i]['nombre'] if i < len(trabajadores) else 'TRABAJADOR DEMO')}</td></tr>" for i in range(10)])
        content=wrap(f"<h2 class='c-title'>Eventos</h2><div class='c-filter'><b>Tipos de Evento:</b><select><option>Renovar Contrato</option><option>Eliminar Contrato</option></select><b>Estados:</b><select><option></option><option>Aprobado</option><option>Pendiente</option></select><b>Código Trabajador</b><input><span></span><span><button class='c-btn'>⌕ Buscar</button> <button class='c-btn gray'>Limpiar</button></span></div><div class='toolbar'>⚙ Acción ▾</div><div class='c-card table-wrap'><table class='c-table'><tr><th></th><th></th><th>No.Operación</th><th>Estado</th><th>Tipo de Evento</th><th>Fecha Registro</th><th>Fecha último Estado</th><th>Trabajador</th></tr>{rows}</table></div>")
    elif sec=='carga':
        emp_opts = "<option>AQUANQA</option><option>AQUANCA II</option>"
        content=wrap(f"""
        <h2 class='c-title'>Carga Masiva</h2>
        <div class='c-card'><div class='tabs'><div class='tab active'>Carga Masiva</div><div class='tab'>Registros de Carga Masiva</div></div></div>
        <div class='c-filter' style='grid-template-columns:1fr 1fr;max-width:980px;margin:auto'><input placeholder='Código/Nombre'><select><option>Grupo</option><option>Actualización de datos</option><option>Baja de trabajador</option></select></div>
        <div class='tile-grid'>
          <div class='c-tile' onclick='abrirCargaModal("actualizar")'><div class='tile-icon'>▤</div><h2>Carga Masiva Actualizar Datos Trabajador</h2><a class='download-corner' onclick='event.stopPropagation()' href='{url_for('descargar_plantilla_carga_contratacion', tipo='actualizar')}'>⬇</a></div>
          <div class='c-tile' onclick='abrirCargaModal("baja")'><div class='tile-icon'>▤</div><h2>Carga Masiva Bajas Trabajador</h2><a class='download-corner' onclick='event.stopPropagation()' href='{url_for('descargar_plantilla_carga_contratacion', tipo='baja')}'>⬇</a></div>
        </div>
        <div id='modalActualizar' class='obs-modal'><div class='obs-box'><div class='obs-head'><h2>Importar registros - Actualización de datos</h2><button type='button' onclick='cerrarCargaModal()'>×</button></div>
          <form method='post' enctype='multipart/form-data' class='obs-form'><input type='hidden' name='accion' value='carga_actualizar_datos'>
            <label>Compañía:</label><select name='empresa'>{emp_opts}</select>
            <label>Archivo Excel:</label><input type='file' name='archivo' accept='.xlsx,.xls' required>
            <div></div><div class='obs-actions'><a class='c-btn gray' href='{url_for('descargar_plantilla_carga_contratacion', tipo='actualizar')}'>⬇ Plantilla</a><button class='c-btn'>Subir</button><button type='button' class='c-btn gray' onclick='cerrarCargaModal()'>Cerrar</button></div>
          </form></div></div>
        <div id='modalBaja' class='obs-modal'><div class='obs-box'><div class='obs-head'><h2>Importar registros - Baja de trabajador</h2><button type='button' onclick='cerrarCargaModal()'>×</button></div>
          <form method='post' enctype='multipart/form-data' class='obs-form'><input type='hidden' name='accion' value='carga_baja_trabajador'>
            <label>Compañía:</label><select name='empresa'>{emp_opts}</select>
            <label>Archivo Excel:</label><input type='file' name='archivo' accept='.xlsx,.xls' required>
            <div></div><div class='obs-actions'><a class='c-btn gray' href='{url_for('descargar_plantilla_carga_contratacion', tipo='baja')}'>⬇ Plantilla</a><button class='c-btn'>Subir</button><button type='button' class='c-btn gray' onclick='cerrarCargaModal()'>Cerrar</button></div>
          </form></div></div>
        <style>.obs-modal{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:99999;align-items:center;justify-content:center;padding:18px}}.obs-modal.show{{display:flex}}.obs-box{{width:min(620px,96vw);background:#fff;color:#111827;border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.35);overflow:hidden}}.obs-head{{display:flex;justify-content:space-between;align-items:center;padding:18px 24px;border-bottom:1px solid #dce2ea}}.obs-head h2{{color:#111827!important;margin:0}}.obs-head button{{border:0;background:transparent;font-size:32px;color:#8a8f98;cursor:pointer}}.obs-form{{display:grid;grid-template-columns:150px 1fr;gap:14px 12px;padding:24px;align-items:center}}.obs-form label{{color:#374151!important;font-weight:700;text-align:right}}.obs-form input,.obs-form select{{background:#fff!important;color:#111827!important;border:1px solid #cdd5df!important;border-radius:8px!important;padding:10px!important}}.obs-actions{{display:flex;gap:10px;justify-content:flex-end;flex-wrap:wrap}}@media(max-width:760px){{.obs-form{{grid-template-columns:1fr}}.obs-form label{{text-align:left}}}}

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
        <script>function abrirCargaModal(t){{cerrarCargaModal();document.getElementById(t==='baja'?'modalBaja':'modalActualizar').classList.add('show');}}function cerrarCargaModal(){{document.querySelectorAll('.obs-modal').forEach(m=>m.classList.remove('show'));}}</script>
        """)
    elif sec=='reportes':
        content=wrap(f"<h2 class='c-title'>Reportes</h2><div class='c-card table-wrap'><table class='c-table'><tr><th></th><th>Código</th><th>Nombre</th><th>Nombre</th><th>Componente</th><th>Creado por</th></tr>{report_rows}</table></div>")
    elif sec=='actualizar':
        content=wrap(f"<h2 class='c-title'>Actualizar Trabajador / Estado laboral</h2><div class='c-card' style='padding:16px'><p class='muted2'>Aquí se controla la base de trabajadores activos. Al cesar se cambia el estado a inactivo, pero NO se elimina ningún documento ni contrato archivado.</p><div class='c-filter' style='grid-template-columns:1fr 1fr'><input placeholder='Buscar por DNI o apellidos'><button class='c-btn'>⌕ Buscar</button></div></div><div class='c-card table-wrap'><table class='c-table'><tr><th>DNI</th><th>Trabajador</th><th>Empresa</th><th>Cargo</th><th>Estado</th><th>Acción</th></tr>{trabajadores_estado_rows}</table></div>")
    elif sec=='tipos_etapa':
        etapa_opts = ''.join([f"<option>{x}</option>" for x in ['Atraccion','Incorporacion','Induccion','Renovacion','Retencion','Separacion']])
        content=wrap(f"""
        <h2 class='c-title'>Tipos de Documento por Etapa</h2>
        <div class='c-bar'><div><select id='filtroStage' class='input' style='background:#fff!important;color:#111!important;width:360px' onchange='filtrarTipoStage(this.value)'><option value=''>Stage</option>{etapa_opts}</select></div><button type='button' class='c-btn' onclick='abrirTipoModal()'>+ Crear Tipo Doc</button></div>
        <div class='c-card table-wrap'><table id='tablaTiposEtapa' class='c-table'><tr><th>Acción</th><th>Estado</th><th>Código Tipo Doc</th><th>Tipo Doc</th><th>Stage</th><th>Requiere Firma Digital</th><th>Requiere Firma Electrónica</th><th>Mandatorio</th><th>Req. confirmación de recepción</th></tr>{tipos_rows}</table></div>
        <div id='tipoModal' class='obs-modal'>
          <div class='obs-box tipo-box'>
            <div class='obs-head'><h2 id='tipoModalTitulo'>Editar Tipo de Documento por Etapa</h2><button type='button' onclick='cerrarTipoModal()'>×</button></div>
            <form method='post' class='obs-form'>
              <input type='hidden' name='accion' value='guardar_tipo_etapa'><input type='hidden' name='tipo_id' id='tipo_id'>
              <label>Código Tipo Doc:</label><input name='codigo' id='tipo_codigo' required>
              <label>Tipo Documento:</label><input name='descripcion' id='tipo_descripcion' required>
              <label>Stage:</label><select name='etapa' id='tipo_etapa'>{etapa_opts}</select>
              <label>Estado:</label><select name='activo' id='tipo_activo'><option value='1'>Active</option><option value='0'>Inactive</option></select>
              <label>Requiere Firma Digital</label><input type='checkbox' name='requiere_firma_digital' id='tipo_fd' value='1'>
              <label>Requiere Firma Electrónica</label><input type='checkbox' name='requiere_firma_electronica' id='tipo_fe' value='1'>
              <label>Mandatorio</label><input type='checkbox' name='obligatorio' id='tipo_obligatorio' value='1'>
              <label>Req. confirmación de recepción</label><input type='checkbox' name='requiere_confirmacion_recepcion' id='tipo_cr' value='1'>
              <div></div><div class='obs-actions'><button class='c-btn'>Guardar</button><button type='button' class='c-btn gray' onclick='cerrarTipoModal()'>Cerrar</button></div>
            </form>
          </div>
        </div>
        <style>.nowrap{{white-space:nowrap}}.icon-btn{{border:0;background:transparent;font-size:20px;cursor:pointer;margin:0 3px}}.select-mini{{background:#fff!important;color:#111!important;border:1px solid #d1d5db;border-radius:999px;padding:6px 10px}}.c-badge.gray{{background:#e5e7eb!important;color:#6b7280!important}}.obs-modal{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:99999;align-items:center;justify-content:center;padding:18px}}.obs-modal.show{{display:flex}}.obs-box{{width:min(760px,96vw);background:#fff;color:#111827;border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.35);overflow:hidden}}.tipo-box{{width:min(700px,96vw)}}.obs-head{{display:flex;justify-content:space-between;align-items:center;padding:18px 24px;border-bottom:1px solid #dce2ea}}.obs-head h2{{color:#111827!important;margin:0}}.obs-head button{{border:0;background:transparent;font-size:32px;color:#8a8f98;cursor:pointer}}.obs-form{{display:grid;grid-template-columns:190px 1fr;gap:14px 12px;padding:24px;align-items:center}}.obs-form label{{color:#374151!important;font-weight:700;text-align:right}}.obs-form input:not([type=checkbox]),.obs-form select{{background:#fff!important;color:#111827!important;border:1px solid #cdd5df!important;border-radius:8px!important;padding:10px!important}}.obs-form input[type=checkbox]{{width:20px;height:20px;accent-color:#1a73e8}}.obs-actions{{display:flex;gap:10px;justify-content:flex-end}}@media(max-width:760px){{.obs-form{{grid-template-columns:1fr}}.obs-form label{{text-align:left}}}}

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
        <script>
        function abrirTipoModal(){{document.getElementById('tipoModalTitulo').innerText='Crear Tipo de Documento por Etapa';document.getElementById('tipo_id').value='';document.getElementById('tipo_codigo').value='';document.getElementById('tipo_descripcion').value='';document.getElementById('tipo_etapa').value='Incorporacion';document.getElementById('tipo_activo').value='1';document.getElementById('tipo_fd').checked=true;document.getElementById('tipo_fe').checked=true;document.getElementById('tipo_obligatorio').checked=true;document.getElementById('tipo_cr').checked=false;document.getElementById('tipoModal').classList.add('show');}}
        function cerrarTipoModal(){{document.getElementById('tipoModal').classList.remove('show');}}
        function editarTipoEtapa(id,codigo,desc,etapa,activo,oblig,fd,fe,cr){{document.getElementById('tipoModalTitulo').innerText='Editar Tipo de Documento por Etapa';document.getElementById('tipo_id').value=id;document.getElementById('tipo_codigo').value=codigo;document.getElementById('tipo_descripcion').value=desc;document.getElementById('tipo_etapa').value=etapa||'Incorporacion';document.getElementById('tipo_activo').value=activo;document.getElementById('tipo_fd').checked=fd==='1';document.getElementById('tipo_fe').checked=fe==='1';document.getElementById('tipo_obligatorio').checked=oblig==='1';document.getElementById('tipo_cr').checked=cr==='1';document.getElementById('tipoModal').classList.add('show');}}
        function filtrarTipoStage(v){{v=(v||'').toLowerCase();document.querySelectorAll('#tablaTiposEtapa tr').forEach((tr,i)=>{{if(i===0)return;let st=(tr.children[4]?.innerText||'').toLowerCase();tr.style.display=(!v||st===v)?'':'none';}});}}
        </script>
        """)
    elif sec=='tipo_empleado':
        grupo_opts = ''.join(["<option>"+x+"</option>" for x in ['Contratación Trabajadores','Renovaciones de Contrato']])
        html_tipo = """
        <h2 class='c-title'>Tipos de Documento de Empleado</h2>
        <div class='c-bar'><div><input id='filtroTipoEmp' class='input' style='background:#fff!important;color:#111!important;max-width:520px' placeholder='Nombre / Código' onkeyup='filtrarTabla(&quot;tablaTipoEmpleado&quot;, this.value)'><br><br><button class='c-btn' type='button'>⌕ Buscar</button> <button class='c-btn gray' type='button' onclick='document.getElementById(&quot;filtroTipoEmp&quot;).value=&quot;&quot;;filtrarTabla(&quot;tablaTipoEmpleado&quot;,&quot;&quot;)'>Limpiar</button></div><button type='button' class='c-btn' onclick='abrirTipoEmpleado()'>+ Crear Tipo Doc</button></div>
        <div class='c-card table-wrap'><table id='tablaTipoEmpleado' class='c-table'><tr><th>Acción</th><th>Estado</th><th>Código</th><th>Descripción</th><th>ShortName</th><th>Grupo</th></tr>__ROWS__</table></div>
        <div id='tipoEmpleadoModal' class='obs-modal'>
          <div class='obs-box tipo-box'><div class='obs-head'><h2 id='tipoEmpleadoTitulo'>Crear tipo documento de empleado</h2><button type='button' onclick='cerrarTipoEmpleado()'>×</button></div>
            <form method='post' class='obs-form'>
              <input type='hidden' name='accion' value='guardar_tipo_empleado'><input type='hidden' name='doc_emp_id' id='doc_emp_id'>
              <label>Code</label><input name='codigo' id='doc_emp_codigo'>
              <label>Descripción</label><input name='descripcion' id='doc_emp_desc' required>
              <label>Nombre Corto</label><input name='nombre_corto' id='doc_emp_corto' required>
              <label>Grupo</label><select name='grupo' id='doc_emp_grupo'>__GRUPOS__</select>
              <label>Estado</label><select name='activo' id='doc_emp_activo'><option value='1'>Active</option><option value='0'>Inactive</option></select>
              <label>Permite generar y/o adjuntar doc. sin cargo</label><input type='checkbox' name='permite_doc_sin_cargo' id='doc_emp_sin_cargo' value='1'>
              <label>Requiere validación</label><input type='checkbox' name='requiere_validacion' id='doc_emp_validacion' value='1'>
              <label>El documento tiene fecha de vencimiento</label><input type='checkbox' name='tiene_fecha_vencimiento' id='doc_emp_vencimiento' value='1'>
              <label>El documento es un contrato</label><input type='checkbox' name='es_contrato' id='doc_emp_contrato' value='1'>
              <div></div><div class='obs-actions'><button class='c-btn'>Guardar</button><button type='reset' class='c-btn gray'>Limpiar</button><button type='button' onclick='cerrarTipoEmpleado()' class='c-btn gray'>Atrás</button></div>
            </form></div></div>
        <style>.nowrap{white-space:nowrap}.icon-btn{border:0;background:transparent;font-size:20px;cursor:pointer;margin:0 3px}.select-mini{background:#fff!important;color:#111!important;border:1px solid #d1d5db;border-radius:999px;padding:6px 10px}.obs-modal{display:none;position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:99999;align-items:center;justify-content:center;padding:18px}.obs-modal.show{display:flex}.obs-box{width:min(900px,96vw);background:#fff;color:#111827;border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.35);overflow:hidden}.tipo-box{width:min(920px,96vw)}.obs-head{display:flex;justify-content:space-between;align-items:center;padding:18px 24px;border-bottom:1px solid #dce2ea}.obs-head h2{margin:0;color:#111827!important}.obs-head button{border:0;background:transparent;font-size:32px;color:#8a8f98;cursor:pointer}.obs-form{display:grid;grid-template-columns:260px 1fr;gap:18px 14px;padding:24px;align-items:center}.obs-form label{font-weight:700;text-align:right;color:#111827!important}.obs-form input:not([type=checkbox]),.obs-form select,.obs-form textarea{background:#fff!important;color:#111827!important;border:1px solid #cdd5df!important;border-radius:8px!important;padding:10px!important}.obs-form input[type=checkbox]{width:20px;height:20px;accent-color:#1a73e8}.obs-actions{display:flex;gap:10px;justify-content:flex-end}@media(max-width:760px){.obs-form{grid-template-columns:1fr}.obs-form label{text-align:left}}

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
        <script>
        function abrirTipoEmpleado(){document.getElementById('tipoEmpleadoTitulo').innerText='Crear tipo documento de empleado';document.getElementById('doc_emp_id').value='';document.getElementById('doc_emp_codigo').value='';document.getElementById('doc_emp_desc').value='';document.getElementById('doc_emp_corto').value='';document.getElementById('doc_emp_grupo').value='Contratación Trabajadores';document.getElementById('doc_emp_activo').value='1';['doc_emp_sin_cargo','doc_emp_validacion','doc_emp_vencimiento','doc_emp_contrato'].forEach(id=>document.getElementById(id).checked=false);document.getElementById('tipoEmpleadoModal').classList.add('show');}
        function cerrarTipoEmpleado(){document.getElementById('tipoEmpleadoModal').classList.remove('show');}
        function editarTipoEmpleado(id,cod,desc,corto,grupo,activo,sinCargo,validacion,vencimiento,contrato){document.getElementById('tipoEmpleadoTitulo').innerText='Editar: '+desc;document.getElementById('doc_emp_id').value=id;document.getElementById('doc_emp_codigo').value=cod;document.getElementById('doc_emp_desc').value=desc;document.getElementById('doc_emp_corto').value=corto;document.getElementById('doc_emp_grupo').value=grupo||'Contratación Trabajadores';document.getElementById('doc_emp_activo').value=activo;document.getElementById('doc_emp_sin_cargo').checked=sinCargo==='1';document.getElementById('doc_emp_validacion').checked=validacion==='1';document.getElementById('doc_emp_vencimiento').checked=vencimiento==='1';document.getElementById('doc_emp_contrato').checked=contrato==='1';document.getElementById('tipoEmpleadoModal').classList.add('show');}
        function filtrarTabla(id,q){q=(q||'').toLowerCase();document.querySelectorAll('#'+id+' tr').forEach((tr,i)=>{if(i===0)return;tr.style.display=tr.innerText.toLowerCase().includes(q)?'':'none';});}
        </script>
        """.replace('__ROWS__', tipo_emp_rows).replace('__GRUPOS__', grupo_opts)
        content=wrap(html_tipo)
    elif sec=='cargo':
        html_cargo = """
        <h2 class='c-title'>Cargos</h2>
        <div class='c-bar'><div><input id='filtroCargo' class='input' style='background:#fff!important;color:#111!important;max-width:520px' placeholder='Nombre' onkeyup='filtrarTabla(&quot;tablaCargos&quot;, this.value)'><br><br><button class='c-btn' type='button'>⌕ Buscar</button> <button class='c-btn gray' type='button' onclick='document.getElementById(&quot;filtroCargo&quot;).value=&quot;&quot;;filtrarTabla(&quot;tablaCargos&quot;,&quot;&quot;)'>Limpiar</button></div><button type='button' class='c-btn' onclick='abrirCargo()'>+ Crear Cargo</button></div>
        <div class='c-card table-wrap'><table id='tablaCargos' class='c-table'><tr><th>Acción</th><th>Estado</th><th>Código</th><th>Descripción</th><th>NombreCorto</th></tr>__ROWS__</table></div>
        <div id='cargoModal' class='obs-modal'><div class='obs-box tipo-box'><div class='obs-head'><h2 id='cargoTitulo'>Crear Cargo</h2><button type='button' onclick='cerrarCargo()'>×</button></div>
          <form method='post' class='obs-form'>
            <input type='hidden' name='accion' value='guardar_cargo'><input type='hidden' name='cargo_id' id='cargo_id'>
            <label>Codigo</label><input name='codigo' id='cargo_codigo'>
            <label>Nombre</label><input name='nombre' id='cargo_nombre' required>
            <label>Nombre Corto</label><input name='nombre_corto' id='cargo_corto' required>
            <label>Resumen</label><input name='resumen' id='cargo_resumen'>
            <label>Funciones:</label><textarea name='funciones' id='cargo_funciones' rows='4'></textarea>
            <label>Estado</label><select name='activo' id='cargo_activo'><option value='1'>Active</option><option value='0'>Inactive</option></select>
            <div></div><div class='obs-actions'><button class='c-btn'>Guardar</button><button type='button' onclick='cerrarCargo()' class='c-btn gray'>Atrás</button></div>
          </form></div></div>
        <style>.nowrap{white-space:nowrap}.icon-btn{border:0;background:transparent;font-size:20px;cursor:pointer;margin:0 3px}.select-mini{background:#fff!important;color:#111!important;border:1px solid #d1d5db;border-radius:999px;padding:6px 10px}.obs-modal{display:none;position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:99999;align-items:center;justify-content:center;padding:18px}.obs-modal.show{display:flex}.obs-box{width:min(900px,96vw);background:#fff;color:#111827;border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.35);overflow:hidden}.tipo-box{width:min(920px,96vw)}.obs-head{display:flex;justify-content:space-between;align-items:center;padding:18px 24px;border-bottom:1px solid #dce2ea}.obs-head h2{margin:0;color:#111827!important}.obs-head button{border:0;background:transparent;font-size:32px;color:#8a8f98;cursor:pointer}.obs-form{display:grid;grid-template-columns:220px 1fr;gap:18px 14px;padding:24px;align-items:center}.obs-form label{font-weight:700;text-align:right;color:#111827!important}.obs-form input:not([type=checkbox]),.obs-form select,.obs-form textarea{background:#fff!important;color:#111827!important;border:1px solid #cdd5df!important;border-radius:8px!important;padding:10px!important}.obs-actions{display:flex;gap:10px;justify-content:flex-end}@media(max-width:760px){.obs-form{grid-template-columns:1fr}.obs-form label{text-align:left}}

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
        <script>
        function abrirCargo(){document.getElementById('cargoTitulo').innerText='Crear Cargo';document.getElementById('cargo_id').value='';document.getElementById('cargo_codigo').value='';document.getElementById('cargo_nombre').value='';document.getElementById('cargo_corto').value='';document.getElementById('cargo_resumen').value='';document.getElementById('cargo_funciones').value='';document.getElementById('cargo_activo').value='1';document.getElementById('cargoModal').classList.add('show');}
        function cerrarCargo(){document.getElementById('cargoModal').classList.remove('show');}
        function editarCargo(id,cod,nombre,corto,resumen,funciones,activo){document.getElementById('cargoTitulo').innerText='Editar: '+nombre;document.getElementById('cargo_id').value=id;document.getElementById('cargo_codigo').value=cod;document.getElementById('cargo_nombre').value=nombre;document.getElementById('cargo_corto').value=corto;document.getElementById('cargo_resumen').value=resumen;document.getElementById('cargo_funciones').value=funciones;document.getElementById('cargo_activo').value=activo;document.getElementById('cargoModal').classList.add('show');}
        function filtrarTabla(id,q){q=(q||'').toLowerCase();document.querySelectorAll('#'+id+' tr').forEach((tr,i)=>{if(i===0)return;tr.style.display=tr.innerText.toLowerCase().includes(q)?'':'none';});}
        </script>
        """.replace('__ROWS__', cargo_rows)
        content=wrap(html_cargo)
    elif sec in ['maestros','observados','tipos_etapa','tipo_empleado','cargo']:
        content=wrap(f"""
        <h2 class='c-title'>Listas de trabajadores observados</h2>
        <div class='c-bar'>
          <div>
            <input class='input' style='background:#fff!important;color:#111!important;max-width:520px' placeholder='Nombre / Num. Documento' onkeyup='filtrarObs(this.value)'>
            <br><br><button class='c-btn' type='button'>⌕ Buscar</button> <button class='c-btn gray' type='button' onclick='location.href=location.pathname+"?sec=observados"'>Limpiar</button>
          </div>
          <button type='button' class='c-btn' onclick='abrirObsModal()'>+ Crear trabajador observado</button>
        </div>
        <div class='toolbar'>⚙ Acción ▾ &nbsp; ⬇ Descargar ▾</div>

        <div id='obsModal' class='obs-modal'>
          <div class='obs-box'>
            <div class='obs-head'><h2>Crear trabajador observado</h2><button type='button' onclick='cerrarObsModal()'>×</button></div>
            <form method='post' class='obs-form'>
              <input type='hidden' name='accion' value='crear_observado'>
              <label>Tipo Persona:</label><div class='radio-line'><label><input type='radio' name='tipo_persona' value='Trabajador' checked onchange='toggleObsTipo()'> Trabajador</label><label><input type='radio' name='tipo_persona' value='Externo' onchange='toggleObsTipo()'> Externo</label></div>
              <label>Trabajador:</label><input name='trabajador_sel' id='obsTrabajador' list='obsTrabajadores' placeholder='Buscar trabajador / DNI' oninput='copiarDniObs(this.value)'><datalist id='obsTrabajadores'>{opt_trab}</datalist>
              <label>DNI / Documento:</label><input name='numero_documento' id='obsDni' placeholder='DNI obligatorio'>
              <label>Nombre:</label><input name='nombre_externo' id='obsNombre' placeholder='Nombre externo o se toma de trabajador'>
              <label>Motivo:</label><select name='motivo' id='obsMotivo'><option value=''>Motivos</option>{motivo_options}</select>
              <label>Motivo digitado:</label><input name='motivo_otro' placeholder='Opcional: digitar motivo personalizado'>
              <label>Nivel Restricción:</label><select name='nivel_restriccion'><option value=''>Niveles</option>{nivel_options}</select>
              <label>Comentario:</label><input name='comentario' placeholder='Comentario / detalle'>
              <div></div><div class='obs-actions'><button class='c-btn'>Guardar</button><button type='reset' class='c-btn gray'>Limpiar</button><button type='button' onclick='cerrarObsModal()' class='c-btn gray'>Cerrar</button></div>
            </form>
          </div>
        </div>

        <div class='c-card'>
          <div class='tabs'><div class='tab active' onclick='showObsTab("activos",this)'>Trabajadores observados</div><div class='tab' onclick='showObsTab("anulados",this)'>Lista de anulados</div></div>
          <div id='tabObsActivos' class='table-wrap obs-tab'><table id='tablaObs' class='c-table'><tr><th></th><th>Acción</th><th>Estado</th><th>Tipo Documento</th><th>Número Documento</th><th>Nombre</th><th>Motivo</th><th>Nivel Restricción</th><th>Comentario</th><th>Creado Por</th><th>Fecha/Hora Creación</th><th>Editor Por</th><th>Fecha/Hora Edición</th></tr>{obs_rows}</table></div>
          <div id='tabObsAnulados' class='table-wrap obs-tab' style='display:none'><table class='c-table'><tr><th>Log</th><th>Tipo Documento</th><th>Número Documento</th><th>Nombre</th><th>Motivo</th><th>Nivel Restricción</th><th>Comentario</th><th>Creado Por</th><th>Fecha/Hora Creación</th><th>Editor Por</th><th>Fecha/Hora Edición</th></tr>{obs_anulados_rows}</table></div>
        </div>
        <style>.obs-modal{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:99999;align-items:center;justify-content:center;padding:18px}}.obs-modal.show{{display:flex}}.obs-box{{width:min(760px,96vw);background:#fff;color:#111827;border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.35);overflow:hidden}}.obs-head{{display:flex;justify-content:space-between;align-items:center;padding:18px 24px;border-bottom:1px solid #dce2ea}}.obs-head h2{{color:#111827!important;margin:0}}.obs-head button{{border:0;background:transparent;font-size:32px;color:#8a8f98;cursor:pointer}}.obs-form{{display:grid;grid-template-columns:190px 1fr;gap:14px 12px;padding:24px;align-items:center}}.obs-form label{{color:#374151!important;font-weight:700;text-align:right}}.obs-form input,.obs-form select{{background:#fff!important;color:#111827!important;border:1px solid #cdd5df!important;border-radius:8px!important;padding:10px!important}}.radio-line{{display:flex;gap:24px}}.radio-line label{{text-align:left!important;font-weight:500!important}}.obs-actions{{display:flex;gap:10px;justify-content:flex-end}}@media(max-width:760px){{.obs-form{{grid-template-columns:1fr}}.obs-form label{{text-align:left}}}}

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}

</style>
        <script>
        function abrirObsModal(){{document.getElementById('obsModal').classList.add('show');}}
        function cerrarObsModal(){{document.getElementById('obsModal').classList.remove('show');}}
        function showObsTab(tipo,el){{document.getElementById('tabObsActivos').style.display=tipo==='activos'?'block':'none';document.getElementById('tabObsAnulados').style.display=tipo==='anulados'?'block':'none';document.querySelectorAll('.tabs .tab').forEach(t=>t.classList.remove('active'));el.classList.add('active');}}
        function copiarDniObs(v){{const m=(v||'').match(/(\d{{8}})/); if(m) document.getElementById('obsDni').value=m[1];}}
        function toggleObsTipo(){{const tipo=document.querySelector('input[name=tipo_persona]:checked').value;document.getElementById('obsTrabajador').disabled=(tipo==='Externo');}}
        function filtrarObs(q){{q=(q||'').toLowerCase();document.querySelectorAll('#tablaObs tr').forEach((tr,i)=>{{if(i===0)return;tr.style.display=tr.innerText.toLowerCase().includes(q)?'':'none';}});}}
        </script>
        """)
    elif sec=='anuncios':
        content=wrap("<h2 class='c-title'>Anuncios de la empresa</h2><div class='c-bar'><div class='c-form'><b>Fecha Registro</b><span><input placeholder='Desde'> - <input placeholder='Hasta'></span><b>Nombre</b><input></div><a class='c-btn'>+ Crear Anuncio</a></div><button class='c-btn'>⌕ Buscar</button> <button class='c-btn gray'>Limpiar</button><br><br><form class='anuncio-upload' method='post' enctype='multipart/form-data'><input type='hidden' name='accion' value='anuncio'><h2>Subir anuncio multimedia</h2><p class='muted2'>Acepta video MP4, PDF, imagen o documento para comunicar a trabajadores.</p><input name='titulo' placeholder='Título del anuncio'><br><br><input type='file' name='archivo' accept='.mp4,.pdf,.png,.jpg,.jpeg,.doc,.docx' required><br><br><button class='c-btn'>Subir anuncio</button><div class='video-box'><b>Vista previa MP4</b><video controls></video></div></form>")
    elif sec=='renovacion':
        content=wrap(f"<h2 class='c-title'>Renovación masiva de contratos</h2><div class='c-form'><b>Renovar por:</b><span>Meses <input type='checkbox' checked> Fecha Termino</span><b>Fecha Termino:</b><input placeholder='d/MM/yyyy'><b>Meses:</b><input type='number' value='0'></div><div class='toolbar'>🔎 Filtros &nbsp; ⚙ Acción ▾ &nbsp; ⬇ Descargar</div><div class='c-card table-wrap'><table class='c-table'><tr><th></th><th>Código</th><th>Apellidos y Nombres</th><th>Modalidad</th><th>FI Planilla</th><th>Fecha Migración</th><th>FI Contrato</th><th>FF Contrato</th><th>Firmado</th><th>Archivado</th><th>Nro File</th></tr>{renov_rows}</table></div>")
    elif sec=='ficha':
        # FICHA TRABAJADOR MEJORADA: búsqueda real por DNI y pestañas funcionales.
        dni_sel = normalizar_dni(request.args.get('dni')) or (sample_trab['dni'] if sample_trab else '')
        trabajador_sel = get_trabajador(dni_sel) if dni_sel else None
        if not trabajador_sel and sample_trab:
            dni_sel = sample_trab['dni']; trabajador_sel = get_trabajador(dni_sel)

        def rv(row, key, default=''):
            try:
                return row[key] if row and key in row.keys() and row[key] not in (None, '') else default
            except Exception:
                return default

        with db() as con:
            ficha_docs = con.execute("""SELECT id,dni,categoria,tipo,periodo,detalle,estado,archivo_nombre,fecha_subida,uploaded_by,ruta_archivo
                                        FROM documentos WHERE dni=? OR categoria='empresa'
                                        ORDER BY COALESCE(periodo,'' ) DESC, id DESC LIMIT 500""", (dni_sel,)).fetchall() if dni_sel else []
            ficha_contratos = con.execute("""SELECT * FROM contratacion_docs WHERE dni=? ORDER BY id DESC LIMIT 300""", (dni_sel,)).fetchall() if dni_sel else []
            vac_periodos = con.execute("""SELECT periodo_inicio,periodo_fin,fecha_ingreso,saldo,dias_ganados FROM vacaciones_saldos
                                       WHERE dni=? ORDER BY periodo_inicio DESC, periodo_fin DESC""", (dni_sel,)).fetchall() if dni_sel else []

        opt_trab_buscar=''.join([f"<option value='{h(r['dni'])}'>{h(r['dni'])} - {h(r['nombre'])}</option>" for r in trabajadores])
        nombre = h(rv(trabajador_sel,'nombre','SIN TRABAJADOR'))
        correo = h(rv(trabajador_sel,'correo',''))
        celular = h(rv(trabajador_sel,'telefono', rv(trabajador_sel,'celular','')))
        estado_txt = 'Activo' if int(rv(trabajador_sel,'activo',1) or 0)==1 else 'Baja'
        estado_cls = 'ok' if estado_txt == 'Activo' else 'bad'
        fecha_ing = h(fecha_sin_hora(rv(trabajador_sel,'fecha_ingreso','')))
        fecha_reg = h(fecha_sin_hora(rv(trabajador_sel,'fecha_registro','')))
        cargo = h(rv(trabajador_sel,'cargo',''))
        area = h(rv(trabajador_sel,'area','NINGUNO'))
        empresa = h(rv(trabajador_sel,'empresa',''))
        planilla = h(rv(trabajador_sel,'planilla',''))
        jefe = h(rv(trabajador_sel,'jefe_nombre','')) or h(rv(trabajador_sel,'jefe_dni',''))

        docs_rows = ''
        current_period = None
        for d in ficha_docs:
            periodo = h(d['periodo'] or 'SIN PERIODO')
            if periodo != current_period:
                current_period = periodo
                docs_rows += f"<tr class='period-row'><td colspan='7'>▸ Movimiento / Periodo: {periodo}</td></tr>"
            open_link = f"<a class='icon-btn' title='Abrir documento' href='/ver/{d['id']}'>📄</a>" if d['id'] else ''
            docs_rows += f"""<tr><td>{open_link}</td><td>{h(d['dni'] or 'EMPRESA')}</td><td>{h(d['tipo'])}</td><td>{h(d['detalle'])}</td><td>{h(d['estado'] or 'Pendiente')}</td><td>{h(fecha_sin_hora(d['fecha_subida']))}</td><td>{h(d['uploaded_by'])}</td></tr>"""
        if not docs_rows:
            docs_rows = "<tr><td colspan='7'>No hay documentos acumulados para este trabajador. Carga documentos o ejecuta sincronización de carpetas.</td></tr>"

        contratos_rows = ''
        for c in ficha_contratos:
            log_url = url_for('contratacion_doc_log', doc_id=c['id'])
            file_url = url_for('contratacion_doc_archivo', doc_id=c['id']) if c['ruta_archivo'] else '#'
            firmado = "<span class='check-green'>✓</span>" if str(c['estado']).upper() in ['FIRMADO','ARCHIVADO','GENERADO'] else "<span class='check-gray'>-</span>"
            archivado = "<span class='check-green'>✓</span>" if c['ruta_archivo'] else "<span class='check-gray'>-</span>"
            contratos_rows += f"""<tr>
              <td><a class='icon-btn' title='Ver log del contrato' href='{log_url}'>📄</a> <a class='icon-btn' title='Abrir archivo' href='{file_url}'>⬇</a></td>
              <td>{h(fecha_sin_hora(c['fecha_registro']))}</td><td>{h(c['etapa'])}</td><td>{firmado}</td><td>{archivado}</td>
              <td>{h(c['tipo_doc'])}</td><td>{h(c['estado'])}</td><td>{h(c['uploaded_by'])}</td>
            </tr>"""
        if not contratos_rows:
            contratos_rows = "<tr><td colspan='8'>No hay contratos registrados para este DNI.</td></tr>"

        periodos_html = ''.join([f"<span class='mini-chip'>{h(p['periodo_inicio'])}/{h(p['periodo_fin'])} · saldo {h(p['saldo'])}</span>" for p in vac_periodos]) or '<span class="mini-chip">Sin periodos cargados</span>'

        content=wrap(f"""
        <h2 class='c-title'>Ficha Trabajador</h2>
        <form method='get' action='/admin/contratacion' class='ficha-search'>
          <input type='hidden' name='sec' value='ficha'>
          <input name='dni' value='{h(dni_sel)}' list='trabajadores_ficha_list' placeholder='Buscar por DNI'>
          <datalist id='trabajadores_ficha_list'>{opt_trab_buscar}</datalist>
          <button class='c-btn'>Buscar</button>
        </form>
        <div class='ficha-profile'>
          <div class='avatar-panel'><div class='avatar-big'>👤</div><span class='status-dot {estado_cls}'>{estado_txt}</span></div>
          <div class='profile-main'>
            <h2>{nombre}</h2><p><b>DNI:</b> {h(dni_sel)} &nbsp; <b>Empresa:</b> {empresa}</p>
            <p><b>Dirección:</b> {h(rv(trabajador_sel,'direccion','')) or 'Sin dirección registrada'}</p>
            <p>✉ {correo or 'Sin correo'} &nbsp;&nbsp; 📱 {celular or 'Sin celular'}</p>
            <div class='created-box'><b>Fecha de Creación:</b> {fecha_reg or '-'} <b>Creado por:</b> SISTEMA PORTAL HR PRO</div>
          </div>
          <div class='profile-col'><p><b>Gerencia:</b> NINGUNO</p><p><b>Área:</b> {area}</p><p><b>Puesto:</b> {cargo}</p><p><b>Supervisor:</b> {jefe or '-'}</p><p><b>Planilla:</b> {planilla or '-'}</p></div>
          <div class='profile-col'><p><b>Estado:</b> <span class='status-pill {estado_cls}'>{estado_txt}</span></p><p><b>Fecha de Ingreso:</b> {fecha_ing or '-'}</p><p><b>Fecha de Cese:</b> {h(rv(trabajador_sel,'fecha_cese','')) or '-'}</p><p><b>Cargo:</b> {cargo}</p><p><b>Sindicalizado:</b> {h(rv(trabajador_sel,'sindicalizado','NO'))}</p></div>
        </div>
        <div class='ficha-tabs c-card'>
          <button type='button' class='tab active' onclick="openFichaTab(event,'datos')">Datos Laborales</button>
          <button type='button' class='tab' onclick="openFichaTab(event,'documentos')">Documentos</button>
          <button type='button' class='tab' onclick="openFichaTab(event,'contratos')">Contratos</button>
        </div>
        <div id='tab-datos' class='ficha-tab-content active'>
          <div class='c-card laboral-grid'>
            <label>Discapacidad:<input readonly value='{h(rv(trabajador_sel,'discapacidad',''))}'></label>
            <label>Fecha Ingreso Planilla:<input readonly value='{fecha_ing}'></label>
            <label>Código:<input readonly value='{h(dni_sel)}'></label>
            <label>Sindicalizado:<input readonly value='{h(rv(trabajador_sel,'sindicalizado','NO'))}'></label>
            <label>Cargo:<input readonly value='{cargo}'></label>
            <label>Categoría Ocupacional:<input readonly value='{h(rv(trabajador_sel,'categoria_ocupacional',''))}'></label>
            <label>Centro de Costo:<input readonly value='{h(rv(trabajador_sel,'centro_costo',''))}'></label>
            <label>Situación Especial:<input readonly value='{h(rv(trabajador_sel,'situacion_especial',''))}'></label>
            <label>Modalidad de Trabajo:<input readonly value='{h(rv(trabajador_sel,'modalidad',''))}'></label>
            <label>Nivel Educativo:<input readonly value='{h(rv(trabajador_sel,'nivel_educativo',''))}'></label>
            <label>Nivel Jerárquico:<input readonly value='{h(rv(trabajador_sel,'nivel_jerarquico',''))}'></label>
            <label>Ocupación:<input readonly value='{h(rv(trabajador_sel,'ocupacion',''))}'></label>
            <label>Planilla:<input readonly value='{planilla}'></label>
            <label>Zona:<input readonly value='{h(rv(trabajador_sel,'zona',''))}'></label>
          </div>
          <div class='periodos-box'><b>Periodos laborales/vacacionales detectados:</b> {periodos_html}</div>
        </div>
        <div id='tab-documentos' class='ficha-tab-content'>
          <div class='toolbar'><a class='c-btn gray' href='/admin/crear_carpetas'>Refrescar carpetas</a> <span>Documentos acumulados por periodos de ingreso, renovación, baja o carga documental.</span></div>
          <div class='c-card table-wrap'><table class='c-table'><tr><th></th><th>Código</th><th>Tipo Documento</th><th>Detalle</th><th>Estado Doc</th><th>Fecha Envío/Subida</th><th>Creado por</th></tr>{docs_rows}</table></div>
        </div>
        <div id='tab-contratos' class='ficha-tab-content'>
          <div class='toolbar'>📄 Hojita = abrir log del contrato &nbsp; | &nbsp; ⬇ = abrir/descargar archivo</div>
          <div class='c-card table-wrap'><table class='c-table'><tr><th>Acción</th><th>Fecha Registro</th><th>Modalidad/Etapa</th><th>Firmado</th><th>Archivado</th><th>Tipo Contrato</th><th>Estado</th><th>Creado por</th></tr>{contratos_rows}</table></div>
        </div>
        <script>
        function openFichaTab(ev,name){{
          document.querySelectorAll('.ficha-tab-content').forEach(x=>x.classList.remove('active'));
          document.querySelectorAll('.ficha-tabs .tab').forEach(x=>x.classList.remove('active'));
          document.getElementById('tab-'+name).classList.add('active'); ev.currentTarget.classList.add('active');
        }}
        </script>
        """)
    elif sec=='plantillas':
        f_nombre_v = html.escape(clean(request.args.get('f_nombre')))
        f_tipo_v = html.escape(clean(request.args.get('f_tipo')))
        f_esquema_v = html.escape(clean(request.args.get('f_esquema')))
        f_cond_v = html.escape(clean(request.args.get('f_condicion')))
        content=wrap(f"""
        <div class='plantilla-top'>
          <h2 class='c-title'>Plantillas</h2>
          <a class='c-btn crear-btn' href='#crearPlantilla'>+ Crear Plantilla</a>
        </div>
        <div class='c-card filter-card' style='padding:18px'>
          <form method='get' action='/admin/contratacion' class='plantilla-filter'>
            <input type='hidden' name='sec' value='plantillas'>
            <b>Nombre Plantilla:</b><input name='f_nombre' value='{f_nombre_v}'>
            <b>Tipo Documento:</b><input name='f_tipo' value='{f_tipo_v}' list='tipos_doc_list_filter'><datalist id='tipos_doc_list_filter'>{opt_tipo}</datalist>
            <b>Esquema:</b><select name='f_esquema'><option value=''></option><option {'selected' if f_esquema_v=='Trabajador Contrato Laboral' else ''}>Trabajador Contrato Laboral</option><option {'selected' if f_esquema_v=='Trabajador Datos Laborales' else ''}>Trabajador Datos Laborales</option><option {'selected' if f_esquema_v=='Esquema Trabajador Datos Laborales GR' else ''}>Esquema Trabajador Datos Laborales GR</option><option {'selected' if f_esquema_v=='Trabajador Declaración Jurada Datos Personales' else ''}>Trabajador Declaración Jurada Datos Personales</option><option {'selected' if f_esquema_v=='Trabajador Declaración Jurada Parentesco' else ''}>Trabajador Declaración Jurada Parentesco</option></select>
            <b>Condición:</b><select name='f_condicion'><option value=''></option><option value='SIN CONDICIONES' {'selected' if f_cond_v=='SIN CONDICIONES' else ''}>SIN CONDICIONES</option><option value='CONDICIONES' {'selected' if f_cond_v=='CONDICIONES' else ''}>CONDICIONES</option></select>
            <span></span><span><button class='c-btn'>⌕ Buscar</button> <a class='c-btn gray' href='/admin/contratacion?sec=plantillas'>Limpiar</a></span>
          </form>
        </div>
        <div id='crearPlantilla' class='modal-prize'>
          <div class='modal-box'>
            <div class='modal-head'>
              <h2>Crear Plantilla</h2>
              <a class='modal-close' href='#' aria-label='Cerrar'>×</a>
            </div>
            <div class='modal-body'>
              <form method='post' enctype='multipart/form-data' class='modal-form'>
                <input type='hidden' name='accion' value='plantilla'>
                <input type='hidden' name='version' value='Version 01'>

                <label class='req'>Tipo Plantilla</label>
                <select name='tipo_plantilla' class='select-soft' required>
                  <option value=''></option>
                  <option>WORD</option>
                  <option>PDF</option>
                </select>
                <div class='modal-help'>Debe seleccionar Tipo Plantilla</div>

                <label class='req'>Tipo Documento</label>
                <input name='tipo_documento' list='tipos_doc_list' required autocomplete='off'>
                <datalist id='tipos_doc_list'>{opt_tipo}</datalist>
                <div class='modal-help'>Debe seleccionar Tipo Documento</div>

                <label class='req'>Nombre Plantilla</label>
                <input name='nombre_plantilla' required autocomplete='off'>
                <div class='modal-help'>Debe ingresar Nombre Archivo</div>

                <label>Descripcion</label>
                <textarea name='descripcion'></textarea>

                <label>Modo de<br>selección de la<br>plantilla</label>
                <select name='condicion' class='select-soft' required>
                  <option value='SIN CONDICIONES'>Utilizar para todos los trabajadores</option>
                  <option value='CONDICIONES'>Usar criterios de selección que cumplan con los datos del trabajador</option>
                </select>

                <label>Esquema</label>
                <select name='esquema' id='modal_esquema'>
                  <option>Trabajador Contrato Laboral</option>
                  <option>Trabajador Datos Laborales</option>
                  <option>Esquema Trabajador Datos Laborales GR</option>
                  <option>Trabajador Declaración Jurada Datos Personales</option>
                  <option>Trabajador Declaración Jurada Parentesco</option>
                </select>

                <label>Archivo plantilla</label>
                <div class='file-row'><input type='file' name='archivo' accept='.pdf,.doc,.docx'></div>

                <div class='modal-actions'>
                  <button class='c-btn' type='submit'>Guardar</button>
                  <button class='c-btn gray' type='reset'>Limpiar</button>
                  <a class='c-btn gray' href='#'>Cerrar</a>
                </div>
              </form>
            </div>
          </div>
        </div>
        <div class='c-card table-wrap'><table class='c-table plantilla-table'><tr><th>Proceso</th><th>Estado</th><th>Nombre Plantilla</th><th>Tipo Documento</th><th>Esquema</th><th>Descripción</th><th>Versión</th><th>Condición</th><th>Nombre Archivo</th></tr>{plantillas_rows or '<tr><td colspan=9>No hay plantillas registradas.</td></tr>'}</table></div>
        """)
    elif sec=='firma':
        opt_docs = ''.join([f"<option value='{d['id']}' data-dni='{h(d['dni'])}' data-trabajador='{h(d['trabajador'])}' data-tipo='{h(d['tipo_doc'])}'>ID {d['id']} - {h(d['dni'])} - {h(d['trabajador'])} - {h(d['tipo_doc'])}</option>" for d in docs])
        rows_docs_lote = ''
        for d in docs[:250]:
            rows_docs_lote += f"""<tr><td><input type='checkbox' class='chk-doc-firma' value='{d['id']}' data-dni='{h(d['dni'])}' data-trabajador='{h(d['trabajador'])}' data-tipo='{h(d['tipo_doc'])}'></td><td>{d['id']}</td><td>{h(d['dni'])}</td><td>{h(d['trabajador'])}</td><td>{h(d['tipo_doc'])}</td><td><span class='estado-soft'>{h(d['estado'] or '')}</span></td></tr>"""
        doc_cards_firma = ''
        for d in docs[:250]:
            doc_cards_firma += f"""<label class='doc-sign-card'><input type='checkbox' class='chk-doc-firma' value='{d['id']}' data-dni='{h(d['dni'])}' data-trabajador='{h(d['trabajador'])}' data-tipo='{h(d['tipo_doc'])}' checked><span class='doc-icon'>W</span><span class='doc-info'><b>{h(d['tipo_doc'] or 'DOCUMENTO')}</b><small>{h(d['trabajador'] or '')} · DNI {h(d['dni'] or '')}</small><small>Estado: {h(d['estado'] or 'Pendiente')}</small></span></label>"""
        # También mostrar las plantillas Word/PDF activas en la pestaña Firma/Facial/Digital.
        # Si cargaste un contrato en Plantilla Documentos, aparece aquí inmediatamente.
        for plf in plantillas[:250]:
            tiene_archivo = bool(plf['ruta_archivo'] or plf['archivo_nombre'])
            if plf['activo'] and tiene_archivo:
                ext = (plf['archivo_nombre'] or '').split('.')[-1].upper() if (plf['archivo_nombre'] or '') else 'DOCX'
                doc_cards_firma += f"""<label class='doc-sign-card plantilla-sign-card'><input type='checkbox' class='chk-doc-firma chk-plantilla-firma' value='PL{plf['id']}' data-dni='74324033' data-trabajador='TRABAJADOR SELECCIONADO' data-tipo='{h(plf['tipo_documento'] or plf['nombre_plantilla'])}' checked><span class='doc-icon'>{'PDF' if ext == 'PDF' else 'W'}</span><span class='doc-info'><b>{h(plf['archivo_nombre'] or plf['nombre_plantilla'] or 'PLANTILLA')}</b><small>Tipo: {h(plf['tipo_documento'] or 'CONTRATO TRABAJADOR')}</small><small>Estado: ACTIVO · Lista para firma</small></span></label>"""

        rows_plantillas_firma = ''
        for plf in plantillas[:250]:
            estado_pl = 'ACTIVA' if plf['activo'] else 'INACTIVA'
            rows_plantillas_firma += f"""<tr><td><input type='checkbox' name='plantillas_firma' value='{plf['id']}' {'disabled' if not plf['activo'] else ''}></td><td>{plf['id']}</td><td>{h(plf['nombre_plantilla'])}</td><td>{h(plf['tipo_documento'])}</td><td>{h(plf['archivo_nombre'])}</td><td><span class='{'ok-chip' if plf['activo'] else 'pend-chip'}'>{estado_pl}</span></td></tr>"""

        rows_firma = ''
        for r in firma_sols:
            token = r['firma_token'] if 'firma_token' in r.keys() and r['firma_token'] else ''
            link = firma_url_token(token) if token else ''
            evidencia = r['evidencia_ref'] if 'evidencia_ref' in r.keys() and r['evidencia_ref'] else ''
            estado_val = r['validacion_estado'] if 'validacion_estado' in r.keys() and r['validacion_estado'] else ''
            rows_firma += f"""<tr>
              <td>{r['id']}</td><td>{h(r['dni'])}</td><td>{h(r['trabajador'])}</td><td>{h(r['metodo'])}</td>
              <td><span class='estado-pill'>{h(r['estado'])}</span><br><small>{h(estado_val)}</small></td>
              <td>{h(r['fecha_envio'])}</td><td>{h(r['fecha_firma'] or '')}</td>
              <td>{'<a class="c-btn gray mini-btn" href="'+h(link)+'" target="_blank">Abrir móvil/web</a>' if link else '-'}</td>
              <td>{'<span class="ok-chip">Con evidencia</span>' if evidencia else '<span class="pend-chip">Pendiente</span>'}</td>
              <td>{h(r['observacion'] or '')}</td>
            </tr>"""
        camara_demo_url = url_for('firma_camara_demo')
        content=wrap(f"""
        <div class='firma-page firma-boceto-final'>
          <div class='firma-topbar'>
            <div class='title-wrap'><div class='title-icon'>📸</div><div><h1>Firma / Facial / Digital</h1><p>Captura facial en tiempo real y firma los documentos seleccionados.</p></div></div>
            <a class='btn-back' href='/admin/contratacion?sec=bandeja'>← Volver a bandeja</a>
          </div>
          <div class='person-strip'>
            <div class='strip-item'><span class='strip-ico'>👤</span><div><small>Trabajador</small><b id='stripTrabajador'>JOSE QUITO</b></div></div>
            <div class='strip-item'><span class='strip-ico'>🪪</span><div><small>DNI</small><b id='stripDni'>72244462</b></div></div>
            <div class='strip-item'><span class='strip-ico'>🗓️</span><div><small>Fecha y hora</small><b id='stripFecha'>26/05/2026 10:34:22</b></div></div>
            <div class='strip-item'><span class='strip-ico'>🖊️</span><div><small>Método de firma</small><b>FACIAL + FIRMA DIGITAL</b></div></div>
          </div>
          <div class='firma-grid-boceto-main'>
            <div class='firma-card-b camera-card-b'>
              <h2>Activación de cámara</h2><p class='b-muted'>Captura facial en tiempo real. Al detectar rostro se tomará captura automática y sonará confirmación.</p>
              <div class='cam-wrap cam-boceto'><video id='firmaVideo' autoplay playsinline muted></video><canvas id='firmaCanvas' style='display:none'></canvas><img id='firmaPreview' style='display:none'><div class='face-frame'></div><div class='face-mesh'></div><div id='liveBadge' class='live-badge'>● APAGADA</div><div id='captureToast' class='capture-toast'>✅ Rostro reconocido correctamente<br><small>Captura realizada automáticamente</small></div></div>
              <div id='soundBox' class='sound-ok'><span class='sound-icon'>🔊</span><div><b>¡Captura exitosa!</b><small>Rostro reconocido correctamente</small></div><span class='wave'>▂▃▄▅▆▇▆▅▄▃▂▃▄▅▆▇</span></div>
              <div class='firma-actions boceto-actions'>
                <button type='button' id='btnActivarCamara' class='btn-green' onclick='return firmaStartCam(event)'>🎥 Activar cámara</button>
                <button type='button' class='btn-yellow' onclick='return firmaCapture()'>📸 Capturar evidencia</button>
                <button type='button' class='btn-dark' onclick='return firmaStopCam()'>■ Detener</button>
                <label class='btn-dark filecam-label'>📁 Cámara/archivo<input id='firmaFileCam' type='file' accept='image/*' capture='user' onchange='firmaLoadFileCam(this)' style='display:none'></label>
              </div><p id='firmaCamMsg' class='b-muted'></p>
            </div>
            <div class='firma-card-b docs-panel-b'>
              <h2>Documentos a firmar <span id='docsBadge' class='badge-green'>0</span></h2><p class='b-muted'>Se firmarán automáticamente los siguientes documentos:</p>
              <div class='doc-sign-list'>{doc_cards_firma or '<div class="empty-docs">No hay documentos cargados. Sube Word/PDF en Archivos Trabajador o Plantilla Documentos.</div>'}</div>
              <label class='switch-row'><input type='checkbox' id='firmaMasivaSwitch' checked onchange='marcarTodosFirma(this.checked)'><span>Firma masiva (todos los documentos)</span><span class='info-dot'>i</span></label>
              <div class='ready-box'>🛡️ <div><b>Todo listo para firmar</b><small>Los documentos seleccionados serán firmados automáticamente.</small></div></div>
            </div>
          </div>
          <form method='post' class='firma-card firma-form firma-selector-pro sr-only-form'>
            <input type='hidden' name='accion' value='firma_solicitud'><h3>Selección de contrato / documento a firmar</h3>
            <div class='form-row'><label>Contrato / documento principal</label><select name='documento_id' required>{opt_docs}</select></div>
            <div class='form-row'><label>Método de firma</label><select name='metodo'><option>FACIAL + FIRMA DIGITAL</option><option>RECONOCIMIENTO FACIAL</option><option>FIRMA DIGITAL</option><option>OTP + ACEPTACIÓN</option></select></div>
            <div class='form-row wide'><label>Observación / mensaje para el trabajador</label><input name='observacion' placeholder='Tienes un contrato pendiente por firmar.'></div>
            <div class='firma-actions'><button class='btn-green'>Generar enlace individual</button><a class='btn-dark' href='#bandeja'>Ver bandeja</a></div>
          </form>
          <form method='post' class='firma-progress-bar' onsubmit='return prepararFirmaMasiva()'>
            <input type='hidden' name='accion' value='firma_masiva'><input type='hidden' id='documentos_lote' name='documentos_lote'>
            <div class='progress-left'><b>Progreso de firma</b><div class='steps'><span class='done'>1<small>Verificación facial<br>Completado</small></span><i></i><span class='done'>2<small>Validación<br>Completado</small></span><i></i><span class='done'>3<small>Firma de documentos<br>En proceso...</small></span><i></i><span>4<small>Finalizado<br>Pendiente</small></span></div></div>
            <button class='btn-green btn-firmar'>🖊️ Firmar todos los documentos<br><small id='firmaMassCounter'>0 seleccionados</small></button>
          </form>
          <div id='bandeja' class='firma-card'><h3>🧾 Bandeja de Firmas</h3><div class='table-wrap'><table class='c-table firma-table'><tr><th>ID</th><th>DNI</th><th>Trabajador</th><th>Método</th><th>Estado</th><th>Fecha envío</th><th>Fecha firma</th><th>Link cámara</th><th>Evidencia</th><th>Observación</th></tr>{rows_firma or '<tr><td colspan=10>No hay solicitudes de firma.</td></tr>'}</table></div></div>
          <div class='firma-card'><h3>🔐 Trazabilidad</h3><div class='trace-grid'><span>IP y navegador</span><span>Fecha / hora</span><span>Hash de evidencia</span><span>Selfie/captura</span><span>Estado RENIEC/API</span><span>Documento archivado</span></div></div>
        </div>
        <style>
        .firma-boceto-final{{background:#f5f7fb!important;margin:-10px -12px 0;padding:18px 26px 26px;min-height:calc(100vh - 80px);color:#0f172a!important;font-family:Inter,Segoe UI,Arial,sans-serif!important}}.firma-boceto-final *{{box-sizing:border-box!important;text-shadow:none!important}}.firma-topbar{{display:flex;justify-content:space-between;align-items:center;margin:0 0 18px}}.title-wrap{{display:flex;align-items:center;gap:12px}}.title-icon{{width:42px;height:42px;border-radius:12px;background:#edf2f7;display:grid;place-items:center;font-size:22px;box-shadow:inset 0 0 0 1px #e2e8f0}}.firma-topbar h1{{margin:0;font-size:30px;line-height:1;color:#0b1220;font-weight:1000}}.firma-topbar p{{margin:8px 0 0;color:#64748b;font-weight:800}}.btn-back{{background:#eef1f6;border:1px solid #dce3eb;border-radius:10px;color:#111827;text-decoration:none;font-weight:950;padding:12px 18px;box-shadow:0 6px 16px #0f172a0d}}.person-strip{{display:grid;grid-template-columns:repeat(4,1fr);gap:0;background:#fff;border:1px solid #dfe7ef;border-radius:12px;box-shadow:0 10px 26px #0f172a10;margin-bottom:14px;overflow:hidden}}.strip-item{{display:flex;gap:14px;align-items:center;padding:20px 26px;border-right:1px solid #e5eaf0}}.strip-item:last-child{{border-right:0}}.strip-ico{{font-size:28px;color:#3b82f6}}.strip-item small{{display:block;color:#111827;font-size:12px;font-weight:950;margin-bottom:8px}}.strip-item b{{font-size:13px;color:#020617;font-weight:1000}}.firma-grid-boceto-main{{display:grid;grid-template-columns:1fr 1.12fr;gap:16px}}.firma-card-b{{background:#fff;border:1px solid #e0e6ee;border-radius:12px;box-shadow:0 10px 26px #0f172a10;padding:18px}}.firma-card-b h2{{margin:0 0 8px;color:#0b1220;font-size:22px;font-weight:1000}}.b-muted{{color:#64748b;font-weight:750;line-height:1.45;margin:0 0 14px}}.cam-wrap{{position:relative;overflow:hidden;background:#000;border-radius:15px;min-height:525px;display:grid;place-items:center}}.cam-wrap video,.cam-wrap img{{width:100%;height:525px;object-fit:cover;background:#000;position:relative;z-index:3;display:block}}.cam-wrap video{{transform:scaleX(-1)}}.cam-wrap:not(.cam-live) .face-frame,.cam-wrap:not(.cam-live) .face-mesh,.cam-wrap:not(.capture-ok) #captureToast{{display:none!important}}.cam-error{{background:#fff7ed!important;border-color:#fed7aa!important;color:#9a3412!important}}.filecam-label{{position:relative;overflow:hidden}}.face-frame{{position:absolute;z-index:4;inset:17% 26%;border:3px solid #22c55e;border-radius:22px;pointer-events:none;display:none}}.face-frame:before,.face-frame:after{{content:'';position:absolute;inset:-3px;border-color:#22c55e;border-style:solid;border-width:0}}.face-mesh{{position:absolute;z-index:5;inset:23% 32%;opacity:.58;pointer-events:none;background:radial-gradient(circle,#34d399 1.4px,transparent 2px) 0 0/34px 34px,linear-gradient(32deg,transparent 49%,rgba(52,211,153,.55) 50%,transparent 51%) 0 0/70px 70px,linear-gradient(145deg,transparent 49%,rgba(52,211,153,.38) 50%,transparent 51%) 0 0/80px 80px;border-radius:50%;display:none}}.live-badge{{position:absolute;right:18px;top:18px;z-index:6;background:#11823b;color:#fff;border-radius:999px;padding:8px 16px;font-weight:1000;font-size:12px}}.capture-toast{{position:absolute;left:18px;right:18px;bottom:18px;z-index:7;background:linear-gradient(90deg,#0f5132,#0b5d34);color:#fff;border-radius:12px;padding:16px 22px;font-weight:1000;box-shadow:0 8px 22px #0006;display:none}}.capture-toast small{{display:block;color:#dcfce7;font-weight:800;margin-top:4px}}.sound-ok{{margin:16px 0 14px;background:#eafff1;border:1px solid #b9edcc;border-radius:10px;color:#16a34a;display:none;align-items:center;gap:12px;padding:12px 14px;font-weight:1000}}.sound-icon{{font-size:27px}}.sound-ok small{{display:block;color:#16a34a;font-weight:700}}.wave{{margin-left:auto;letter-spacing:2px}}.boceto-actions{{display:flex;gap:10px;flex-wrap:wrap}}.btn-yellow,.btn-green,.btn-dark{{border:0;border-radius:8px;padding:13px 18px;font-weight:1000;cursor:pointer;text-decoration:none;display:inline-flex;align-items:center;justify-content:center;box-shadow:0 8px 16px #0f172a18}}.btn-yellow{{background:#10b981;color:#111827}}.btn-green{{background:#10b84e;color:#fff!important}}.btn-dark{{background:#334155;color:#fff!important}}.doc-sign-list{{display:grid;gap:12px;margin:18px 0;max-height:505px;overflow:auto;padding:0 4px 0 0}}.doc-sign-card{{display:grid;grid-template-columns:32px 56px 1fr;align-items:center;gap:14px;border:1px solid #dbe3ec;border-radius:12px;background:#fafcff;padding:17px;min-height:88px;box-shadow:0 6px 18px #0f172a08;cursor:pointer}}.doc-sign-card input{{width:20px;height:20px;accent-color:#10b84e}}.doc-icon{{width:45px;height:50px;border-radius:8px;background:#2f67c7;color:#fff;display:grid;place-items:center;font-size:20px;font-weight:1000;box-shadow:inset 0 -8px 0 rgba(0,0,0,.08)}}.doc-info b{{display:block;color:#0f172a;font-size:14px;text-transform:uppercase;margin-bottom:7px}}.doc-info small{{display:block;color:#475569;font-weight:850;margin-top:3px}}.badge-green{{display:inline-flex;align-items:center;justify-content:center;min-width:28px;height:28px;border-radius:999px;background:#10b84e;color:#fff;font-size:14px;vertical-align:middle}}.switch-row{{display:flex;align-items:center;gap:10px;margin:18px 0;color:#111827;font-weight:950}}.switch-row input{{width:42px;height:22px;accent-color:#10b84e}}.info-dot{{width:18px;height:18px;border-radius:50%;display:inline-grid;place-items:center;background:#111827;color:#fff;font-size:12px}}.ready-box{{display:flex;gap:12px;align-items:center;background:#eafff1;border:1px solid #b9edcc;border-radius:10px;color:#16a34a;padding:16px;font-weight:1000}}.ready-box small{{display:block;color:#16a34a;font-weight:700;margin-top:4px}}.firma-progress-bar{{display:grid;grid-template-columns:1fr 300px;gap:22px;align-items:center;background:#fff;border:1px solid #dfe7ef;border-radius:12px;box-shadow:0 10px 26px #0f172a10;margin-top:16px;padding:18px}}.progress-left>b{{display:block;font-size:16px;margin-bottom:14px;color:#0f172a}}.steps{{display:flex;align-items:flex-start;gap:10px}}.steps i{{height:1px;background:#cbd5e1;flex:1;margin-top:17px}}.steps span{{width:34px;height:34px;border-radius:50%;display:grid;place-items:center;background:#e5e7eb;color:#111827;font-weight:1000;position:relative;flex:0 0 auto}}.steps span.done{{background:#10b84e;color:#fff}}.steps small{{position:absolute;top:42px;left:50%;transform:translateX(-50%);width:130px;color:#475569;font-size:10px;text-align:center;line-height:1.35}}.steps .done small{{color:#10b84e}}.btn-firmar{{height:54px;font-size:15px;flex-direction:column}}.btn-firmar small{{font-size:11px;color:#eafff1;margin-top:4px}}.sr-only-form{{display:none!important}}.firma-card{{background:#fff;border:1px solid #e0e6ee;border-radius:12px;padding:18px;margin-top:16px;color:#0f172a}}.firma-table th{{background:#0b2135!important;color:#fff!important}}.firma-table td{{background:white!important;color:#1f2937!important}}.trace-grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:10px}}.trace-grid span{{background:#e0f2fe;border:1px solid #7dd3fc;color:#075985;padding:12px;border-radius:10px;font-weight:950}}.estado-pill,.estado-soft,.ok-chip,.pend-chip{{display:inline-flex;border-radius:999px;padding:7px 10px;font-weight:950}}.estado-pill{{background:#fef3c7;color:#92400e}}.estado-soft{{background:#e0f2fe;color:#075985}}.ok-chip{{background:#dcfce7;color:#166534}}.pend-chip{{background:#fee2e2;color:#991b1b}}.docs-panel-b h2,.docs-panel-b p,.docs-panel-b label,.camera-card-b h2,.camera-card-b p{{color:#0f172a!important}}.doc-sign-card{{color:#0f172a!important}}.empty-docs{{background:#fff7ed;border:1px solid #fed7aa;color:#9a3412;border-radius:12px;padding:14px;font-weight:900}}.plantilla-sign-card{{border-color:#bbf7d0!important;background:#f8fffb!important}}.cam-wrap.cam-live .face-frame,.cam-wrap.cam-live .face-mesh{{display:block}}.cam-wrap.capture-ok #captureToast{{display:block}}.boceto-actions button:disabled{{opacity:.65;cursor:wait}}@media(max-width:1180px){{.firma-grid-boceto-main,.firma-progress-bar{{grid-template-columns:1fr}}.person-strip{{grid-template-columns:1fr 1fr}}}}@media(max-width:720px){{.firma-boceto-final{{padding:14px}}.firma-topbar,.steps{{display:grid}}.person-strip{{grid-template-columns:1fr}}.strip-item{{border-right:0;border-bottom:1px solid #e5eaf0}}.cam-wrap,.cam-wrap video,.cam-wrap img{{min-height:360px;height:360px}}}}
        

/* === PORTAL HR PRO - interfaz verde/blanca tipo referencia === */
:root{--txt:#0f172a!important;--mut:#64748b!important;--green:#15965a;--green2:#16a34a;--green3:#0f6f4f;--darkgreen:#0b2d2b;--soft:#eef8f3;--line:#dbe5ee;--panel:#ffffff;--panel2:#f8fafc;--shadow:0 18px 45px rgba(15,23,42,.10)}
body{background:#eef2f5!important;color:#0f172a!important;font-weight:700!important}.main{background:#eef2f5!important;color:#0f172a!important;padding:28px 30px!important}.card,.mini,.metric,.stat,.module-tile,.hero,.login-card{background:#fff!important;border:1px solid #dbe5ee!important;color:#0f172a!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;border-radius:24px!important}.hero{padding:24px!important}.hero h1,.topbar h1,.card h2,.module-tile h2{color:#0f172a!important;font-weight:1000!important}.muted,.subtitle,.muted2{color:#5f6b7a!important}.btn-green,.btn-blue,.btn{background:linear-gradient(135deg,#19aa63,#0f8f55)!important;color:#fff!important;border:0!important;border-radius:14px!important;font-weight:1000!important;box-shadow:0 12px 24px rgba(22,163,74,.20)!important}.btn-red{border:0!important;border-radius:14px!important}.input,select,textarea,input[type=file]{background:#fff!important;color:#111827!important;border:1px solid #d5e1ec!important;border-radius:14px!important}.input:focus,select:focus,textarea:focus{border-color:#19aa63!important;box-shadow:0 0 0 4px rgba(22,163,74,.14)!important}table{color:#243042!important}th{background:#f3f7fb!important;color:#334155!important}td{border-bottom:1px solid #e5edf4!important}.table-wrap{background:#fff!important;border-radius:18px!important;border:1px solid #e3ebf2!important}.flash{background:#dff7e9!important;color:#064e3b!important;border-color:#9adbb8!important}.flash.err{background:#fee2e2!important;color:#991b1b!important;border-color:#fecaca!important}
.side{background:linear-gradient(180deg,#124f34,#0e322d 42%,#091827)!important;border-right:1px solid rgba(255,255,255,.10)!important;box-shadow:16px 0 36px rgba(15,23,42,.22)!important}.side-top{height:86px!important;background:#124f34!important;border-bottom:1px solid rgba(255,255,255,.12)!important}.side-top b{color:#fff!important;font-size:18px!important}.brand{padding:22px 18px 18px!important;text-align:left!important}.brand img{display:none!important}.brand:before{content:'PORTAL HR PRO';display:block;color:#fff;font-size:22px;font-weight:1000;letter-spacing:.3px}.brand p{margin:4px 0 0!important;color:#d1fae5!important;font-size:13px!important;font-weight:900}.menu-title,.menu-item{background:transparent!important;border:0!important;color:#d7e1ea!important;box-shadow:none!important;border-radius:16px!important;margin:5px 8px!important;padding:15px 18px!important;font-weight:1000!important}.menu-title:hover,.menu-item:hover{background:rgba(255,255,255,.09)!important;color:#fff!important}.menu-item.active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:linear-gradient(135deg,#17aa55,#158a48)!important;color:#fff!important;box-shadow:0 14px 28px rgba(0,0,0,.20)!important}.submenu{padding:4px 0 8px!important}.side-user{background:rgba(255,255,255,.10)!important;color:#fff!important;border:1px solid rgba(255,255,255,.12)!important}.side-user small{color:#b7f7d1!important}.mobile-head{background:#124f34!important;color:#fff!important}.app{background:#eef2f5!important}.app.side-collapsed{grid-template-columns:92px 1fr!important}.side.collapsed .brand:before{content:'HR';text-align:center;font-size:22px}.side.collapsed .brand{padding:18px 8px!important;text-align:center!important}
.login-body{background:radial-gradient(circle at 9% 5%,rgba(34,197,94,.18) 0 19%,transparent 19.4%),radial-gradient(circle at 94% 0%,rgba(20,184,166,.18) 0 20%,transparent 20.4%),linear-gradient(180deg,#f8fafc 0%,#ecfdf5 100%)!important}.login-body:after{content:'';position:absolute;left:-4%;right:-4%;bottom:-6%;height:31%;background:linear-gradient(135deg,#22c55e,#064e3b);clip-path:polygon(0 38%,25% 22%,50% 15%,75% 25%,100% 35%,100% 100%,0 100%);z-index:0}.login-card{width:min(92vw,560px)!important;padding:88px 56px 34px!important;overflow:visible!important;text-align:left!important;background:rgba(255,255,255,.96)!important;position:relative!important;z-index:2!important}.login-card:before{content:'👥'!important;position:absolute!important;left:50%!important;top:-68px!important;transform:translateX(-50%)!important;width:136px!important;height:136px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe5ee!important;color:#149459!important;display:grid!important;place-items:center!important;font-size:52px!important;box-shadow:0 22px 55px rgba(15,23,42,.13)!important}.login-card:after{display:none!important}.login-logo{display:none!important}.login-title h1{color:#1f2937!important;font-size:42px!important;letter-spacing:1px!important}.login-title b{color:#64748b!important;font-size:17px!important}.login-title:after{content:'';display:block;margin:20px auto 0;width:62px;height:4px;border-radius:999px;background:#10b981}.login-input{background:#eaf2ff!important;border:1px solid #cfe0f2!important;border-radius:16px!important;padding:0 14px!important;margin-bottom:20px!important;color:#149459!important}.login-input input,.login-input select{background:transparent!important;color:#0f172a!important;border:0!important}.login-input input::placeholder{color:#64748b!important}.login-card .field label{display:block!important;color:#111827!important;margin:10px 0 8px;font-weight:1000}.login-card .btn-green{width:100%!important;margin:0 0 28px!important;font-size:20px!important;border-radius:16px!important;padding:17px!important}.login-links{margin-top:0!important;padding-bottom:0!important}.login-links a{color:#64748b!important}.contract-detail-wrap,.contract-detail-wrap *{color:#0f172a!important}



/* ===== CORRECCIÓN REAL FINAL: HEADER LIMPIO, LOGIN PORTAL HR PRO, DASHBOARD VISIBLE MÓVIL ===== */
.login-card{{
  width:min(90vw,500px)!important;
  max-width:500px!important;
  padding:122px 46px 30px!important;
}}
.login-logo{{top:-78px!important;z-index:30!important;}}
.login-avatar-svg{{width:122px!important;height:122px!important;}}
.login-avatar-svg:before{{width:68px!important;height:68px!important;}}
.login-title{{position:relative!important;z-index:20!important;margin:0 0 24px!important;}}
.login-title h1{{font-size:36px!important;line-height:1.05!important;color:#182233!important;text-transform:uppercase!important;white-space:nowrap!important;}}
.login-title b{{color:#667085!important;}}
.login-body:after{{content:"🛡️  Sistema seguro y confiable\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;white-space:pre!important;position:absolute!important;left:0!important;right:0!important;bottom:18px!important;text-align:center!important;color:#fff!important;font-size:16px!important;line-height:1.9!important;font-weight:650!important;z-index:1!important;background:none!important;height:auto!important;clip-path:none!important;}}

.side{{position:sticky!important;top:0!important;height:100vh!important;padding-top:0!important;background:#073126!important;overflow-y:auto!important;overflow-x:hidden!important;}}
.side::before{{content:""!important;position:sticky!important;top:0!important;display:block!important;height:112px!important;margin-bottom:-112px!important;background:#073126!important;z-index:2147483638!important;pointer-events:none!important;}}
.side-head-pro{{position:sticky!important;top:12px!important;z-index:2147483640!important;margin:12px 12px 18px!important;height:82px!important;min-height:82px!important;padding:0 12px!important;background:#123f33!important;border:1px solid rgba(255,255,255,.16)!important;border-radius:20px!important;box-shadow:0 16px 34px rgba(0,0,0,.34)!important;overflow:hidden!important;isolation:isolate!important;backdrop-filter:none!important;}}
.side-head-pro:before{{content:""!important;position:absolute!important;inset:0!important;background:linear-gradient(90deg,#123f33 0%,#196b48 52%,#123f33 100%)!important;opacity:1!important;z-index:-1!important;}}
.side-head-pro:after{{display:none!important;}}
.side-head-pro *{{position:relative!important;z-index:3!important;}}
.side-brand-text b{{color:#ffffff!important;font-size:16px!important;font-weight:800!important;}}
.side-brand-text small{{color:#d6e9e2!important;font-size:14px!important;font-weight:650!important;}}
nav{{position:relative!important;z-index:1!important;padding-top:0!important;}}

@media(max-width:1000px){{
  .side{{position:fixed!important;top:0!important;left:-340px!important;width:315px!important;height:100dvh!important;padding-top:0!important;background:#073126!important;}}
  .side.open{{left:0!important;}}
  .side-head-pro{{top:10px!important;margin:10px 10px 18px!important;}}
  .mobile-head{{background:#0f4a35!important;color:#fff!important;z-index:999999!important;height:56px!important;}}
  .menu-title,.menu-item{{color:#f1f7f5!important;opacity:1!important;text-shadow:none!important;font-weight:760!important;}}
  .menu-title .label,.menu-item .label{{color:inherit!important;opacity:1!important;visibility:visible!important;}}
  .menu-item.active,.menu-item.parent-active,.menu-title.active,.menu-group.force-open>.menu-title.active{{background:rgba(22,185,120,.30)!important;color:#ffffff!important;box-shadow:inset 4px 0 0 #34d399!important;border:1px solid rgba(134,239,172,.22)!important;}}
  .menu-item.active .label,.menu-item.parent-active .label,.menu-title.active .label,.menu-group.force-open>.menu-title.active .label{{color:#ffffff!important;}}
  .menu-title i,.menu-item i{{color:#d7eee5!important;background:rgba(255,255,255,.08)!important;}}
  .login-card{{padding:116px 26px 30px!important;}}
  .login-logo{{top:-70px!important;}}
  .login-avatar-svg{{width:116px!important;height:116px!important;}}
  .login-title h1{{font-size:31px!important;}}
}}

</style>
        <script>
        let firmaStream=null;
        let firmaCaptured=false;
        let firmaStarting=false;
        function firmaSetMsg(txt, ok=false, error=false){{
          const msg=document.getElementById('firmaCamMsg');
          if(msg){{ msg.innerHTML=txt; msg.style.color=error?'#b91c1c':(ok?'#059669':'#475569'); msg.style.fontWeight='900'; }}
        }}
        function firmaBadge(txt,bg){{ const badge=document.getElementById('liveBadge'); if(badge){{ badge.textContent=txt; badge.style.background=bg; }} }}
        function firmaBeep(){{ try{{ const A=window.AudioContext||window.webkitAudioContext; const ctx=new A(); const osc=ctx.createOscillator(); const gain=ctx.createGain(); osc.type='sine'; osc.frequency.value=880; gain.gain.setValueAtTime(0.001,ctx.currentTime); gain.gain.exponentialRampToValueAtTime(0.25,ctx.currentTime+0.02); gain.gain.exponentialRampToValueAtTime(0.001,ctx.currentTime+0.28); osc.connect(gain); gain.connect(ctx.destination); osc.start(); osc.stop(ctx.currentTime+0.30); }}catch(e){{}} }}
        function firmaResetVisual(){{ const wrap=document.querySelector('.cam-wrap'); const preview=document.getElementById('firmaPreview'); const sound=document.getElementById('soundBox'); if(preview){{ preview.removeAttribute('src'); preview.style.display='none'; }} if(sound) sound.style.display='none'; if(wrap){{ wrap.classList.remove('capture-ok','cam-live','cam-error'); }} firmaBadge('● APAGADA','#334155'); }}
        function firmaEsContextoSeguro(){{ return window.isSecureContext || location.protocol==='https:' || ['localhost','127.0.0.1','::1'].includes(location.hostname); }}
        async function firmaEsperarVideo(video, ms=12000){{ const inicio=Date.now(); while(Date.now()-inicio < ms){{ if(video && video.videoWidth && video.videoHeight && video.readyState>=2) return true; await new Promise(r=>setTimeout(r,150)); }} return false; }}
        async function firmaStartCam(ev){{
          if(ev) ev.preventDefault();
          if(firmaStarting) return false;
          firmaStarting=true;

          const wrap=document.querySelector('.cam-wrap');
          const btnActivar=document.getElementById('btnActivarCamara');
          const v=document.getElementById('firmaVideo');

          try{{
            if(btnActivar){{ btnActivar.disabled=true; btnActivar.innerHTML='⏳ Activando...'; }}
            firmaCaptured=false;
            firmaResetVisual();
            firmaBadge('● ACTIVANDO','#16a34a');
            firmaSetMsg('Activando cámara real... cuando el navegador pregunte, presiona <b>Permitir</b>.');

            if(!firmaEsContextoSeguro()) throw new Error('CONTEXTO_NO_SEGURO');
            if(!navigator.mediaDevices || !navigator.mediaDevices.getUserMedia) throw new Error('MEDIADEVICES_NO_DISPONIBLE');

            if(firmaStream){{ firmaStream.getTracks().forEach(t=>t.stop()); firmaStream=null; }}
            if(!v) throw new Error('VIDEO_NO_ENCONTRADO');

            v.pause();
            v.removeAttribute('src');
            v.removeAttribute('poster');
            v.srcObject=null;
            v.autoplay=true;
            v.muted=true;
            v.playsInline=true;
            v.setAttribute('playsinline','');
            v.setAttribute('webkit-playsinline','');
            v.style.display='block';
            v.style.background='#000';

            const intentos=[
              {{video:{{facingMode:{{ideal:'user'}},width:{{ideal:1280}},height:{{ideal:720}}}},audio:false}},
              {{video:{{facingMode:'user'}},audio:false}},
              {{video:{{width:{{ideal:960}},height:{{ideal:540}}}},audio:false}},
              {{video:true,audio:false}}
            ];

            let ultimoError=null;
            for(const cfg of intentos){{
              try{{
                firmaStream=await navigator.mediaDevices.getUserMedia(cfg);
                break;
              }}catch(e){{
                ultimoError=e;
              }}
            }}
            if(!firmaStream) throw (ultimoError || new Error('SIN_CAMARA'));

            v.srcObject=firmaStream;
            await new Promise(r=>setTimeout(r,250));
            try{{ await v.play(); }}catch(e){{ await v.play().catch(()=>{{}}); }}

            if(!await firmaEsperarVideo(v,15000)) throw new Error('VIDEO_NEGRO_O_SIN_IMAGEN');

            if(wrap){{ wrap.classList.add('cam-live'); wrap.classList.remove('cam-error'); }}
            firmaBadge('● EN VIVO','#16a34a');
            firmaSetMsg('✅ Cámara activada correctamente. Ya puedes presionar <b>Capturar evidencia</b>.',true);
            if(btnActivar) btnActivar.innerHTML='✅ Cámara activa';
            return true;
          }}catch(e){{
            if(firmaStream){{ firmaStream.getTracks().forEach(t=>t.stop()); firmaStream=null; }}
            if(v) v.srcObject=null;
            if(wrap) wrap.classList.add('cam-error');
            const nombre=(e&&e.name)?e.name:((e&&e.message)?e.message:'Error');
            let ayuda='';
            if(nombre==='CONTEXTO_NO_SEGURO') ayuda=' Abre el sistema en http://127.0.0.1:5000, localhost o en HTTPS. En celular con IP local HTTP el navegador bloquea la cámara.';
            else if(nombre==='NotAllowedError'||nombre==='PermissionDeniedError') ayuda=' Permite Cámara desde el candado del navegador y vuelve a presionar Activar cámara.';
            else if(nombre==='NotFoundError'||nombre==='DevicesNotFoundError') ayuda=' No se encontró cámara conectada.';
            else if(nombre==='NotReadableError'||nombre==='TrackStartError') ayuda=' La cámara está ocupada por otra app. Cierra Zoom/Meet/Teams/Cámara de Windows.';
            else if(nombre==='VIDEO_NEGRO_O_SIN_IMAGEN') ayuda=' El permiso fue aceptado, pero no llegó imagen. Cierra otras apps de cámara y vuelve a intentar.';
            else ayuda=' Revisa permisos de Windows/Android/iPhone y del navegador.';
            firmaBadge('● APAGADA','#dc2626');
            firmaSetMsg('❌ No se pudo activar cámara: '+nombre+'.'+ayuda+' Como respaldo puedes usar <b>Cámara/archivo</b>.',false,true);
            if(btnActivar) btnActivar.innerHTML='🎥 Activar cámara';
            return false;
          }}finally{{
            firmaStarting=false;
            if(btnActivar) btnActivar.disabled=false;
          }}
        }}
                async function firmaCapture(){{ const v=document.getElementById('firmaVideo'), c=document.getElementById('firmaCanvas'), img=document.getElementById('firmaPreview'), wrap=document.querySelector('.cam-wrap'), sound=document.getElementById('soundBox'); if(!v || !v.srcObject || !v.videoWidth){{ firmaSetMsg('Primero se activará la cámara. Acepta el permiso del navegador.',false,false); const ok=await firmaStartCam(); if(!ok || !v || !v.videoWidth){{ firmaSetMsg('No hay imagen de cámara. Presiona Activar cámara o usa Cámara/archivo como respaldo.',false,true); return false; }} }} c.width=v.videoWidth; c.height=v.videoHeight; c.getContext('2d').drawImage(v,0,0,c.width,c.height); const data=c.toDataURL('image/jpeg',0.88); img.src=data; img.style.display='block'; firmaCaptured=true; if(wrap){{ wrap.classList.add('capture-ok','cam-live'); wrap.classList.remove('cam-error'); }} if(sound) sound.style.display='flex'; firmaBeep(); firmaSetMsg('✅ Evidencia facial capturada correctamente. Lista para firmar los documentos seleccionados.',true); return true; }}
        function firmaLoadFileCam(input){{ const file=input && input.files ? input.files[0] : null; if(!file) return; const img=document.getElementById('firmaPreview'), wrap=document.querySelector('.cam-wrap'), sound=document.getElementById('soundBox'); const reader=new FileReader(); reader.onload=function(){{ if(img){{ img.src=reader.result; img.style.display='block'; }} firmaCaptured=true; if(wrap){{ wrap.classList.add('cam-live','capture-ok'); wrap.classList.remove('cam-error'); }} if(sound) sound.style.display='flex'; firmaBadge('● EVIDENCIA','#16a34a'); firmaSetMsg('✅ Evidencia cargada correctamente desde cámara/archivo.',true); firmaBeep(); }}; reader.readAsDataURL(file); }}
        function firmaStopCam(){{ if(firmaStream){{ firmaStream.getTracks().forEach(t=>t.stop()); firmaStream=null; }} const v=document.getElementById('firmaVideo'); if(v){{ v.pause(); v.srcObject=null; }} const wrap=document.querySelector('.cam-wrap'); if(wrap) wrap.classList.remove('cam-live'); firmaBadge('● DETENIDA','#334155'); firmaSetMsg('Cámara detenida.'); }}
        function updateFirmaCounter(){{ const checks=[...document.querySelectorAll('.doc-sign-list .chk-doc-firma:checked')]; const n=checks.length; const el=document.getElementById('firmaMassCounter'); if(el) el.textContent='Se firmarán '+n+' documentos'; const b=document.getElementById('docsBadge'); if(b) b.textContent=n; const first=checks[0]; if(first){{ const dni=document.getElementById('stripDni'), trab=document.getElementById('stripTrabajador'); if(dni) dni.textContent=first.dataset.dni||''; if(trab) trab.textContent=(first.dataset.trabajador||'').toUpperCase(); }} const f=document.getElementById('stripFecha'); if(f){{ const d=new Date(); f.textContent=d.toLocaleDateString('es-PE')+' '+d.toLocaleTimeString('es-PE'); }} }}
        function marcarTodosFirma(on){{ document.querySelectorAll('.chk-doc-firma').forEach(x=>x.checked=on); updateFirmaCounter(); }}
        function prepararFirmaMasiva(){{ const ids=[...new Set([...document.querySelectorAll('.doc-sign-list .chk-doc-firma:checked')].map(x=>x.value))]; if(ids.length===0){{ alert('Selecciona al menos un contrato para firmar.'); return false; }} if(!firmaCaptured){{ const continuar=confirm('Aún no se capturó evidencia facial. ¿Deseas continuar igual?'); if(!continuar) return false; }} const lote=document.getElementById('documentos_lote'); if(lote) lote.value=ids.join(','); return confirm('Se firmarán/generarán '+ids.length+' documento(s). ¿Continuar?'); }}
        window.firmaStartCam=firmaStartCam; window.firmaCapture=firmaCapture; window.firmaStopCam=firmaStopCam; window.firmaLoadFileCam=firmaLoadFileCam; window.marcarTodosFirma=marcarTodosFirma; window.prepararFirmaMasiva=prepararFirmaMasiva;
        document.addEventListener('change',e=>{{ if(e.target.classList && e.target.classList.contains('chk-doc-firma')) updateFirmaCounter(); }});
        document.addEventListener('DOMContentLoaded',()=>{{ updateFirmaCounter(); firmaResetVisual(); firmaSetMsg('Presiona <b>🎥 Activar cámara</b>. Cuando el navegador pregunte, elige <b>Permitir</b>. Debe abrirse en localhost/127.0.0.1 o HTTPS.'); }});
        </script>
        """)
    elif sec=='nisira':
        content=wrap("<h2 class='c-title'>Contratación NISIRA</h2><div class='c-card' style='padding:22px'><p class='muted2'>Sección preparada para importar contratos / altas desde NISIRA y cruzar por DNI.</p><button class='c-btn'>Sincronizar NISIRA</button></div>")
    elif sec=='descargas':
        content=wrap(f"<h2 class='c-title'>Descargas</h2><div class='c-card table-wrap'><table class='c-table'><tr><th></th><th>Código</th><th>Apellidos y Nombres</th><th>Tipo Documento</th><th>Estado Doc</th><th>Fecha Envío</th></tr>{docs_rows or '<tr><td colspan=6>No hay archivos.</td></tr>'}</table></div>")
    else:
        content=wrap(f"<h2 class='c-title'>Archivos Trabajador</h2><div class='toolbar'>🔎 Filtros &nbsp; ⚙ Acción ▾ &nbsp; ⬇ Descargar ▾</div><form method='post' enctype='multipart/form-data' class='c-card c-form' style='padding:18px'><b>Trabajador</b><input name='dni' list='trabajadores_list' required><datalist id='trabajadores_list'>{opt_trab}</datalist><b>Etapa</b><select name='etapa'><option>Incorporación</option><option>Renovación</option><option>Cese</option></select><b>Tipo documento</b><select name='tipo_doc'>{opt_tipo}</select><b>Archivo</b><input type='file' name='archivo' required><span></span><button class='c-btn'>⬆ Subir Docs Individual</button></form><div class='c-card table-wrap'><table class='c-table'><tr><th></th><th></th><th>Código</th><th>Apellidos y Nombres</th><th>Tipo Documento</th><th>Estado Doc</th><th>Fecha Envío</th></tr>{docs_rows or '<tr><td colspan=7>No hay archivos.</td></tr>'}</table></div>")
    return render_page(content, active=f'Gestion Contratacion:{sec}')


@app.route('/admin/crear_carpetas')
@admin_required
def admin_crear_carpetas():
    asegurar_carpetas_documentales()
    total = sincronizar_documentos_carpeta()
    flash(f'Carpeta local creada/actualizada: {DOCUMENTOS_BASE_DIR}. Coloca allí los PDFs y presiona Sincronizar. Documentos detectados automáticamente: {total}.', 'ok')
    return redirect(request.referrer or url_for('admin'))

@app.route('/admin/sincronizar')
@admin_required
def admin_sincronizar():
    resumen = sincronizar_documentos_carpeta(devolver_resumen=True)
    flash(f"Sincronización completada. Nuevos: {resumen['nuevos']} | Revisados: {resumen['revisados']} | Duplicados: {resumen['duplicados']} | Sin DNI: {resumen['sin_dni']} | Sin trabajador activo: {resumen['sin_trabajador']} | Ruta: {DOCUMENTOS_BASE_DIR}", 'ok')
    return redirect(url_for('admin_documentos'))


@app.route('/admin/modo_prueba')
@admin_required
def admin_modo_prueba():
    modo_txt = 'ACTIVO' if modo_prueba_activo() else 'INACTIVO'
    content = f"""
    <div class='hero'><div class='topbar'><div><h1>Modo prueba <span class='accent'>y limpieza general</span></h1><div class='subtitle'>Esta opción aplica a todas las gestiones. Actívala solo cuando vayas a probar con usuarios y administrador.</div></div></div></div>
    <section class='grid'>
      <div class='card span-12'><h2>🧪 Control general de pruebas</h2><p class='muted'>Todo lo cargado quedará marcado como [MODO PRUEBA]. Luego puedes limpiarlo sin tocar la información real.</p><div class='actions'><a class='btn-green' href='/admin/modo_prueba/toggle'>Modo prueba: {modo_txt}</a><a class='btn-danger' onclick='return confirm("¿Borrar documentos y eventos de MODO PRUEBA?")' href='/admin/modo_prueba/limpiar'>Limpiar pruebas</a><a class='btn-blue' href='/admin/desbloquear_usuarios'>Desbloquear usuarios</a></div></div>
    </section>"""
    return render_page(content, active='Modo prueba')

@app.route('/admin/modo_prueba/toggle')
@admin_required
def admin_modo_prueba_toggle():
    set_config('modo_prueba', '0' if modo_prueba_activo() else '1')
    flash('Modo prueba actualizado.', 'ok')
    return redirect(url_for('admin'))

@app.route('/admin/modo_prueba/limpiar')
@admin_required
def admin_modo_prueba_limpiar():
    borrados = 0
    with db() as con:
        rows = con.execute("SELECT id,ruta_archivo FROM documentos WHERE uploaded_by LIKE '%MODO PRUEBA%'").fetchall()
        ids = [r['id'] for r in rows]
        for r in rows:
            try:
                p = Path(r['ruta_archivo'])
                if p.exists() and str(p).startswith(str(UPLOAD_DIR)):
                    p.unlink()
            except Exception:
                pass
        if ids:
            q = ','.join(['?']*len(ids))
            con.execute(f'DELETE FROM eventos_documento WHERE documento_id IN ({q})', ids)
            con.execute(f'DELETE FROM documentos WHERE id IN ({q})', ids)
            borrados = len(ids)
        con.commit()
    flash(f'Modo prueba limpiado. Documentos de prueba borrados: {borrados}.', 'ok')
    return redirect(url_for('admin'))

@app.route('/admin/desbloquear_usuarios')
@admin_required
def admin_desbloquear_usuarios():
    with db() as con:
        con.execute('DELETE FROM login_intentos')
        con.commit()
    flash('Usuarios desbloqueados. Los intentos fallidos fueron reiniciados.', 'ok')
    return redirect(url_for('admin'))


@app.route('/admin/contratacion/doc/<int:doc_id>/archivo')
@admin_required
def contratacion_doc_archivo(doc_id):
    with db() as con:
        r = con.execute('SELECT * FROM contratacion_docs WHERE id=?', (doc_id,)).fetchone()
    if not r or not r['ruta_archivo']:
        abort(404)
    path = Path(r['ruta_archivo'])
    if not path.exists():
        abort(404)
    return send_file(path, as_attachment=False, download_name=r['archivo_nombre'] or 'contrato')

@app.route('/admin/contratacion/doc/<int:doc_id>/log')
@admin_required
def contratacion_doc_log(doc_id):
    with db() as con:
        r = con.execute('SELECT * FROM contratacion_docs WHERE id=?', (doc_id,)).fetchone()
    if not r:
        abort(404)
    eventos = [
        ('Fecha Registro', fecha_sin_hora(r['fecha_registro']), r['uploaded_by'] or 'SISTEMA', 'Documento registrado en Gestión Contratación'),
        ('Firmado', fecha_sin_hora(r['fecha_registro']), r['uploaded_by'] or 'SISTEMA', 'Marcado como firmado / reconocido facialmente o cargado por RR.HH.' if str(r['estado']).upper() in ['FIRMADO','ARCHIVADO','GENERADO'] else 'Pendiente de firma'),
        ('Archivado', fecha_sin_hora(r['fecha_registro']), r['uploaded_by'] or 'SISTEMA', 'Archivo almacenado correctamente' if r['ruta_archivo'] else 'Sin archivo adjunto'),
    ]
    rows = ''.join([f"<tr><td>{html.escape(a)}</td><td>{html.escape(b or '-')}</td><td>{html.escape(c or '-')}</td><td>{html.escape(d or '')}</td></tr>" for a,b,c,d in eventos])
    content = f"""
    <div class='c-card' style='padding:24px;max-width:1100px;margin:auto'>
      <h1 class='c-title'>Log de Contrato</h1>
      <p><b>DNI:</b> {html.escape(r['dni'] or '')} &nbsp; <b>Trabajador:</b> {html.escape(r['trabajador'] or '')} &nbsp; <b>Tipo:</b> {html.escape(r['tipo_doc'] or '')}</p>
      <div class='table-wrap'><table class='c-table'><tr><th>Evento</th><th>Fecha Registro</th><th>Creado por</th><th>Comentario</th></tr>{rows}</table></div>
      <br><a class='c-btn gray' href='{url_for('admin_contratacion', sec='ficha', dni=r['dni'])}'>Cerrar / volver a ficha</a>
    </div>"""
    return render_page(content, active='Gestion Contratacion:ficha')

# API compatibles
@app.route('/api/health')
def api_health(): return jsonify({'ok': True, 'mensaje': 'PORTAL HR PRO activo - optimizado Render Free'})
@app.route('/api/boleta/<dni>')
def api_boleta(dni):
    docs = listar_documentos(dni=dni, categoria='pago', limit=20)
    t = get_trabajador(dni)
    return jsonify({'ok': bool(t), 'trabajador': dict(t) if t else None, 'documentos': [dict(x) for x in docs]})


@app.route('/admin/firma_digital')
@admin_required
def admin_firma_digital():
    return redirect(url_for('admin_contratacion', sec='firma'))



# ===== PARCHE FINAL REAL: evita que el icono tape el título y corrige header/menu móvil =====
@app.after_request
def parche_final_login_header(response):
    try:
        ct = response.headers.get('Content-Type','')
        if 'text/html' not in ct.lower():
            return response
        html_text = response.get_data(as_text=True)
        css = '\n<style id="fix-final-real-login-header">\n.login-body{min-height:100vh!important;display:flex!important;align-items:center!important;justify-content:center!important;padding:58px 18px 118px!important;overflow:hidden!important;background:radial-gradient(circle at 8% -4%,rgba(34,197,94,.17) 0 20%,transparent 20.4%),radial-gradient(circle at 94% -2%,rgba(45,212,191,.17) 0 22%,transparent 22.4%),linear-gradient(180deg,#fbfdff 0%,#effaf5 100%)!important;}\n.login-body:before{content:""!important;position:absolute!important;left:0!important;right:0!important;bottom:0!important;height:30%!important;background:linear-gradient(135deg,#22c55e 0%,#065f46 100%)!important;clip-path:polygon(0 36%,22% 26%,48% 17%,73% 28%,100% 22%,100% 100%,0 100%)!important;z-index:0!important;}\n.login-body:after{content:"🛡️  Sistema seguro y confiable\\A © 2025 Portal HR Pro. Todos los derechos reservados."!important;white-space:pre!important;position:absolute!important;left:0!important;right:0!important;bottom:18px!important;text-align:center!important;color:#fff!important;font-size:16px!important;line-height:1.9!important;font-weight:650!important;z-index:1!important;background:none!important;height:auto!important;clip-path:none!important;}\n.login-card{width:min(90vw,500px)!important;max-width:500px!important;margin:0 auto!important;padding:52px 46px 30px!important;border-radius:28px!important;background:rgba(255,255,255,.975)!important;border:1px solid #dce8ef!important;box-shadow:0 18px 48px rgba(15,23,42,.13)!important;overflow:visible!important;position:relative!important;z-index:2!important;text-align:left!important;}\n.login-card:before,.login-card:after{display:none!important;content:none!important;}\n.login-inner{position:relative!important;z-index:3!important;}\n.login-logo{display:flex!important;position:relative!important;top:auto!important;left:auto!important;right:auto!important;bottom:auto!important;transform:none!important;margin:-94px auto 26px!important;width:124px!important;height:124px!important;z-index:6!important;justify-content:center!important;align-items:center!important;text-align:center!important;}\n.login-logo img{display:none!important;}\n.login-avatar-svg{width:124px!important;height:124px!important;border-radius:50%!important;background:#fff!important;border:1px solid #dbe7ef!important;box-shadow:0 14px 36px rgba(15,23,42,.10)!important;display:grid!important;place-items:center!important;}\n.login-avatar-svg:before{content:""!important;width:68px!important;height:68px!important;display:block!important;background:no-repeat center/contain url("data:image/svg+xml,%3Csvg xmlns=\'http://www.w3.org/2000/svg\' viewBox=\'0 0 96 96\'%3E%3Cdefs%3E%3ClinearGradient id=\'g\' x1=\'0\' y1=\'0\' x2=\'1\' y2=\'1\'%3E%3Cstop stop-color=\'%2327c77b\'/%3E%3Cstop offset=\'1\' stop-color=\'%230b8f5a\'/%3E%3C/linearGradient%3E%3C/defs%3E%3Ccircle cx=\'35\' cy=\'30\' r=\'13\' fill=\'url(%23g)\'/%3E%3Ccircle cx=\'61\' cy=\'28\' r=\'15\' fill=\'url(%23g)\'/%3E%3Cpath d=\'M13 72c1-17 13-28 29-28 5 0 9 1 13 3-7 5-12 13-13 25H13z\' fill=\'url(%23g)\'/%3E%3Cpath d=\'M38 72c1-19 14-31 32-31 15 0 25 12 25 31H38z\' fill=\'url(%23g)\'/%3E%3C/svg%3E")!important;}\n.login-title{text-align:center!important;margin:0 0 24px!important;position:relative!important;z-index:5!important;}\n.login-title h1{margin:0 0 10px!important;font-size:36px!important;line-height:1.05!important;font-weight:800!important;letter-spacing:-.035em!important;color:#182233!important;text-transform:uppercase!important;white-space:nowrap!important;position:relative!important;z-index:5!important;}\n.login-title b{display:block!important;color:#667085!important;font-size:16px!important;font-weight:520!important;line-height:1.32!important;}\n.login-title:after{content:""!important;display:block!important;width:60px!important;height:4px!important;border-radius:999px!important;background:#16b978!important;margin:18px auto 0!important;}\n.login-input input,.login-input select{color:#172033!important;-webkit-text-fill-color:#172033!important;background:#fff!important;}\n.login-input input::placeholder{color:#6b7787!important;opacity:1!important;-webkit-text-fill-color:#6b7787!important;}\n.side{background:#073126!important;overflow-y:auto!important;overflow-x:hidden!important;}\n.side-head-pro{position:sticky!important;top:0!important;z-index:2147483640!important;margin:0 12px 16px!important;height:82px!important;min-height:82px!important;padding:0 12px!important;display:flex!important;align-items:center!important;justify-content:space-between!important;background:#123f33!important;border:1px solid rgba(255,255,255,.16)!important;border-radius:0 0 20px 20px!important;box-shadow:0 16px 34px rgba(0,0,0,.34)!important;overflow:hidden!important;isolation:isolate!important;backdrop-filter:none!important;}\n.side-head-pro:before{content:""!important;position:absolute!important;left:-16px!important;right:-16px!important;top:-26px!important;bottom:0!important;background:#073126!important;z-index:-3!important;opacity:1!important;}\n.side-head-pro:after{content:""!important;position:absolute!important;inset:0!important;background:linear-gradient(90deg,#123f33 0%,#196b48 52%,#123f33 100%)!important;z-index:-2!important;border-radius:0 0 20px 20px!important;opacity:1!important;}\n.side-head-pro *{position:relative!important;z-index:3!important;}\n.side-brand-text b{color:#fff!important;font-size:16px!important;font-weight:800!important;white-space:nowrap!important;}\n.side-brand-text small{color:#d6e9e2!important;font-size:14px!important;font-weight:650!important;display:block!important;}\n.menu-title,.menu-item{color:#eaf7f1!important;opacity:1!important;text-shadow:none!important;font-weight:760!important;}\n.menu-title .label,.menu-item .label{color:inherit!important;opacity:1!important;visibility:visible!important;display:inline!important;}\n.menu-item.active,.menu-item.parent-active,.menu-title.active,.menu-group.force-open>.menu-title.active{background:rgba(22,185,120,.30)!important;color:#fff!important;box-shadow:inset 4px 0 0 #34d399!important;border:1px solid rgba(134,239,172,.22)!important;}\n.menu-item.active .label,.menu-item.parent-active .label,.menu-title.active .label,.menu-group.force-open>.menu-title.active .label{color:#fff!important;}\n@media(max-width:1000px){.side{position:fixed!important;top:0!important;left:-340px!important;width:315px!important;height:100dvh!important;background:#073126!important;z-index:2147483644!important;}.side.open{left:0!important;}.side-head-pro{top:0!important;margin:0 10px 16px!important;border-radius:0 0 18px 18px!important;}.mobile-head{background:#0f4a35!important;color:#fff!important;z-index:2147483645!important;}.mobile-head b,.mobile-head a{color:#fff!important;opacity:1!important;visibility:visible!important;}.menu-title,.menu-item{font-size:13.5px!important;color:#f8fffc!important;font-weight:760!important;}.menu-title .label,.menu-item .label{color:#f8fffc!important;}.login-body{padding:82px 12px 108px!important;align-items:center!important;overflow-y:auto!important;}.login-card{width:calc(100vw - 26px)!important;max-width:420px!important;padding:50px 26px 30px!important;border-radius:26px!important;}.login-logo{width:116px!important;height:116px!important;margin:-92px auto 24px!important;}.login-avatar-svg{width:116px!important;height:116px!important;}.login-avatar-svg:before{width:64px!important;height:64px!important;}.login-title h1{font-size:31px!important;}.login-title b{font-size:17px!important;max-width:310px!important;margin:auto!important;}.login-body:after{font-size:12px!important;bottom:14px!important;}}\n</style>\n'
        html_text = html_text.replace('</head>', css + '</head>')
        response.set_data(html_text)
        response.headers['Content-Length'] = str(len(response.get_data()))
    except Exception:
        pass
    return response



# ===== PARCHE FINAL PEDIDO: PORTAL HR + CELULAR COMPACTO + SIDEBAR SIN TAPA DASHBOARD =====
@app.after_request
def parche_portal_hr_movil_compacto(response):
    try:
        ct = response.headers.get('Content-Type','')
        if 'text/html' not in ct.lower():
            return response
        html_text = response.get_data(as_text=True)
        html_text = html_text.replace('PRIZE PRO', 'PORTAL HR')
        html_text = html_text.replace('Prize Pro', 'Portal HR')
        html_text = html_text.replace('PRIZE', 'PORTAL HR')
        css = '\n<style id="portal-hr-final-mobile-fix">\n.login-title h1{font-size:36px!important;line-height:1.05!important;font-weight:800!important;color:#182233!important;letter-spacing:-.035em!important;white-space:nowrap!important;text-transform:uppercase!important;}\n.login-logo{z-index:10!important;}.login-card{overflow:visible!important;}.login-body:after{content:"🛡️  Sistema seguro y confiable\\\\A © 2025 Portal HR. Todos los derechos reservados."!important;white-space:pre!important;}\n.login-input input,.login-input select{background:#fff!important;color:#172033!important;-webkit-text-fill-color:#172033!important;caret-color:#172033!important;opacity:1!important;}\n.login-input input::placeholder{color:#6b7787!important;-webkit-text-fill-color:#6b7787!important;opacity:1!important;}\n@media(max-width:700px){\n  .login-body{padding:76px 12px 104px!important;align-items:center!important;justify-content:center!important;overflow-y:auto!important;}\n  .login-card{width:calc(100vw - 34px)!important;max-width:390px!important;padding:78px 24px 28px!important;border-radius:26px!important;}\n  .login-logo{width:112px!important;height:112px!important;margin:-88px auto 22px!important;top:auto!important;position:relative!important;left:auto!important;transform:none!important;display:flex!important;justify-content:center!important;}\n  .login-avatar-svg{width:112px!important;height:112px!important;}.login-avatar-svg:before{width:62px!important;height:62px!important;}\n  .login-title{margin:0 0 22px!important;}.login-title h1{font-size:31px!important;}.login-title b{font-size:16px!important;line-height:1.28!important;max-width:300px!important;margin:auto!important;}\n  .login-title:after{margin-top:16px!important;width:58px!important;}.login-card .field{margin-bottom:12px!important;}.login-card .field label{font-size:14px!important;margin-bottom:6px!important;}\n  .login-input{height:54px!important;border-radius:16px!important;}.login-input i{width:30px!important;height:30px!important;font-size:15px!important;}\n  .login-input input,.login-input select{height:42px!important;min-height:42px!important;font-size:16px!important;border-radius:14px!important;}\n  .login-card .btn-green{height:56px!important;font-size:19px!important;border-radius:16px!important;margin:15px 0 18px!important;}.login-links a{font-size:14px!important;}\n  .login-body:after{font-size:12px!important;bottom:12px!important;line-height:1.7!important;}\n}\n@media(max-width:1000px){\n  .mobile-head{height:56px!important;min-height:56px!important;position:sticky!important;top:0!important;z-index:2147483647!important;background:#0f4a35!important;color:#fff!important;padding:0 14px!important;display:flex!important;align-items:center!important;}\n  .mobile-head b{font-size:15px!important;font-weight:800!important;letter-spacing:.02em!important;color:#fff!important;}.mobile-head a{font-size:15px!important;font-weight:800!important;color:#fff!important;text-decoration:none!important;}\n  .mobile-head .toggle{width:42px!important;height:42px!important;border-radius:12px!important;background:rgba(255,255,255,.12)!important;}\n  .side{top:56px!important;height:calc(100dvh - 56px)!important;width:286px!important;left:-306px!important;background:#073126!important;padding-top:8px!important;z-index:2147483644!important;}\n  .side.open{left:0!important;}.side-head-pro{display:none!important;}nav{padding:8px 14px 18px!important;}.menu-group{margin:7px 0!important;}\n  .menu-title{min-height:52px!important;padding:11px 13px!important;border-radius:16px!important;font-size:14px!important;line-height:1.22!important;}\n  .menu-item{min-height:44px!important;padding:10px 12px!important;border-radius:14px!important;font-size:13.5px!important;line-height:1.22!important;}\n  .menu-title i,.menu-item i{width:28px!important;height:28px!important;font-size:16px!important;}.menu-title .label,.menu-item .label{color:#f8fffc!important;opacity:1!important;visibility:visible!important;font-weight:760!important;}\n  .submenu{padding:6px 0!important;}.menu-group.nested{margin-left:10px!important;padding-left:10px!important;}.side-user{margin:14px 14px 18px!important;padding:12px!important;border-radius:16px!important;}\n  .main{padding-top:14px!important;}.hero,.topbar,.card{scroll-margin-top:70px!important;}\n}\n</style>\n'
        html_text = html_text.replace('</head>', css + '</head>')
        response.set_data(html_text)
        response.headers['Content-Length'] = str(len(response.get_data()))
    except Exception:
        pass
    return response

if __name__ == '__main__':
    port = int(os.getenv('PORT', '5000'))
    host = os.getenv('HOST', '0.0.0.0')
    debug = os.getenv('FLASK_DEBUG', '0') == '1'
    ssl_context = None
    # Para probar desde CELULAR en red local, el navegador exige HTTPS.
    # Ejecuta en Windows: set APP_SSL=1 && python app.py
    # Luego abre en el celular: https://IP-DE-TU-PC:5000/firma/<token>
    if os.getenv('APP_SSL', '0') == '1':
        cert_file = os.getenv('SSL_CERT', 'cert.pem')
        key_file = os.getenv('SSL_KEY', 'key.pem')
        if Path(cert_file).exists() and Path(key_file).exists():
            ssl_context = (cert_file, key_file)
        else:
            ssl_context = 'adhoc'
    try:
        print('============================================================')
        print('PORTAL HR PRO iniciado')
        print('PC local:  http://127.0.0.1:%s' % port)
        if ssl_context:
            print('Celular:   https://IP-DE-TU-PC:%s  (aceptar certificado)' % port)
        else:
            print('Celular:   usar enlace HTTPS de Render o iniciar con APP_SSL=1')
        print('Nota cámara:', contexto_camara_seguro_texto())
        print('============================================================')
        app.run(host=host, port=port, debug=debug, ssl_context=ssl_context)
    except Exception as e:
        print('No se pudo iniciar con SSL:', e)
        print('Reintentando en HTTP solo para PC localhost...')
        app.run(host=host, port=port, debug=debug)
